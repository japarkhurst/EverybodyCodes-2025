# -*- coding: utf-8 -*-
"""
Created on Wed Aug 11 11:12:28 2021

@author: japarkhu
"""

# -*- coding: utf-8 -*-
"""
Created on Tue Apr 6 10:47:50 2021

@author: japarkhu

Change Log
4/19/22:
All Tools:
    -Limited layer list dropdown to pertitent subset of map layers
Submit Resource:
    -Validated Resource Name for length and invalid characters
Append State Resources:
    -Removed Points and Lines from dropdown
Exporting Reports:
    -Removed email address blank if none is provided.
    -Converted Project Lead, Project Reviewer, Field Director, Principal Investigator, Task Lead, TVA Monitor, Staff Reviewer names from domain description (last name first) to report format (first and last name)
    -Entered "N/A" for fields that have no value specified
    -Document name for Activities changed to Activity ID + timestamp (removed project name)
    -Changed Activity Name to be free text; added check to ensure the value is unique

"""

'''
References:
Value Table Object
https://pro.arcgis.com/en/pro-app/latest/arcpy/classes/valuetable.htm
Arcade Conditional Expressions
https://community.esri.com/t5/arcgis-online-documents/conditional-field-display-with-arcade-in-pop-ups-revisited/ta-p/920869
Value Table with Drop Downs
https://community.esri.com/t5/python-documents/creating-a-valuetable-with-columns-that-have-drop-down-choices/ta-p/907702
Python Toolboxes:
https://pro.arcgis.com/en/pro-app/latest/arcpy/geoprocessing_and_python/defining-parameter-data-types-in-a-python-toolbox.htm
ArcGIS API for Python:
Feature Layers and Features: https://developers.arcgis.com/python/guide/working-with-feature-layers-and-features

Variables in ALL_CAPITALS are coming from the Constants.py file.
'''

'''Dynamic Parameters'''
# TODO - add to CONSTANTS file.
# FEATURE_LAYER_NAME_RESOURCE_TYPE_DICT = {'archaeol_survey_activity_poly':'Archaeological_Survey',
# 'archaeology_resource_poly':'Archaeological_Resources',
# 'cemetery_polygon':'Cemeteries',
# 'cultural_landscape_poly':'Cultural_Landscapes',
# 'district_facilities_poly':'Districts_and_Facilities',
# 'har_poly':'Historic_Architectural_Resources',
# 'object_poly':'Objects',
# 'other_resources_poly':'Other'}


''' Define Tools Visible in Toolbox '''
### Create Toolbox and Define Tools within the Toolbox
class Toolbox(object):
    def __init__(self):
        """Define the toolbox (the name of the toolbox is the name of the
        .pyt file)."""
        self.label = "Managed Task CRMS Toolbox"
        self.alias = "Managed Task CRMS Toolbox"

        # List of tool classes associated with this toolbox
        toolList = []
        toolList.append(exportResourcesForManagedTask)
        toolList.append(analyzeMTDifferences)
        # # # toolList.append(copyResourcesIntoPreviousBoundaries)
        # # # toolList.append(placeholderExportReport)
        # # # toolList.append(placeholderManuallyUpdateResources)
        # # # toolList.append(appendNewMTtoGIS)
        # # # toolList.append(populateMTGISGlobalID)
        # # # toolList.append(appendMTtoAGOL)
        # toolList.append(testTool)
        toolList.append(UpdateAGOLLists)
        self.tools = toolList

''' Import Modules '''
import arcpy
import ast
from collections import namedtuple
from copy import deepcopy
import os
import datetime
import glob
import importlib
import json
import pandas
import shutil
import sys
import traceback

def messageTimestamp():
    import datetime
    return str('{:%Y%m%d %H:%M:%S}'.format(datetime.datetime.utcnow()))

printLevel = 'verbose'
messageLog = ''
def msg(text,level=''):
    global messageLog
    timestamp_time = '{:%Y%m%d %H:%M:%S}'.format(datetime.datetime.utcnow())
    text = '{}\t{}'.format(timestamp_time,text)
    arcpy.AddMessage(text)
    print(text)
    messageLog = messageLog + '\n' + str(text)

def plain_msg(text,level=''):
    global messageLog
    arcpy.AddMessage(text)
    print(text)
    messageLog = messageLog + '\n' + str(text)

def vmsg(text,printLevel):
    global messageLog
    timestamp_time = '{:%Y%m%d %H:%M:%S}'.format(datetime.datetime.utcnow())
    text = '{}\t{}'.format(timestamp_time,text)
    if printLevel == 'verbose':
        arcpy.AddMessage(text)
        print(text)
    messageLog = messageLog + '\n' + text

def wrn(text):
    global messageLog
    timestamp_time = '{:%Y%m%d %H:%M:%S}'.format(datetime.datetime.utcnow())
    text = "{}\tWarning: {}".format(timestamp_time,text)
    arcpy.AddWarning(text)
    print(text)
    messageLog = messageLog + '\n' + text

def err(text):
    global messageLog
    timestamp_time = '{:%Y%m%d %H:%M:%S}'.format(datetime.datetime.utcnow())
    text = "{}\tWarning: {}".format(timestamp_time,text)
    arcpy.AddError('> ')
    arcpy.AddError('>>')
    arcpy.AddError(">>>Error: {}<<<".format(text))
    arcpy.AddError('>>')
    arcpy.AddError('>')
    print("\n\nError: {}\n\n".format(text))
    messageLog = messageLog + '\n' + text

exceptionsTxt = ''
def printException():
    """
    requires global text string variable titled exceptionsTxt
    requires import traceback, sys
    exceptionTxt = ""
    """
    try:
        # Get the traceback object
        tb = sys.exc_info()[2]
        tbinfo = traceback.format_tb(tb)[0]
        # Concatenate information together concerning the error into a message string
        pymsg = "Traceback info: " + tbinfo + "Error Info: " + str(sys.exc_info()[1])
        # msgs = "ArcPy ERRORS:\n" + arcpy.GetMessages(2) + "\n"
        wrn("!!!{}".format(pymsg))
        global exceptionsTxt
        exceptionsTxt = exceptionsTxt + ";" + pymsg
    except Exception as e:
        msg("Error with printException function: {}".format(e))

def createFolder(folderPath):
    """function to wrap around folder locations so that they are generated if they don't exist"""
    if not os.path.exists(folderPath):
        os.makedirs(folderPath)
    return folderPath

def createFGDB(path):
    """function to wrap around FGDB locations so that they are generated if they don't exist"""
    import os
    if not arcpy.Exists(path):
        arcpy.CreateFileGDB_management(os.path.dirname(path),os.path.basename(path))
    return path

# def checkIfFileIsOpen(path):
#     try:
#         tempPath =
#         os.rename(path, 'tempfile.xls')
#         os.rename('tempfile.xls', 'file.xls')
#     except OSError:
#         wrn('File is open: {}'.format(path))

def returnFileList(folderPath,wildcard=''):
    """ This function returns the list of files in a specified folder
    wildcard should include astrisk if desired; e.g. *.pdf """
    if wildcard:
        fileList = glob.glob(r'{}\{}'.format(folderPath,wildcard))
    else:
        fileList = glob.glob(r'{}\*'.format(folderPath))
    print('Files in Folder: {}'.format(fileList))
    return fileList

'''Environment Parameters'''
arcpy.env.parallelProcessingFactor = "100%" #environment variable for using 100% of processor when running
arcpy.env.overwriteOutput = True # environment variable to overwrite output dataset
arcpy.env.preserveGlobalIds = True # environment variable to preserve GlobalIDs

timestamp_date = str('{:%Y%m%d}'.format(datetime.datetime.utcnow()))
day = str('{:%m/%d/%Y}'.format(datetime.datetime.utcnow()))
timestamp_time = str('{:%Y%m%d_%H%M%S}'.format(datetime.datetime.utcnow()))
reportDateText = str('{:%B %#d, %Y}'.format(datetime.datetime.now()))
startingDate = '2021-08-01' # starting date for epoch time difference for record id creation

''' Reference Links
https://pro.arcgis.com/en/pro-app/arcpy/geoprocessing_and_python/defining-parameters-in-a-python-toolbox.htm
https://pro.arcgis.com/en/pro-app/arcpy/mapping/map-class.htm
https://pro.arcgis.com/en/pro-app/arcpy/mapping/arcgisproject-class.htm
https://pro.arcgis.com/en/pro-app/help/mapping/text/text-formatting-tags.htm
https://pro.arcgis.com/en/pro-app/tool-reference/conversion/layer-to-kml.htm
All KML and KMZ files are created in the WGS84 coordinate system. You need to be sure your layers will properly project if they aren't already in WGS84. You can use the Project tool to reproject your data prior to KML conversion if your projection requires a transformation.
'''

''' Global Parameters Across All Tools in the Python Toolbox '''
### Define Project Folders
aprx = arcpy.mp.ArcGISProject("CURRENT")
aprxFilePath = aprx.filePath
aprxFolder = os.path.dirname(aprxFilePath)
localProjectFolder = os.path.dirname(aprxFilePath) # set local project folder as parent of APRX location
localProjectTempFolder = createFolder(os.path.join(localProjectFolder,'temp'))

''' Folder where external python modules are located '''
#pythonModulesFolder = r'\\tva\egis\EGIS-Projects\Environment\Cultural\Tools\APRX\Modules' # TODO - update when deploying
pythonModulesFolder = os.path.join(localProjectFolder,'Tools','Modules')
sys.path.append(pythonModulesFolder)
# from docx import Document
# from docx.shared import Pt

''' Import Modules from other TVA Python Files '''
#constantsPath = os.path.join(localProjectFolder,'Tools','Config')
try:
    pytFilePath = os.path.realpath(__file__)
    constantsPath = os.path.dirname(pytFilePath) # os.path.join(pytFilePath,'Config')
except:
    constantsPath = r'\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS_Development\Dev_Toolbox\MT_Testing'
sys.path.append(constantsPath)

# importlib.reload(sys.modules['Constants'])
constantsModule = sys.modules.get('Constants',None)
if constantsModule:
    importlib.reload(sys.modules['Constants']) # reload to refresh changes in Constants.py file

#import Constants
from Constants import GET_ITEM_ID_REFERENCE_LIST, GET_EXISTING_CULTURAL_RESOURCE_SOURCE_LIST # function to get item IDs and existing cultural resource source list
from Constants import SDE_ENVIRONMENT, SDE_CONNECTION_PATH, PARENT_PROJECT_FOLDER, PARENT_RESOURCE_FOLDER, PARENT_ARPA_FOLDER
from Constants import ENVIRONMENTAL_SETTING_PDF_FOLDER, HISTORIC_SETTING_PDF_FOLDER, PREHISTORIC_SETTING_PDF, PRECONTACT_SETTING_PDF
from Constants import CONTRIBUTING_RESOURCES_LAYER_NAME, NON_CONTRIBUTING_RESOURCES_LAYER_NAME, RESOURCES_CONTRIBUTED_TO_LAYER_NAME
from Constants import CONTRIBUTING_RESOURCE_FIELD_SUFFIX, NON_CONTRIBUTING_RESOURCE_FIELD_SUFFIX, RESOURCES_CONTRIBUTED_TO_FIELD_SUFFIX
from Constants import ACTIVITIES_LAYER_NAME, PREVIOUSLY_IDENTIFIED_RESOURCES_LAYER_NAME, NEWLY_IDENTIFIED_RESOURCES_LAYER_NAME, CONSULTING_PARTY_RESPONSES_LAYER_NAME
from Constants import ARPA_ACTIVITIES_LAYER_NAME, RELATED_RESOURCE_LAYER_NAME
from Constants import PROJECT_LEAD_CONTACTS_CSV, AGOL_CONTACTS_CSV_ITEM_ID
from Constants import MAX_PROJECT_NAME_LENGTH, MAX_RESOURCE_NAME_LENGTH
from Constants import RESOURCE_FIELD_SUFFIX_DICT, ACTIVITY_TYPE_ABBREV_DICT, RESOURCE_TYPE_ABBREV_DICT
from Constants import SHAREPOINT_SITE, SHAREPOINT_FOLDER_DICT, TRIBE_NAME_DOMAIN_DICT, TRIBE_DOMAIN_NAME_DICT, EXCLUDED_TRIBES, CONSULTING_PARTY_CONTACT_LIST_XLSX
from Constants import LEAD_FED_AGENCY_DOMAIN_DICT, STATE_ABBREV_DICT
from Constants import PROJECT_TYPE_DOMAIN_DICT, ACTIVITY_TYPE_DOMAIN_DICT, RESOURCE_TYPE_DOMAIN_DICT
from Constants import PREVIOUSLY_IDENTIFIED_RESOURCE_FIELD_PREFIX, NEWLY_IDENTIFIED_RESOURCE_FIELD_PREFIX, GET_SERVICE_FIELD_MAPPING_DICT, ICD_RESOURCE_DATASET_NAME
from Constants import ACTIVITIES_WITH_RESOURCE
from Constants import TVA_AGOL_URL,SURVEY123_WKID
from Constants import WEB_SOIL_URL
from Constants import EXTERNAL_FEATURE_SERVICES_FOR_BACKGROUND_RESEARCH_LYR, EXTERNAL_MAP_SERVICES_FOR_BACKGROUND_RESEARCH_LYR
from Constants import MT_GDB_TEMPLATE_FOLDER, FEATURE_LAYER_NAME_RESOURCE_TYPE_DICT
from Constants import CRMS_FOLDER_ON_SERVER, LAYER_FILES_FOLDER_ON_SERVER
from Constants import CONTRACTOR_AND_ACTIVITY_NAMES_CSV_LIST_NAME, CONTRACTOR_CSV_LIST_NAME
#from Constants import * # import all constants at once

# alternative: https://gis.tva.gov/arcgis/rest/services/SQLFS/EGIS_AllLayers/FeatureServer/63
TOPO_POLYGONS_URL = 'https://gis.tva.gov/arcgis/rest/services/SQLFS/EGIS_AllLayers/FeatureServer/94' # TODO - update if Topo Polygon Service URL Changes

# TODO - remove after Constants cache is refreshed
pcUser = os.getenv('username')
CONSULTING_PARTY_CONTACT_LIST_XLSX = "C:\\Users\\{}\\Tennessee Valley Authority\\TVACulturalResources - CRMS\\TVA\\Contact_List.xlsx".format(pcUser) # contact list with consulting party email addresses

importlib.reload(sys.modules['Constants'])

# CONTRIBUTING_RESOURCE_FIELD_SUFFIX = '_cr'
# NON_CONTRIBUTING_RESOURCE_FIELD_SUFFIX = '_ncr'
# RESOURCES_CONTRIBUTED_TO_FIELD_SUFFIX = '_rct'
RELATED_RESOURCES_FIELD_PREFIX = 'related_'
RELATED_RESOURCE_LAYER_NAME = 'related_resource'
GENERALIZE_THRESHOLD = "0.5 Feet"

ITEM_ID_REFERENCE_LIST = GET_ITEM_ID_REFERENCE_LIST()
EXISTING_CULTURAL_RESOURCE_SOURCE_LIST = GET_EXISTING_CULTURAL_RESOURCE_SOURCE_LIST()
SERVICE_FIELD_MAPPING_DICT = GET_SERVICE_FIELD_MAPPING_DICT(EXISTING_CULTURAL_RESOURCE_SOURCE_LIST)

CONTRACTOR_LIST = ['TVAR','New South Associates','WSP'] # TODO - move to Constants file

templateDatasetFolder = createFolder(f'{localProjectFolder}\\Tools\\Data\\Templates') # local folder for template datasets
localLayerFileFolder = createFolder(f'{localProjectFolder}\\Tools\\Layers') # local folder for layer files

# connectionFileFolder = createFolder(f'{localProjectFolder}\\Connections') # local folder for sde connection files
# sde_connection = os.path.join(connectionFileFolder,SDE_CONNECTION_PATH)
sde_connection = SDE_CONNECTION_PATH
ecoregions_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.Reference_Ecoregions','ENVIRONMENT.GIS.EPA_Level_IV_Ecoregions') #'ICD.ICD.EPA_Level_IV_Ecoregions'
tribes_by_county_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.Tribal_Areas_By_County') # 'ICD.ICD.Tribal_Areas_By_County'
counties_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ENVIRONMENT.GIS.Counties') # 'ICD.ICD.Counties'
cultural_resource_polygons_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS', 'ENVIRONMENT.GIS.Cultural_Resources') #'ICD.ICD.Cultural_Resources'
cultural_activity_polygons_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS', 'ENVIRONMENT.GIS.Cultural_Activities') #'ICD.ICD.Cultural_Activities'
cultural_resource_polygons_unverified_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ENVIRONMENT.GIS.Cultural_Resources_Unverified') #'ICD.ICD.Cultural_Resources_Unverified'
cultural_resource_polygons_previous_boundaries_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ENVIRONMENT.GIS.Cultural_Resources_Previous_Boundaries') #'ICD.ICD.Cultural_Resources_Previous_Boundaries'
cultural_project_boundary_polygons_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ENVIRONMENT.GIS.Cultural_Project_Boundary_Polygons') #'ICD.ICD.Cultural_Project_Boundary_Polygons'
#cultural_project_ape_polygons_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ICD.ICD.Cultural_Project_APE_Polygons') #'ICD.ICD.Cultural_Project_APE_Polygons'
cultural_project_footprint_polygons_sde = os.path.join(sde_connection,'ENVIRONMENT.GIS.CRMS','ENVIRONMENT.GIS.Cultural_Project_Footprint_Polygons') # 'ICD.ICD.Cultural_Project_Footprint_Polygons'
# cultural_project_polygons_sde = os.path.join(sde_connection,'ICD.ICD.Cultural_Projects')

templateDatasetsGDB = templateDatasetFolder + os.path.sep + 'template_datasets.gdb'
viewshed_points_template = templateDatasetsGDB + os.path.sep + 'Viewshed_Points'
viewshed_polygons_template = templateDatasetsGDB + os.path.sep + 'Viewshed_Polygons'
viewshed_polygons_temp = templateDatasetsGDB + os.path.sep + 'Viewshed_Polygons_Temp'

# TODO - change if needed
individualDocumentsFolderName = 'Individual_Documents'
mergedReportsFolderName = 'Merged_Reports'

### Template Datasets

'''
### code to create blank template datasets ###
shpFileList = [(projectFlightLinesTemplate,'POLYLINE'),(projectLimitsTemplate,'POLYGON'),(requestedProjectAreasTemplates,'POLYGON')]
shpList = []
for shpFile in shpFileList:
    shp = shpFile[0]
    shpList.append(shp)
    geom = shpFile[1]
    arcpy.CreateFeatureclass_management(os.path.dirname(shp),os.path.basename(shp), geom)
'''

### Parameter Defaults
projectLyrParamDefault = 'Cultural Project Boundary Polygons' # set default cultural project layer
activityLyrParamDefault = 'Cultural Activities' # set default cultural activity layer
resourceLyrParamDefault = 'Cultural Resources' # set default cultural resources layer
resourcePrevBoundariesLyrParamDefault = 'Cultural Resources Previous Boundaries' # set default cultural resources layer
mtResourceLyrParamDefault = 'cultural_landscape_poly'
mtResourceLyrParamDefault = 'object_poly'
unverifiedLyrParamDefault = 'Cultural Resources Unverified' # set default cultural resources unverified layer
projectFootprintLyrParamDefault = 'Cultural Project Footprint Polygons' # set default project footprint layer
viewshedUnitsParamDefault = ['Feet','Meters','Yards','Kilometers','Miles']
maxViewshedDistanceParamDefault = '0.5 Miles'
observerHeightParamDefault = '5.75 Feet'
surfaceOffsetParamDefault = '0 Feet'
projectTypeParamList = ['Section_106','Section_110','ARPA']
#projectTypeParamList = [x for x in PROJECT_TYPE_DOMAIN_DICT.values()]
viewshedPointsLyrParamDefault = 'Viewshed Points'
apeProjectBufferParamDefault = '0.5 Miles'
unverifiedResourceSearchDistance = '50 FEET' # '500 FEET'
contributingResourceSearchDistanceDefault = '500 FEET'
contributingStatusToAssignParamList = ['Contributing Resource','Non Contributing Resource','Contributes To This Resource','Not Included','Related Resource',' '] #,'Unknown','Remove',' ']
# submissionTypeParamList = ['Newly Identified Resource','Resource from External Service']
activityTypeDomainList = [k for k in ACTIVITY_TYPE_DOMAIN_DICT.keys()]
retainResourceDistanceParamDefault = '0.5 Miles'

### Group Layers and Group Layer Names
blankGroupLayer = os.path.join(localLayerFileFolder,'Blank_Group_Layer.lyrx')

### Field Name Parameters
projectIDField = 'PROJECT_ID' # project identifier
projectNameField = 'PROJECT_NAME' # project name
projectTypeField = 'PROJECT_TYPE' # Type of Project (Section 106, Section 110, ARPA)
activityIDField = 'ACTIVITY_ID' # activity identifier
activityNameField = 'ACTIVITY_NAME' # activity name
activityTypeField = 'ACTIVITY_TYPE' # Type of activity (Background Research, Findings & Effects, etc)
resourceIDField = 'RESOURCE_ID' # resource identifier
resourceStateIDField = 'RES_STATE_ID' # state id for cultural resource
resourceNameField = 'RESOURCE_NAME' # name for cultural resource
resourceTypeField = 'RESOURCE_TYPE' # type of cultural resource
otherResourceTypeField = 'OTHER_RESOURCE_TYPE' # text field for custom resource type
cemeteryNumberField = 'CEMETERY_NUMBER' # cemetery number
legacyObjectIDField = 'LEGACY_OBJECT_ID' #  legacy id for Object resources
fieldSiteNumberField = 'FIELD_SITE_NUMBER' #  field site number for cultural resources
relatedResourceIDField = 'RELATED_RESOURCE_ID' # related resource id (used for Interments and Reinterments)
# projectDirField = 'PROJ_DIR' # directory for the project (search for ProjectID and find the folder that corresponds to that project)
maxViewshedDistanceField = 'MAX_VIEWSHED_DISTANCE'
observerHeightField = 'OBSERVER_HEIGHT'
surfaceOffsetField = 'SURFACE_OFFSET'
formItemIDField = 'FORM_ITEM_ID'
formRecordGlobalIDField = 'FORM_RECORD_GLOBALID'
featureLayerItemIDField = 'FEATURE_LAYER_ITEM_ID'
formURLField = 'FORM_URL'
gisCommentsField = 'GIS_COMMENTS'

# resource fields in common for syncing with managed tasks GDB # TODO*** update if new resource fields are added
COMMON_RESOURCE_FIELDS_FOR_SYNC = ['resource_id', 'res_state_id', 'resource_name', 'cemetery_number', 'legacy_object_id', 'field_site_number', 'resource_type', 'other_resource_type', 'activity_id', 'activity_type', 'project_id']

COORDINATE_ROUNDING = 6 # decimal places for comparing geometry between MT and AGOL

def addFields(datasetList,fieldNameList_):
    """
        Adds fields to datasets in a dataset list based on a fieldname list
        - datasetList (list) - list of datasets to process
        - fieldNameList_ (list) - list of fields to process that belong in each dataset
    """
    for dataset in datasetList:
        for fieldName in fieldNameList_:
            if fieldName.lower() not in [f.name.lower() for f in arcpy.ListFields(dataset)]:
                if fieldName == projectIDField: arcpy.AddField_management(dataset,projectIDField,"TEXT",0,0,20,"Project ID","NULLABLE","NON_REQUIRED","")
                elif fieldName == projectNameField: arcpy.AddField_management(dataset,projectNameField,"TEXT",0,0,100,"Project Name","NULLABLE","NON_REQUIRED","")
                elif fieldName == activityTypeField: arcpy.AddField_management(dataset,activityTypeField,"TEXT",0,0,100,"Activity Type","NULLABLE","NON_REQUIRED","")
                # elif fieldName == projectTypeField: arcpy.AddField_management(dataset,projectTypeField,"TEXT",0,0,100,"Project Name","NULLABLE","NON_REQUIRED","Cultural_Project_Types")
                ### add new fields here
                else: err('{} not defined in addFields function'.format(fieldName))

                msg('{} field added to {}'.format(fieldName,dataset))
            else:
                # msg('{} field already in {}'.format(fieldName,dataset))
                pass
            '''
            AddField(in_table, field_name, field_type, {field_precision}, {field_scale}, {field_length}, {field_alias}, {field_is_nullable}, {field_is_required}, {field_domain})
            TEXT —Any string of characters.
            FLOAT — Fractional numbers between -3.4E38 and 1.2E38.
            DOUBLE — Fractional numbers between -2.2E308 and 1.8E308.
            SHORT — Whole numbers between -32,768 and 32,767.
            LONG — Whole numbers between -2,147,483,648 and 2,147,483,647.
            DATE —Date and/or time.
            BLOB —Long sequence of binary numbers. You need a custom loader or viewer or a third-party application to load items into a BLOB field or view the contents of a BLOB field.
            RASTER —Raster images. All ArcGIS software-supported raster dataset formats can be stored, but it is highly recommended that only small images be used.
            GUID —Globally unique identifier.
            '''

def generateAddFieldSyntax(dataset):
    """
        Generates syntax for arcpy.AddField() based on a datatset input parameter
        - dataset - input dataset for which to generate the arcpy.AddField() syntax
    """
    for field in arcpy.ListFields(dataset):
        if "SHAPE" not in field.name and "OBJECTID" not in field.name and not field.required:
            fieldType = field.type
            if "Blob" in field.type or "Date" in field.type or "Double" in field.type:
                fieldType = field.type
                fieldType = fieldType.upper()
            if "Integer" in field.type:
                fieldType = "LONG"
            if "Single" in field.type:
                fieldType = "FLOAT"
            if "SmallInteger" in field.type:
                fieldType = "SHORT"
            if "String" in field.type:
                fieldType = "TEXT"
            if field.isNullable == True:
                isNullable = "NULLABLE"
            else:
                isNullable = "NON_NULLABLE"
            if field.required == True:
                isRequired = "REQUIRED"
            else:
                isRequired = "NON_REQUIRED"
            if field.domain is None:
                domain = ''
            else:
                domain = field.domain
            print('arcpy.AddField_management(dataset,"{1}","{2}",{3},{4},{5},"{6}","{7}","{8}","{9}")'.format(dataset,field.name,fieldType,field.precision,field.scale,field.length,field.aliasName,isNullable,isRequired,domain))
            #print('{1}\t{2}\t{3}\t{4}\t{5}\t{6}\t{7}\t{8}\t{9}'.format(dataset,field.name,fieldType,field.precision,field.scale,field.length,field.aliasName,isNullable,isRequired,field.domain))

excludedExportFieldList = ['OBJECTID','OBJECTID_1','Shape','SHAPE','Shape_STAr','Shape_STLe','Shape_Length','Shape_Leng','Shape_Area','created_user','created_date','last_edited_user','last_edited_date']  # used in sdeToLocalShp function to remove invalid shapefile fields

### Global Text, Buffer and Distance Parameters

### Text Value Parameters

### Parameter Domains for Tools
# used for getParameterInfo(self)

### Default Parameter Values
# used for getParameterInfo(self)

# def getStateList(egisConnectionFile):
#     stateList = []
#     sql = 'select distinct STATE from [EGIS].[EGIS].[TGR_County_2000_Polygons_evw] order by STATE'
#     egisConnection = arcpy.ArcSDESQLExecute(egisConnectionFile)
#     sqlReturn = egisConnection.execute(sql)
#     for row in sqlReturn:
#         stateList.append(row[0])
#     return stateList

# def getCountyListForState(egisConnectionFile,state):
#     countyList = []
#     sql = '''select distinct COUNTY from [EGIS].[EGIS].[TGR_County_2000_Polygons_evw] where STATE = '{}' order by COUNTY'''.format(state)
#     egisConnection = arcpy.ArcSDESQLExecute(egisConnectionFile)
#     sqlReturn = egisConnection.execute(sql)
#     for row in sqlReturn:
#         countyList.append(row[0])
#     return countyList


# def getStateListAndStateCountyDict(connectionFile):
#     countyList = []
#     #sql = '''select distinct STATE,COUNTY from [EGIS].[EGIS].[TGR_County_2000_Polygons_evw]'''
#     sql = '''select distinct STATE,COUNTY from ICD.ICD.Counties'''
#     connection = arcpy.ArcSDESQLExecute(connectionFile)
#     sqlReturn = connection.execute(sql)
#     state_list = []
#     state_county_dict = {}
#     for row in sqlReturn:
#         state = row[0]
#         county = row[1]
#         state_list.append(state)
#         state_county_dict[state] = state_county_dict.get(state,[]) + [county]

#     return sorted(list(set(state_list))),state_county_dict

# # egisConnectionFile = r'\\tva\egis\EGIS-TeamShare\Doc\Connections\Production\EGIS@UDCPWGISSQL1@OSA.sde'
# # county_fc = r'{}\EGIS.EGIS.TGR_County_2000_Polygons'.format(egisConnectionFile)
# icdConnectionFile = r'\\tva\egis\EGIS-TeamShare\Doc\Connections\Production\ICD@UDCPWGISSQL1@OSA.sde'
# icdConnectionFile = r'\\tva\egis\EGIS-TeamShare\Doc\Connections\Acceptance\ICD@UDKAWGISSQL1@OSA.sde'
# county_fc = r'{}\ICD.ICD.Counties'.format(icdConnectionFile)
# global state_list
# global state_county_dict
# state_list, state_county_dict = getStateListAndStateCountyDict(connectionFile=icdConnectionFile)



# state = pValues['state']
# if state:
#     countyList = getCountyListForState(egisConnectionFile,state)
#     parameters[pNum['county']].filter.list = countyList

defineGlobalFunctions = True
if defineGlobalFunctions:
    ''' Global Functions '''
    ### Messaging Functions


    ### Functions to retrieve parameter values in validation or execution
    def paramDictValues(inputParameters):
        """ returns a dictionary of the parameter name and the parameter value """
        paramDict = {}
        i = 0
        msg('Input Parameter Values:')
        for param in inputParameters:
            paramDict[param.name]=param.value
            if 'password' not in param.name.lower():
                msg('   {} = {}'.format(param.name,param.value))
            i+=1
        return paramDict

    def paramDictValuesAsText(inputParameters):
        paramDict = {}
        i = 0
        for param in inputParameters:
            paramDict[param.name]=param.valueAsText
            i+=1
        return paramDict

    def paramDictNum(inputParameters):
        """ returns a dictionary of the parameter name and relative position - useful in updateParameter functions"""
        paramDict = {}
        i = 0
        for param in inputParameters:
            paramDict[param.name]= i
            i+=1
        return paramDict

    def addBrackets(txt):
        # function to add brackets for the globalid field
        try:
            if not txt.startswith('{'):
                txt = '{' + txt
            if not txt.endswith('}'):
                txt += '}'
            return txt
        except:
            return txt

    def removeBrackets(txt):
        # function to removes brackets from the globalid field
        try:
            if txt.startswith('{') and txt.endswith('}'):
                return txt[1:-1]
            else:
                return txt
        except:
            return txt

    def tryUpper(txt):
        # function to try converting a text string to upper
        try:
            return txt.upper()
        except:
            return txt

    def sendEmail(fromAddress,toAddress,subjectText,bodyText):
        import smtplib
        from email.mime.text import MIMEText
        try:
            smtpObj = smtplib.SMTP('mailhost.cha.tva.gov')
            emsg = MIMEText(bodyText)
            emsg['Subject'] = subjectText
            emsg['From'] = fromAddress
            emsg['To'] = ', '.join(toAddress)
            smtpObj.sendmail(fromAddress, toAddress, emsg.as_string())
            smtpObj.quit()
        except Exception as e:
            msg(e.message)
    # def connectToAGOL():
    #     # TODO - set to where it will prompt for username/password if not your active portal
    #     from arcgis.gis import GIS
    #     try:
    #         gis = GIS("pro",verify_cert=False)
    #         if gis.url != TVA_AGOL_URL:
    #             err('Currently signed into {}. Sign in to ArcGIS Online: {}'.format(gis.url,TVA_AGOL_URL))
    #         return gis
    #     except Exception as e:
    #         err('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
    #         return None

    # def checkGIS_URL():
    #     # TODO - set to where it will prompt for username/password if not your active portal
    #     from arcgis.gis import GIS
    #     try:
    #         gis = GIS("pro",verify_cert=False)
    #         return gis.url
    #     except Exception as e:
    #         err('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
    #         return None

    def datetimeUTCToEpoch(datetimeObject):
        if isinstance(datetimeObject,datetime.datetime):
            return int(datetimeObject.replace(tzinfo=datetime.timezone.utc).timestamp())
        else:
            return datetimeObject

    def datetimeUTCToEpochMilli(datetimeObject):
        if isinstance(datetimeObject,datetime.datetime):
            return int(datetimeObject.replace(tzinfo=datetime.timezone.utc).timestamp()*1000)
        else:
            return datetimeObject

    def epochToDatetimeUTC(epochObject):
        if isinstance(epochObject,(int,float)):
            if len(str(epochObject)) == 13:
                return datetime.datetime.fromtimestamp(int(epochObject/1000), datetime.timezone.utc)
            else:
                return datetime.datetime.fromtimestamp(int(epochObject), datetime.timezone.utc)
        else:
            return epochObject

    def xlsxToList(xlsxPath,sheetName=None):
        """ Function for converting an XLSX worksheet into corresponding columnList and rows (list of value lists)
        sample call: columns,rows = xlsxToList(xlsxPath,sheetName='GEO')
        # https://stackoverflow.com/questions/53965596/python-3-openpyxl-userwarning-data-validation-extension-not-supported/72559583#72559583
        """
        import pandas
        import warnings
        warnings.simplefilter(action='ignore', category=UserWarning)
        if sheetName:
            df = pandas.read_excel(xlsxPath,sheet_name=sheetName,keep_default_na=False) # keep_default_na = False prevents pandas from converting NA and N/A as NaN values
        else:
            df = pandas.read_excel(xlsxPath)
        columns = df.columns.tolist()
        df1 = df.where(pandas.notnull(df), None)
        rows = df1.values.tolist()
        return columns,rows

    def saveSheets(dfDict, xlsxFile, index=False, header=True):
        """
        Save a dictionary of dataframes to an excel file, with each dataframe as a seperate page
        dfDict is a dictionary where the sheet name is the key and the dataframe is the value
        # must be run in python 3.x
        """
        from pandas import ExcelWriter
        import warnings
        warnings.simplefilter(action='ignore', category=UserWarning)
        writer = ExcelWriter(xlsxFile)
        for sheet in dfDict:
            dfDict[sheet].to_excel(writer, sheet, index=index, header=header)
        writer.close()

    def saveSheets(dfDict, xlsxFile, index=False, header=True, style=None):
        """
        Save a dictionary of dataframes to an excel file, with each dataframe as a seperate page
        dfDict is a dictionary where the sheet name is the key and the dataframe is the value
        # must be run in python 3.x
        """
        from pandas import ExcelWriter
        import warnings
        with ExcelWriter(xlsxFile) as writer:
            for sheet,df in dfDict.items():
                try:
                    starttime = datetime.datetime.now()
                    if style:
                        df = df.style.applymap(style)
                    df.to_excel(writer, sheet, index=index, header=header)
                    #msg(f'{sheet} added to ExcelWriter object. Duration: {datetime.datetime.now()-starttime}')
                except Exception as e:
                    wrn(f'Error exporting {sheet}: {e}')
                    singlexlsx = xlsxFile.replace('.xlsx',f'_{sheet}.xlsx')
                    try:
                        df.to_excel(singlexlsx,index=index,header=header)
                    except Exception as e:
                        wrn(f'Error exporting individual xlsx for {sheet}: {e}')

    def connectToAGOL():
        # TODO - set to where it will prompt for username/password if not your active portal
        from arcgis.gis import GIS
        try:
            gis = GIS("pro",verify_cert=False)
            if gis.url != TVA_AGOL_URL:
                err('Currently signed into {}. Sign in to ArcGIS Online: {}'.format(gis.url,TVA_AGOL_URL))
                return
            return gis
        except Exception as e:
            err('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
            raise Exception('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
            return

    def checkGIS_URL():
        # TODO - set to where it will prompt for username/password if not your active portal
        from arcgis.gis import GIS
        try:
            gis = GIS("pro",verify_cert=False)
            return gis.url
        except Exception as e:
            err('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
            raise Exception('Sign in to ArcGIS Online: {}'.format(TVA_AGOL_URL))
            return None

    def checkGIS(parameters):
        from arcgis.gis import GIS
        gis = GIS("pro",verify_cert=False)
        gisURL = gis.url
        if gisURL != TVA_AGOL_URL:
            parameters[0].setErrorMessage('Currently signed into {}. Sign in to ArcGIS Online: {}'.format(gisURL,TVA_AGOL_URL))
        return parameters

    def deleteDatasets(datasetList):
        """ list of datasets to delete """
        for dataset in datasetList:
            if arcpy.Exists(dataset):
                arcpy.Delete_management(dataset)

    def deleteRows(datasetList):
        """ list of datasets to delete """
        for dataset in datasetList:
            if arcpy.Exists(dataset):
                arcpy.DeleteRows_management(dataset)

    def returnShapefilesInFolder(folder):
        """ This function returns the list of shapefiles in a specified folder """
        fileList = glob.glob(r'{}\*.shp'.format(folder))
        msg('Shapefiles in Folder: {}'.format(fileList))
        return fileList

    def listToString(list_):
        """ converts list to a string; if values of list are of type string, the quotations will be removed; see the listToStringWithQuotes function to maintain quotations
        ['abc','def'] returns "abc,def"
        [123,456] return '123,456'
         """
        return ",".join(list_)

    def listToStringWithQuotes(list_):
        """ converts list to a string; quotations will be added; to remove quotations, see the listToString function
        ['abc','def'] returns "'abc','def'"
        [123,456] return "'123','456'"
         """
        return "'" + "','".join(list_) + "'"

    def generateFieldMappings(inFC,fieldList):
        """ Function to generate field mappings used in FeatureClassToFeatureClass_Conversion and fc2fc function;
        inFC - input feature class
        fieldList - list of fields to include in export
            # field_mapping = generateFieldMappings(fc,fieldList)
            # arcpy.FeatureClassToFeatureClass_conversion(in_features, out_path, out_name, {where_clause}, {field_mapping}, config_keyword)
        """
        fms = arcpy.FieldMappings()
        for field in fieldList:
            if field in [f.name for f in arcpy.ListFields(inFC)]:
                pass
            else:
                truncatedField = field[:10]
                wrn('{} field is not found in {} in generateFieldMapping function. Trying to match truncated field name {}'.format(field,inFC,truncatedField))
                if truncatedField in [f.name for f in arcpy.ListFields(inFC)]:
                    field = truncatedField
                else:
                    wrn('{} field is not found in {} in generateFieldMapping function. {} truncated field is not found'.format(field,inFC,truncatedField))

            fm = arcpy.FieldMap()
            fm.addInputField(inFC,field)
            fms.addFieldMap(fm)
        return fms

    def sdeToLocalShp(sdeInputFC_,outputShp_,fieldList=[],query='',excludedExportfieldList=excludedExportFieldList):
        """ Function to copy one feature class to a shapefile
        sdeInputFC_ - input feature class
        fieldList - list of fields to copy
        outputShp_ - path to shapefile
        query - definition query used to limit the output
        # sample call:
            lpcFC = r'\\tva\egis\EGIS-TeamShare\Doc\Connections\Production\EGIS@UDCPWGISSQL1@OSA.sde\EGIS.CUSTOMER.Local_Power_Company'
            lpcShp = r'c:\temp\surveys_tools\LPC_output'
            lpcFieldList = ['DISTR_NAME','last_edited_date'] # could be [] if all fields should be exported
            ftCampbellQuery = "DISTR_NAME = 'Fort Campbell'" # could be "" if not filter is desired
            sdeToLocalShp(lpcFC,lpcShp,lpcFieldList,ftCampbellQuery,excludedExportFieldList)
        """
        # option to export multiple shapefiles at once, but no way to specify fields or query: https://pro.arcgis.com/en/pro-app/tool-reference/conversion/feature-class-to-shapefile.htm #

        if len(fieldList) == 0: fms = '' # fieldList = [f.name for f in arcpy.ListFields(sdeInputFC_) if f.name not in excludedExportfieldList]
        else: fms = generateFieldMappings(sdeInputFC_,fieldList)
        # msg(fieldList)
        # fms = generateFieldMappings(sdeInputFC_,fieldList)
        msg(fms)
        shpFolder = os.path.dirname(outputShp_)
        if not os.path.exists(shpFolder): os.makedirs(shpFolder)
        shpName = os.path.basename(outputShp_)
        try:
            arcpy.conversion.FeatureClassToFeatureClass(sdeInputFC_, shpFolder, shpName, query, fms, None)
            msg(f'{outputShp_} created')
        except:
            if os.path.exists(outputShp_):
                try:
                    arcpy.Delete_management(outputShp_)
                    msg(f'{outputShp_} deleted')
                    arcpy.conversion.FeatureClassToFeatureClass(sdeInputFC_, shpFolder, shpName, query, fms, None)
                except:
                    try:
                        os.remove(outputShp_)
                        msg(f'{outputShp_} removed')
                        arcpy.conversion.FeatureClassToFeatureClass(sdeInputFC_, shpFolder, shpName, query, fms, None)
                    except:
                        try:
                            arcpy.TruncateTable_management(outputShp_)
                        except:
                            arcpy.DeleteRows_management(outputShp_)
                        arcpy.Append_management(sdeInputFC_,outputShp_, 'NO_TEST')
                        msg(f'{outputShp_} appended')
            else:
                printException()

    def fc2fc(inputFC,outputFC,fieldList=None,query='',excludedExportfieldList=None):
        """Function to copy one feature class to another; also works on shapefiles
        inputFC - input feature class
        fieldList - list of fields to copy
        outputFC - path to output feature class
        query - definition query used in output
        """

        if not fieldList:
            fieldList = [f.name for f in arcpy.ListFields(inputFC) if f.name not in excludedExportfieldList]
        fms = generateFieldMappings(inputFC,fieldList)
        try:
            arcpy.conversion.FeatureClassToFeatureClass(inputFC, os.path.dirname(outputFC), os.path.basename(outputFC), query, fms, None)
            msg(f'{outputFC} created')
        except:
            try:
                arcpy.Delete_management(outputFC)
                msg(f'{outputFC} deleted')
                arcpy.conversion.FeatureClassToFeatureClass(inputFC, os.path.dirname(outputFC), os.path.basename(outputFC), query, fms, None)
            except:
                try:
                    os.remove(outputFC) #if os.path.exists(outputFC) in [True,'true']
                    msg(f'{outputFC} removed')
                    arcpy.conversion.FeatureClassToFeatureClass(inputFC, os.path.dirname(outputFC), os.path.basename(outputFC), query, fms, None)
                except:
                    try:
                        arcpy.TruncateTable_management(outputFC)
                    except:
                        arcpy.DeleteRows_management(outputFC)
                    arcpy.Append_management(inputFC,outputFC, 'NO_TEST')
                    msg(f'{outputFC} appended')
            return

    def inputToList(item):
        """Function to convert the input value to a list, if the user provides a single string or number
        enables function to work with both a list and single value
        https://www.w3schools.com/python/ref_func_isinstance.asp
        """
        if isinstance(item,list): # check to see if lineNameValue is single value or list of values
            listInput = item
        else:
            listInput = [item]
        return listInput

    def returnMapObject(aprxFile='CURRENT',mapName=None):
        """Return the map object for a specified aprx and specified map name; defaults are provided so you can actually call the function using aprxMap=returnMapObject()"""
        aprx =arcpy.mp.ArcGISProject(aprxFile)
        try:
            if not mapName:
                mapName = aprx.activeMap.name
            aprxMap = aprx.listMaps(mapName)[0]
            return aprxMap
        except:
            wrn('!!!Active Map must be open. Open or import a map to continue!!!')
            return None

    def getValidLayers(aprxMap=None,searchLyrName=None):
        if not aprxMap:
            aprxMap = returnMapObject()
        if searchLyrName:
            lyrList= aprxMap.listLayers(searchLyrName)
        else:
            lyrList = aprxMap.listLayers()
        validLayerList = []
        for lyr in lyrList:
            try:
                #msg(lyr.name)
                if lyr.isFeatureLayer: # only return feature layers
                    x = lyr.name # only return if they have a name property
                    validLayerList.append(lyr)
            except Exception as e:
                pass
                #wrn('Unable to process {} layer'.format(lyr))
        return validLayerList

    def getValidTables(aprxMap=None,searchLyrName=None):
        if not aprxMap:
            aprxMap = Object()
        if searchLyrName:
            lyrList= aprxMap.listTables(searchLyrName)
        else:
            lyrList = aprxMap.listTables()
        validLayerList = []
        for lyr in lyrList:
            try:
                #msg(lyr.name)
                if lyr.isFeatureLayer: # only return feature layers
                    x = lyr.name # only return if they have a name property
                    validLayerList.append(lyr)
            except Exception as e:
                pass
                #wrn('Unable to process {} layer'.format(lyr))
        return validLayerList

    def getActiveViewType(activeMapView):
        """
        aprx = arcpy.mp.ArcGISProject("CURRENT")
        mv = aprx.activeView # get active map view
        """
        mvTypeStr = str(type(activeMapView))
        if 'MapView' in mvTypeStr:
            return 'Map'
        elif 'Layout' in mvTypeStr:
            return 'Layout'
        else:
            return None

    def turnLyrsOn(datasetList):
        """Turns on layers in aprx corresponding with the datasets provided in the input list"""
        datasetList = inputToList(datasetList)
        aprxMap_ = returnMapObject()
        lyrList = aprxMap_.listLayers()
        for dataset in datasetList:
            for lyr in lyrList:
                if lyr.supports('DATASOURCE'):
                    if lyr.dataSource == dataset:
                        lyr.visible = True

    def getVersionForMapLayerFromLyrName(lyrName):
        """ function to get the version used in a map layer
        connection properties example:
        {'dataset': 'ICD.ICD.Cultural_Resources', 'workspace_factory': 'SDE', 'connection_info': {'authentication_mode': 'OSA', 'database': 'ICD', 'dbclient': 'sqlserver', 'db_connection_properties': 'udkawgissql1', 'instance': 'sde:sqlserver:udkawgissql1', 'server': 'udkawgissql1', 'version': 'ICD_EDITOR.Cultural_Collector'}}
        """
        aprxMap = returnMapObject()
        lyrList = getValidLayers(aprxMap)
        lyrs = [lyr for lyr in lyrList if lyr.name == lyrName]
        if lyrs:
            lyr = lyrs[0]
            connProp = lyr.connectionProperties
            if connProp['workspace_factory'] == 'SDE':
                version = connProp['connection_info']['version']
                msg('Layer {} is using version: {}'.format(lyrName,version))
                return version
            else:
                msg('Layer {} is not an SDE dataset and has no version'.format(lyrName))
        else:
            wrn('No layer with name {}'.format(lyrName))

    def getVersionForMapLayerFromLyr(lyr):
        """ function to get the version used in a map layer
        connection properties example:
        {'dataset': 'ICD.ICD.Cultural_Resources', 'workspace_factory': 'SDE', 'connection_info': {'authentication_mode': 'OSA', 'database': 'ICD', 'dbclient': 'sqlserver', 'db_connection_properties': 'udkawgissql1', 'instance': 'sde:sqlserver:udkawgissql1', 'server': 'udkawgissql1', 'version': 'ICD_EDITOR.Cultural_Collector'}}
        """
        lyrName = lyr.name
        connProp = lyr.connectionProperties
        if connProp['workspace_factory'] == 'SDE':
            version = connProp['connection_info']['version']
            msg('Layer {} is using version: {}'.format(lyrName,version))
            return version
        else:
            msg('Layer {} is not an SDE dataset and has no version'.format(lyrName))

    def getVersionAndParentVersionForMapLayerFromLyr(lyr):
        """ function to get the version used in a map layer
        connection properties example:
        {'dataset': 'ICD.ICD.Cultural_Resources', 'workspace_factory': 'SDE', 'connection_info': {'authentication_mode': 'OSA', 'database': 'ICD', 'dbclient': 'sqlserver', 'db_connection_properties': 'udkawgissql1', 'instance': 'sde:sqlserver:udkawgissql1', 'server': 'udkawgissql1', 'version': 'ICD_EDITOR.Cultural_Collector'}}
        """
        lyrName = lyr.name
        connProp = lyr.connectionProperties
        if connProp['workspace_factory'] == 'SDE':
            version = connProp['connection_info']['version']
            msg('Layer {} is using version: {}'.format(lyrName,version))
            workspace = os.path.dirname(arcpy.Describe(lyr).catalogPath)
            conn = arcpy.ArcSDESQLExecute(workspace)
            owner,name = version.split('.')
            query = "select parent_owner + '.' + parent_name from sde.sde_versions where name = '{}' and owner = '{}'".format(name,owner)
            sqlResult = conn.execute(query)
            parentVersion = sqlResult
            return version,parentVersion
        else:
            msg('Layer {} is not an SDE dataset and has no version'.format(lyrName))

    def turnLyrsOff(datasetList):
        """Turns on layers in aprx corresponding with the datasets provided in the input list"""
        datasetList = inputToList(datasetList)
        aprxMap_ = returnMapObject()
        lyrList = aprxMap_.listLayers()
        msg(str(datasetList))
        for dataset in datasetList:
            msg(dataset)
            for lyr in lyrList:
                if lyr.supports('DATASOURCE'):
                    if lyr.dataSource == dataset:
                        msg(lyr.dataSource)
                        lyr.visible = False

    def addDatasetToAPRX(aprxMap_='CURRENT',dataset='',layerName_='',layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery=''):
        """ Add Dataset To APRX
        aprxMap_ - map item to which you want to add the dataset; example: aprxMap = aprx.listMaps('Mission Estimation')[0]
        dataset - path to the dataset (can be left blank if you use an empty group layer file)
        layerName_ - what you want the layer to be called; if left blank, it will default to the dataset name
        layerFile_ - .lyrx file containing symbology, or group layer (can be left blank and it will use default symbology)
        replaceExisting_ - parameter to determine how to handle existing layers that may have the same name
            '' - multiple layers will exist with the same name
            'REMOVE' - will remove any other layers that have the same name
            'OVERWRITE' - will change the datasource of the existing layer to the new dataset specified
        # sample call:
        lineOfInterestLyr = addDatasetToAPRX(aprxMap_=aprxMap,dataset=txLinesWithinDistance,layerName_=f'{lineName}',layerFile_='',replaceExisting_='OVERWRITE')
        or
        lineOfInterestLyr = addDatasetToAPRX(aprxMap,txLinesWithinDistance,f'{lineName}','','OVERWRITE')
        """
        if aprxMap_ == '' or aprxMap_ == 'CURRENT':
            aprxMap_ = returnMapObject()

        if layerName_ == '': layerName_ = os.path.basename(dataset)

        if replaceExisting_ == 'REMOVE':
            existingLyrList = aprxMap_.listLayers(layerName_)
            if len(existingLyrList) > 0:
                for existingLyr in existingLyrList:
                    if existingLyr.name == layerName_:
                        aprxMap_.removeLayer(existingLyr)
            add = True
        elif replaceExisting_ == 'OVERWRITE' and dataset != '':
            existingLyrList = aprxMap_.listLayers(layerName_)
            if len(existingLyrList) > 0:
                for existingLyr in existingLyrList:
                    tempFeatureLyr = arcpy.MakeFeatureLayer_management(dataset,'tempFeatureLyr')
                    tempLyr = aprxMap_.listLayers('tempFeatureLyr')[0]
                    existingLyr.updateConnectionProperties(existingLyr.connectionProperties,tempLyr.connectionProperties)
                    aprxMap_.removeLayer(tempLyr)
                    add = False
            else: add = True
        else: add = True

        if add == True:
            if layerFile_ == '':
                lyr = aprxMap_.addDataFromPath(dataset)
                lyr.name = layerName_
            else:
                if dataset == '': # use layer file as is
                    lyrFile = arcpy.mp.LayerFile(layerFile_)
                    lyrFileLyr = aprxMap_.addLayer(lyrFile)[0]
                    lyr = lyrFileLyr
                else: # use original datasource, update symbology
                    lyrFile = arcpy.mp.LayerFile(layerFile_)
                    lyrFileLyr = aprxMap_.addLayer(lyrFile)[0]
                    datasetLyr = aprxMap_.addDataFromPath(dataset)
                    # datasetLyr.symbology = lyrFileLyr.symbology
                    # aprxMap_.removeLayer(lyrFileLyr)
                    # lyr = datasetLyr
                    lyrFileLyr.updateConnectionProperties(lyrFileLyr.connectionProperties,datasetLyr.connectionProperties)
                    aprxMap_.removeLayer(datasetLyr)
                    lyr = lyrFileLyr
                    lyr.name = layerName_
            lyr.visible = visibility_
            if defQuery:
                lyr.definitionQuery = defQuery

            if groupLayer_ != '':
                groupLyrSearch = aprxMap_.listLayers(groupLayer_)
                if len(groupLyrSearch)>0:
                    groupLyr = groupLyrSearch[0] # return layer object for group layer of interest
                else:
                    blankGroupLayerLyrObject = arcpy.mp.LayerFile(blankGroupLayer)
                    aprxMap_.addLayer(blankGroupLayerLyrObject)
                    blankGroupLyr = aprxMap_.listLayers("Blank Group Layer")[0]
                    blankGroupLyr.name = groupLayer_
                    groupLyr = blankGroupLyr
                aprxMap_.addLayerToGroup(groupLyr,lyr)
                aprxMap_.removeLayer(lyr)
            return lyr

    def removeLayers(inputLayerList_):
        """Removes from aprx those layers corresponding to the datasets provided in the input list, or names provided in the input list"""
        if isinstance(inputLayerList_,list): # if input is a list, pass
            pass
        else: # if input is single value, load into a list
            inputLayerList_ = [inputLayerList_]
        aprxMap_ = returnMapObject()
        lyrList = aprxMap_.listLayers()
        for inputLyr in inputLayerList_:
            for lyr in lyrList:
                if lyr.supports('DATASOURCE'):
                    if lyr.dataSource == inputLyr or lyr.name == inputLyr:
                        aprxMap_.removeLayer(lyr)

    def calcFields(datasetList,fieldList,fieldValueList,whereClause='1=1'):
        """calculate multiple fields with a value in a values list for multiple datasets; this will check to make sure the field exists; if it does not, it will add the field
        * this requires the addFields function to be included in the script
        datasetList - list of datasets to update
        fieldList - fields to calculate
        fieldValueList - values to calculate
        whereClause - optional parameter to specify a filter for what will be updated
        """
        if isinstance(datasetList,list): # if input is a list, pass
            pass
        else: # if input is single value, load into a list
            datasetList = [datasetList]
        # addFields(datasetList,fieldList)
        for dataset in datasetList:
            desc = arcpy.Describe(dataset)
            versioned = desc.isVersioned
            #workspace = os.path.dirname(dataset) # TODO - how to account for feature datasets, layers???
            #workspace = arcpy.management.CreateDatabaseConnectionString('SQL_SERVER','UDCPWGISSQL1', 'OPERATING_SYSTEM_AUTH',database='ICD') # doesn't work
            if versioned:
                workspace = os.path.dirname(desc.catalogPath)
                # account for feature datasets
                if (('.sde' in workspace.lower() and not workspace.lower().endswith('.sde'))
                or ('.gdb' in workspace.lower() and not workspace.lower().endswith('.gdb'))):
                    workspace = os.path.dirname(workspace)
                try:
                    editSession = startEditSession(workspace,versioned)
                    msg('Edit session opened for {} workspace'.format(workspace))
                except Exception as e:
                    wrn('Unable to open edit session for {} workspace: {}'.format(workspace,e))

            datasetType = str(type(dataset))
            if datasetType == "<class 'MappingLayerObject'>":
                msg('Attempting to update {} Layer: {} with {}'.format(dataset.name,fieldList,fieldValueList))
            else:
                msg('Attempting to update {}: {} with {}'.format(dataset,fieldList,fieldValueList))
            with arcpy.da.UpdateCursor(dataset,fieldList,whereClause) as cursor:
                for row in cursor:
                    cursor.updateRow(fieldValueList)
            if versioned:
                try:
                    stopEditSession(editSession)
                except:
                    msg('Unable to close edit session')

            # for field,value in dict(zip(fieldList,fieldValueList)).items():
            #     msg('Attempting to update {}: {} with {}'.format(dataset,field,value))
            #     arcpy.CalculateField_management(dataset,field,"'{}'".format(value),'PYTHON3')

    def calcArea(polygonDataset_,sqMilesField_='',acresField_=''):
        """ fuction to calculate area for a polygon dataset, in square miles and/or acres """
        # option to verify geometry as enhancement
        if sqMilesField_ != '':
            addFields(datasetList=[polygonDataset_],fieldNameList_=[sqMilesField_])
            arcpy.CalculateGeometryAttributes_management(polygonDataset_, [[sqMilesField_,'AREA']],'','SQUARE_MILES_US''')
        if acresField_ != '':
            addFields(datasetList=[polygonDataset_],fieldNameList_=[acresField_])
            arcpy.CalculateGeometryAttributes_management(polygonDataset_, [[acresField_,'AREA']],'','ACRES','')

    def calcDistance(lineDataset_,distanceField=''):
        """ fuction to calculate distance for a line dataset, in miles """
        # option to verify geometry as enhancement
        if distanceField != '':
            addFields(datasetList=[lineDataset_],fieldNameList_=[distanceField])
            arcpy.CalculateGeometryAttributes_management(lineDataset_, [[distanceField,'LENGTH']],'MILES_US','','')
            msg('Distance calculated line feature.')

    def returnRowValues(dataset,fieldList,whereClause=None):
        """ Returns a list of lists containing the field values for the fields specified in the fieldList for the specified dataset """
        rowValues = []
        with arcpy.da.SearchCursor(dataset,fieldList,where_clause=whereClause) as cursor:
            for row in cursor:
                rowValues.append(list(row))
        return rowValues

    def returnDistinctFieldValues(table,field,whereClause=''):
        """ Return distinct values for field in table
        table - table to search
        field - field to search
        whereClause - filter to use in search # will default to '' if no parameter is entered
        valueList = returnDistinctFieldValues(table,field,{whereClause})
        ### Note that shapefiles do not support the "DISTINCT" clause; add in logic to search cursor
        """
        with arcpy.da.SearchCursor(table,field,where_clause=whereClause,sql_clause=('DISTINCT','')) as cursor:
            result = []
            for row in cursor:
                if row[0] in result:
                    continue
                result.append(row[0])

        return result
        # #msg('Field: {}; Result: {}'.format(field,result))
        # if len(result) == 1:
        #     return result[0]
        # elif len(result)>1:
        #     return result
        # else:
        #     # err('No values found for {} field in {}'.format(field,table))
        #     # return 'No values found for {} field in {}'.format(field,table)
        #     return

    def returnMultipleDistinctFieldValues(table,fieldList,whereClause=''):
        """
            Return distinct values for fields in table ; same as the returnDistinctFieldValues function, but allows for returning multiple values in one function
            + table - table to search
            + fieldList - list of fields to search
            + whereClause - filter to use in search # will default to '' if no parameter is entered
            returns a list of value lists - ex: a,b,c = returnMultipleDistinctFieldValues(table,['aField','bField','cField'],{whereClause})
            ### Note that shapefiles do not support the "DISTINCT" clause; add in logic to search cursor
        """
        masterResult = []
        for field in fieldList:
            with arcpy.da.SearchCursor(table,field,where_clause=whereClause,sql_clause=('DISTINCT','')) as cursor:
                result = []
                for row in cursor:
                    if row[0] in result:
                        continue
                    result.append(row[0])
            msg('Field: {}; Result: {}'.format(field,result))
            if len(result) == 1:
                masterResult.append(result[0])
            elif len(result)>1:
                masterResult.append(result)
            else:
                # masterResult.append('No values found for {} field in {}'.format(field,table))
                err('No values found for {} field in {}'.format(field,table))

        if len(masterResult) == 1:
            masterResultReturn = masterResult[0]
        elif len(masterResult)>1:
            masterResultReturn = masterResult
        else:
            masterResultReturn = 'Error finding values for {} fields in {}'.format(fieldList,table)

        return masterResultReturn

    def returnMultipleDistinctFieldValuesForOneRecord(table,fieldList,whereClause=''):
        with arcpy.da.SearchCursor(table,fieldList,where_clause=whereClause) as cursor:
            for row in cursor:
                msg(row)
                return row

    def returnDistinctFieldValueForOneRecord(table,field,whereClause=''):
        with arcpy.da.SearchCursor(table,field,where_clause=whereClause) as cursor:
            for row in cursor:
                msg(row[0])
                return row[0]

    def returnUniqueFieldValuesMultipleFields(table, fields, whereClause=None):
        '''
        Return concatenated list of unique values for multiple field combinations
        table = table, feature class, ect.
        fields = [field0.name,field1.name]
        whereClause = optional definition query
        '''
        with arcpy.da.SearchCursor(table,fields,where_clause=whereClause,sql_clause=('DISTINCT','')) as cursor:
            results = []
            for row in cursor:
                if row in results:
                    continue
                results.append(row)
        return results

    def returnProjectFolders(projectID_,parentFolder_):
        """
            Returns folder list that matches project ID
            - projectID_ (string) - Project ID wildcard to parse folders with
            - parentFolder_ (directory) - parent directory to search for matching project ID results
        """
        matchingFolderList = []
        # walk through folders from parentFolder_
        folderList = glob.glob(r'{}\*'.format(parentFolder_))
        for folder in folderList:
            if projectID_ in folder:
                matchingFolderList.append(folder)
        if len(matchingFolderList) != 1:
            if len(matchingFolderList) > 1 :
                wrn("More than one folder ({}) was found for Project {}:{}".format(len(matchingFolderList),projectID_,matchingFolderList))
            else:
                wrn("No Project Folder exists for {}".format(projectID_))
                newFolder = r'{}\{}_New_Project_Folder'.format(parentFolder_,projectID_)
                wrn("New folder will be created for {}: {}".format(projectID_,newFolder))
                matchingFolderList.append(newFolder)
        return matchingFolderList

    def generateCopyOfShapefilesWithProjectID(inputShpList_,projectID_,whereClause='',placeholderTxt_='projectIDPlaceholder'):
        """ Generates a copy of the shapefiles, replacing the "projectIDPlaceholder" text with the ProjectID
        if one input, returns shapefile with project id; if multiple input shapefiles, returns list of new shapefiles """
        outputShpList = []
        for inputShp in inputShpList_:
            outputShp = inputShp.replace(placeholderTxt_,projectID_)
            arcpy.Copy_management(inputShp,outputShp)
            #arcpy.FeatureClassToFeatureClass_conversion(inputShp,os.path.dirname(outputShp),os.path.basename(outputShp),whereClause)
            outputShpList.append(outputShp)
        if len(outputShpList)==1:
            return outputShpList[0]
        else:
            return outputShpList

    def fcToPandasDF(featureClass_, fieldList='', query=''):
        """
        Load data into a Pandas Data Frame for subsequent analysis.
        :param featureClass_: Input ArcGIS Feature Class.
        :param fieldList: Fields for input (optional) - if not specified, it will include all fields
        :param query: subset selection query (optional)
        :return: Pandas DataFrame object.
        """
        if fieldList == '':
            fieldList = [f.name for f in arcpy.ListFields(featureClass_)]
        dataList = []
        with arcpy.da.SearchCursor(featureClass_,fieldList,where_clause=query) as cursor:
            for row in cursor:
                dataList.append(row)

        df = pandas.DataFrame(dataList)
        df.columns = fieldList
        return df

    def startEditSession(workspace,versioned):
        """
        Function to create an edit session within a workspace
        workspace: the workspace to edit
        versioned: True if versions, False if not versioned
        Return from function is the edit session, which would be used as the input for the stopEditSession function
        Sample Call:
        updateLinename_EditSession = startEditSession(editConnection)
        """
        #initiate Edit Session
        edit = arcpy.da.Editor(workspace)
        if edit.isEditing:
            wrn("{} is currently in an edit session please try again later.".format(workspace))
        else:
            try:
                # Start Editing update feature with da.Editor
                ## Start an edit session. Must provide the workspace.
                ## Edit session is started without an undo/redo stack for versioned data
                ##  (for second argument, use False for un-versioned data)
                edit.startEditing(False, versioned)
                ## Start an edit operation
                edit.startOperation()
                msg("    {} - Editing Session Opened for {}".format(datetime.datetime.today().strftime('%H:%M:%S.%f'),workspace))
            except Exception as e:
                wrn("    Error opening edit session for {}: {}".format(workspace,e))
        return edit

    def stopEditSession(editSession):
        """
        Function to stop an edit session
        editSession: the edit session variable defined during the startEditSession function
        Sample Call:
        stopEditSession(updateLinename_EditSession)
        """
        edit = editSession
        #Stop Operations and
        if editSession.isEditing:
            try:
                edit.stopOperation()
                edit.stopEditing(True)
                msg("    {} - Editing Session Closed".format(datetime.datetime.today().strftime('%H:%M:%S.%f')))
            except Exception as e:
                wrn("    Error closing edit session: {}".format(e))
        else:
            wrn("        Not in an edit session.")

    def startEditingLyr(lyr):
        desc = arcpy.Describe(lyr)
        versioned = desc.isVersioned
        msg('Versioned? {}'.format(versioned))
        workspace = os.path.dirname(desc.catalogPath)
        msg('workspace: {}'.format(workspace))
        if versioned or '.gdb' in workspace.lower():
            # account for feature datasets
            if (('.sde' in workspace.lower() and not workspace.lower().endswith('.sde'))
            or ('.gdb' in workspace.lower() and not workspace.lower().endswith('.gdb'))):
                workspace = os.path.dirname(workspace)
            editSession = startEditSession(workspace,versioned)
            return editSession

    def stopEditing(editSession):
        if editSession:
            stopEditSession(editSession)

    def populateEmptyDomainValuesCursor(dataset):
        """ Function to populate domain values in SDE if they are blank in the shapefile,
        as this causes issues if you try to edit the SDE feature class through the attribute table
        """
        import os
        datasetName = os.path.basename(dataset)
        domainFields = [f.name for f in arcpy.ListFields(dataset) if f.domain and f.isNullable] # get list of fields that have domains and are nullable
        msg('Populating empty domain fields: {}'.format(domainFields))
        workspace = os.path.dirname(dataset)
        # determine if dataset is versioned
        desc = arcpy.Describe(dataset)
        versioned = desc.isVersioned
        editSession = startEditSession(workspace,versioned)
        for field in domainFields:
            whereClause = "{} IN ('',' ')".format(field)
            with arcpy.da.UpdateCursor(dataset,[field,'OBJECTID'],whereClause) as cursor:
                for row in cursor:
                    initialValue = row[0]
                    newValue = None
                    row[0] = newValue
                    cursor.updateRow(row)
                    msg('{} changed from {} to {} for OBJECTID {} in {}'.format(field,initialValue,newValue,row[1],datasetName))
        stopEditSession(editSession)

    def populateEmptyDomainValuesScratchGDB(dataset,scratchFC,query=''):
        """ Function to populate domain values in SDE if they are blank in the shapefile,
        as this causes issues if you try to edit the SDE feature class through the attribute table
        """
        import os
        datasetName = os.path.basename(dataset).replace('.shp','')
        msg('scratchFC: {}'.format(scratchFC))
        arcpy.FeatureClassToFeatureClass_conversion(dataset,os.path.dirname(scratchFC),os.path.basename(scratchFC),query)
        domainFields = [f.name for f in arcpy.ListFields(dataset) if f.domain and f.isNullable] # get list of fields that have domains and are nullable
        msg('domain fields: {}'.format(domainFields))
        for field in domainFields:
            whereClause = "{} IN ('',' ')".format(field)
            msg(whereClause)
            with arcpy.da.UpdateCursor(scratchFC,[field,'OBJECTID'],whereClause) as cursor:
                for row in cursor:
                    initialValue = row[0]
                    newValue = None
                    row[0] = newValue
                    msg('{} changed from {} to {} for OBJECTID {} in {} Scratch FC'.format(field,initialValue,newValue,row[1],datasetName))
                    cursor.updateRow(row)
            msg('Update cursor finished for {}'.format(field))

    def populateEmptyDomainValuesCalcField(dataset,lyr=None):
        """ Function to populate domain values in SDE if they are blank in the shapefile,
        as this causes issues if you try to edit the SDE feature class through the attribute table.
        This function uses Calculate Field because it is more stable that opening an edit session
        and running update cursor on versioned datasets
        """
        import os
        datasetName = os.path.basename(dataset)
        if lyr is None:
            lyr = dataset
        domainFields = [f.name for f in arcpy.ListFields(dataset) if f.domain and f.isNullable] # get list of fields that have domains and are nullable
        for field in domainFields:
            whereClause = "{} IN ('',' ')".format(field)
            arcpy.SelectLayerByAttribute_management(lyr,"NEW_SELECTION",whereClause)
            emptyCount = int(str(arcpy.GetCount_management(lyr)))
            if emptyCount > 0:
                arcpy.CalculateField_management(lyr,field,"None","PYTHON3",None)
                msg('{} field changed to None for {} record(s) in {} due to empty domain values'.format(field,emptyCount,datasetName))
            else:
                msg('No records contain empty values for {} in {}'.format(field,datasetName))

    def getMaxFieldValue(dataset,fieldName):
        with arcpy.da.SearchCursor(dataset,[fieldName]) as cursor:
            valueList = [row[0] for row in cursor]
            maxValue = max(valueList)
        return maxValue

    def returnMaxValueSDE(dataset,fieldName,sdeConnectionFile=None,format='NUMBER',whereClause=''):
        """ Function to return the maximum value of a field in a SQL table
        dataset = r'\\tva\egis\EGIS-ARCGISSERVER\connections\Development\SURVEYS@KNXDWGISSQL2@SURVEYS.sde\Transmission_Structures_Active'
        fieldName = 'OBJECTID'
        sdeConnectionFile = r'\\tva\egis\EGIS-ARCGISSERVER\connections\Development\SURVEYS@KNXDWGISSQL2@SURVEYS.sde'
        format = 'NUMBER'
        whereClause = "WHERE VOLTAGE = '161'"
        returnMaxValueSDE(dataset,fieldName,sdeConnectionFile=None,format='NUMBER',whereClause='')
        """
        try:
            import os
            if sdeConnectionFile is None:
                sdeConnectionFile = os.path.dirname(dataset)
            datasetName = os.path.basename(dataset)
            sqlConnection = arcpy.ArcSDESQLExecute(sdeConnectionFile)
            if format == 'NUMBER':
                sqlQuery = '''select max({}) FROM {}_evw {}'''.format(fieldName,datasetName,whereClause)
                returnVal = sqlConnection.execute(sqlQuery)
            elif format == 'DATE':
                sqlQuery = '''select CAST(max({}) as varchar(10)) FROM {}_evw {}'''.format(fieldName,datasetName,whereClause)
                returnVal = sqlConnection.execute(sqlQuery)
            elif format == 'DATETIME':
                sqlQuery = '''select CAST(max({}) as varchar) FROM {}_evw {}'''.format(fieldName,datasetName,whereClause)
                sqlReturn = sqlConnection.execute(sqlQuery)
                if sqlReturn != None:
                    returnVal = sqlReturn[:sqlReturn.find('.')] # round milliseconds off
                else:
                    returnVal = sqlReturn
            else:
                msg('Other Format for {}'.format(datasetName))
                sqlQuery = '''select max({}) FROM {}_evw {}'''.format(fieldName,datasetName,whereClause)
                returnVal = sqlConnection.execute(sqlQuery)
            return returnVal
        except:
            try:
                # msg("{}_evw does not exist; trying {} (unversioned)".format(datasetName,datasetName))
                if format == 'NUMBER':
                    sqlQuery = '''select max({}) FROM {} {}'''.format(fieldName,datasetName,whereClause)
                    returnVal = sqlConnection.execute(sqlQuery)
                elif format == 'DATE':
                    sqlQuery = '''select CAST(max({}) as varchar(10)) FROM {} {}'''.format(fieldName,datasetName,whereClause)
                    returnVal = sqlConnection.execute(sqlQuery)
                elif format == 'DATETIME':
                    sqlQuery = '''select CAST(max({}) as varchar) FROM {} {}'''.format(fieldName,datasetName,whereClause)
                    sqlReturn = sqlConnection.execute(sqlQuery)
                    if sqlReturn != None:
                        returnVal = sqlReturn[:sqlReturn.find('.')] # round milliseconds off
                    else:
                        returnVal = sqlReturn
                else:
                    msg('Other Format for {}'.format(datasetName))
                    sqlQuery = '''select max({}) FROM {} {}'''.format(fieldName,datasetName,whereClause)
                    returnVal = sqlConnection.execute(sqlQuery)
                return returnVal
            except:
                wrn("Error with returnMaxValue function for {},{},{},{}".format(dataset,fieldName,format,whereClause))
                printException()

        '''
        for SQL queries, date fields must be converted to string using CAST(field as varchar) or CAST(field as varchar(10)) syntax
        '''

    def replaceOrAppendFeatures(inputDataset_,targetDataset_,uidfieldList,whereClause=''):
        """Function to either replace (if corresponding record exists) or append (if corresponding record does not exist) features based on a unique id, with the option of adding query to the input dataset
        uidfieldList is the list of fields used to identify uniquely the records being compared between the input and target datasets
        examples: subFootprintUIDFieldList = [assetNoField]
        flightlineUIDFieldList = [projectIDField,acqTypeField,flightlineTypeField]
        projectLimitUIDFieldList = [projectIDField,acqTypeField]
        *** define uidfieldList options at top of script
        """
        def returnDistinctList(table,fieldList,oidField_,whereClause=''):
            """ Returns a dictionary with key being the unique field values and the value being a list of OIDs
            fieldList is the unique field being queried
            oidField_ is the name of the object id field (e.g. OBJECTID, FID, etc) """
            cursorFieldList = [oidField_] + fieldList
            resultDict = {}
            dataList = []
            with arcpy.da.SearchCursor(table,cursorFieldList,where_clause=whereClause,sql_clause=('DISTINCT','')) as cursor:
                for row in cursor:
                    dataList.append(row)
            if len(dataList) != 0:
                df = pandas.DataFrame(dataList)
                df.columns = cursorFieldList
                dfRowValues = df.values.tolist()
                for row in dfRowValues:
                    oid = row[0]
                    rowValueList = row.copy()
                    rowValueList.remove(rowValueList[0])
                    rowValueTuples = tuple(rowValueList)
                    if rowValueTuples in resultDict:
                        existingOIDList = resultDict[rowValueTuples]
                        newOIDList = existingOIDList.copy()
                        newOIDList.append(oid)
                        resultDict[rowValueTuples]=newOIDList
                    else:
                        resultDict[rowValueTuples]=[oid]
            #msg(resultDict)
            return resultDict


        if arcpy.Exists(inputDataset_):
            inputRecordCount = int(str(arcpy.GetCount_management(inputDataset_)))
            if inputRecordCount != 0:
                msg('Attempting to update {} with records from {}'.format(targetDataset_,inputDataset_))
                inputOIDField = getOIDField(inputDataset_) # name of the object id field (e.g. OBJECTID, FID, etc)
                targetOIDField = getOIDField(targetDataset_) # name of the object id field (e.g. OBJECTID, FID, etc)
                inputDistinctValuesDict = returnDistinctList(inputDataset_,uidfieldList,inputOIDField,whereClause)
                targetDistinctValuesDict = returnDistinctList(targetDataset_,uidfieldList,targetOIDField,whereClause)

                combinedInputOIDAppendList = []
                combinedTargetOIDDeleteList = []
                for valueList in inputDistinctValuesDict:
                    if valueList in targetDistinctValuesDict:
                        combinedTargetOIDDeleteList = combinedTargetOIDDeleteList + targetDistinctValuesDict[valueList] # list of OIDs to delete from target
                        combinedInputOIDAppendList = combinedInputOIDAppendList + inputDistinctValuesDict[valueList] # list of OIDs in input that match records in target and need to be appended to target
                    else:
                        combinedInputOIDAppendList = combinedInputOIDAppendList + inputDistinctValuesDict[valueList] # list of OIDs in input but not in target that need to be appended to target

                # append records
                msg('{} OIDs to Append from Input to Target: {}'.format(os.path.basename(inputDataset_),combinedInputOIDAppendList))
                countOIDAppendList = len(combinedInputOIDAppendList)
                if countOIDAppendList > 0:
                    inputAppendOIDQuery = "{} in ({})".format(inputOIDField,listToString(combinedInputOIDAppendList))
                    msg('Append OID Query: {}'.format(inputAppendOIDQuery))

                    useCursor = True
                    if useCursor:
                        msg('Appending to {}'.format(targetDataset_))
                        arcpy.Append_management(inputDataset_, targetDataset_, 'NO_TEST', '', '',inputAppendOIDQuery) # python 3 append has an additional parameter for a query expression
                        # populate domain values in SDE if they are blank from the shapefile
                        populateEmptyDomainValuesCursor(targetDataset_)

                    useScratchGDB = False
                    if useScratchGDB:
                        scratchGDB = arcpy.env.scratchGDB
                        datasetName = os.path.basename(inputDataset_).replace('.shp','').replace('-','_')
                        scratchFC = r'{}\scratchfc_{}'.format(scratchGDB,datasetName)
                        populateEmptyDomainValuesScratchGDB(inputDataset_,scratchFC,inputAppendOIDQuery)
                        arcpy.Append_management(scratchFC, targetDataset_, 'NO_TEST', '', '','') # python 3 append has an additional parameter for a query expression

                    useCalcField = False
                    if useCalcField:
                        def getListOfLastOIDs(maxOID_,numOIDs_):
                            """ Function to get the list of the last N number of OIDs, given the max OID and number of OIDs requested """
                            oidList = []
                            for x in range(0,numOIDs_):
                                oid = maxOID_ - x
                                oidList.append(oid)
                            return oidList
                        arcpy.Append_management(inputDataset_, targetDataset_, 'NO_TEST', '', '',inputAppendOIDQuery) # python 3 append has an additional parameter for a query expression
                        maxOID = returnMaxValueSDE(targetDataset_,'OBJECTID')
                        newOIDList = getListOfLastOIDs(maxOID,countOIDAppendList)
                        newOIDQuery = "{} in ({})".format(targetOIDField,listToString(newOIDList))
                        newRecordsLyr = arcpy.MakeFeatureLayer_management(targetDataset_,'newRecordsLyr',newOIDQuery)
                        populateEmptyDomainValuesCalcField(targetDataset_,newRecordsLyr) # populate domain values in SDE if they are blank from the shapefile


                # delete records
                msg('{} OIDs to Delete from Target to make way for Input Records: {}'.format(os.path.basename(targetDataset_),combinedTargetOIDDeleteList))
                if len(combinedTargetOIDDeleteList) > 0:
                    targetDeleteOIDQuery = "{} in ({})".format(targetOIDField,listToString(combinedTargetOIDDeleteList))
                    targetDeleteLyr = arcpy.MakeFeatureLayer_management(targetDataset_,'deleteLyr',targetDeleteOIDQuery)
                    arcpy.DeleteRows_management(targetDeleteLyr)
            else:
                msg('No records exist in {}. {} will not be updated'.format(inputDataset_,targetDataset_))
        else:
            msg('{} does not exist. {} will not be updated'.format(inputDataset_,targetDataset_))

    def getOIDListFromLyrSelection(lyr):
        # determine selection through FIDset
        try:
            fidSet = arcpy.Describe(lyr).FIDset
        except AttributeError as e:
            wrn('Error getting OID List from Layer Selection. Ensure that Map View is Active andlayer name exists.')
            return []
        msg('FID Set: {}'.format(fidSet))
        oidListString = fidSet.replace(' ','')
        oidList = oidListString.split(';')
        oidList = [oid for oid in oidList if oid] # account for fact that when no records are selected, it shows ['']
        oidList = [int(oid) for oid in oidList]
        msg(oidList)
        return oidList

    def getOIDField(table):
        """ Function to determine the name of the object id field (e.g. OBJECTID, FID, etc) """
        oidField = [f.name for f in arcpy.ListFields(table) if f.type == 'OID'][0]
        return oidField

    def getOIDWhereClause(lyr,oidList):
        oidField = getOIDField(lyr)
        oidWhereClause = "{} in ({})".format(oidField,",".join([str(oid) for oid in oidList]))
        msg(oidWhereClause)
        return oidWhereClause

    def getValueUnit(valueUnit):
        """ Function to split the value and unit portions from a GPLinearUnit Arcpy Parameter """
        value,unit = valueUnit.split(' ',1)
        if '.' in value:
            value = float(value)
        else:
            value = int(value)
        return value,unit

    def returnProjectIDsFromLayerNames():
        """ Function to return Project IDs based on layer names in an APRX
        Sample Call: projectIDs = returnProjectIDsFromLayerNames() """
        aprxMap = returnMapObject()
        lyrList = aprxMap.listLayers()
        lyrNames = [lyr.name for lyr in lyrList if lyr.supports('DATASOURCE') and '.shp' in lyr.name and lyr.name[:4].isnumeric() and lyr.name[5:8].isnumeric() and lyr.name[4] == '-']
        projectIDs = list(set([lyrName[:8] for lyrName in lyrNames]))
        return projectIDs

    def clearMapLayers():
        """Code to clear out all map layers, leaving group layers and basemaps intact"""
        aprx =arcpy.mp.ArcGISProject('CURRENT')
        mapName = aprx.activeMap.name
        aprxMap = aprx.listMaps(mapName)[0]
        for lyr in aprxMap.listLayers(): #ctc
            if lyr.isWebLayer or lyr.isBasemapLayer:
                pass
            elif lyr.supports('DATASOURCE'):
                aprxMap.removeLayer(lyr)

    def get_count_through_search_cursor(dataset,field='OBJECTID'):
        """
        function to get the record cound of a feature class or table
        this is a faster alternative to arcpy.GetCount_management
        default field is OBJECTID, but can be changed if needed
        record_count = get_count_cursor(fc)
        """
        def get_count(dataset,field):
            with arcpy.da.SearchCursor(dataset,[field]) as cursor:
                cnt = len([row for row in cursor])
                return cnt
        try:
            cnt = get_count(dataset,field)
            return cnt
        except Exception as e:
            desc = arcpy.Describe(dataset)
            try:
                field = desc.OIDFieldName
                cnt = get_count(dataset,field)
                return cnt
            except AttributeError as e:
                all_index_fields = []
                for index in desc.indexes:
                    index_fields = [f.name for f in index.fields]
                    for index_field in index_fields:
                        all_index_fields.append(index_field)
                field = all_index_fields[0] # get first index field
                try:
                    cnt = get_count(dataset,field)
                    return cnt
                except Exception as e:
                    field_list = [f.name for f in desc.fields]
                    raise ValueError('\nError using {} field; try another field ({})\n{}\n'.format(field,field_list,e))

    def convertParameterMultiValueTextToList(parameterName,pValuesAsText):
        initialInput = pValuesAsText[parameterName]
        if initialInput:
            valueList = initialInput.replace("'","").split(';')
            return valueList
        else:
            return []

    def getEcoregionList(ecoregions_sde,project_polygon_lyr):
        lyr = arcpy.MakeFeatureLayer_management(ecoregions_sde,'ecoregions')
        selection = arcpy.SelectLayerByLocation_management(lyr, 'INTERSECT', project_polygon_lyr)
        ecoregions_list = []
        with arcpy.da.SearchCursor(selection,['STATE_NAME','US_L4CODE','US_L4NAME']) as cursor:
            for row in cursor:
                state, code, name = row[0], row[1], row[2]
                ecoregion_name = state + ' ' + code + '-' + name
                ecoregions_list.append(ecoregion_name)
        ecoregions_list = sorted(ecoregions_list)
        msg('Ecoregions: {}: {}'.format(len(ecoregions_list),ecoregions_list))
        return sorted(ecoregions_list)

    def getCountyList(counties_sde,project_polygon_lyr):
        lyr = arcpy.MakeFeatureLayer_management(counties_sde,'counties')
        selection = arcpy.SelectLayerByLocation_management(lyr, 'INTERSECT', project_polygon_lyr)
        county_state_list = []
        stcntyfips_list = []
        state_list = []
        state_county_list = []
        with arcpy.da.SearchCursor(selection,['COUNTY','STATE','FIPSSTCO','STATE_COUNTY']) as cursor:
            for row in cursor:
                county, state, fipsstco, state_county = row[0], row[1], row[2], row[3]
                county_state_name = county + ', ' + state
                stcntyfips_list.append(fipsstco)
                county_state_list.append(county_state_name)
                state_list.append(state)
                state_county_list.append(state_county)
        msg('Counties: {}: {}'.format(len(county_state_list),county_state_list))
        state_list = sorted(list(set(state_list)))
        msg('State List: {}'.format(state_list))
        msg('State County List: {}'.format(state_county_list))
        return county_state_list,stcntyfips_list,state_list,state_county_list

    def getTribesList(tribes_by_county_sde,stcntyfips_list):
        query = "FIPSSTCO IN ({})".format(listToStringWithQuotes(stcntyfips_list))
        tribes_list = []
        with arcpy.da.SearchCursor(tribes_by_county_sde,['TRIBE','FIPSSTCO'],where_clause=query) as cursor:
            for row in cursor:
                if row[1] in stcntyfips_list:
                    tribes_list.append(row[0])
        if tribes_list:
            tribes_list = sorted(list(set(tribes_list)))
            msg('Tribes: {}: {}'.format(len(tribes_list),tribes_list))
            return tribes_list
        else:
            return ''

    def get_State_Cultural_Resources(state_list,EXISTING_CULTURAL_RESOURCE_SOURCE_LIST):
        state_resource_sources = []
        for state in state_list:
            msg('Evaluating resources for {}'.format(state))
            source_list = [x for x in EXISTING_CULTURAL_RESOURCE_SOURCE_LIST if x.state == state] # get those resource sources for that state
            state_resource_sources += source_list
        return state_resource_sources

    def getEpochTime():
        import time
        epoch = time.time()
        # epoch = epoch *1000
        return int(epoch)

    def getStartingEpoch(startingDate):
        import time
        from calendar import timegm
        start_time = time.strptime("{}T00:00:00.000Z".format(startingDate), "%Y-%m-%dT%H:%M:%S.%fZ")
        epoch_time = timegm(start_time) * 1000
        return epoch_time

    def getEpochTimeMilli(startingEpochTime):
        import time
        epoch = time.time()
        epoch = epoch * 1000
        time.sleep(.001)
        #msg(epoch)
        epochDifference = epoch-startingEpochTime
        #msg('epochDifference: {}'.format(epochDifference))
        return str(int(epochDifference))

    def epochToDatetimeUTC(epochObject):
        if isinstance(epochObject,(int,float)):
            if len(str(epochObject)) == 13:
                return datetime.datetime.fromtimestamp(int(epochObject/1000), datetime.timezone.utc)
            else:
                return datetime.datetime.fromtimestamp(int(epochObject), datetime.timezone.utc)
        else:
            return epochObject

    # def getFeatures(itemTitle,outputFields,query,relativeLyrNum,layerType,item_id=''):
    #     lyr = returnLyrObjectFromAGOLItem(gis=gis,itemTitle=itemTitle,relativeLyrNum=relativeLyrNum,outputFieldList=outputFields,layerType=layerType,query=query)
    #     fset = lyr.query(where=query,out_fields=listToString(outputFields))
    #     features = fset.features
    #     return features

    # def getFieldListForLayer(lyrObject_):
    #     """ Function that returns the list of fields in a layer object """
    #     fieldList = [f.name for f in lyrObject_.properties.fields]
    #     return fieldList

    def returnLyrObjectFromAGOLItem(gis,item_id='',itemTitle='',relativeLyrNum=0,layerName=None):
        """ Function that searches AGOL / Portal for the Feature Service item and returns the layer object for the specified relative layer number
        inputs:
            gis - GIS session for Portal or AGOL
            item_id - the item ID for the feature layer
            itemTitle - the title of the item
            relativeLyrNum - the relative layer number within the feature service
            layerName - the name of the layer within the feature service

        # you must specify either an item id or item title
        # you must specify either a relative layer number or layer name
         """
        #search_result = GISConnection_.content.search(f"title:{itemTitle_}", item_type = "Feature Service")
        if item_id: # get item directly from item id
            item = gis.content.get(item_id)
        else: # get item by searching for item title
            search_result = gis.content.search(query="title:{}".format(itemTitle),item_type = "Feature Layer", max_items=1000)
            search_result = [x for x in search_result if x.title == itemTitle]
            print(search_result)
            if len(search_result) == 1:
                item = search_result[0]
            else:
                err('No Item with {} Title'.format(itemTitle))
                item = None
        if item:
            msg('Item: {}'.format(item.title))
            lyrs = item.layers + item.tables
            msg('Layers: {}'.format([lyr.properties.name for lyr in lyrs]))
            if layerName:
                selectedLyrs = [lyr for lyr in lyrs if lyr.properties.name==layerName.lower().replace(' ','_')] # TODO remove if domains standardized with form
                if selectedLyrs:
                    if len(selectedLyrs) == 1:
                        selectedLyr = selectedLyrs[0]
                        msg('Selected Layer: {}: {}'.format(layerName,selectedLyr))
                        return selectedLyr
                    else:
                        wrn('More than 1 layer with name {} in item {}'.format(layerName,item))
                        return None
                else:
                    wrn('No layer with name {} in item {}'.format(layerName,item))
                    return None
            if relativeLyrNum == 0:
                selectedLyr = lyrs[0]
                msg('Selected Layer: {}: {}'.format(selectedLyr.properties.name,selectedLyr))
                return selectedLyr
            elif relativeLyrNum:
                for lyr in lyrs:
                    url = lyr.url
                    print(url)
                    url_parts = [x for x in url.split('/')]
                    print(url_parts)
                    rel_num = url_parts[-1] # get last part of url path
                    print(rel_num)
                    print(relativeLyrNum)
                    if int(rel_num) == int(relativeLyrNum):
                        selectedLyr = lyr
                        break
                msg('Selected Layer: {}'.format(selectedLyr))
                return selectedLyr
            else:
                return None
        else:
            return None

    def getFeaturesFromLyr(lyr,query='1=1',outputFields='*',returnGeometry=True):
        """ Function to return the features (in list of dictionaries) from a feature service layer """
        import json
        lyrName = lyr.properties.name

        msg('Query for getting features from {}: {} with {} output fields requested'.format(lyrName,query,outputFields))
        fset = lyr.query(where=query,out_fields=",".join(outputFields),return_geometry=returnGeometry)
        # features_json = fset.features
        # features = json.loads(str(features_json))
        features = fset.features
        if not features:
            msg('No features found in {}'.format(lyrName))
        return features

    def getShapeJSON(layer):
        # https://pro.arcgis.com/en/pro-app/2.7/arcpy/get-started/reading-geometries.htm
        cnt = 0
        with arcpy.da.SearchCursor(layer,['SHAPE@JSON']) as cursor:
            for row in cursor:
                shapeJSON = row[0]
                #msg(shapeJSON)
                cnt+=1
        if cnt == 1:
            return shapeJSON
        else:
            msg('No selection or multiple records selected. Please select one feature.')

    def getXY(layer,mapLyrOID):
        # https://pro.arcgis.com/en/pro-app/2.7/arcpy/get-started/reading-geometries.htm
        cnt = 0
        with arcpy.da.SearchCursor(layer,['SHAPE@XY'],"OBJECTID = {}".format(mapLyrOID)) as cursor:
            for row in cursor:
                x,y = row[0]
                #msg('{},{}'.format(x,y))
                cnt+=1
        if cnt == 1:
            return (x,y)
        else:
            msg('No selection or multiple records selected. Please select one feature.')

    def convert_coordinates(pt_x_easting_long,pt_y_northing_lat,source_WKID,target_WKID=SURVEY123_WKID,transformation=''):
        """Convert source projection SHAPE XY to WGS 1984 (WKID 4326) or another coordinate system
        Note process to get source_WKID from source feature class:
            desc = arcpy.Describe(sourceFC)
            source_WKID = desc.SpatialReference.factoryCode
            target_WKID = desc.SpatialReference.factoryCode # default set to WGS 1984 (4326)
            transformation - e.g. 'NAD_1927_To_NAD_1983_NADCON'
        WKID for State Plane NAD83 US Feet:
            TN - 2274
            AL East - 102629
            AL West - 102630
            GA East - 2239
            GA West - 2240
            KY North - 2246
            KY South - 2247
            MS East - 2254
            MS West - 2255
            NC - 2264
            VA South - 2284
        WKID for WGS 1984 Web Mercator (auxiliary sphere) - 3857
        sample call:
        x,y = return_decimal_degrees_coords(417243.02,1530416.49,source_WKID=2274,target_WKID=4326,transformation='')
        """
        pointGeometry = arcpy.PointGeometry(arcpy.Point(pt_x_easting_long,pt_y_northing_lat),arcpy.SpatialReference(source_WKID)).projectAs(arcpy.SpatialReference(target_WKID),transformation)
        x = pointGeometry.firstPoint.X #'LONGITUDE' [X]
        y = pointGeometry.firstPoint.Y #'LATITUDE' [Y]
        #print(lat,lon)
        return(x,y)


    def getFieldValuesAndShapeJSON(mapLyr,field_names):
        cnt = 0
        if field_names:
            fields = field_names + ['SHAPE@JSON']
        else:
            fields = ['SHAPE@JSON']
        with arcpy.da.SearchCursor(mapLyr,fields) as cursor:
            for row in cursor:
                shapeJSON = row[-1]
                fieldValues = None
                if fieldValues:
                    fieldValues = row[:-1]
                    msg(fieldValues)
                #print(shapeJSON)
                cnt+=1
        if cnt == 1:
            return fieldValues,shapeJSON
        else:
            msg('No selection or multiple records selected. Please select one project polygon.')

    def addFormRecord(lyr,adds_dict_list):
        add_result = lyr.edit_features(adds=adds_dict_list)
        msg(add_result)
        result = add_result['addResults'][0]
        form_globalid = result['globalId']
        form_objectid = result['objectId']
        if form_globalid:
            return tryUpper(addBrackets(form_globalid)), form_objectid # add brackets if they don't already exists
        else:
            return add_result

    def deleteFormRecord(lyr,deletes_oid_list=None,deletes_where_clause=None):
        # https://gis.stackexchange.com/questions/278263/arcgis-api-for-python-delete-multiple-features
        # remove_features = [f for f in test_features if f.attributes['Biologist'] == 'jah']
        # object_ids = [f.get_value('OBJECTID') for f in remove_features]
        # object_ids_str = ",".join(object_ids)
        # delete_result = test_flayer.edit_features(deletes=object_ids_str)
        #oid_list_string = None
        #where = None
        if deletes_where_clause:
            #deletes_dict = {'where':deletes_where_clause}
            delete_result = lyr.delete_features(where=deletes_where_clause)# ,returnDeleteResults=True)
        if deletes_oid_list:
            #deletes_dict = {'objectIds':",".join([str(oid) for oid in deletes_oid_list])}
            oid_list_string = ",".join([str(oid) for oid in deletes_oid_list])
            delete_result = lyr.delete_features(deletes=oid_list_string)# ,returnDeleteResults=True)
        msg(delete_result)
        result = delete_result['deleteResults'][0]
        form_globalid = result['globalId']
        form_objectid = result['objectId']        
        if form_globalid:
            return tryUpper(addBrackets(form_globalid)), form_objectid # add brackets if they don't already exists
        else:
            return delete_result

    def updateFormRecord(lyr,updates_dict_list):
        update_result = lyr.edit_features(updates=updates_dict_list)
        msg(update_result)
        result = update_result['updateResults'][0]
        form_globalid = result['globalId']
        form_objectid = result['objectId']
        if form_globalid:
            return tryUpper(addBrackets(form_globalid)), form_objectid # add brackets if they don't already exists
        else:
            return update_result

    # def calcProjectID(project_polygon_lyr): # old process
    #     if int(str(arcpy.GetCount_management(project_polygon_lyr))) == 1:
    #         with arcpy.da.UpdateCursor(project_polygon_lyr,[projectIDField]) as cursor:
    #             for row in cursor:
    #                 currentValue = row[0]
    #                 if not currentValue:
    #                     epochTime = getEpochTimeMilli()
    #                     row[0] = epochTime
    #                     print(epochTime)
    #                     cursor.updateRow(row)
    #     else:
    #         msg('No selection or multiple records selected. Please select one project polygon.')

    def calcSurvey123URL(project_polygon_lyr,form_url):
        if int(str(arcpy.GetCount_management(project_polygon_lyr))) == 1:
            with arcpy.da.UpdateCursor(project_polygon_lyr,['FORM_URL']) as cursor:
                for row in cursor:
                    row[0] = form_url
                    print(form_url)
                    cursor.updateRow(row)
        else:
            msg('No selection or multiple records selected. Please select one project polygon.')

    def calcProjectID(project_polygon_lyr,projectID,projectIDField):
        """ Function to calculate Project ID field
        This checks to see if Project ID is already calculated before updating it.
        """
        if int(str(arcpy.GetCount_management(project_polygon_lyr))) == 1:
            with arcpy.da.UpdateCursor(project_polygon_lyr,[projectIDField]) as cursor:
                for row in cursor:
                    #epochTime = getEpochTime()
                    epochTimeMilli = getEpochTimeMilli()
                    #msg(epochTimeMilli)
                    epochTime = int(str(epochTimeMilli)[:10]) # formerly 10
                    #msg(epochTime)
                    currentProjectID = row[0]
                    if not currentProjectID:
                        row[0] = epochTime
                        currentProjectID = epochTime
                    #print(epochTime)
                    projectID = currentProjectID
                    cursor.updateRow(row)
            return projectID
        else:
            msg('No selection or multiple records selected. Please select one project polygon.')
            return

    def getProjectIDFromActivityID(activityID):
        """ Function to search the activity ID for the first non-numeric character and return everything to the left
        this should be the project ID because the activity ID is the project ID plus the Activity abbreviation
        """
        if activityID:
            return activityID.split('-')[0]
        else:
            return activityID
        # abbrevStart = 0
        # if activityID:
        #     for i,x in enumerate(list(activityID)):
        #         if not x.isnumeric():
        #             abbrevStart = i-1
        # if abbrevStart == 0:
        #     return None
        # else:
        #     return activityID[:abbrevStart]

    def convert_ValuesList_to_StringList(valuesList):
        stringList = []
        for x in valuesList:
            xStringList = [str(y) for y in x]
            xString = "-".join(xStringList)
            print(xString)
            stringList.append(xString)
        return stringList

    def getDuplicatesInList(valueList):
        # find duplicates in a given list
        # returns a list of those values which have duplicates
        from collections import Counter
        count_dict = dict(Counter([txt for txt in valueList]))
        duplicates_list = [txt for txt,count in count_dict.items() if count > 1]
        return duplicates_list

    def getResourceValueList(source,search_extent_lyr,resourceLyr=None,def_query=None):
        existing_Cultural_Resources_Value_List = []
        if source in ('Internal ICD','Internal CRMS'):
            # path = resourceLyr
            path = arcpy.Describe(resourceLyr).catalogPath
            source_name = resourceLyr.name
        else:
            source_name = source.source_name
            path = source.path
            def_query = source.def_query
            resource_type = source.resource_type
        msg('Path for evaluating resources: {}'.format(path))
        lyr = arcpy.MakeFeatureLayer_management(in_features=path,out_layer=source_name + '_temp_lyr',where_clause=def_query)
        lyr_count = arcpy.GetCount_management(lyr)
        msg('Lyr Count: {}'.format(lyr_count))
        selection = arcpy.SelectLayerByLocation_management(lyr, 'INTERSECT', search_extent_lyr)
        selection_count = arcpy.GetCount_management(selection)
        msg('Selection Count: {}'.format(selection_count))
        cnt = 0

        if source == 'Internal CRMS':
            with arcpy.da.SearchCursor(selection,[resourceIDField,resourceTypeField,resourceStateIDField,resourceNameField,otherResourceTypeField,cemeteryNumberField,legacyObjectIDField,fieldSiteNumberField,activityIDField,projectIDField]) as cursor:
                for row in cursor:
                    existing_Cultural_Resources_Value_List.append(list(row))
                    cnt+=1
        else:
            with arcpy.da.SearchCursor(selection,[source.resource_state_id_field,source.resource_name_field]) as cursor:
                for row in cursor:
                    state_id = row[0]
                    name = row[1]
                    existing_Cultural_Resources_Value_List.append(['',resource_type,state_id,name,'','','','',''])
                    cnt+=1
        msg('{} records for {}'.format(cnt,source_name))
        return existing_Cultural_Resources_Value_List
        # return list has the following information: resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number

    def getEnvironmentalSettingPDF_List(ecoregions_list,ENVIRONMENTAL_SETTING_PDF_FOLDER):
        msg("Searching for environmentalSetting PDFs")
        environmentalSettingPDF_List = []
        for enviroDesc in ecoregions_list:
            state_num,desc = enviroDesc.split('-')
            state_num_split = state_num.split(' ')
            state = state_num_split[:-1]
            state = " ".join(state)
            num = state_num_split[-1]
            msg(state)
            msg(num)
            msg(desc)
            num_desc = num + '_' + desc.replace('/','_')
            folder = os.path.join(ENVIRONMENTAL_SETTING_PDF_FOLDER,state,num_desc)
            filename = 'Description_{}.pdf'.format(num)
            filepath = os.path.join(folder,filename)
            # check to see if filepath exists; if it doesn't, search for the same file in all other state folders
            # if not os.path.exists(filepath):
            #     filelist = getFileListWithSubfolders(inputPath,wildcard=filename)
            #     if filelist:
            #         filepath = filelist[0]
            environmentalSettingPDF_List.append(filepath)
        msg(environmentalSettingPDF_List)
        return environmentalSettingPDF_List

    def getHistoricSettingPDF_List(county_state_list,HISTORIC_SETTING_PDF_FOLDER):
        msg("Searching for County Histories PDFs")
        historicSettingPDF_List = []
        for countyState in county_state_list:
            county,state = countyState.split(',')
            state = state.strip()
            print(county)
            print(state)
            folder = HISTORIC_SETTING_PDF_FOLDER + '\\' + state
            filename = '{} County.pdf'.format(county)
            filepath = folder + '\\' + filename
            historicSettingPDF_List.append(filepath)
        msg(historicSettingPDF_List)
        return historicSettingPDF_List

    def getProjectFolderFromProjectIDAndParentProjectFolder(parent_project_folder,projectID):
        folders = os.listdir(parent_project_folder)
        try:
            projectFolderName = [x for x in folders if projectID in x][-1] # get last one created
            projectFolder = createFolder(os.path.join(parent_project_folder,projectFolderName))
            return projectFolder
        except:
            wrn('No folder found for ProjectID {} in {}'.format(projectID,parent_project_folder))
            # exit()

    def getProjectFolder(parentProjectFolder,projectID,projectName):
        if len(projectID) == 5: # do not add project name as suffix for existing projects that were migrated from ICD
            projectFolder = os.path.join(parentProjectFolder,projectID)
        else:
            projectFolder = os.path.join(parentProjectFolder,'{}_{}'.format(projectID,projectName))
        return projectFolder

    def getResourceFolder(parentResourceFolder,resourceType,resourceID):
        resourceFolder = os.path.join(parentResourceFolder,resourceType,resourceID)
        return resourceFolder

    def combinePDFs(pdfList,combinedPDFPath,overwrite=False):
        msg('Combining PDFs: {}'.format(pdfList))
        import os
        from PyPDF2 import PdfFileMerger, PdfFileReader
        # Call the PdfFileMerger
        mergedPDF = PdfFileMerger()
        if not os.path.exists(combinedPDFPath) or overwrite==True:
            for inputPDF in pdfList:
                if os.path.exists(inputPDF) and inputPDF.endswith(".pdf"):
                    mergedPDF.append(PdfFileReader(inputPDF, 'rb'))
                    msg('Input PDF {} added to merged PDF'.format(os.path.basename(inputPDF)))
                else:
                    wrn('Input PDF {} does not exist or is not a PDF'.format(inputPDF))
            # Write all the files into a file which is named as shown below
            try:
                mergedPDF.write(combinedPDFPath)
                msg('Combined PDF Created: {}'.format(combinedPDFPath))
            except PermissionError as e:
                err('Output PDF already open or inaccessible. Close PDF or re-name Output PDF Name\n{}:{}'.format(combinedPDFPath,e))
        else:
            wrn('Combined PDF Already Exists: {}'.format(combinedPDFPath))
        return combinedPDFPath

    def pdf2txt(path, pages=None):
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.layout import LAParams
        from pdfminer.pdfpage import PDFPage
        from io import StringIO
        if not pages:
            pagenums = set()
        else:
            pagenums = set(pages)
        output = StringIO()
        manager = PDFResourceManager()
        converter = TextConverter(manager, output, laparams=LAParams())
        interpreter = PDFPageInterpreter(manager, converter)

        infile = open(path, 'rb')
        for page in PDFPage.get_pages(infile, pagenums):
            interpreter.process_page(page)
        infile.close()
        converter.close()
        text = output.getvalue()
        output.close()
        return text

    def word2PDF(docPath,pdfPath=None,overwrite=True):
        import os
        if not pdfPath:
            pdfPath = docPath.replace('.docx','.pdf')
        if not os.path.exists(pdfPath) or overwrite:
            import sys
            sys.path.append(pythonModulesFolder)
            #import comtypes
            from comtypes import client
            word = client.CreateObject('Word.Application')
            doc = word.Documents.Open(docPath)
            doc.SaveAs(pdfPath, FileFormat=17) # file format for pdf
            doc.Close()
            word.Quit()
            msg('Word Doc {} converted to PDF {}'.format(docPath,pdfPath))
        else:
            msg('Word Doc {} has already been converted to PDF {}'.format(docPath,pdfPath))
        return pdfPath

    def pdf2Word(pdfPath,docPath=None):
        # https://stackoverflow.com/questions/26358281/convert-pdf-to-doc-python-bash
        import sys
        sys.path.append(pythonModulesFolder)
        import comtypes
        from comtypes import client
        word = comtypes.client.CreateObject('Word.Application')
        pdf = word.Documents.Open(pdfPath)
        if not docPath:
            docPath = pdfPath.replace('.pdf','.docx')
        pdf.SaveAs2(docPath, FileFormat=16) # file format for docx
        pdf.Close()
        word.Quit()
        msg('PDF {} converted to Word Doc {}'.format(pdfPath,docPath))
        return docPath

    def createTableOfContents(pdfInputList,tempFolder,tocRelativeIndex=0,projectName=None,projectID=None):
        import sys
        sys.path.append(pythonModulesFolder)
        from docx import Document
        from docx.shared import Pt

        def findReplaceChar(txt,findChar,replaceChar):
            while findChar in txt:
                txt = txt.replace(findChar,replaceChar)
            return txt

        def createLine(delimiter,count):
            startingString = ''
            for _ in range(1,count):
                startingString = startingString + delimiter
            return startingString

        from PyPDF2 import PdfFileMerger, PdfFileReader

        ### Generate Page Dictionary
        pageDict = {}
        pdfNameList = []
        for pdfPath in pdfInputList:
            basename = os.path.basename(pdfPath)
            #correspondingDoc = os.path.dirname(os.path.dirname(pdfPath)) + '\\' + os.path.basename(pdfPath).replace('.pdf','.docx')
            #print(correspondingDoc)
            #pdfName = docPathpdfDict[correspondingDoc]
            pdfName = basename
            pdfNameList.append(pdfName)
            numPages = PdfFileReader(open(pdfPath, "rb")).getNumPages()
            pageDict[pdfName] = numPages

        ### Print out Table of Contents
        maxLength = 1
        for pdfName in pdfNameList:
            if len(pdfName)>maxLength:
                maxLength=len(pdfName)
        delimiter = '.'
        maxDelimiter = maxLength + 15
        tocTxt = 'Table of Contents'
        print(tocTxt)
        pdfPathTOCTextDict = {}

        tocDoc = os.path.join(tempFolder,'Table_of_Contents.docx')

        includePageNumbers = False # TODO - confirm with CC that we do not need page numbers
        if includePageNumbers:
            if tocRelativeIndex:
                pdfInputList = pdfInputList[:tocRelativeIndex] + [tocDoc] + pdfInputList[tocRelativeIndex:]
                lastPage = 1
                for i,pdfPath in enumerate(pdfInputList):
                    if i == tocRelativeIndex:
                        lastPage += 1 # add page number to account for TOC (assume 1 page; if more, will update later)
                        continue
                    pdfName = os.path.basename(pdfPath)
                    pdfPage = lastPage + pageDict[pdfName]
                    pdfPage = lastPage + 1
                    lastPage = pdfPage + pageDict[pdfName] - 1
                    if i > tocRelativeIndex:
                        pdfName = findReplaceChar(pdfName,'  ',' ')
                        pdfName = findReplaceChar(pdfName,'\t','')
                        pdfLength = len(pdfName)
                        delimiterLength = maxDelimiter - pdfLength
                        lastPage = pdfPage + pageDict[pdfName] - 1
                        # lineTxt = pdfName + createLine(delimiter,delimiterLength) + str(pdfPage)
                        # #lineTxt = pdfName + ' - ' + str(pdfPage)
                        # print('\t' + lineTxt)
                        # lineTxt = lineTxt.replace('_',' ').replace('.pdf','')
                        # tocTxt = tocTxt + '\n' + lineTxt
                        # pdfPathTOCTextDict[pdfPath] = (lineTxt)
                        lineTxt = pdfName.replace('_',' ').replace('.pdf','')
                        lineDelimiter = createLine(delimiter,delimiterLength)
                        pdfPathTOCTextDict[pdfPath] = [lineTxt,lineDelimiter,pdfPage]
                pdfInputList.remove(tocDoc) # remove from list so that it is not included in next section

            else:
                lastPage = 1
                for pdfPath in pdfInputList:
                    pdfName = os.path.basename(pdfPath)
                    pdfPage = lastPage + pageDict[pdfName]
                    pdfPage = lastPage + 1
                    pdfName = findReplaceChar(pdfName,'  ',' ')
                    pdfName = findReplaceChar(pdfName,'\t','')
                    pdfLength = len(pdfName)
                    delimiterLength = maxDelimiter - pdfLength
                    lastPage = pdfPage + pageDict[pdfName] - 1
                    # lineTxt = pdfName + createLine(delimiter,delimiterLength) + str(pdfPage)
                    # #lineTxt = pdfName + ' - ' + str(pdfPage)
                    # print('\t' + lineTxt)
                    # lineTxt = lineTxt.replace('_',' ').replace('.pdf','')
                    # tocTxt = tocTxt + '\n' + lineTxt
                    # pdfPathTOCTextDict[pdfPath] = (lineTxt)
                    lineTxt = pdfName.replace('_',' ').replace('.pdf','')
                    lineDelimiter = createLine(delimiter,delimiterLength)
                    pdfPathTOCTextDict[pdfPath] = [lineTxt,lineDelimiter,pdfPage]
        else:
            pdfInputListForTOC = pdfInputList[tocRelativeIndex:]
            for pdfPath in pdfInputListForTOC:
                pdfName = os.path.basename(pdfPath)
                lineTxt = pdfName.replace('_',' ').replace('.pdf','')
                pdfPathTOCTextDict[pdfPath] = [lineTxt,'','']

        def createTOC(tocDoc,pdfInputList,pdfPathTOCTextDict,projectName,projectID):
            document = Document()
            if projectName:
                p = document.add_paragraph() #'A plain paragraph having some ')
                p.add_run(projectName).bold = True

            if projectID:
                p = document.add_paragraph() #'A plain paragraph having some ')
                p.add_run('Project ID {}'.format(projectID)).bold = True

            # add Table of Contents
            p = document.add_paragraph() #'A plain paragraph having some ')
            p.add_run('Table of Contents').bold = True
            for pdfPath in pdfInputList:
                #tocText = pdfPathTOCTextDict.get(pdfPath,None)
                pdfPathTOCText = pdfPathTOCTextDict.get(pdfPath,None)
                if pdfPathTOCText:
                    msg('pdfPathTOCText: {}'.format(pdfPathTOCText))
                    lineTxt,lineDelimiter,pdfPage = pdfPathTOCText
                    tocText = lineTxt + lineDelimiter + str(pdfPage)
                    p = document.add_paragraph() #'A plain paragraph having some ')
                    p.add_run(tocText).bold = True
            document.save(tocDoc)

            tocPDF = tocDoc.replace('.docx','.pdf')
            word2PDF(tocDoc,tocPDF)
            os.remove(tocDoc)
            return tocPDF

        tocPDF = createTOC(tocDoc,pdfInputList,pdfPathTOCTextDict,projectName,projectID)
        tocPageNum = PdfFileReader(open(tocPDF, "rb")).getNumPages()
        if tocPageNum > 1: # regenerate TOC if it is more than 1 page in length
            addedPages = tocPageNum - 1
            new_pdfPathTOCTextDict = {}
            for pdfPath,TOCText in pdfPathTOCTextDict.items():
                lineTxt,lineDelimiter,pdfPage = TOCText
                new_pdfPathTOCTextDict[pdfPath] = (lineTxt,lineDelimiter,pdfPage+addedPages)
                os.remove(tocDoc)
                os.remove(tocPDF)
                tocPDF = createTOC(tocDoc,pdfInputList,pdfPathTOCTextDict,projectName,projectID)

        return tocPDF

    def splitPDF(pdf, splitList, tempFolder):
        import os
        import PyPDF2
        # creating input pdf file object
        pdfFileObj = open(pdf, 'rb')
        # creating pdf reader object
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        splits=[]
        for split in splitList:
            splits.append(int(split))
        # starting index of first slice
        start = 0
        # starting index of last slice
        end = splits[0]
        #msg(splits)
        outputPDFList = []
        for i in range(len(splits)+1):
            # creating pdf writer object for (i+1)th split
            pdfWriter = PyPDF2.PdfFileWriter()
            # output pdf file name
            name = os.path.basename(pdf).replace('.pdf','')
            outputpdf = tempFolder + os.path.sep + name + '_' + str(start) + '.pdf'
            #msg(outputpdf)
            # adding pages to pdf writer object
            for page in range(start,end):
                msg('Adding page {} to the pdf writer object'.format(page))
                pdfWriter.addPage(pdfReader.getPage(page))
            # writing split pdf pages to pdf file
            with open(outputpdf, "wb") as f:
                pdfWriter.write(f)
            # interchanging page split start position for next split
            start = end
            try:
                # setting split end position for next split
                end = splits[i+1]
            except IndexError:
                # setting split end position for last split
                end = pdfReader.numPages
            if outputpdf not in outputPDFList:
                outputPDFList.append(outputpdf)
        # closing the input pdf file object
        pdfFileObj.close()
        return outputPDFList

    def removeExtraHeaderPages(inputPDF):
        """ Function to remove pages that were exported that only have the header on them """
        import os
        import PyPDF2
        import shutil
        totalNumPages = PyPDF2.PdfFileReader(open(inputPDF, "rb")).getNumPages()
        splitList = [i for i in range(0,totalNumPages)]
        pdfFolder = os.path.dirname(inputPDF)
        tempFolder = createFolder(os.path.join(pdfFolder,'splitTemp')) # created splitTemp folder if it doesn't already exist
        tempPDFList = splitPDF(inputPDF, splitList, tempFolder) # split each page into its own PDF
        outputPDFList = []
        for tempPDF in tempPDFList:
            txt = pdf2txt(tempPDF) # convert page to text
            #msg(txt)
            txtList = txt.split('\n') # convert each line to its own item in a list
            txtListMinusBlankLines = [x.strip() for x in txtList if x.strip()] # remove blank lines
                # example of change: 'Findings and Effects \n\n \n\n \n\n\x0c' >>> ['Findings and Effects']
            msg(txtListMinusBlankLines)
            if len(txtListMinusBlankLines) > 1: # only include those pages that have more than just a header
                outputPDFList.append(tempPDF)
                msg('{} added to output list'.format(tempPDF))

        #shutil.copy2(inputPDF,inputPDF.replace(pdfFolder,tempFolder))
        os.rename(inputPDF,inputPDF.replace(pdfFolder,tempFolder))
        outputPDF = combinePDFs(outputPDFList,inputPDF) # overwrite original PDF
        shutil.rmtree(tempFolder)
        return outputPDF

    def listFieldsFromService(mapLyr):
        """ take map layer, export to temp layer file; use layer file to get list of fields
        this is needed because you cannot run arcpy.ListFields() on the map layer in the aprx
        """
        import os
        folder = arcpy.env.scratchFolder
        testLyrFile = os.path.join(folder,'testLyrFile.lyrx')
        arcpy.management.SaveToLayerFile(mapLyr,testLyrFile)
        lyrFile = arcpy.mp.LayerFile(testLyrFile)
        lyrFileList = lyrFile.listLayers()
        lyr = lyrFileList[0]
        fieldList = [f for f in arcpy.ListFields(lyr)]
        # ['PROP_ID', 'MAPCODE', 'QUAD_ID', 'COUNTY', 'GIS_EASTIN', 'GIS_NORTHI', 'Shape']
        arcpy.Delete_management(testLyrFile)
        return fieldList

    def getLayerDataSource(mapLyr):
        """ take map layer, export to temp layer file; use layer file to get list of fields
        this is needed because you cannot run arcpy.ListFields() on the map layer in the aprx
        """
        import os
        folder = arcpy.env.scratchFolder
        testLyrFile = os.path.join(folder,'testLyrFile.lyrx')
        arcpy.management.SaveToLayerFile(mapLyr,testLyrFile)
        lyrFile = arcpy.mp.LayerFile(testLyrFile)
        lyrFileList = lyrFile.listLayers()
        lyr = lyrFileList[0]
        if lyr.supports('DATASOURCE'):
            dataSource = lyr.dataSource
            if 'http' not in dataSource:
                connProp = lyr.connectionProperties
                datasetName = connProp['dataset']
                dataSource = datasetName
        else:
            dataSource = None
        arcpy.Delete_management(testLyrFile)
        return dataSource

    def getGeometryDictFromShapeJSON(shapeJSON,target_WKID=4326):
        # target WKID
        # WGS 1984 - 4326
        # TNSP NAD 83 - 2274
        def convertRings(shapeJSON,target_WKID):
            # https://developers.arcgis.com/documentation/common-data-types/geometry-objects.htm
            #msg('shapeJSONDict: {}'.format(shapeJSONDict))
            #msg('shapeJSONDictKeys: {}'.format(shapeJSONDict.keys()))
            #msg('shapeJSONDict: {}'.format(shapeJSONDict))
            source_WKID = shapeJSONDict["spatialReference"]["latestWkid"]
            rings = shapeJSONDict.get("rings",None)
            if rings:
                if source_WKID != target_WKID:
                    msg('Converting coordinates from {} to {}'.format(source_WKID,target_WKID))
                    converted_rings = []
                    for ring in rings:
                        converted_ring = []
                        for x,y in ring:
                            new_x,new_y = convert_coordinates(pt_x_easting_long=x,pt_y_northing_lat=y,source_WKID=source_WKID,target_WKID=target_WKID,transformation='')
                            converted_ring.append([new_x,new_y])
                        converted_rings.append(converted_ring)
                else:
                    msg('No conversion of coordinates needed')
                    converted_rings = rings
            else:
                converted_rings = None

            curveRings = shapeJSONDict.get("curveRings",None)
            if curveRings:
                if source_WKID != target_WKID:
                    msg('Converting coordinates from {} to {}'.format(source_WKID,target_WKID))
                   #msg(curveRings)
                    converted_curve_rings = []
                    for curveRing in curveRings:
                        newCurveRing = []
                        for ring in curveRing:
                            if isinstance(ring,list) and len(ring) == 2:
                                new_x,new_y = convert_coordinates(pt_x_easting_long=ring[0],pt_y_northing_lat=ring[1],source_WKID=source_WKID,target_WKID=target_WKID,transformation='')
                                newCurveRing.append([new_x,new_y])
                            elif isinstance(ring,dict):
                                new_ring_dict = {}
                                for k,v_list in ring.items():
                                    new_v_list = []
                                    for v in v_list:
                                        print(v)
                                        if isinstance(v,list) and len(v) == 2:
                                            new_y,new_x = convert_coordinates(pt_x_easting_long=v[0],pt_y_northing_lat=v[1],source_WKID=source_WKID,target_WKID=target_WKID,transformation='')
                                            new_v_list.append([new_x,new_y])
                                        else:
                                            new_v_list.append(v)
                                    #msg(new_v_list)
                                    new_ring_dict[k] = new_v_list
                                newCurveRing.append(new_ring_dict)
                        converted_curve_rings.append(newCurveRing)
                    #msg(converted_curve_rings)
                else:
                    converted_curve_rings = curveRings
                    msg('No conversion of coordinates needed')
            else:
                converted_curve_rings = None

            return converted_rings,converted_curve_rings

        import json
        shapeJSONDict = json.loads(shapeJSON)
        #msg('shapeJSONDict: {}'.format(shapeJSONDict))
        #msg('shapeJSONDictKeys: {}'.format(shapeJSONDict.keys()))
        #msg('shapeJSONDict: {}'.format(shapeJSONDict))
        source_WKID = shapeJSONDict["spatialReference"]["latestWkid"]
        if source_WKID != target_WKID:
            msg('Attempting to convert geometry coordinates from {} to {}'.format(source_WKID,target_WKID))
            rings,curve_rings = convertRings(shapeJSONDict,target_WKID)
        else:
            msg('No geometry conversion needed')
            rings = shapeJSONDict.get('rings')
            curve_rings = shapeJSONDict.get('curveRings')
            # shapeJSONDict.pop('spatialReference')
            # return shapeJSONDict
        geometryDict = {}
        if rings:
            geometryDict['rings'] = rings
        if curve_rings:
            geometryDict['curveRings'] = curve_rings
        #msg('Geometry Dict: {}'.format(geometryDict))
        return geometryDict

    # def getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle,gis):
    #     # msg(ITEM_ID_REFERENCE_LIST)
    #     # feature_layer_item_id = [x.feature_layer_item_id for x in ITEM_ID_REFERENCE_LIST if x.title == itemTitle][0]
    #     form_item_id = [x.form_item_id for x in ITEM_ID_REFERENCE_LIST if x.title == itemTitle][0]
    #     form_item = gis.content.get(form_item_id)
    #     msg('Form ID for {}: {}'.format(itemTitle,form_item_id))
    #     related_items = form_item.related_items(rel_type='Survey2Service',direction='forward')
    #     if len(related_items) == 1:
    #         related_item = related_items[0]
    #         feature_layer_item_id = related_item.id
    #     elif len(related_items) > 1:
    #         wrn('More than 1 related feature layer for form id {}'.format(form_item_id))
    #         feature_layer_item_id = None
    #     else:
    #         wrn('No related feature layer for form id {}'.format(form_item_id))
    #         feature_layer_item_id = None

    #     return form_item_id, feature_layer_item_id

    def getFeatureLayerIDFromFormID(form_item_id,gis):
        try:
            form_item = gis.content.get(form_item_id)
        except Exception as e:
            err('{}\n\n\nUnable to access ArcGIS Online. Check Active Portal.\n\n\n'.format(e))
            exit()
        related_items = form_item.related_items(rel_type='Survey2Service',direction='forward')
        if len(related_items) == 1:
            related_item = related_items[0]
            feature_layer_item_id = related_item.id
        elif len(related_items) > 1:
            wrn('More than 1 related feature layer for form id {}'.format(form_item_id))
            feature_layer_item_id = None
        else:
            wrn('No related feature layer for form id {}'.format(form_item_id))
            feature_layer_item_id = None
        return feature_layer_item_id

    def getFeatureUrlFromFeatureLayerID(feature_layer_item_id,gis):
        try:
            feature_layer_item = gis.content.get(feature_layer_item_id)
        except Exception as e:
            err('{}\n\n\nUnable to access ArcGIS Online. Check Active Portal.\n\n\n'.format(e))
            exit()
        return feature_layer_item.url

    def getFormIDFromFeatureLayerID(feature_layer_item_id,gis):
        try:
            feature_layer_item = gis.content.get(feature_layer_item_id)
        except Exception as e:
            err('{}\n\n\nUnable to access ArcGIS Online. Check Active Portal.\n\n\n'.format(e))
            exit()
        related_items = feature_layer_item.related_items(rel_type='Survey2Service',direction='reverse')
        if len(related_items) == 1:
            related_item = related_items[0]
            form_item_id = related_item.id
        elif len(related_items) > 1:
            wrn('More than 1 related form for feature layer id {}'.format(feature_layer_item_id))
            form_item_id = None
        else:
            wrn('No related feature form for feature layer id {}'.format(feature_layer_item_id))
            form_item_id = None
        return form_item_id

    def getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle,gis,formType=None):
        # msg(ITEM_ID_REFERENCE_LIST)
        # feature_layer_item_id = [x.feature_layer_item_id for x in ITEM_ID_REFERENCE_LIST if x.title == itemTitle][0]
        # TODO - update if Constants.py is modified
        if formType == 'Activity' and itemTitle == 'Other':
            itemTitle = 'Other_Activity'
        elif formType == 'Resource' and itemTitle == 'Other':
            itemTitle = 'Other_Resource'
        form_item_ids = [x.form_item_id for x in ITEM_ID_REFERENCE_LIST if x.title == itemTitle]
        if form_item_ids:
            if len(form_item_ids) == 1:
                form_item_id = form_item_ids[0]
                feature_layer_item_id = getFeatureLayerIDFromFormID(form_item_id,gis)
                msg('Form ID for {}: {}'.format(itemTitle,form_item_id))
                msg('Feature Layer ID for {}: {}'.format(itemTitle,feature_layer_item_id))
                return form_item_id, feature_layer_item_id
            else:
                err('More than 1 form for {} in item ID Reference List: {}'.format(itemTitle,ITEM_ID_REFERENCE_LIST))
                return None,None
        else:
            err('No form for {} in item ID Reference List: {}'.format(itemTitle,ITEM_ID_REFERENCE_LIST))
            return None,None

    def generateFormURL(form_item_id,form_globalid):
        form_url = 'https://survey123.arcgis.com/share/{}?mode=edit&globalId={}'.format(form_item_id,form_globalid)
        return form_url

    def addMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict=None,shapeJSON=None):
        if not shapeJSON:
            fieldValues, shapeJSON = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
        else:
            fieldValues, _ = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
        # msg('shapeJSON: {}'.format(shapeJSON))
        msg('{} >>> {}'.format(sde_fields,fieldValues))
        if survey123_fields:
            fieldValuesDict = dict(zip(survey123_fields,fieldValues))
        else:
            fieldValuesDict = {}
        gis_globalid = returnDistinctFieldValueForOneRecord(mapLyr,'GLOBALID')
        fieldValuesDict['gis_globalid'] = gis_globalid # .replace('{','').replace('}','') # added GLOBALID in order to populate gis_globalid in survey123
        currentDateEpoch = getEpochTime()
        msg(currentDateEpoch)
        fieldValuesDict['begin_date'] = currentDateEpoch * 1000
        if attribute_dict:
            attributes_dict = {**spatial_analysis_dict, **fieldValuesDict, **attribute_dict} # combined spatial analysis dict with field values dict and any custom attribute dictionary provided
        else:
            attributes_dict = {**spatial_analysis_dict, **fieldValuesDict} # combined spatial analysis dict with field values dict
        mainSurveyLyrName = mainSurveyLyr.properties.name
        msg('Attributes Dict for {}: {}'.format(mainSurveyLyrName,attributes_dict))
        geometryDict = getGeometryDictFromShapeJSON(shapeJSON,target_WKID=SURVEY123_WKID)
        adds_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
        form_globalid, form_objectid = addFormRecord(mainSurveyLyr,adds_dict_list)
        return form_globalid # returns the global id of the new survey123 form record

    def updateMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,existingGlobalID,existingObjectid,gis_globalid,attribute_dict=None,shapeJSON=None):
        # if a shape is provided, use that for the geometry (in cases where updating based on another feature); otherwise, use the geometry from the main record
        if not shapeJSON:
            fieldValues, shapeJSON = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
        else:
            fieldValues, _ = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
        if survey123_fields:
            fieldValuesDict = dict(zip(survey123_fields,fieldValues))
        else:
            fieldValuesDict = {}
        fieldValuesDict['gis_globalid'] = gis_globalid # added GLOBALID in order to populate gis_globalid in survey123
        if attribute_dict:
            attributes_dict = {**spatial_analysis_dict, **fieldValuesDict, **attribute_dict} # combined spatial analysis dict with field values dict and any custom attribute dictionary provided
        else:
            attributes_dict = {**spatial_analysis_dict, **fieldValuesDict} # combined spatial analysis dict with field values dict
        attributes_dict['globalId'] = existingGlobalID
        attributes_dict['objectid'] = existingObjectid
        msg('Attributes Dict: {}'.format(attributes_dict))
        geometryDict = getGeometryDictFromShapeJSON(shapeJSON,target_WKID=SURVEY123_WKID)
        updates_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
        form_globalid, form_objectid = updateFormRecord(mainSurveyLyr,updates_dict_list)
        return form_globalid # returns the global id of the new survey123 form record

    def createOrUpdateSurvey123Record(mainSurveyLyr,mapLyr,gis_globalid,itemTitle,form_item_id,feature_layer_item_id,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict=None,agolQuery=None,shapeJSON=None):
        """
        Creates a new Survey123 Record or Updates an Existing Records

        Args:
            mainSurveyLyr (AGOL Feature Layer Object): [description]
            mapLyr (ArcGIS Pro Map Layer): [description]
            gis_globalid (str): [description]
            itemTitle (str): [description]
            form_item_id (str): [description]
            feature_layer_item_id (str): [description]
            sde_fields (list): [description]
            survey123_fields (list): [description]
            spatial_analysis_dict (dict): [description]
            attribute_dict (dict, optional): [description]. Defaults to None.
            agolQuery (str, optional): [description]. Defaults to None.

        Returns:
            [type]: [description]
            form_globalid,form_url,new_record_TF
            (str,str,boolean)
        """
        #existingFeatures = getFeaturesFromLyr(mainSurveyLyr,query="gis_globalid = '{}'".format(gis_globalid),outputFields='*')
        if not agolQuery:
            agolQuery = "gis_globalid = '{}'".format(gis_globalid)
        existingFeatures = getFeaturesFromLyr(mainSurveyLyr,query=agolQuery,outputFields=['globalid','objectid','gis_globalid'],returnGeometry=False)
        if existingFeatures:
            msg('existingFeatures: {}'.format(existingFeatures))
            # TODO - ask how they want to handle multiple Survey123 records with the same gis_globalid
            existingFeature = existingFeatures[0]
            # existingFeatureAttributes = existingFeature['attributes']
            # existingGlobalID = existingFeatureAttributes['globalid']
            # existingObjectid = existingFeatureAttributes['objectid']
            existingGlobalID = existingFeature.attributes['globalid']
            existingObjectid = existingFeature.attributes['objectid']
            existingFeatureCount = len(existingFeatures)
            if existingFeatureCount > 1:
                msg('Existing Features Attributes: {}'.format([feature.attributes for feature in existingFeatures]))
                multipleGlobalIDs = [feature.attributes['globalid'] for feature in existingFeatures]
                #form_url = generateFormURL(form_item_id,existingGlobalID)
                #wrn('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,multipleGlobalIDs,existingObjectid,existingGlobalID,form_url))
                form_url_list = [generateFormURL(form_item_id,globalid) for globalid in multipleGlobalIDs]
                err('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,"\n".join(form_url_list)))
            msg('Survey Record for GIS GlobalID {} already exists in {}: ObjectID: {}, GlobalID: {}'.format(gis_globalid,itemTitle,existingObjectid,existingGlobalID))
            updateExistingRecord = True # TODO - make parameter to determine if update is needed
            if updateExistingRecord:
                form_globalid = updateMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,existingGlobalID,existingObjectid,gis_globalid,attribute_dict,shapeJSON=shapeJSON)
            else:
                msg('No change to existing record')
                form_globalid = existingGlobalID
            # update GIS dataset based on record in Survey123
            form_url = generateFormURL(form_item_id,form_globalid)
            calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
            new_record_TF = False
        else:
            form_globalid = addMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict,shapeJSON=shapeJSON)
            msg('Survey Record for GIS GlobalID {} does not yet exist in {}.'.format(gis_globalid,itemTitle))
            if isinstance(form_globalid,str): # check to see if the resulting globalid is a string
                form_url = generateFormURL(form_item_id,form_globalid)
                calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
                new_record_TF = True
            else:
                err('Error creating new Survey123 record: {}'.format(form_globalid))
                form_url = None
                form_globalid = None
                new_record_TF = False

        return form_globalid,form_url,new_record_TF

    def getBackgroundResearchFeatures(gis,projectID,activity_polygon_lyr,background_research_feature_layer_item_id):
        """ Function to find the Background Research Features associated with a specified Activity
        Inputs:
        projectID
        activity_polygon_layer
        background_research_feature_layer_item_id
        """
        background_research_activity_global_id_query = "PROJECT_ID='{}' AND ACTIVITY_TYPE='Background_Research'".format(projectID)
        arcpy.SelectLayerByAttribute_management(activity_polygon_lyr,where_clause=background_research_activity_global_id_query)
        background_research_activity_global_id = returnDistinctFieldValueForOneRecord(activity_polygon_lyr,'GLOBALID',whereClause=background_research_activity_global_id_query)

        msg('background_research_activity_global_id: {}'.format(background_research_activity_global_id))
        # find corresponding Background Research form
        backgroundResearchSurveyLyr = returnLyrObjectFromAGOLItem(gis,item_id=background_research_feature_layer_item_id,relativeLyrNum=0)
        backgroundResearchSurveyLyrQuery = "gis_globalid='{}'".format(background_research_activity_global_id)
        backgroundResearchFeatures = getFeaturesFromLyr(backgroundResearchSurveyLyr,query=backgroundResearchSurveyLyrQuery)
        if backgroundResearchFeatures:
            if len(backgroundResearchFeatures)>1:
                wrn('More than 1 Background Research Record for GIS GlobalID {}'.format(background_research_activity_global_id))
                return None
            else:
                return backgroundResearchFeatures
        else:
            wrn('No Background Research Features for GIS GlobalID {}, Project ID {}'.format(background_research_activity_global_id,projectID))
            return None

    def getValueListFromService(gis,mainSurveyLayerItemID,mainSurveyLayerQuery,relatedTableName=None,relatedTableQuery=None,fieldList=None,returnGeometry=False):
        """ Function to return a list of values from a Feature Layer or related Feature Layer Table
        gis - the connection to ArcGIS Online or Portal
        mainSurveyLayerItemID - the itemID for the feature layer of interest
        mainSurveyLayerQuery - query used to filter the main feature layer
            example: query="gis_globalid='{}'".format(background_research_activity_global_id)
        relatedTableName - name of the related table to be searched (if any)
        relatedTableQuery - query used to filter the related table
            example: "parentglobalid='{}'".format(backgroundResearchSurveyGlobalID)
        fieldList - list of fields to return
            example: ['previous_resource_id','previous_resource_state_id','previously_id_res_state_id','previous_resource_name','previously_id_resources_type','previous_other_resource_type','previous_cemetery_number','previous_legacy_object_id','previous_field_site_number']
        returnGeometry - boolean - indicates whether geometry should be returned

        valueList = getValueListFromService(gis,mainSurveyLayerItemID=feature_layer_item_id,mainSurveyLayerQuery="1=1",relatedTableName='add_related_records',relatedTableQuery="CEC = '42342'",fieldList=['parentglobalids'])
        """

        # find corresponding form
        mainSurveyLayer = returnLyrObjectFromAGOLItem(gis,item_id=mainSurveyLayerItemID,relativeLyrNum=0)
        if relatedTableName:
            features = getFeaturesFromLyr(mainSurveyLayer,query=mainSurveyLayerQuery,outputFields=['globalid'],return_geometry=returnGeometry)
        else:
            if fieldList:
                features = getFeaturesFromLyr(mainSurveyLayer,query=mainSurveyLayerQuery,outputFields=fieldList,return_geometry=returnGeometry)
            else:
                features = getFeaturesFromLyr(mainSurveyLayer,query=mainSurveyLayerQuery,outputFields='*',return_geometry=returnGeometry)
        if features:
            if relatedTableName:
                if len(features)>1:
                    wrn('More than 1 Record for Query {}'.format(mainSurveyLayerQuery))
                    return None
                else:
                    pass
            else: # if no related table specified, return fields from main layer
                featuresValList = []
                for feature in features:
                    valList = []
                    if not fieldList:
                        fieldList = [k for k in feature.attributes.keys()]
                    for field in fieldList:
                        val = feature.attributes[field]
                        valList.append(val)
                    featuresValList.append(valList)
                # msg('featuresValList: {}'.format(featuresValList))
                return featuresValList
                # if len(result) == 1:
                #     return featuresValList[0]
                # elif len(result)>1:
                #     return featuresValList

        else:
            wrn('No Record for Query {}'.format(mainSurveyLayerQuery))
            return None

        featureGlobalID = features[0].attributes['globalid']

        # relTable = getRelatedTable(background_research_feature_layer_item_id,relationship_name='Background_Research_Survey_Form_previous_identified_resource')
        relTable = returnLyrObjectFromAGOLItem(gis,item_id=mainSurveyLayerItemID,layerName=relatedTableName)
        if relatedTableQuery:
            relatedTableQuery = "parentglobalid='{}' AND {}".format(featureGlobalID,relatedTableQuery)
        else:
            relatedTableQuery = "parentglobalid='{}'".format(featureGlobalID)

        if not fieldList:
            outputFields = '*'
        else:
            outputFields = fieldList
        relatedTableRecords = getFeaturesFromLyr(relTable,query=relatedTableQuery,outputFields=outputFields,returnGeometry=False)

        if relatedTableRecords:
            msg('{} Related Records: {}'.format(len(relatedTableRecords),relatedTableRecords))
            relatedRecordsValueList = []
            for relatedRecord in relatedTableRecords:
                valList = []
                if not fieldList:
                    fieldList = [k for k in relatedRecord.attributes.keys()]
                for field in fieldList:
                    val = relatedRecord.attributes[field]
                    valList.append(val)
                relatedRecordsValueList.append(valList)
            return relatedRecordsValueList
        else:
            msg('No Related Records for Query: {}'.format(relatedTableQuery))
            # return None
            return []

    def getValueListFromRelatedTable(gis,mainSurveyLayerItemID,mainSurveyLayerQuery="1=1",relatedTableName=None,relatedTableQuery=None,fieldList=None,returnGeometry=False):
        """ Function to return a list of values from a Feature Layer or related Feature Layer Table
        gis - the connection to ArcGIS Online or Portal
        mainSurveyLayerItemID - the itemID for the feature layer of interest
        mainSurveyLayerQuery - query used to filter the main feature layer
            example: query="gis_globalid='{}'".format(background_research_activity_global_id)
        relatedTableName - name of the related table to be searched (if any)
        relatedTableQuery - query used to filter the related table
            example: "parentglobalid='{}'".format(backgroundResearchSurveyGlobalID)
        fieldList - list of fields to return
            example: ['previous_resource_id','previous_resource_state_id','previously_id_res_state_id','previous_resource_name','previously_id_resources_type','previous_other_resource_type','previous_cemetery_number','previous_legacy_object_id','previous_field_site_number']
        returnGeometry - boolean - indicates whether geometry should be returned

        valueList = getValueListFromService(gis,mainSurveyLayerItemID=feature_layer_item_id,mainSurveyLayerQuery="1=1",relatedTableName='add_related_records',relatedTableQuery="CEC = '42342'",fieldList=['parentglobalids'])
        """

        # find corresponding form
        # relTable = getRelatedTable(background_research_feature_layer_item_id,relationship_name='Background_Research_Survey_Form_previous_identified_resource')
        relTable = returnLyrObjectFromAGOLItem(gis,item_id=mainSurveyLayerItemID,layerName=relatedTableName)
        if relatedTableQuery:
            relatedTableQuery = "{}".format(relatedTableQuery)
        else:
            relatedTableQuery = "1=1"

        if not fieldList:
            outputFields = '*'
        else:
            outputFields = fieldList
        relatedTableRecords = getFeaturesFromLyr(relTable,query=relatedTableQuery,outputFields=outputFields,returnGeometry=False)

        if relatedTableRecords:
            msg('{} Related Records: {}'.format(len(relatedTableRecords),relatedTableRecords))
            relatedRecordsValueList = []
            for relatedRecord in relatedTableRecords:
                valList = []
                if not fieldList:
                    fieldList = [k for k in relatedRecord.attributes.keys()]
                for field in fieldList:
                    val = relatedRecord.attributes[field]
                    valList.append(val)
                relatedRecordsValueList.append(valList)
            return relatedRecordsValueList
        else:
            msg('No Related Records for Query: {}'.format(relatedTableQuery))
            # return None
            return []

    def getBackgroundResearchResourceValueList(backgroundResearchFeatures,background_research_feature_layer_item_id,gis):
        backgroundResearchSurveyGlobalID = backgroundResearchFeatures[0].attributes['globalid']

        # relTable = getRelatedTable(background_research_feature_layer_item_id,relationship_name='Background_Research_Survey_Form_previous_identified_resource')
        relTable = returnLyrObjectFromAGOLItem(gis,item_id=background_research_feature_layer_item_id,layerName=PREVIOUSLY_IDENTIFIED_RESOURCES_LAYER_NAME)
        backgroundResearchPrevIdentifiedResources = getFeaturesFromLyr(relTable,query="parentglobalid='{}'".format(backgroundResearchSurveyGlobalID))
        if backgroundResearchPrevIdentifiedResources:
            msg('{} Background Research Previously Identified Resources: {}'.format(len(backgroundResearchPrevIdentifiedResources),backgroundResearchPrevIdentifiedResources))
            backgroundResearchResourceValueList = []
            #backgroundResearchResourceIDList = []
            for resource in backgroundResearchPrevIdentifiedResources:
                resource_id = resource.attributes['previous_resource_id']
                resource_type = resource.attributes['previous_resource_type']
                res_state_id = resource.attributes['previous_res_state_id']
                resource_name = resource.attributes['previous_resource_name']
                other_resource_type = resource.attributes['previous_other_resource_type']
                cemetery_number = resource.attributes['previous_cemetery_number']
                legacy_object_id = resource.attributes['previous_legacy_object_id']
                field_site_number = resource.attributes['previous_field_site_number']
                # TODO - get other required fields
                # resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id
                backgroundResearchResourceValueList.append([resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number])
                #backgroundResearchResourceIDList.append(resource_id)
            return backgroundResearchResourceValueList
        else:
            msg('No Background Research Previously Identified Resources')
            return None

    # def getRelatedTable(feature_layer_item_id,relationship_name='Background_Research_Survey_Form_previous_identified_resource'):
    #     lyr = gis.content.get(feature_layer_item_id)
    #     #tbl = lyr.tables[0]
    #     #tblProperties = tbl.properties
    #     #lyrProperties = lyr.properties
    #     flyr = lyr.layers[0] # TODO - change if main feature layer is not the first feature layer

    #     relationships = flyr.properties['relationships']
    #     relationship_dict_list = [d for d in relationships if d['name'] == relationship_name]
    #     if relationship_dict_list:
    #         related_table_id = relationship_dict_list[0]['relatedTableId']
    #         related_table_index = related_table_id - len(lyr.layers)
    #         #wrn(locals())
    #         relatedTable = lyr.tables[related_table_index]
    #         return relatedTable
    #     else:
    #         wrn('Unable to find {} relationship for Item {}'.format(relationship_name,feature_layer_item_id))

    def addPreviouslyIdentifiedResourcesToRelatedTable(relTable,resourceValueList,form_globalid):
        ''' Function to add a list of resources to a related table
        relTable - the table object for the feature service
        resourceValueList - the list of values for the resources to be added
            # contains the following fields:resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id
        form_globalid - the globalid of the related main form record
        '''
        # check to see if the record already exists
        existingPrevIDResources = getFeaturesFromLyr(relTable,query="parentglobalid='{}'".format(form_globalid),outputFields=['previous_resource_id'],returnGeometry=False)
        # existingPrevIDResourcesIDs = [feature['attributes']['previous_resource_id'] for feature in existingPrevIDResources]
        existingPrevIDResourcesIDs = [feature.attributes['previous_resource_id'] for feature in existingPrevIDResources]

        for resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number in resourceValueList:
            if resource_id not in existingPrevIDResourcesIDs:
                attributes_dict = {"previous_resource_type" : resource_type,
                        "previous_resource_name" : resource_name,
                        "previous_resource_id" : resource_id,
                        "previous_res_state_id" : res_state_id,
                        "previous_other_resource_type" : other_resource_type,
                        "previous_cemetery_number" : cemetery_number,
                        "previous_legacy_object_id" : legacy_object_id,
                        "previous_field_site_number" : field_site_number,
                        "parentglobalid" : form_globalid}
                adds_dict_list = [{"attributes":attributes_dict}]
                rel_globalid = addFormRecord(relTable,adds_dict_list)
                msg('{} added as Previously Identified Resource; related globalid: {}'.format(attributes_dict,rel_globalid))
            else:
                msg('{} already in Previously Identified Resources'.format(resource_id))
            # TODO - ask if they want to update the existing form record, or leave as is

    def addNewlyIdentifiedResourcesToRelatedTable(relTable,resourceValueList,form_globalid):
        ''' Function to add a list of resources to a related table
        relTable - the table object for the feature service
        resourceValueList - the list of values for the resources to be added
            # contains the following fields:resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number
        form_globalid - the globalid of the related main form record
        '''
        # check to see if the record already exists
        existingNewIDResources = getFeaturesFromLyr(relTable,query="parentglobalid='{}'".format(form_globalid),outputFields=['new_resource_id'],returnGeometry=False)
        # existingPrevIDResourcesIDs = [feature['attributes']['new_resource_id'] for feature in existingNewIDResources]
        existingNewIDResourcesIDs = [feature.attributes['new_resource_id'] for feature in existingNewIDResources]

        msg('Adding to Newly Identified Resources in {}: {} for form_globalid {}'.format(relTable,resourceValueList,form_globalid))
        for resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number in resourceValueList:
            if resource_id not in existingNewIDResourcesIDs:
                attributes_dict = {"new_resource_type" : resource_type,
                        "new_resource_name" : resource_name,
                        "new_resource_id" : resource_id,
                        "new_res_state_id" : res_state_id,
                        "new_other_resource_type" : other_resource_type,
                        "new_cemetery_number" : cemetery_number,
                        "new_legacy_object_id" : legacy_object_id,
                        "new_field_site_number" : field_site_number,
                        "parentglobalid" : form_globalid}
                adds_dict_list = [{"attributes":attributes_dict}]
                rel_globalid = addFormRecord(relTable,adds_dict_list)
                msg('{} added as Newly Identified Resource; related globalid: {}'.format(attributes_dict,rel_globalid))
            else:
                msg('{} already in Newly Identified Resources'.format(resource_id))

    def addContributingResourcesToRelatedTable(relTable,resourceValueList,form_globalid,fieldNameSuffix):
        ''' Function to add a list of resources to a related table
        relTable - the table object for the feature service
        resourceValueList - the list of values for the resources to be added
        fieldNamePrefix - prefix for each field name (e.g. related_, previous_, new_)
        fieldNameSuffix - suffix for each field name (e.g. _cr, _ncr, _rct)
            # contains the following fields:resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number,
        form_globalid - the globalid of the related main form record
        '''
        msg('Adding to Contributing Resources in {}: {} for form_globalid {}'.format(relTable,resourceValueList,form_globalid))
        resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number = resourceValueList
        attributes_dict = {"resource_type{}".format(fieldNameSuffix) : resource_type,
                "resource_name{}".format(fieldNameSuffix) : resource_name,
                "resource_id{}".format(fieldNameSuffix) : resource_id,
                "res_state_id{}".format(fieldNameSuffix) : res_state_id,
                "other_resource_type{}".format(fieldNameSuffix) : other_resource_type,
                "cemetery_number{}".format(fieldNameSuffix) : cemetery_number,
                "legacy_object_id{}".format(fieldNameSuffix) : legacy_object_id,
                "field_site_number{}".format(fieldNameSuffix) : field_site_number,
                "parentglobalid" : form_globalid}
        adds_dict_list = [{"attributes":attributes_dict}]
        rel_globalid = addFormRecord(relTable,adds_dict_list)
        msg('{} added as resource; related globalid: {}'.format(attributes_dict,rel_globalid))

    def addResourceValueListToRelatedTable(relTable,resourceValueList,form_globalid,fieldNamePrefix='',fieldNameSuffix=''):
        ''' Function to add a list of resources to a related table
        relTable - the table object for the feature service
        resourceValueList - the list of values for the resources to be added
        fieldNamePrefix - prefix for each field name (e.g. related_, previous_, new_)
        fieldNameSuffix - suffix for each field name (e.g. _cr, _ncr, _rct)
            # contains the following fields:resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number,
        form_globalid - the globalid of the related main form record
        '''
        msg('Adding to Contributing Resources in {}: {} for form_globalid {}'.format(relTable,resourceValueList,form_globalid))
        resource_id,resource_type,res_state_id,resource_name,other_resource_type,cemetery_number,legacy_object_id,field_site_number = resourceValueList
        attributes_dict = {"{}resource_type{}".format(fieldNamePrefix,fieldNameSuffix) : resource_type,
                "{}resource_name{}".format(fieldNamePrefix,fieldNameSuffix) : resource_name,
                "{}resource_id{}".format(fieldNamePrefix,fieldNameSuffix) : resource_id,
                "{}res_state_id{}".format(fieldNamePrefix,fieldNameSuffix) : res_state_id,
                "{}other_resource_type{}".format(fieldNamePrefix,fieldNameSuffix) : other_resource_type,
                "{}cemetery_number{}".format(fieldNamePrefix,fieldNameSuffix) : cemetery_number,
                "{}legacy_object_id{}".format(fieldNamePrefix,fieldNameSuffix) : legacy_object_id,
                "{}field_site_number{}".format(fieldNamePrefix,fieldNameSuffix) : field_site_number,
                "parentglobalid" : form_globalid}
        adds_dict_list = [{"attributes":attributes_dict}]
        rel_globalid = addFormRecord(relTable,adds_dict_list)
        msg('{} added as resource; related globalid: {}'.format(attributes_dict,rel_globalid))

    def resourceToFindingsAndEffectsFieldMapDict(resourceFields,fe_resourceFields,resource_field_suffix):
        ''' Function to map the fields in a Resource form with the corresponding
        field in the Findings and Effects form
        example: {'resource_id_har':'resource_id'} '''
        excludedFields = ['creationdate', 'creator', 'editdate', 'editor','objectid','globalid']
        field_map_dict = {}
        for field in [x for x in fe_resourceFields if x.lower() not in excludedFields]:
            if field in resourceFields:
                field_map_dict[field] = field
            elif field.endswith(resource_field_suffix):
                mod_field = field[:field.rfind(resource_field_suffix)]
                if mod_field in resourceFields:
                    field_map_dict[mod_field] = field
        return field_map_dict

    def resourceToContributingFieldMapDict(resourceFields,contributing_resourceFields,resource_field_suffix):
        ''' Function to map the fields in a Resource form with the corresponding
        field in the Contributing or Non-Contributing or Resource Contributes To
        example: {'resource_id_cr':'resource_id'}
        example: {'resource_id_ncr':'resource_id'}
        example: {'resource_id_rct':'resource_id'}
        '''
        excludedFields = ['creationdate', 'creator', 'editdate', 'editor','objectid','globalid']
        field_map_dict = {}
        for field in [x for x in contributing_resourceFields if x.lower() not in excludedFields]:
            if field in resourceFields:
                field_map_dict[field] = field
            elif field.endswith(resource_field_suffix):
                mod_field = field[:field.rfind(resource_field_suffix)]
                if mod_field in resourceFields:
                    field_map_dict[mod_field] = field
        return field_map_dict

    def resourceToNewlyIdentifiedFieldMapDict(resourceFields,act_resourceFields,newlyIdentifiedSuffix=NEWLY_IDENTIFIED_RESOURCE_FIELD_PREFIX):
        ''' Function to map the fields in a Resource form with the corresponding
        field in the Newly Identified Resource section of an Activity form
        example: {'new_resource_id':'resource_id} '''
        excludedFields = ['creationdate', 'creator', 'editdate', 'editor','objectid','globalid']
        prefixLength = len(newlyIdentifiedSuffix)
        field_map_dict = {}
        for field in [x for x in act_resourceFields if x.lower() not in excludedFields]:
            if field in resourceFields:
                field_map_dict[field] = field
            elif field.startswith(newlyIdentifiedSuffix):
                mod_field = field[prefixLength:]
                if mod_field in resourceFields:
                    field_map_dict[mod_field] = field
        return field_map_dict

    def resourceToPreviouslyIdentifiedFieldMapDict(resourceFields,act_resourceFields,previouslyIdentifiedSuffix=PREVIOUSLY_IDENTIFIED_RESOURCE_FIELD_PREFIX):
        ''' Function to map the fields in a Resource form with the corresponding
        field in the Previously Identified Resource section of an Activity form
        example: {'new_resource_id':'resource_id} '''
        excludedFields = ['creationdate', 'creator', 'editdate', 'editor','objectid','globalid']
        prefixLength = len(previouslyIdentifiedSuffix)
        field_map_dict = {}
        for field in [x for x in act_resourceFields if x.lower() not in excludedFields]:
            if field in resourceFields:
                field_map_dict[field] = field
            elif field.startswith(previouslyIdentifiedSuffix):
                mod_field = field[prefixLength:]
                if mod_field in resourceFields:
                    field_map_dict[mod_field] = field
        return field_map_dict

    def resourceToNewOrPreviousResourceFieldMapDict(resourceFields,act_resourceFields,fieldPrefix=PREVIOUSLY_IDENTIFIED_RESOURCE_FIELD_PREFIX):
        ''' Function to map the fields in a Resource form with the corresponding
        field in the Previously Identified or Newly Identified Resource section of an Activity form
        example: {'new_resource_id':'resource_id}
        fieldPrefix=PREVIOUSLY_IDENTIFIED_RESOURCE_FIELD_PREFIX
        fieldPrefix=NEWLY_IDENTIFIED_RESOURCE_FIELD_PREFIX
        '''
        excludedFields = ['creationdate', 'creator', 'editdate', 'editor','objectid','globalid']
        prefixLength = len(fieldPrefix)
        field_map_dict = {}
        for field in [x for x in act_resourceFields if x.lower() not in excludedFields]:
            if field in resourceFields:
                field_map_dict[field] = field
            elif field.startswith(fieldPrefix):
                mod_field = field[prefixLength:]
                if mod_field in resourceFields:
                    field_map_dict[mod_field] = field
        return field_map_dict

    def getFileListWithSubfolders(inputPath,wildcard=None):
        ''' This function returns the list of files in a specified folder and in all subfolders'''
        import os
        fileList = []
        for root, dirs, files in os.walk(inputPath):
            for f in files:
                if not wildcard or wildcard.lower() in f.lower() :
                    filePath=os.path.join(root,f)
                    fileList.append(filePath)
                    # msg(filePath)
        return fileList

    def unzip(zipFilePath,outputFolderPath=None):
        '''
        Function to unzip a zip file
        zipFilePath - required - file path for zip file
        outputFolderPath - optional - full path for resulting folder
            if not provided, the resulting folder will be in the same directory as the zip file
            and have the same name
        returns - outputFolderPath
        sample use:
        zipFilePath = "\\\\tva\\egis\\EGIS-Temp\\Shared\\ALIS_Inventory\\New_Archive_20210831T141714.zip"
        outputFolderPath = unzip(zipFilePath)
        returns "\\\\tva\\egis\\EGIS-Temp\\Shared\\ALIS_Inventory\\New_Archive_20210831T141714"
        '''

        import zipfile
        msg("Unzipping file {}".format(zipFilePath))
        filename = os.path.basename(zipFilePath)
        root = os.path.dirname(zipFilePath)
        if not outputFolderPath:
            outputFolderPath = os.path.join(root, os.path.splitext(filename)[0])
        zipfile.ZipFile(zipFilePath).extractall(outputFolderPath)
        msg("{} unzipped to {}".format(zipFilePath,outputFolderPath))
        return outputFolderPath

    def saveXLSXForSurvey123Form(itemID,outputFolder=None,gis=None):
        ''' Function to open the XLSX configuration for a Survey123 Form
        xlsxPath = saveXLSXForSurvey123Form('7b9d7199cbd9485da34a141eb31edc11') '''
        if not gis:
            gis = connectToAGOL()
        item = gis.content.get(itemID)
        zipFile = item.get_data()
        unzippedFolder = unzip(zipFile,outputFolder)
        fileList = getFileListWithSubfolders(unzippedFolder,wildcard='.xlsx')

        os.remove(zipFile)

        # if fileList:
        #     xlsxFile = fileList[0]
        #     return xlsxFile
        return fileList[0] if fileList else None

    def downloadSurveyGDB(form_item_id,outputFolder,gis=None):
        ''' Function to download the file geodatabase for a Survey123 Form '''
        from arcgis.apps.survey123._survey import SurveyManager, Survey
        if not gis:
            gis = connectToAGOL()
        survey_mgr = SurveyManager(gis)
        survey = survey_mgr.get(form_item_id)
        zipFile = survey.download('File Geodatabase',outputFolder)
        unzippedFolder = unzip(zipFile)
        fileList = getFileListWithSubfolders(unzippedFolder,wildcard='.gdb')
        os.remove(zipFile)

        # if fileList:
        #     gdb = os.path.dirname(fileList[0])
        #     return gdb
        return os.path.dirname(fileList[0]) if fileList else None

    # def xlsxToList(xlsxPath,sheetName=None):
    #     # surveyColumnList,surveyRows = xlsxToList(xlsxPath,sheetName='survey')
    #     # surveyColumnDict = {name:i for i,name in enumerate(surveyColumnList)}

    #     # choiceColumnList,choiceRows = xlsxToList(xlsxPath,sheetName='choice')
    #     # choiceColumnDict = {name:i for i,name in enumerate(choiceColumnList)}
    #     import os
    #     import xlrd
    #     print('Workbook path: {}'.format(xlsxPath))
    #     workbookLocation = os.path.dirname(xlsxPath)
    #     workbookName,workbookExtension = os.path.basename(xlsxPath).split('.')
    #     workbook = xlrd.open_workbook(xlsxPath, on_demand=True)
    #     if sheetName:
    #         sheet = workbook.sheet_by_name(sheetName)
    #     else:
    #         sheet = workbook.sheet_by_index(0)

    #     columnList = sheet.row_values(0, start_colx=0, end_colx=None)
    #     columnList = [x if x else i for i,x in enumerate(columnList)] # convert any empty column names to the relative index
    #     rowLists = []
    #     row_count = sheet.nrows
    #     col_count = sheet.ncols
    #     for cur_row in range(1, row_count):
    #         rowValues = sheet.row_values(cur_row, start_colx=0, end_colx=None)
    #         rowLists.append(rowValues)

    #     return columnList,rowLists

    def validateActivityIDandName(activity_polygon_lyr,projectID,activityID,activityName):
        oidList = getOIDListFromLyrSelection(activity_polygon_lyr)
        oidWhereClause = getOIDWhereClause(activity_polygon_lyr,oidList)
        msg('Initial Count: {}'.format(arcpy.GetCount_management(activity_polygon_lyr)))
        arcpy.SelectLayerByAttribute_management(activity_polygon_lyr,where_clause="{} = '{}'".format(projectIDField,projectID))
        existingActivityPolygonCount = int(str(arcpy.GetCount_management(activity_polygon_lyr)))
        msg('Count for Project {}: {}'.format(projectID,existingActivityPolygonCount))
        existingActivityIDs,existingActivityNames = [],[] # set initial values to empty lists
        if existingActivityPolygonCount:
            whereClause="{} = '{}' AND {} IS NOT NULL".format(projectIDField,projectID,projectIDField)
            results = returnUniqueFieldValuesMultipleFields(activity_polygon_lyr,[activityIDField,activityNameField],whereClause=whereClause)
            existingActivityIDs = [x[0] for x in results if x[0] != activityID]
            existingActivityNames = [x[1] for x in results if x[0] != activityName]
            msg('Existing Activity IDs: {}'.format(existingActivityIDs))
            msg('Existing Activity Names: {}'.format(existingActivityNames))

            # if existingActivityNames:
            #     if activityName in existingActivityNames:
            #         wrn('{} Activity Name already exists for Project {}'.format(activityName,projectID))
            #         #exit()
            #         return False

        arcpy.SelectLayerByAttribute_management(activity_polygon_lyr,where_clause=oidWhereClause)
        msg('Count for Processing: {}'.format(arcpy.GetCount_management(activity_polygon_lyr)))
        return existingActivityIDs,existingActivityNames

    def validateProjectName(project_polygon_lyr,projectName):
        oidList = getOIDListFromLyrSelection(project_polygon_lyr)
        oidWhereClause = getOIDWhereClause(project_polygon_lyr,oidList)
        msg('Initial Count: {}'.format(arcpy.GetCount_management(project_polygon_lyr)))
        arcpy.SelectLayerByAttribute_management(project_polygon_lyr,where_clause="{} = '{}'".format(projectNameField,projectName))
        existingProjectPolygonCount = int(str(arcpy.GetCount_management(project_polygon_lyr)))
        msg('Count for Project {}: {}'.format(projectName,existingProjectPolygonCount))
        if existingProjectPolygonCount:
            wrn('{} Project Name already exists'.format(projectName))

        arcpy.SelectLayerByAttribute_management(project_polygon_lyr,where_clause=oidWhereClause)
        msg('Count for Processing: {}'.format(arcpy.GetCount_management(project_polygon_lyr)))

        return existingProjectPolygonCount

        #return True

    def validateResourceName(resourceLyr,resourceName):
        oidList = getOIDListFromLyrSelection(resourceLyr)
        oidWhereClause = getOIDWhereClause(resourceLyr,oidList)
        msg('Initial Count: {}'.format(arcpy.GetCount_management(resourceLyr)))
        arcpy.SelectLayerByAttribute_management(resourceLyr,where_clause="{} = '{}'".format(resourceNameField,resourceName))
        existingResourcePolygonCount = int(str(arcpy.GetCount_management(resourceLyr)))
        msg('Count for Project {}: {}'.format(resourceName,existingResourcePolygonCount))
        if existingResourcePolygonCount>1:
            wrn('{} Resource Name already exists'.format(resourceName))

        arcpy.SelectLayerByAttribute_management(resourceLyr,where_clause=oidWhereClause)
        msg('Count for Processing: {}'.format(arcpy.GetCount_management(resourceLyr)))

        return existingResourcePolygonCount

    def addLayerFileToMap(layerFile,position='TOP'):
        #https://pro.arcgis.com/en/pro-app/latest/arcpy/mapping/map-class.htm
        #'AUTO_ARRANGE', 'BOTTOM','TOP'
        lyrObject = arcpy.mp.LayerFile(layerFile)
        aprxMap = returnMapObject()
        lyr = aprxMap.addLayer(lyrObject,position)
        return lyr

    def checkVersion(parameters,pNum,lyr,paramName):
            """ Function to check which version a layer is referencing
            This is used in the updateMessages validation to ensure
            that the user is editing a personal version """
            # import os
            version = getVersionForMapLayerFromLyr(lyr)
            validVersion_TF = True
            if 'TVA\\' not in version:
                parameters[pNum[paramName]].setErrorMessage('Editing {} version. Switch to personal version (CRMS_Editing_{}).'.format(version,os.getenv('username')))
                validVersion_TF = False
            return parameters,validVersion_TF

    def csvToList(xlsxPath,sheetName=None):
        """ Function for converting an XLSX worksheet into corresponding columnList and rows (list of value lists)
        :param str xlsxPath - path to the XLSX file
        :param str sheetName - optional - name of sheet in XLSX file; if empty, will default to first sheet
        :return tuple - two lists - one for column names and one for row values
        """
        import pandas
        if sheetName:
            df = pandas.read_csv(xlsxPath,sheet_name=sheetName,keep_default_na=False)
        else:
            df = pandas.read_csv(xlsxPath,keep_default_na=False)
        columns = df.columns.tolist()
        df1 = df.where(pandas.notnull(df), None)
        rows = df1.values.tolist()
        return columns,rows

    def getProjectLeadContactDict(PROJECT_LEAD_CONTACTS_CSV):
        """ Function to generate dictionary of contact information. Return format:
        {'WELLS_EDWARD_W':('Edward W. Wells','ewwells@tva.gov','865-632-2259'),
        'WELLS, EDWARD W.':('Edward W. Wells','ewwells@tva.gov','865-632-2259')
        """

        columns,rows = csvToList(xlsxPath=PROJECT_LEAD_CONTACTS_CSV)
        colDict = {col:i for i,col in enumerate(columns)}
        PROJECT_LEAD_CONTACTS_DICT = {}
        for row in rows:
            # ['name', 'label', 'reportname', 'email']
            PROJECT_LEAD_CONTACTS_DICT[row[colDict['name']]] = (row[colDict['reportname']],row[colDict['email']])
            PROJECT_LEAD_CONTACTS_DICT[row[colDict['label']]] = (row[colDict['reportname']],row[colDict['email']])
        return PROJECT_LEAD_CONTACTS_DICT

    def shutilCopy2WithWarning(inputFile,outputFile):
        try:
            shutil.copy2(inputFile,outputFile)
        except PermissionError as e:
            wrn('Output File already open: {}\nPlease close and try again.'.format(outputFile))

    def getAGOLContactsCSV(itemID):
        ''' Function to open the XLSX configuration for a Survey123 Form
        csvFile = getAGOLContactsCSV('87101c0fcc86444491c4356ff86f8145') '''
        import arcgis
        from arcgis import GIS

        gis = GIS("pro",verify_cert=False)
        item = gis.content.get(itemID)
        csvFile = item.download()
        return csvFile

    def getSettingDocumentsForExport(settingDocumentFolder):
        """
        Function to return setting documents
        envSettingDocList,preContactDocList,prehistoricDocList,historicDocList,soilDocList = getSettingDocumentsForExport(settingDocumentFolder)
        """
        settingDocuments = getFileListWithSubfolders(settingDocumentFolder,wildcard='.pdf')
        envSettingDocList = [x for x in settingDocuments if 'Environmental_Setting' in x]
        preContactDocList = [x for x in settingDocuments if 'Pre-contact_Setting' in x]
        prehistoricDocList = [x for x in settingDocuments if 'Prehistoric_Setting' in x]
        historicDocList = [x for x in settingDocuments if 'Historic_Setting' in x]
        soilDocList = [x for x in settingDocuments if 'Soil' in x]

        return envSettingDocList,preContactDocList,prehistoricDocList,historicDocList,soilDocList

    def returnAllLayerObjectsFromAGOLItem(gis,item_id='',itemTitle='',relativeLyrNum=0,layerName=None):
        """ Function that searches AGOL / Portal for the Feature Service item and returns the layer object for the specified relative layer number
        inputs:
            gis - GIS session for Portal or AGOL
            item_id - the item ID for the feature layer
            itemTitle - the title of the item
            relativeLyrNum - the relative layer number within the feature service
            layerName - the name of the layer within the feature service

        # you must specify either an item id or item title
        # you must specify either a relative layer number or layer name
            """
        #search_result = GISConnection_.content.search(f"title:{itemTitle_}", item_type = "Feature Service")
        if item_id: # get item directly from item id
            item = gis.content.get(item_id)
        else: # get item by searching for item title
            search_result = gis.content.search(query="title:{}".format(itemTitle),item_type = "Feature Layer", max_items=1000)
            print(search_result)
            if len(search_result) == 1:
                item = search_result[0]
            else:
                err('No Item with {} Title'.format(itemTitle))
                item = None
        if item:
            msg('Item: {}'.format(item.title))
            lyrs = item.layers + item.tables
            msg('Layers: {}'.format([lyr.properties.name for lyr in lyrs]))
            return lyrs
        else:
            return None

    def compare_GIS_and_MT_Resources(gis_dict,gis_fields,gis_wkid,managed_task_dict,mt_fields,mt_wkid):
        agol_different_values_dict = {}
        changed_GlobalIDs = []
        unchanged_GlobalIDs = []
        sync_changed_GlobalIDs = []
        for gis_GlobalID,gisValueDict in gis_dict.items():
            msg('---')
            different_values = []
            gis_GlobalID = gis_dict[gis_GlobalID]['globalid'].upper()

            common_fields = [field for field in gis_fields if field.lower() in mt_fields and field.lower() not in ('objectid','creationdate','creator','editdate','editor','globalid','shape@json','shape')]
            msg('common_fields: {}'.format(common_fields))

            if 'shape' in mt_fields:
                gis_geo_json = gis_dict[gis_GlobalID]['SHAPE@JSON']
                gis_geo = json.loads(gis_geo_json) if gis_geo_json else None
                gis_geo_rings = gis_geo['rings'] if gis_geo else [[]]

                mt_geo_json = managed_task_dict[gis_GlobalID]['SHAPE@JSON']
                mt_geo = json.loads(mt_geo_json) if mt_geo_json else None
                mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

                gis_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in gis_geo_rings for coord in ring]
                mt_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in mt_geo_rings for coord in ring]

                #mt_coord_lists = [sorted([convert_coordinates(coord[0],coord[1],source_WKID=mt_wkid,target_WKID=gis_wkid,transformation='') for ring in mt_geo_rings for coord in ring])]

                if sorted(gis_coord_lists) != sorted(mt_coord_lists):
                    msg('Different Geometry for {}'.format(gis_GlobalID))
                    msg(f'{gis_coord_lists=}\n{mt_coord_lists=}')
                    different_values += [('Geometry','','')]
                else:
                    msg('Same Geometry for {}'.format(gis_GlobalID))

            different_attribute_values = [(field,gis_dict[gis_GlobalID][field],managed_task_dict[gis_GlobalID][field]) for field in common_fields if tryUpper(removeBrackets(gis_dict[gis_GlobalID][field])) != tryUpper(removeBrackets(datetimeUTCToEpochMilli(managed_task_dict[gis_GlobalID][field])))] # must attempt to convert datetime to epoch in order to compare feature class to AGOL values; must convert globalids to uppercase; must remove '{' and '}' for parentglobalid

            if different_attribute_values:
                msg('Different Attributes for {}'.format(gis_GlobalID))
            else:
                msg('Same Attributes for {}'.format(gis_GlobalID))

            different_values += different_attribute_values

            if different_values:
                msg(different_values)
                msg('Different Values for {}:\nField\tGIS\tManaged Task\n{}'.format(gis_GlobalID,"\n".join(['{}\t{}\t{}'.format(field,gis,mt_value) for field,gis,mt_value in different_values])))
                agol_different_values_dict[gis_GlobalID] = different_values
                changed_GlobalIDs.append(gis_GlobalID)
                # determine if field differences apply when syncing with GIS
                # if sync_gis_fields:
                #     sync_value_changes = [x for x in different_values if x[2] in sync_gis_fields]
                #     if sync_value_changes:
                #         sync_changed_GlobalIDs.append(gis_GlobalID)
            else:
                unchanged_GlobalIDs.append(gis_GlobalID)
                msg('Same Values for {}'.format(gis_GlobalID))

        sync_changed_GlobalIDs = list(set(sync_changed_GlobalIDs))
        return agol_different_values_dict

    def compare_AGOL_and_MT_Resources(agol_features,managed_task_dict,mt_fields,sync_gis_fields):
        agol_different_values_dict = {}
        changed_GlobalIDs = []
        unchanged_GlobalIDs = []
        sync_changed_GlobalIDs = []
        for agol_feature in agol_features:
            msg('---')
            different_values = []
            agol_globalid = agol_feature.attributes['globalid'].upper() # convert to upper in order to compare with MT and GIS

            common_fields = [field for field in agol_feature.attributes.keys() if field.lower() in mt_fields and field.lower() not in ('objectid','creationdate','creator','editdate','editor','globalid')]

            if 'shape' in mt_fields:
                agol_geo = agol_feature.geometry
                # if agol_geo:
                #     agol_geo_rings = agol_geo['rings']
                agol_geo_rings = agol_geo['rings'] if agol_geo else [[]]
                mt_geo_json = managed_task_dict['{' + agol_globalid + '}']['SHAPE@JSON']
                mt_geo = json.loads(mt_geo_json) if mt_geo_json else None
                # if mt_geo:
                #     mt_geo_rings = mt_geo['rings']
                mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

                agol_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in agol_geo_rings for coord in ring]
                mt_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in mt_geo_rings for coord in ring]

                if sorted(agol_coord_lists) != sorted(mt_coord_lists):
                    msg('Different Geometry for {}'.format(agol_globalid))
                    msg(f'{agol_coord_lists=}\n{mt_coord_lists=}')
                    different_values += [('','','Geometry','','')]
                    if sync_gis_fields:
                        sync_changed_GlobalIDs.append(agol_globalid)
                else:
                    msg('Same Geometry for {}'.format(agol_globalid))

            if 'resource_id' in common_fields:
                agol_resourceID = agol_feature.attributes['resource_id']
                if agol_resourceID:
                    agol_resourceID = agol_resourceID.upper()
            else:
                agol_resourceID = ''

            if 'res_state_id' in common_fields:
                agol_res_state_id = agol_feature.attributes['res_state_id']
            else:
                agol_res_state_id = ''

            different_attribute_values = [(agol_resourceID,agol_res_state_id,field,agol_feature.attributes[field],managed_task_dict['{' + agol_globalid + '}'][field]) for field in common_fields if tryUpper(removeBrackets(agol_feature.attributes[field])) != tryUpper(removeBrackets(datetimeUTCToEpochMilli(managed_task_dict['{' + agol_globalid + '}'][field])))] # must attempt to convert datetime to epoch in order to compare feature class to AGOL values; must convert globalids to uppercase; must remove '{' and '}' for parentglobalid

            if different_attribute_values:
                msg('Different Attributes for {}'.format(agol_globalid))
            else:
                msg('Same Attributes for {}'.format(agol_globalid))

            different_values += different_attribute_values

            if different_values:
                msg('Different Values for {}:\nField\tAGOL\tManaged Task\n{}'.format(agol_globalid,"\n".join(['{}\t{}\t{}\t{}\t{}'.format(resourceID,resStateID,field,agol_value,mt_value) for resourceID,resStateID,field,agol_value,mt_value in different_values])))
                agol_different_values_dict[agol_globalid] = different_values
                changed_GlobalIDs.append(agol_globalid)
                # determine if field differences apply when syncing with GIS
                if sync_gis_fields:
                    sync_value_changes = [x for x in different_values if x[2] in sync_gis_fields]
                    if sync_value_changes:
                        sync_changed_GlobalIDs.append(agol_globalid)
            else:
                unchanged_GlobalIDs.append(agol_globalid)
                msg('Same Values for {}'.format(agol_globalid))

        sync_changed_GlobalIDs = list(set(sync_changed_GlobalIDs))
        agol_different_values_dict = {addBrackets(k):v for k,v in agol_different_values_dict.items()}
        return agol_different_values_dict,sync_changed_GlobalIDs

    def updateContractorAndActivityLists(contractor,activityName,gis,CONTRACTOR_AND_ACTIVITY_NAMES_CSV_LIST_NAME=CONTRACTOR_AND_ACTIVITY_NAMES_CSV_LIST_NAME,CONTRACTOR_CSV_LIST_NAME=CONTRACTOR_CSV_LIST_NAME):
        """
        Function to add new Contractor and Activity Name to the AGOL Picklist for Managed Task work
        """
        import pandas
        contractor_success_TF = True
        contractor_activity_success_TF = True
        msg(CONTRACTOR_AND_ACTIVITY_NAMES_CSV_LIST_NAME)
        msg(f'Checking for {contractor} and {activityName} in Contractor and Activity Name List')
        contractorAndActivityListItem = getItemFromTitle(gis=gis,title=CONTRACTOR_AND_ACTIVITY_NAMES_CSV_LIST_NAME,item_type="CSV")
        contractorAndActivityListFile = contractorAndActivityListItem.download()
        df = pandas.read_csv(contractorAndActivityListFile)
        df['contractorAndActivity'] = df['contractor'] + ';' + df['name']
        existingValuesList = df.values.tolist()
        dfColumns = df.columns.tolist()
        existingContractors = list(df.contractor.unique())
        existingContractorActivities = list(df.contractorAndActivity.unique())
        if f'{contractor};{activityName}' in existingContractorActivities:
            msg(f"'{contractor}' and '{activityName}' already exist in Contractor and Activity Name List.")
        else:
            newValuesDict = {'contractor':contractor,'name':activityName,'label':activityName,'contractorAndActivity': contractor + ';' + activityName}
            newRow = [newValuesDict[field] for field in dfColumns]
            #df = df.append(newRow, ignore_index=True)
            existingValuesList.append(newRow)
            df2 = pandas.DataFrame(existingValuesList)
            df2.columns = dfColumns
            df2.drop_duplicates(inplace=True)
            df2.sort_values(['contractor','name'],inplace=True)
            df2.to_csv(contractorAndActivityListFile,index=False)
            msg(f"Adding '{contractor}' and '{activityName}' to Contractor and Activity Name List...")
            try:
                result = contractorAndActivityListItem.update({}, contractorAndActivityListFile)
                if result != True:
                    wrn(f'Error updating CSV: {result}')
                    contractor_activity_success_TF = False
            except Exception as e:
                error_msg = str(e)
                if error_msg.startswith('You do not have permissions to access this resource or perform this operation.'):
                    wrn(f'Error updating CSV: {e}')
                    contractor_activity_success_TF = False
                else:
                    raise Exception(f'Error updating CSV: {e}')
        if contractor in existingContractors:
            msg(f"'{contractor}' already exists in Contractor List.")
        else:
            # also add to Contractor List
            contractorListItem = getItemFromTitle(gis=gis,title=CONTRACTOR_CSV_LIST_NAME,item_type="CSV")
            contractorListFile = contractorListItem.download()
            df = pandas.read_csv(contractorListFile)
            existingValuesList = df.values.tolist()
            dfColumns = df.columns.tolist()
            newValuesDict = {'name':contractor,'label':contractor}
            newRow = [newValuesDict[field] for field in dfColumns]
            #df2 = df2.append(newRow, ignore_index=True)
            existingValuesList.append(newRow)
            df2 = pandas.DataFrame(existingValuesList)
            df2.columns = dfColumns
            df2.drop_duplicates(inplace=True)
            df2.sort_values(['name'],inplace=True)
            df2.to_csv(contractorListFile,index=False)
            msg(f"Adding '{contractor}' to Contractor List...")
            try:
                result = contractorListItem.update({}, contractorListFile)
                if result != True:
                    wrn(f'Error updating CSV {CONTRACTOR_CSV_LIST_NAME}: {result}')
                    contractor_success_TF = False
            except Exception as e:
                error_msg = str(e)
                if error_msg.startswith('You do not have permissions to access this resource or perform this operation.'):
                    wrn(f'Error updating CSV {CONTRACTOR_CSV_LIST_NAME}: {e}')
                    contractor_success_TF = False
                else:
                    raise Exception(f'Error updating CSV {CONTRACTOR_CSV_LIST_NAME}: {e}')
        return contractor_success_TF, contractor_activity_success_TF

    def savePickle(object, filepath):
        import pickle
        if not filepath.endswith('.pkl'): # ensure file has proper extension
            filepath = filepath + '.pkl'
        with open(filepath, 'wb') as f:
            pickle.dump(object, f, protocol=2)

    def openPickle(filepath):
        import pickle
        with open(filepath, 'rb') as f:
            object = pickle.load(f)
        return object

    def getItemFromTitle(gis,title,item_type="Feature Layer"):
        items = gis.content.search(title,item_type,max_items=1000)
        items = [item for item in items if item.title == title]
        if len(items) != 1:
            wrn(f'Unable to uniquely find "{title}": {items}')
            return
        return items[0]

    def getFeatureLayerIDFromFormItem(form_item):
        related_items = form_item.related_items(rel_type='Survey2Service',direction='forward')
        if len(related_items) == 1:
            related_item = related_items[0]
            return related_item.id
        elif len(related_items) > 1:
            wrn('More than 1 related feature layer for form id {}'.format(form_item))
            return None
        else:
            wrn('No related feature layer for form id {}'.format(form_item))
            return None

def getLyrGlobalIDandFeaturesDict(lyrs,globalIDList,formType='Resource'):
    globalIDQuery = "globalid IN ('{}')".format("','".join(globalIDList))
    parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(globalIDList))
    globalIDobjectIDDict = {}
    lyrGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
    lyrMissingGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
    lyrFeaturesDict = {lyr.properties.id:[] for lyr in lyrs}
    parentGlobalIDDict = {} # dictionary to store relationship between related record globalid and the parent globalid
    baseGlobalIDDict = {} # dictionary to store relationship between related record globalid and the base polygon feature globalid

    outputFields = "*"
    returnGeometry = True
    if not globalIDList:
        return lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict, baseGlobalIDDict

    full_globalid_list = []
    base_globalid_list = []
    lyrsToReprocess = [] # variable to hold the table layers that are not directly related to the polygon feature (i.e. related to a related table)
    maxNon0Rel = 0 # variable to hold the number of relationships that are not directly with the polygon feature
    for i,lyr in enumerate(lyrs):
        lyrId = lyr.properties.id
        lyrName = lyrs[lyrId].properties.name
        print(f'\n{lyrId}: {lyrName}')
        if i == 0:
            agol_features = getFeaturesFromLyr(lyr=lyr, query=globalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
            #msg(agol_features)
            gid_list = [x.attributes['globalid'] for x in agol_features]
            globalIDobjectIDDict.update({x.attributes['globalid']:x.attributes['objectid'] for x in agol_features})
            msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
            if formType == 'Activity' and not gid_list:
                err(f'No ArcGIS Online record found for Activity ID {activityID}.')
                exit()

            if len(globalIDList) != len(gid_list):
                missing = [x for x in globalIDList if x.lower().strip('{').strip('}') not in gid_list]
                err(f'{len(missing)} not found in ArcGIS Online: {missing}')
                msg(f'{globalIDList=}')
                msg(f'{gid_list=}')
                lyrMissingGlobalIDDict[lyrId] += missing
                #if formType == 'Resource':
                    #exit()
            parentGlobalIDDict.update({gid:gid for gid in gid_list})
            full_globalid_list += gid_list
            base_globalid_list += gid_list
            lyrGlobalIDDict[lyrId] += gid_list
            lyrFeaturesDict[lyrId] += agol_features
            continue
        relationships = lyr.properties['relationships']
        #wrn(f'Relationships: {relationships}')
        non0Rel = [x for x in relationships if x['relatedTableId'] != 0] # get list of layers that have relationships other than with the main polygon features
        if non0Rel:
            lyrsToReprocess.append(lyr)

        agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
        gid_list = [x.attributes['globalid'] for x in agol_features]
        globalIDobjectIDDict.update({x.attributes['globalid']:x.attributes['objectid'] for x in agol_features})
        if i != 0:
            parentGlobalIDDict.update({x.attributes['globalid']:x.attributes['parentglobalid'] for x in agol_features})
        msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
        full_globalid_list += gid_list
        lyrGlobalIDDict[lyrId] += gid_list
        lyrFeaturesDict[lyrId] += agol_features

        if lyrsToReprocess:
            found_new_TF = True
            pass_count = 1
            while found_new_TF:
                pass_count+=1
                msg(f'Making Pass {pass_count} for related records...')
                found_new_TF = False
                for lyr in lyrsToReprocess:
                    lyrId = lyr.properties.id
                    newParentGlobalIDList = []
                    newParentGlobalIDList += full_globalid_list
                    if not newParentGlobalIDList:
                        continue
                    parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(newParentGlobalIDList))
                    agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                    gid_list = [x.attributes['globalid'] for x in agol_features]
                    new_gid_list = [x for x in gid_list if x not in full_globalid_list]
                    if not new_gid_list:
                        continue
                    msg(f'{len(new_gid_list)} found during pass {pass_count} on layer {lyrId} - {lyr.properties.name}')
                    full_globalid_list += new_gid_list
                    lyrGlobalIDDict[lyrId] += new_gid_list
                    found_new_TF = True
                    lyrFeaturesDict[lyrId] += agol_features

    for globalid, parent_globalid in parentGlobalIDDict.items():
        max_iter = 20
        i = 0
        if parent_globalid in base_globalid_list:
            baseGlobalIDDict[globalid] = parent_globalid
            continue
        while parent_globalid not in base_globalid_list:
            wrn(f'Searching for {parent_globalid} for {globalid}')
            parent_globalid = parentGlobalIDDict[parent_globalid]
        baseGlobalIDDict[globalid] = parent_globalid  
    
    return lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict, baseGlobalIDDict, globalIDobjectIDDict



### TOOL CLASSES ###
''' Export Resources for Managed Task Activities '''
class exportResourcesForManagedTask(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "Export Resources for Managed Task Activities"
        self.description = "Tool to Export Resources for Managed Task Activities"
        self.canRunInBackground = False
        self.category = "Managed Task Export"

    def getParameterInfo(self):
        """ Export Resource Form Templates for Managed Task """
        """
            Defines the tool's parameters
        """
        resourceLyrParam = arcpy.Parameter(
            displayName='Resource Layer',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        activityLyrParam = arcpy.Parameter(
            displayName='Cultural Activities Layer',
            name='activityLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Optional',
            direction='Input')
        activityLyrParam.value = activityLyrParamDefault # set default cultural activity layer

        activityLyrParam = arcpy.Parameter(
            displayName='Cultural Activities Layer',
            name='activityLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Optional',
            direction='Input')
        activityLyrParam.value = activityLyrParamDefault # set default cultural activity layer

        activityTypeParam = arcpy.Parameter(
            displayName='Activity Type',
            name='activityType',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        activityTypeParam.filter.type = 'ValueList' # sets input type in tool
        activityTypeParam.filter.list = activityTypeDomainList
        # activityTypeParam.value = 'Test Activity'

        activityNameParam = arcpy.Parameter(
            displayName='Activity Name',
            name='activityName',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        #activityNameParam.value = 'Test Activity'

        activityIDParam = arcpy.Parameter(
            displayName='Activity ID',
            name='activityID',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        #activityIDParam.value = 'Test Activity'

        projectIDParam = arcpy.Parameter(
            displayName='Project ID',
            name='projectID',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        #projectIDParam.value = 'Test Project'

        resourceTypesParam = arcpy.Parameter(
            displayName='Resource Types',
            name='resourceTypes',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input',
            multiValue=True)
        resourceTypesParam.filter.type = 'ValueList'
        resourceTypesParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])
        #resourceTypesParam.values = list(sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys() if k not in ('Interment','Reinterment')]))

        verifyInterment_TFParam = arcpy.Parameter(
            displayName='Only authorized staff can export Interments and Reinterments. Check box to confirm.',
            name='verifyInterment_TF', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        verifyInterment_TFParam.enabled = False

        # downloadGDB_TFParam = arcpy.Parameter(
        #     displayName='Download Geodatabase?',
        #     name='downloadGDB_TF', ### Name the same as parameter in main script
        #     datatype='GPBoolean',
        #     parameterType='Optional',
        #     direction='Input')
        # # set to true by default
        # downloadGDB_TFParam.value = True

        mtProjectNameParam = arcpy.Parameter(
            displayName='Managed Task Project Name',
            name='mtProjectName',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')

        """
        outputFolderParam = arcpy.Parameter(
            displayName='Output Folder',
            name='outputFolder', ### Name the same as parameter in main script
            datatype='DEFolder',
            parameterType='Optional',
            direction='Input')
        """

        retainResourceDistanceParam = arcpy.Parameter(
            displayName='Buffer Distance for Exporting Resources',
            name='retainResourceDistance',
            datatype='GPLinearUnit',
            parameterType='Optional',
            direction='Input')
        retainResourceDistanceParam.value = retainResourceDistanceParamDefault # '0.5 Miles' # set default maximum distance for viewshed

        resourceSelectionBufferLyrParam = arcpy.Parameter(
            displayName='Resource Selection Buffer Layer',
            name='resourceSelectionBufferLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Optional',
            direction='Input')

        overwriteExports_TFParam = arcpy.Parameter(
            displayName='Overwrite Previous Exports for this Activity?',
            name='overwriteExports_TF', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        # set to true by default
        overwriteExports_TFParam.value = True

        emailAddressParam = arcpy.Parameter(
            displayName='Email Address (optional)\nTo be notified when process completes',
            name='emailAddress',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        emailAddressParam.value = '' # @tva.gov'

        contractorParam = arcpy.Parameter(
            displayName='Managed Task Contractor',
            name='contractor',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        contractorParam.filter.type = 'ValueList'
        contractorParam.filter.list = ['TVAR','New South Associates'] # TODO - get from AGOL CSV?

        params = [resourceLyrParam,activityLyrParam,activityTypeParam,activityNameParam,activityIDParam,projectIDParam,resourceTypesParam,verifyInterment_TFParam,retainResourceDistanceParam,resourceSelectionBufferLyrParam,mtProjectNameParam,overwriteExports_TFParam,emailAddressParam,contractorParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Export Resource Form Templates for Managed Task """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)
        # resourceType = pValues['resourceType']

        projectID = pValues['projectID']
        activityID = pValues['activityID']
        activityLyr = pValues['activityLyr']
        oidList = getOIDListFromLyrSelection(activityLyr)
        emailAddress = pValues['emailAddress']

        # must have only one activity selected
        if len(oidList) != 1:
            return

        if not emailAddress:
            gis = connectToAGOL()
            parameters[pNum['emailAddress']].value = gis.users.me.email

        if activityLyr:
            activityName,activityType,activityID,projectID = returnMultipleDistinctFieldValuesForOneRecord(activityLyr,[activityNameField,activityTypeField,activityIDField,projectIDField],whereClause='')
            parameters[pNum['activityName']].value = activityName
            parameters[pNum['activityType']].value = activityType
            parameters[pNum['activityID']].value = activityID
            parameters[pNum['projectID']].value = projectID

            pValuesAsText = paramDictValuesAsText(parameters)
            resourceTypes = convertParameterMultiValueTextToList('resourceTypes',pValuesAsText)
            if not resourceTypes:
                archaeologicalSurveyResourceTypes = ['Archaeological_Resources','Cemeteries','Other']
                historicArchitecturalSurveyResourceTypes = ['Cemeteries','Cultural_Landscapes','Districts_and_Facilities','Historic_Architectural_Resources','Objects','Other']
                if activityType == 'Archaeological_Survey':
                    parameters[pNum['resourceTypes']].values = archaeologicalSurveyResourceTypes
                elif activityType == 'Historic_Architectural_Survey':
                    parameters[pNum['resourceTypes']].values = historicArchitecturalSurveyResourceTypes


        mtProjectName = pValues['mtProjectName']
        if not mtProjectName and activityName:
            parameters[pNum['mtProjectName']].value = activityName.replace(' ','_')
            mtProjectName = activityName.replace(' ','_')

        """
        # populate outputFolder based on activityFolder
        outputFolder = pValuesAsText['outputFolder']
        if activityLyr and mtProjectName and activityID and projectID and not outputFolder:
            projectFolder = getProjectFolderFromProjectIDAndParentProjectFolder(PARENT_PROJECT_FOLDER,projectID)
            if projectFolder:
                activityFolder = createFolder(os.path.join(projectFolder,activityID))
                parameters[pNum['outputFolder']].value = activityFolder
        """

        resourceTypes = convertParameterMultiValueTextToList('resourceTypes',pValuesAsText)
        if 'Interment' in resourceTypes or 'Reinterment' in resourceTypes:
            parameters[pNum['verifyInterment_TF']].enabled = True
        else:
            parameters[pNum['verifyInterment_TF']].enabled = False

        return

    def updateMessages(self, parameters):
        """ Export Resource Form Templates for Managed Task """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        pValuesAsText = paramDictValuesAsText(parameters)

        gisURL = checkGIS_URL()
        if gisURL != TVA_AGOL_URL:
            parameters[0].setErrorMessage('Currently signed into {} Portal. Sign in to ArcGIS Online: {}'.format(gisURL,TVA_AGOL_URL))
            return

        activityLyr = pValues['activityLyr']

        a_oidList = getOIDListFromLyrSelection(activityLyr)
        if not a_oidList:
            parameters[pNum['activityLyr']].setErrorMessage('No Activity feature selected.')
        elif len(a_oidList) > 1:
            parameters[pNum['activityLyr']].setErrorMessage('More than 1 Activity feature selected.')

        resourceTypes = convertParameterMultiValueTextToList('resourceTypes',pValuesAsText)
        resourceSelectionBufferLyr = pValues['resourceSelectionBufferLyr']
        resourceWarningTxt = ''
        if activityLyr and resourceTypes:
            aprxMap = returnMapObject()
            retainResourceDistance = pValuesAsText['retainResourceDistance']

            if not resourceSelectionBufferLyr or not parameters[pNum['retainResourceDistance']].hasBeenValidated:
                # bufferLyrs = [x for x in lyrList if x == lyrName]
                lyrName = 'Resource_Selection_Buffer'
                bufferPath= os.path.join(arcpy.env.scratchGDB,lyrName)
                bufferLyr = arcpy.analysis.Buffer(
                    in_features=activityLyr,
                    out_feature_class=bufferPath,
                    buffer_distance_or_field=retainResourceDistance,
                    line_side="FULL",
                    line_end_type="ROUND",
                    dissolve_option="NONE",
                    dissolve_field=None,
                    method="PLANAR"
                )
                addDatasetToAPRX(aprxMap_=aprxMap,dataset=bufferPath,layerName_=lyrName,layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery='')
                parameters[pNum['resourceSelectionBufferLyr']].value = lyrName

            if not parameters[pNum['resourceTypes']].hasBeenValidated or not parameters[pNum['retainResourceDistance']].hasBeenValidated:
                #lyrList = getValidLayers(aprxMap)
                resourceLyr = pValues['resourceLyr']
                #arcpy.management.SelectLayerByLocation(resourceLyr, 'INTERSECT', activityLyr, retainResourceDistance)
                # r_oidList = getOIDListFromLyrSelection(resourceLyr)
                # r_count = arcpy.management.GetCount(resourceLyr)
                lyrList = getValidLayers(aprxMap,searchLyrName=resourceLyr)
                rLyr = lyrList[0]
                #arcpy.SelectLayerByLocation_management(rLyr,'INTERSECT',activityLyr,retainResourceDistance)
                #arcpy.SelectLayerByLocation_management(rLyr,'WITHIN_A_DISTANCE',activityLyr,retainResourceDistance)
                if not resourceSelectionBufferLyr:
                    resourceSelectionBufferLyr = bufferLyr
                arcpy.SelectLayerByLocation_management(rLyr,'INTERSECT',resourceSelectionBufferLyr)
                # r_oidList = getOIDListFromLyrSelection(rLyr)
                # lyrList = getValidLayers(aprxMap)
                resourceTypeList = convertParameterMultiValueTextToList('resourceTypes',pValuesAsText)
                resourceTypeQuery = "resource_type in ('{}')".format("','".join(resourceTypeList))
                arcpy.SelectLayerByAttribute_management(rLyr,'SUBSET_SELECTION',where_clause=resourceTypeQuery)

                r_oidList = getOIDListFromLyrSelection(rLyr)
                if r_oidList:
                    with arcpy.da.SearchCursor(rLyr,['RESOURCE_TYPE','RESOURCE_ID','FORM_RECORD_GLOBALID']) as cursor:
                        for row in cursor:
                            if not row[2]:
                                resourceWarningTxt+=f'\n{row[0]} with Resource ID {row[1]} has no FORM RECORD GLOBALID. Submit Resource to CRMS or exclude from MT Export Resource Selection'
                if resourceWarningTxt:
                    parameters[pNum['resourceLyr']].setWarningMessage(resourceWarningTxt)
                # parameters[pNum['resourceLyr']].setWarningMessage('{},{} resources to be selected using {} radius'.format(len(r_oidList),r_count,retainResourceDistance))

        verifyInterment_TF = pValues['verifyInterment_TF']
        if ('Interment' in resourceTypes or 'Reinterment' in resourceTypes) and not verifyInterment_TF:
            parameters[pNum['verifyInterment_TF']].setErrorMessage('Only authorized staff can export Interments and Reinterments. Check box to confirm.')

        emailAddress = pValues['emailAddress']
        if emailAddress == 'egis@tva.gov':
            parameters[pNum['emailAddress']].setWarningMessage('Update email address if notification is desired.')

        """
        projectID = pValues['projectID']
        if projectID:
            projectFolder = getProjectFolderFromProjectIDAndParentProjectFolder(PARENT_PROJECT_FOLDER,projectID)
            if not projectFolder:
                parameters[pNum['outputFolder']].setErrorMessage('No folder found for ProjectID {} in {}'.format(projectID,PARENT_PROJECT_FOLDER))
        """

        return

    def execute(self, parameters, messages):
        """ Export Resource Form Templates for Managed Task """
        # Create GIS Connection to AGOL / Portal
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        #//outputFolder = r'{}'.format(pValuesAsText['outputFolder'])
        #// create output folder if it doesn't exist
        #//createFolder(outputFolder)
        #//mtImportFolder = createFolder(os.path.join(outputFolder,'Import_from_MT')) # go ahead and create folder to store returned geodatabase
        #//outputFolderName = os.path.basename(outputFolder)
        #//tempOutputFolder = createFolder(r'C:\Users\{}\Documents\CRMS_Managed_Task\{}'.format(pcUser,outputFolderName))
        #//mtExportFolder = createFolder(os.path.join(tempOutputFolder,'Export_to_MT'))
        #//finalMTExportFolder = mtExportFolder.replace(tempOutputFolder,outputFolder)

        #// msg(f'{outputFolder=}')
        #// msg(f'{tempOutputFolder=}')
        #// msg(f'{mtExportFolder=}')
        #// msg(f'{finalMTExportFolder=}')
        #// msg(f'Would copy {mtExportFolder} to {finalMTExportFolder}')
        #// return



        #downloadGDB_TF = pValues['downloadGDB_TF']
        #updateDefaults_TF = pValues['updateDefaults_TF']
        resourceTypeList = convertParameterMultiValueTextToList('resourceTypes',pValuesAsText)

        start_time = datetime.datetime.now()
        aprxMap = returnMapObject()
        if not aprxMap:
            exit()

        gis = connectToAGOL()

        retainResourceRadius = pValuesAsText['retainResourceDistance'] # '0.5 Miles'
        resourceLyr = pValues['resourceLyr']
        activityLyr = pValues['activityLyr']
        activityName = pValues['activityName']
        activityType = pValues['activityType']
        activityID = pValues['activityID']
        projectID = pValues['projectID']
        mtProjectName = pValues['mtProjectName']
        mtProjectName = mtProjectName.replace(' ' ,'_')
        overwriteExports_TF = pValues['overwriteExports_TF']
        emailAddress = pValues['emailAddress']
        contractor = pValues['contractor']

        msg(f'{resourceTypeList=}')
        getCount = int(str(arcpy.management.GetCount(resourceLyr)))
        fidSetList = getOIDListFromLyrSelection(resourceLyr)
        msg(f'Resource layer count: {getCount}')
        msg(f'Resource layer selection FID Set: {len(fidSetList)}')


        #resourceSelection = arcpy.management.SelectLayerByLocation(resourceLyr, 'INTERSECT', activityLyr, retainResourceRadius, 'NEW_SELECTION')
        resourceTypeCountDict = {x:0 for x in resourceTypeList}
        resourceFormGlobalidListDict = {x:[] for x in resourceTypeList}
        #with arcpy.da.SearchCursor(resourceSelection,['resource_type','form_record_globalid'],"resource_type in ('{}')".format("','".join(resourceTypeList))) as cursor:
        if fidSetList:
            invalidResources = []
            invalidResourceTypes = set()
            with arcpy.da.SearchCursor(resourceLyr,['resource_type','form_record_globalid','resource_id'],"resource_type in ('{}')".format("','".join(resourceTypeList))) as cursor:
                for row in cursor:
                    resourceType = row[0]
                    formglobalid = row[1]
                    resourceID = row[2]
                    if not formglobalid:
                        wrn('No Form GlobalID for Resource ID "{}" in {}'.format(resourceID,resourceType))
                        invalidResources.append(f'{resourceID} ({resourceType})')
                        #wrn('No Form GlobalID for Resource ID "{}" in {}. Check to see if a Survey123 Record exists. If so, populate the Form Record Global ID in the GIS. If not, submit the Resource to create a Survey123 Record and then populate the Form Record Global ID in the GIS'.format(resourceID,resourceType))
                        continue
                    resourceTypeCountDict[resourceType]+=1
                    resourceFormGlobalidListDict[resourceType].append(formglobalid)
            if invalidResources:
                invalidResources = [str(x) for x in invalidResources]
                err('No Form GlobalID for {} Resource IDs ("{}"). Check to see if a Survey123 Record exists. If so, populate the Form Record Global ID in the GIS. If not, submit the Resource to create a Survey123 Record and then populate the Form Record Global ID in the GIS'.format(len(invalidResources),'","'.join(invalidResources)))
                #exit()
        msg('resourceTypeCountDict: {}'.format(resourceTypeCountDict))
        msg('resourceCount: {}'.format(sum(resourceTypeCountDict.values())))
        msg(f'{resourceFormGlobalidListDict=}')

        # get form record globalid for selected activity
        activityGlobalIDList = []
        with arcpy.da.SearchCursor(activityLyr,['form_record_globalid'],"activity_id = '{}'".format(activityID)) as cursor:
            activityGlobalIDList = [row[0] for row in cursor]
        msg(f'{activityGlobalIDList=}')

        if len(activityGlobalIDList) > 1:
            err(f'More than 1 ArcGIS Online record for Activity ID {activityID}: {activityGlobalIDList}.')
            exit()


        surveyExportList = []
        surveyExportList.append([activityType,'Activity'])
        # surveyExportList.append(['Management_Summary','Project'])
        for resourceType in resourceTypeList:
            surveyExportList.append([resourceType,'Resource'])

        new_item_id_dict = {}

        summaryTextList = []
        for surveyInfo in surveyExportList:
            formTitle,formType = surveyInfo
            starttime = datetime.datetime.now()
            msg('\nProcessing {} {}'.format(formTitle,formType))
            summaryTextList.append('\n{}'.format(formTitle))
            form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=formTitle,formType=formType,gis=gis)


            # xlsxPath = saveXLSXForSurvey123Form(form_item_id,gis=gis)
            # file_properties = {'title':os.path.basename(xlsxPath),'type':'Microsoft Excel','tags':'template'}
            # file_item = gis.content.add(data=xlsxPath,item_properties=file_properties)
            # agol_folder = gis.content.create_folder(folder='CRMS_Survey_Templates')
            # file_item.move(agol_folder)
            # new_item_id_dict[file_item.id] = file_item.title

            if formType == 'Resource':
                globalIDList = resourceFormGlobalidListDict[formTitle]
                resourceFeatureLayerItemID = feature_layer_item_id
            elif formType == 'Activity':
                globalIDList = activityGlobalIDList
                activityFeatureLayerItemID = feature_layer_item_id
                resourceFeatureLayerItemID = None


            lyrs = returnAllLayerObjectsFromAGOLItem(gis,item_id=feature_layer_item_id)

            def getLyrGlobalIDandFeaturesDict(lyrs,globalIDList,formType='Resource'):
                globalIDQuery = "globalid IN ('{}')".format("','".join(globalIDList))
                parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(globalIDList))
                lyrGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
                lyrMissingGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
                lyrFeaturesDict = {lyr.properties.id:[] for lyr in lyrs}
                parentGlobalIDDict = {} # dictionary to store relationship between related record globalid and the parent globalid
                baseGlobalIDDict = {} # dictionary to store relationship between related record globalid and the base polygon feature globalid

                outputFields = "*"
                returnGeometry = True
                if not globalIDList:
                    return lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict, baseGlobalIDDict

                full_globalid_list = []
                base_globalid_list = []
                lyrsToReprocess = [] # variable to hold the table layers that are not directly related to the polygon feature (i.e. related to a related table)
                maxNon0Rel = 0 # variable to hold the number of relationships that are not directly with the polygon feature
                for i,lyr in enumerate(lyrs):
                    lyrId = lyr.properties.id
                    lyrName = lyrs[lyrId].properties.name
                    print(f'\n{lyrId}: {lyrName}')
                    if i == 0:
                        agol_features = getFeaturesFromLyr(lyr=lyr, query=globalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                        #msg(agol_features)
                        gid_list = [x.attributes['globalid'] for x in agol_features]
                        msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
                        if formType == 'Activity' and not gid_list:
                            err(f'No ArcGIS Online record found for Activity ID {activityID}.')
                            exit()

                        if len(globalIDList) != len(gid_list):
                            missing = [x for x in globalIDList if x.lower().strip('{').strip('}') not in gid_list]
                            err(f'{len(missing)} not found in ArcGIS Online: {missing}')
                            msg(f'{globalIDList=}')
                            msg(f'{gid_list=}')
                            lyrMissingGlobalIDDict[lyrId] += missing
                            #if formType == 'Resource':
                                #exit()
                        parentGlobalIDDict.update({gid:gid for gid in gid_list})
                        full_globalid_list += gid_list
                        base_globalid_list += gid_list
                        lyrGlobalIDDict[lyrId] += gid_list
                        lyrFeaturesDict[lyrId] += agol_features
                        continue
                    relationships = lyr.properties['relationships']
                    #wrn(f'Relationships: {relationships}')
                    non0Rel = [x for x in relationships if x['relatedTableId'] != 0] # get list of layers that have relationships other than with the main polygon features
                    if non0Rel:
                        lyrsToReprocess.append(lyr)

                    agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                    gid_list = [x.attributes['globalid'] for x in agol_features]
                    if i != 0:
                        parentGlobalIDDict.update({x.attributes['globalid']:x.attributes['parentglobalid'] for x in agol_features})
                    msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
                    full_globalid_list += gid_list
                    lyrGlobalIDDict[lyrId] += gid_list
                    lyrFeaturesDict[lyrId] += agol_features

                    if lyrsToReprocess:
                        found_new_TF = True
                        pass_count = 1
                        while found_new_TF:
                            pass_count+=1
                            msg(f'Making Pass {pass_count} for related records...')
                            found_new_TF = False
                            for lyr in lyrsToReprocess:
                                lyrId = lyr.properties.id
                                newParentGlobalIDList = []
                                newParentGlobalIDList += full_globalid_list
                                if not newParentGlobalIDList:
                                    continue
                                parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(newParentGlobalIDList))
                                agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                                gid_list = [x.attributes['globalid'] for x in agol_features]
                                new_gid_list = [x for x in gid_list if x not in full_globalid_list]
                                if not new_gid_list:
                                    continue
                                msg(f'{len(new_gid_list)} found during pass {pass_count} on layer {lyrId} - {lyr.properties.name}')
                                full_globalid_list += new_gid_list
                                lyrGlobalIDDict[lyrId] += new_gid_list
                                found_new_TF = True
                                lyrFeaturesDict[lyrId] += agol_features

                for globalid, parent_globalid in parentGlobalIDDict.items():
                    max_iter = 20
                    i = 0
                    if parent_globalid in base_globalid_list:
                        baseGlobalIDDict[globalid] = parent_globalid
                        continue
                    while parent_globalid not in base_globalid_list:
                        wrn(f'Searching for {parent_globalid} for {globalid}')
                        parent_globalid = parentGlobalIDDict[parent_globalid]
                    baseGlobalIDDict[globalid] = parent_globalid  
                
                return lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict, baseGlobalIDDict

            lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict, baseGlobalIDDict = getLyrGlobalIDandFeaturesDict(lyrs,globalIDList,formType)
            msg('---')
            msg(f'{lyrGlobalIDDict=}')
            msg(f'{lyrMissingGlobalIDDict=}')
            #msg(lyrFeaturesDict)

            # compare agol features with unique globalid list to identify potential duplicates in AGOL
            full_list = []
            for i,lst in lyrFeaturesDict.items():
                full_list+=[x.attributes['globalid'] for x in lst]
                if len(full_list) != len(set(full_list)):
                    err(f'Duplicates detected in AGOL features for {formTitle}')

            if formTitle == 'Other':
                formTitle = 'Other_Resources'
            elif formTitle in ('Archaeological_Survey','Historic_Architectural_Survey'):
                formTitle = 'Managed_Task_Activities'

            savePickle(lyrFeaturesDict,filepath=fr'\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS_Development\Dev_Toolbox\MT_Testing\pickles\{formTitle}_Test.pkl')

            #lyrFeaturesDict = openPickle(r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS_Development\Dev_Toolbox\MT_Testing\pickles\HAR_Test.pkl")
            # create dictionary to store relationship between original local globalid and new globalid for new records
            orig_globalid_new_globalid_dict = {}

            # # for each related table, get new globalid and update parentglobalid as appropriate
            # orig_globalid_current_parentglobalid_dict = {}
            # orig_globalid_new_parentglobalid_dict = {}

            # # populate initial values for globalid to parentglobalid relationship
            # for lyrId, featuresList in lyrFeaturesDict.items():
            #     for feature in featuresList:
            #         if lyrId == 0:
            #             continue
            #         orig_globalid_current_parentglobalid_dict[feature.attributes['globalid']] = feature.attributes['parentglobalid']

            # def addFormRecord(f):
            #     import random
            #     return random.randint(0,10000)

            tableAliasDict = {'non_contributing_resources': 'Non Contributing Resources', 'resources_contributed_to': 'Resources Contributed To', 'significant_events': 'Significant Events', 'significant_persons': 'Significant Persons', 'resource_design_features': 'Resource Design Features', 'contributing_resources': 'Contributing Resources', 'additional_resources_group': 'Additional Resources Group', 'exterior_component_group': 'Exterior Component Group', 'dates_significant_repeat': 'Significant Dates'}

            mtFormTitle = formTitle.replace('_',' ') + ' - MT Master'
            msg(f'Updating {mtFormTitle}...')
            formItem = getItemFromTitle(gis,mtFormTitle,item_type='Form')
            if not formItem:
                wrn(f'No form {mtFormTitle}')
                continue
            featureLyrItemId = getFeatureLayerIDFromFormItem(formItem)
            featureLyrItem = gis.content.get(featureLyrItemId)
            lyrs = featureLyrItem.layers + featureLyrItem.tables
            #for lyrId, featuresList in lyrFeaturesDict.items():
            for lyrId, lyr in enumerate(lyrs):
                if lyrId == 0 and formType == 'Resource':
                    lyrName = FEATURE_LAYER_NAME_RESOURCE_TYPE_DICT.get(lyr.properties['name']).replace('_',' ') + ' Polygons'
                else:
                    lyrName = tableAliasDict.get(lyr.properties['name'])
                    if not lyrName:
                        lyrName = lyr.properties['name']
                summaryTextList.append(lyrName)
                featuresList = lyrFeaturesDict.get(lyrId)
                if not featuresList:
                    msg(f'No features for {lyrId}')
                    continue
                # set field names based on layer id
                contractorField = f'contractor_{lyrId}'
                origGlobalidField = f'orig_globalid_{lyrId}'
                origParentGlobalidField = f'orig_parentglobalid_{lyrId}'
                tvaAssignedActivityNameField = f'tva_assigned_activity_name_{lyrId}'
                tvaAssignedActivityIdField = f'tva_assigned_activity_id_{lyrId}'
                contractorReviewStatusField = 'contractor_review_status'
                resourceIdentificationStatusField = 'resource_identification_status'


                # TODO - check to see if records already exist for that orig_globalid and contractor and activity combination
                globalidList = [f.attributes['globalid'] for f in featuresList]
                globalidListString = "','".join(globalidList)
                if formType == 'Resource':
                    query = f"{contractorField} = '{contractor}' and {tvaAssignedActivityNameField} = '{activityName}' and {tvaAssignedActivityIdField} = '{activityID}' and {origGlobalidField} in ('{globalidListString}')"
                else:
                    query = f"activity_id = '{activityID}' and {contractorField} = '{contractor}' and {origGlobalidField} in ('{globalidListString}')"
                existingFeatures = getFeaturesFromLyr(lyr=lyr, query=query, outputFields=[origGlobalidField,'globalid'], returnGeometry=False)
                existingOrigGlobalidDict = {f.attributes[origGlobalidField]:f.attributes['globalid'] for f in existingFeatures}
                orig_globalid_new_globalid_dict.update(existingOrigGlobalidDict)
                msg(f'Globalid Mapping dictionary updated with existing records in layer {lyrId}')
                # do we want to overwrite or add duplicate
                # do we need to check all values and geometry to find out if there are any differences

                # report resources or activities that are missing
                missingGlobalidList = lyrMissingGlobalIDDict[lyrId]
                for globalid in missingGlobalidList:
                    summaryTextList.append(f"\tNot found in layer {lyrId}: {globalid.lower()}")

                # populate orig_globalid and orig_parentglobalid and contractor
                countAdded = 0
                for f in featuresList:
                    current_globalid = f.attributes['globalid']
                    if current_globalid in existingOrigGlobalidDict:
                        wrn(f'{current_globalid} already in layer {lyrId} with {contractor}')
                        if lyrId == 0:
                            if formType == 'Resource':
                                summaryText = f"\tAlready in layer {lyrId} with {contractor}: {f.attributes['globalid']}\t{f.attributes['resource_id']}\t{f.attributes['resource_name']}"
                            else:
                                summaryText = f"\tAlready in layer {lyrId} with {contractor}: {f.attributes['globalid']}\t{f.attributes['activity_id']}\t{f.attributes['activity_name']}"
                        else:
                            summaryText = f"\tAlready in layer {lyrId} with {contractor}: {f.attributes['globalid']}"
                        summaryTextList.append(summaryText)
                        continue
                    f.attributes[contractorField] = contractor
                    f.attributes[origGlobalidField] = f.attributes['globalid']
                    f.attributes[contractorReviewStatusField] = 'Not Started' # set initial status to Not Started
                    if formType == 'Resource':
                        f.attributes[tvaAssignedActivityNameField] = activityName
                        f.attributes[tvaAssignedActivityIdField] = activityID
                        if lyrId == 0:
                            f.attributes[resourceIdentificationStatusField] = 'Previously Identified' # set resource identification status to Previously Identified
                    if lyrId == 0:
                        new_globalid = addFormRecord(lyr,[f])
                        orig_globalid_new_globalid_dict[current_globalid] = new_globalid
                        countAdded+=1
                    else:
                        current_parentglobalid = f.attributes['parentglobalid']
                        f.attributes[origParentGlobalidField] = current_parentglobalid
                        new_parentglobalid = orig_globalid_new_globalid_dict.get(current_parentglobalid)
                        if not new_parentglobalid:
                            summaryText = f"\tNo resource found for parent globalid {current_parentglobalid} in layer {lyrId}"
                            summaryTextList.append(summaryText)
                            err(f'No new globalid found for {current_parentglobalid} in layer {lyrId}')
                            exit()
                        # update parent globalid
                        f.attributes['parentglobalid'] = new_parentglobalid
                        #del f.attributes['globalid']
                        new_globalid = addFormRecord(lyr,[f])
                        orig_globalid_new_globalid_dict[current_globalid] = new_globalid
                        countAdded+=1
                    if lyrId == 0:
                        if formType == 'Resource':
                            summaryText = f"\t{f.attributes['globalid']}\t{f.attributes['resource_id']}\t{f.attributes['resource_name']}"
                        else:
                            summaryText = f"\t{f.attributes['globalid']}\t{f.attributes['activity_id']}\t{f.attributes['activity_name']}"
                    else:
                        summaryText = f"\t{f.attributes['globalid']}"
                    summaryTextList.append(summaryText)


                msg(f'{countAdded} Records added to layer {lyrId}')
            msg(f'Process complete for {featureLyrItem.title}')

        # update AGOL Contractor and Activity Lists
        contractor_success_TF, contractor_activity_success_TF = updateContractorAndActivityLists(contractor,activityName,gis)

        # send email to notify EGIS to update the CSV List in ArcGIS Online
        emailAddresses = ['japarkhurst@tva.gov']
        if emailAddress and emailAddress != '@tva.gov':
            emailAddresses.append(emailAddress)
        subjectText = 'Managed Task Export Complete for {}'.format(activityID)
        contractorActivityNameLink = 'https://tva.maps.arcgis.com/home/item.html?id=2270f2491b3a4d37813640dcdc73b255'
        contractorListLink = 'https://tva.maps.arcgis.com/home/item.html?id=054a6ce213d448c3beae339c14778ce5'
        bodyText = f'Managed Task Export Complete for {activityID}\nActivity Name: {activityName}\nContractor: {contractor}'
        if not contractor_success_TF or not contractor_activity_success_TF:
            bodyText += r'\n\nNote for GIS and Mapping Team:'
            if not contractor_activity_success_TF:
                bodyText += f'\nPlease update Activity and Contractor CSV List: {contractorActivityNameLink}'
            elif not contractor_success_TF:
                bodyText += f'\nPlease update Contractor CSV List: {contractorListLink}'
        bodyText += f'\nDuration: {datetime.datetime.now()-start_time}'
        bodyText+="\n\n".join(summaryTextList)
        sendEmail(fromAddress='egis@tva.gov',toAddress=emailAddresses,subjectText=subjectText,bodyText=bodyText)

        return





        # templateResourceTypes = []
        # exportResourceTypes = []
        # for resourceType,count in resourceTypeCountDict.items():
        #     msg('{}:{}'.format(resourceType,count))
        #     if count == 0:
        #         templateResourceTypes.append(resourceType)
        #     else:
        #         templateResourceTypes.append(resourceType)
        #         #exportResourceTypes.append(resourceType)
        # msg('Resource types to use template: {}'.format(templateResourceTypes))
        # msg('Resource types to export: {}'.format(exportResourceTypes))

        # return
        # field_name_change_dict = { 'other_tva_property_affiliation_orig':'other_tva_property_affiliation_','other_tva_property_affiliation_curren':'other_tva_property_affiliation1','participating_parties_sacred_si_other':'participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o','significant_persons_associate_oth':'significant_persons_associate_o','dates_of_construction_approximate':'dates_of_construction_approxima','Exterior_Component_Material_Type':'Exterior_Component_Material_Typ','Other_Exterior_Component_Material_Type':'Other_Exterior_Component_Materi','Exterior_Component_Material_Treatment':'Exterior_Component_Material_Tre','Other_Exterior_Component_Material_Treatment':'Other_Exterior_Component_Mate_1','other_alterations_additions_type':'other_alterations_additions_typ'}

        # """
        # # arch resource field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # cemetery field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # cultural landscape field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # districts and facilities field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # objects field changes
        # { 'other_tva_property_affiliation_orig':'other_tva_property_affiliation_','other_tva_property_affiliation_curren':'other_tva_property_affiliation1','participating_parties_sacred_si_other':'participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # other field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o'}
        # # historic architectural resrouce field changes
        # {'participating_parties_sacred_si_other':'	participating_parties_sacred__1','consulting_parties_tcp_tcl_other':'consulting_parties_tcp_tcl_othe','participating_parties_tcp_tcl_other':'participating_parties_tcp_tcl_o','significant_persons_associate_oth':'significant_persons_associate_o','dates_of_construction_approximate':'dates_of_construction_approxima','Exterior_Component_Material_Type':'Exterior_Component_Material_Typ','Other_Exterior_Component_Material_Type':'Other_Exterior_Component_Materi','Exterior_Component_Material_Treatment':'Exterior_Component_Material_Tre','Other_Exterior_Component_Material_Treatment':'Other_Exterior_Component_Mate_1','other_alterations_additions_type':'other_alterations_additions_typ'}
        # """

        # def assignDefaultValues(out_features,activityID,activityName,activityType,projectID,formTitle):
        #     # assign default values to fields for activity and project information
        #     fieldValueDict = {'activity_id':activityID,'activity_name':activityName,'activity_type':activityType,'project_id':projectID,'resource_type':formTitle}
        #     for field,value in fieldValueDict.items():
        #         try:
        #             arcpy.management.AssignDefaultToField(out_features, field, value)
        #             msg('Default value {} assigned to {} field'.format(value,field))
        #         except Exception as e:
        #             msg('Unable to assign default value {} to {} field in {}'.format(value,field,os.path.basename(out_features)))


        # gis = connectToAGOL()
        # # surveyExportList = deepcopy(resourceTypeList)
        # surveyExportList = []
        # surveyExportList.append([activityType,'Activity'])
        # # surveyExportList.append(['Management_Summary','Project'])
        # for resourceType in resourceTypeList:
        #     surveyExportList.append([resourceType,'Resource'])

        # surveyTableRecordCountLists = []
        # gdbList = []
        # xlsxList = []
        # for surveyInfo in surveyExportList:
        #     formTitle,formType = surveyInfo
        #     msg('\nProcessing {} {}'.format(formTitle,formType))
        #     form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=formTitle,formType=formType,gis=gis)

        #     form_xlsx = saveXLSXForSurvey123Form(form_item_id,gis=gis)

        #     if (formTitle in exportResourceTypes or formTitle in templateResourceTypes):
        #         form_xlsx = updateFormTemplateDefaults(form_xlsx,activityID,activityName,activityType,projectID,formTitle)

        #     xlsxFolder = os.path.dirname(form_xlsx)
        #     # agolUpdateFolder = createFolder(os.path.join(outputFolder,'Update_AGOL')) # go ahead and create folder to store returned geodatabase

        #     new_xlsx = form_xlsx.replace(xlsxFolder,mtExportFolder)
        #     new_xlsx = os.path.join(os.path.dirname(new_xlsx),'{}_{}.xlsx'.format(formTitle,mtProjectName))
        #     if os.path.exists(new_xlsx):
        #         os.remove(new_xlsx)
        #     shutil.copy2(form_xlsx,new_xlsx)

        #     msg('Template Survey XLSX File: {}'.format(new_xlsx))
        #     xlsxList.append(new_xlsx)
        #     # os.startfile(mtExportFolder)

        #     if os.path.exists(form_xlsx):
        #         os.remove(form_xlsx)

        #     # still need to assign default value, even if no records
        #     templateGDB = os.path.join(MT_GDB_TEMPLATE_FOLDER,'{}.gdb'.format(formTitle))
        #     if formTitle != activityType:
        #         globalIDList = resourceFormGlobalidListDict[formTitle]
        #         if not globalIDList :
        #             msg('No records for {}'.format(formTitle))
        #             appendStatus = ''
        #             # for item in items:
        #             #     itemID,datasetName = item
        #             surveyTableRecordCountLists.append([formTitle,'','',0,appendStatus])
        #     #         continue
        #     # if formTitle == activityType:
        #     #     continue

        #     if (formTitle in exportResourceTypes or formTitle in templateResourceTypes or formTitle == activityType):
        #         if formTitle == activityType:
        #             gdb_form_abbrev = ACTIVITY_TYPE_ABBREV_DICT.get(formTitle)
        #         else:
        #             gdb_form_abbrev = RESOURCE_TYPE_ABBREV_DICT.get(formTitle)
        #         # new_gdb_folder = createFolder(os.path.join(mtExportFolder,'{}_GDB'.format(formTitle)))
        #         # new_gdb = os.path.join(new_gdb_folder,'{}_{}.gdb'.format(formTitle,mtProjectName))
        #         new_gdb_folder = createFolder(os.path.join(mtExportFolder,'{}_GDB'.format(gdb_form_abbrev)))
        #         new_gdb = os.path.join(new_gdb_folder,'{}_{}.gdb'.format(gdb_form_abbrev,mtProjectName))
        #         #new_gdb = os.path.join(mtExportFolder,'{}_{}.gdb'.format(formTitle,mtProjectName))
        #         gdbList.append(new_gdb)
        #         if overwriteExports_TF or not arcpy.Exists(new_gdb):
        #             if arcpy.Exists(new_gdb):
        #                 msg('Attempting to delete {}'.format(new_gdb))
        #                 arcpy.Delete_management(new_gdb)
        #             try:
        #                 if formTitle in exportResourceTypes:
        #                     msg('Downloading Survey123 GDB...')
        #                     surveyGDB = downloadSurveyGDB(form_item_id,outputFolder=xlsxFolder,gis=gis)
        #                     msg('surveyGDB download: {}'.format(surveyGDB))
        #                     os.rename(surveyGDB,new_gdb)
        #                 else:
        #                     surveyGDB = templateGDB
        #                     msg('Copying empty template from {} to {}'.format(surveyGDB,new_gdb))
        #                     arcpy.management.Copy(surveyGDB,new_gdb)
        #                     # arcpy.management.Delete(surveyGDB)

        #                 msg('Template Survey GDB: {}'.format(new_gdb))

        #             except Exception as e:
        #                 wrn('Error exporting {} survey: {}'.format(formTitle,e))
        #                 continue
        #         else:
        #             msg('{} already exists'.format(new_gdb))
        #         #https://services.arcgis.com/w8auYAijfGK1Mydj/arcgis/rest/services/service_e5c4c15f908a40ce8c1a2ca9475a2711/FeatureServer
        #         arcpy.env.workspace = new_gdb

        #         arcpy.env.preserveGlobalIds = True
        #         arcpy.env.overwriteOutput = True
        #         #outputGDB = r"C:\Users\japarkhu\Documents\CRMS_Managed_Task\Kingsport_202310\Objects.gdb"
        #         serviceItemsTable = os.path.join(new_gdb,'GDB_ServiceItems')
        #         #serviceUrl = 'https://services.arcgis.com/w8auYAijfGK1Mydj/arcgis/rest/services/service_4cec60597f194b169fe91d6d6eec833a/FeatureServer'
        #         #serviceUrl = 'https://services.arcgis.com/w8auYAijfGK1Mydj/arcgis/rest/services/service_{}/FeatureServer'.format(feature_layer_item_id)
        #         items = sorted(returnUniqueFieldValuesMultipleFields(serviceItemsTable,['ItemID','DatasetName'],whereClause='ItemID>=0'))
        #         main_layer_name = [datasetName for itemID,datasetName in items if itemID == 0][0]
        #         if formTitle == activityType:
        #             globalIDList = activityGlobalIDList
        #         else:
        #             globalIDList = resourceFormGlobalidListDict[formTitle]
        #         if not globalIDList:
        #             msg('No records for {}'.format(formTitle))
        #             appendStatus = ''
        #             for item in items:
        #                 itemID,datasetName = item
        #                 surveyTableRecordCountLists.append([formTitle,itemID,datasetName,0,appendStatus]) # get list of layers for summary report
        #             # update default values
        #             out_features = os.path.join(new_gdb,main_layer_name)
        #             assignDefaultValues(out_features,activityID,activityName,activityType,projectID,formTitle)
        #             continue # TODO - this assumes there are no orphaned related records
        #         globalIDCount = len(globalIDList)
        #         msg('{} globalids to export'.format(globalIDCount))
        #         serviceUrl = getFeatureUrlFromFeatureLayerID(feature_layer_item_id,gis)
        #         msg('GlobalID List: {}'.format(globalIDList))
        #         globalIDQuery = "globalid IN ('{}')".format("','".join(globalIDList))
        #         parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(globalIDList))
        #         for itemID,datasetName in items:
        #             msg('itemID: {},datasetName: {}'.format(itemID,datasetName))
        #             in_features = serviceUrl + '/' + str(itemID)
        #             out_features = os.path.join(new_gdb,datasetName)
        #             if itemID == 0:
        #                 where_clause = globalIDQuery
        #                 #msg('Count before append: {}'.format(arcpy.management.GetCount(out_features)))
        #                 msg('in_features: {}\nout_features: {}\nwhere_clause: {}'.format(in_features,out_features,where_clause))
        #                 arcpy.conversion.ExportFeatures(in_features,out_features,where_clause,use_field_alias_as_name="NOT_USE_ALIAS")
        #                 msg('Features exported for {}'.format(datasetName))
        #                 # cannot maintain GlobalIDs with file geodatabase
        #                 # https://pro.arcgis.com/en/pro-app/latest/tool-reference/tool-errors-and-warnings/001001-010000/tool-errors-and-warnings-03326-03350-003340.htm
        #                 # arcpy.management.Append(
        #                 #     inputs=in_features,
        #                 #     target=out_features,
        #                 #     schema_type="NO_TEST",
        #                 #     subtype="",
        #                 #     expression=where_clause,
        #                 #     match_fields=None,
        #                 #     update_geometry="NOT_UPDATE_GEOMETRY"
        #                 #     )

        #                 # msg('Features appended for {}'.format(datasetName))
        #                 exportCount = int(str(arcpy.management.GetCount(out_features)))
        #                 appendStatus = ''
        #                 if exportCount:
        #                     appendStatus = 'Y'
        #                 msg('Count after export: {}'.format(exportCount))
        #                 surveyTableRecordCountLists.append([formTitle,itemID,datasetName,exportCount,appendStatus]) # get list of layers for summary report
        #                 if exportCount != globalIDCount:
        #                     wrn('Some features from Resource layer not found in AGOL layer')

        #                 # assign default values to fields for activity and project information
        #                 assignDefaultValues(out_features,activityID,activityName,activityType,projectID,formTitle)
        #             else: # TODO - account for cases where the parentGlobalID is the parent of a related record and not the polygon record
        #                 where_clause = parentGlobalIDQuery
        #                 msg('in_features: {}\nout_features: {}\nwhere_clause: {}'.format(in_features,out_features,where_clause))
        #                 #msg('Count before append: {}'.format(arcpy.management.GetCount(out_features)))
        #                 arcpy.conversion.ExportTable(in_features,out_features,where_clause,use_field_alias_as_name="NOT_USE_ALIAS")
        #                 msg('Record exported for {}'.format(datasetName))
        #                 # arcpy.management.Append(
        #                 #     inputs=in_features,
        #                 #     target=out_features,
        #                 #     schema_type="NO_TEST",
        #                 #     subtype="",
        #                 #     expression=where_clause,
        #                 #     match_fields=None,
        #                 #     update_geometry="NOT_UPDATE_GEOMETRY"
        #                 #     )
        #                 #msg('Count after append: {}'.format(arcpy.management.GetCount(out_features)))
        #                 exportCount = int(str(arcpy.management.GetCount(out_features)))
        #                 appendStatus = ''
        #                 if exportCount:
        #                     appendStatus = 'Y'
        #                 msg('Count after export: {}'.format(exportCount))
        #                 surveyTableRecordCountLists.append([formTitle,itemID,datasetName,exportCount,appendStatus]) # get list of layers for summary report

        #             # try:
        #             #     addDatasetToAPRX(aprxMap_='CURRENT',dataset=out_features,layerName_=datasetName,layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery='')
        #             # except Exception as e:
        #             #     wrn('Unable to add {} layer to map: {}'.format(datasetName,out_features))

        #         def getRelationshipClassPropertiesDict(templateGDB):
        #             arcpy.env.workspace = templateGDB
        #             fcs = []
        #             #root of workspace
        #             for item in arcpy.ListFeatureClasses("*"):    fcs.append(item)
        #             for item in arcpy.ListTables("*"):    fcs.append(item)

        #             fds = arcpy.ListDatasets("*","Feature")
        #             for fd in fds:
        #                 env.workspace = inGDB +'\\'+fd
        #                 for fc in arcpy.ListFeatureClasses("*"):
        #                     fcs.append(fd+'/'+fc)
        #                 for tb in arcpy.ListTables("*"):
        #                     fcs.append(fd+'/'+tb)

        #             arcpy.env.workspace = templateGDB
        #             relClassPropertiesDict = {}
        #             for fc in fcs:
        #                 print('Processing {}'.format(fc))
        #                 desc = arcpy.Describe(os.path.join(templateGDB,fc))
        #                 #desc = arcpy.Describe(fc)
        #                 for relClassName in desc.relationshipClassNames:
        #                     relClassDesc = arcpy.Describe(os.path.join(templateGDB,relClassName))
        #                     #relClassDesc = arcpy.Describe(relClassName)
        #                     relClassPropertiesDict[relClassDesc.name] = relClassDesc

        #             return relClassPropertiesDict

        #         def createRelationshipClasses(targetGDB,relClassPropertiesDict):
        #             arcpy.env.workspace = targetGDB
        #             msg(f'Creating relationship classes: {relClassPropertiesDict.keys()}')
        #             for relClassName,relClassDesc in relClassPropertiesDict.items():
        #                 msg(relClassName)
        #                 if relClassDesc.cardinality == 'OneToMany':
        #                     rel_card = 'ONE_TO_MANY'
        #                 elif relClassDesc.cardinality == 'OneToOne':
        #                     rel_card = 'ONE_TO_ONE'
        #                 elif relClassDesc.cardinality == 'ManyToMany':
        #                     rel_card = 'MANY_TO_MANY'

        #                 if relClassDesc.isComposite:
        #                     rel_type = 'COMPOSITE'
        #                 else:
        #                     rel_type = 'SIMPLE'

        #                 if relClassDesc.isAttributed:
        #                     rel_attributed = 'ATTRIBUTED'
        #                 else:
        #                     rel_attributed = 'NONE'


        #                 originClassKeys = relClassDesc.originClassKeys
        #                 origin_primary,origin_foreign = None,None
        #                 if originClassKeys:
        #                     for field,key,_ in originClassKeys:
        #                         if key == 'OriginPrimary':
        #                             origin_primary = field
        #                         elif key == 'OriginForeign':
        #                             origin_foreign = field

        #                 destClassKeys = relClassDesc.destinationClassKeys
        #                 dest_primary,dest_foreign = None,None
        #                 if destClassKeys:
        #                     for field,key,_ in destClassKeys:
        #                         if key == 'DestinationPrimary':
        #                             dest_primary = field
        #                         elif key == 'DestinationForeign':
        #                             dest_foreign = field

        #                 try:
        #                     arcpy.management.CreateRelationshipClass(relClassDesc.originClassNames[0],relClassDesc.destinationClassNames[0],relClassName,rel_type,relClassDesc.forwardPathLabel,relClassDesc.backwardPathLabel,relClassDesc.notification,rel_card,rel_attributed,origin_primary,origin_foreign,dest_primary,dest_foreign)
        #                     msg('Relationship class created: {}'.format(relClassName))
        #                 except Exception as e:
        #                     wrn('Unable to create relationship class: {}: {}'.format(relClassName,e))

        #         relClassPropertiesDict = getRelationshipClassPropertiesDict(templateGDB)

        #         createRelationshipClasses(new_gdb,relClassPropertiesDict)

        #         continue


        #         # fc_name = arcpy.ListFeatureClasses()[0]
        #         # fc_path = os.path.join(new_gdb,fc_name)
        #         # if formType == 'Activity':
        #         #     activityFC = os.path.join(new_gdb,fc_name)
        #         #     activityIDQuery = "ACTIVITY_ID = '{}'".format(activityID)
        #         #     lyr = arcpy.management.MakeFeatureLayer(fc_name,'Activity')
        #         #     arcpy.management.SelectLayerByAttribute(lyr, 'NEW_SELECTION', activityIDQuery, 'INVERT')
        #         #     arcpy.management.DeleteFeatures(lyr)
        #         #     msg('Outside Activities have been deleted')
        #         # elif formType == 'Resource':
        #         #     lyr = arcpy.management.MakeFeatureLayer(fc_name,'Resource')
        #         #     # selectLyr = arcpy.management.MakeFeatureLayer(activityFC,'Activity')
        #         #     selectLyr = activityLyr
        #         #     arcpy.management.SelectLayerByLocation(lyr, 'INTERSECT', selectLyr, retainResourceRadius, 'NEW_SELECTION', 'INVERT')
        #         #     arcpy.management.DeleteFeatures(lyr)
        #         #     msg('{} outside of {} have been deleted. {} remain.'.format(formTitle,retainResourceRadius,arcpy.management.GetCount(lyr)))
        #         #     msg(fc_path)
        #         #     #addDatasetToAPRX(fc_path)
        #         #     try:
        #         #         addDatasetToAPRX(aprxMap_='CURRENT',dataset=fc_path,layerName_=fc_name,layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery='')
        #         #     except Exception as e:
        #         #         wrn('Unable to add {} layer to map: {}'.format(fc_name,fc_path))

        #         # elif formType == 'Project':
        #         #     projectIDQuery = "PROJECT_ID = '{}'".format(projectID)
        #         #     lyr = arcpy.management.MakeFeatureLayer(fc_name,'Project')
        #         #     arcpy.management.SelectLayerByAttribute(lyr, 'NEW_SELECTION', projectIDQuery, 'INVERT')
        #         #     arcpy.management.DeleteFeatures(lyr)
        #         #     msg('Outside Projects have been deleted')

        # if surveyTableRecordCountLists:
        #     df = pandas.DataFrame(surveyTableRecordCountLists)
        #     df.columns = ['Form_Title','Layer_ID','Layer_Name','Record_Count','Need to Update AGOL Form?']
        #     xlsxPath = r'{}\Summary_for_{}.xlsx'.format(mtExportFolder,mtProjectName)
        #     if os.path.exists(xlsxPath):
        #         existingDF = pandas.read_excel(xlsxPath)
        #         combinedDF = pandas.concat([existingDF,df],axis=0)
        #         combinedDF.drop_duplicates(inplace=True)
        #         try:
        #             combinedDF.to_excel(xlsxPath,index=False)
        #         except (PermissionError,OSError) as e:
        #             wrn('Error exporting to {}; {}; trying new filename'.format(xlsxPath,e))
        #             combinedDF.to_excel(xlsxPath.replace('.xlsx','_{}.xlsx'.format(timestamp_time)),index=False)

        #     else:
        #         df.to_excel(xlsxPath,index=False)
        #     msg('Summary XLSX Exported: {}'.format(xlsxPath))
        #     xlsxList.append(xlsxPath)

        # def zip_gdb_folder(infolder,outfile):
        #     # infolder - folder containing .gdb
        #     # outfile - output zip file path
        #     import zipfile
        #     with zipfile.ZipFile(outfile, 'w', zipfile.ZIP_DEFLATED) as myzip:
        #         infolder = os.path.normpath(infolder)
        #         # os.walk visits every subdirectory, returning a 3-tuple
        #         #  of directory name, subdirectories in it, and file names
        #         #  in it.
        #         for (dirpath, dirnames, filenames) in os.walk(infolder):
        #             # Iterate over every file name
        #             for file in filenames:
        #                 # Ignore .lock files
        #                 if not file.endswith('.lock'):
        #                     print("Adding %s..." % os.path.join(infolder, dirpath, file))
        #                     try:
        #                         myzip.write(os.path.join(dirpath, file),
        #                         #os.path.join(os.path.basename(infolder), os.path.join(dirpath, file)[len(infolder)+len(os.sep):]))
        #                         os.path.join(dirpath, file)[len(infolder)+len(os.sep):])
        #                     except Exception as e:
        #                         print("Error adding %s: %s" % (file, e))
        #     msg('{} zip file created'.format(outfile))

        # def zip_up_files(infiles,outfile):
        #     # infolder - folder containing .gdb
        #     # outfile - output zip file path
        #     import zipfile
        #     with zipfile.ZipFile(outfile, 'w', zipfile.ZIP_DEFLATED) as myzip:
        #         for file in infiles:
        #             myzip.write(file,file[len(os.path.dirname(file))+len(os.sep):])
        #             msg('{} added to zip file'.format(file))

        #     msg('{} zip file created'.format(outfile))

        # # # zip up GDB
        # gdbZipList = []
        # for new_gdb in gdbList:
        #     gdbZip = os.path.dirname(new_gdb)+'.zip'
        #     zip_gdb_folder(infolder=os.path.dirname(new_gdb),outfile=gdbZip)
        #     gdbZipList.append(gdbZip)

        # filesToZip = gdbZipList + xlsxList
        # zip_up_files(filesToZip,outfile=os.path.join(mtExportFolder,'Export_for_{}.zip'.format(mtProjectName)))

        # if os.path.exists(finalMTExportFolder):
        #     wrn('{} already exists. Manually reconcile with local folder {}'.format(finalMTExportFolder,mtExportFolder))
        #     os.startfile(mtExportFolder)
        #     os.startfile(finalMTExportFolder)
        # else:
        #     arcpy.management.Copy(mtExportFolder,finalMTExportFolder)
        #     msg('{} copied to {}'.format(mtExportFolder,finalMTExportFolder))
        #     os.startfile(finalMTExportFolder)

        # # open output folder


        # if emailAddress and emailAddress != '@tva.gov':
        #     subjectText = 'Managed Task Export Complete for {}'.format(activityID)
        #     bodyText = 'Managed Task Export Complete for {}\nLocal Folder: {}\nDuration: {}'.format(activityID,finalMTExportFolder,datetime.datetime.now()-start_time)
        #     sendEmail(fromAddress='egis@tva.gov',toAddress=[emailAddress],subjectText=subjectText,bodyText=bodyText)
        # return

''' Analyze Differences for Managed Task Deliverables '''
new_cells = set()
updated_cells = set()
class analyzeMTDifferences(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "1. Analyze Differences for Managed Task Deliverables"
        self.description = "Tool to Analyze Differences for Managed Task Deliverables"
        self.canRunInBackground = False
        #self.category = "Managed Task Import"

    def getParameterInfo(self):
        """ Analyze Differences for Managed Task Deliverables """
        '''
            Defines the tool's parameters
        '''
        # fgdbParam = arcpy.Parameter(
        #     displayName='Managed Task FGDB',
        #     name='fgdb', ### Name the same as parameter in main script
        #     datatype='DEWorkspace',
        #     parameterType='Required',
        #     direction='Input')
        #fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb"
        # fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\Projects\70568\70568-AS-1\MT\Import_from_MT\Historic_Architectural_Resources_70568-AS-1\2bd66427-6ea8-4c6a-8823-7a6e8381fff7.gdb"
        # \\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\Projects\21573447673_26a Kim-Geyer\21573447673-AS-1\Export_to_MT
        # resourceTypeParam = arcpy.Parameter(
        #     displayName='Resource Type',
        #     name='resourceType',
        #     datatype='GPSTRING',
        #     parameterType='Required',
        #     direction='Input')
        # resourceTypeParam.filter.type = 'ValueList'
        # resourceTypeParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])
        # #resourceTypeParam.value = 'Cultural_Landscapes'

        contractorParam = arcpy.Parameter(
            displayName='Contractor Filter (Optional)',
            name='contractor',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        contractorParam.filter.type = 'ValueList'
        contractorParam.filter.list = CONTRACTOR_LIST

        activityStatusParam = arcpy.Parameter(
            displayName='Activity Status Filter (Optional)',
            name='activityStatus',
            datatype='GPSTRING',
            parameterType='Optional',
            direction='Input')
        activityStatusParam.filter.type = 'ValueList'
        activityStatusParam.filter.list = ['Not Started','Ready for TVA Review']

        activityParam = arcpy.Parameter(
            displayName='Activity To Process',
            name='activity',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        activityParam.value = "contractor='TVAR'; activity_id='128529793407-AS-1'; activity_name='TVAR Test Archaeological Survey'; review_status='Not Started'" # TODO - remove default after testing

        searchParam = arcpy.Parameter(
            displayName='Search for Activities',
            name='search',
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')

        # resourceLyrParam = arcpy.Parameter(
        #     displayName='Resource Layer',
        #     name='resourceLyr', ### Name the same as parameter in main script
        #     datatype='GPFeatureLayer',
        #     parameterType='Required',
        #     direction='Input')
        # resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        # prevResourceLyrParam = arcpy.Parameter(
        #     displayName='Resource Previous Boundary Layer',
        #     name='prevResourceLyr', ### Name the same as parameter in main script
        #     datatype='GPFeatureLayer',
        #     parameterType='Required',
        #     direction='Input')
        # prevResourceLyrParam.value = resourcePrevBoundariesLyrParamDefault # set default cultural resources layer

        # mtResourceLyrParam = arcpy.Parameter(
        #     displayName='Managed Task Resource Layer',
        #     name='mtResourceLyr', ### Name the same as parameter in main script
        #     datatype='GPFeatureLayer',
        #     parameterType='Required',
        #     direction='Input')
        #mtResourceLyrParam.value = mtResourceLyrParamDefault # set default cultural resources layer

        # params = [resourceLyrParam,prevResourceLyrParam,mtResourceLyrParam,queryForDifferencesParam,queryForNewRecordsParam]

        params = [contractorParam,activityStatusParam,activityParam,searchParam] # [resourceLyrParam,fgdbParam,mtResourceLyrParam,resourceTypeParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Analyze Differences for Managed Task Deliverables """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        pValuesAsText = paramDictValuesAsText(parameters)
        contractor = pValues['contractor']
        activityStatus = pValues['activityStatus']
        activity = pValues['activity']
        search = pValues['search']

        if search:
            gis = connectToAGOL()
            # get activity records
            mt_activityTitle = 'Managed Task Activities - MT Master'
            mt_activityFeatureLyrItem = getItemFromTitle(gis,mt_activityTitle,item_type='Feature Layer')
            mt_lyr = mt_activityFeatureLyrItem.layers[0]
            if contractor and activityStatus:
                mt_query = "contractor_0 = '{}' and contractor_review_status = '{}'".format(contractor,activityStatus)
            elif contractor:
                mt_query = "contractor_0 = '{}'".format(contractor)
            elif activityStatus:
                mt_query = "contractor_review_status = '{}'".format(activityStatus)
            else:
                mt_query = "1=1"
            mt_activities_to_process = getFeaturesFromLyr(lyr=mt_lyr, query=mt_query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'activity_id', 'project_id', 'contractor_review_status'],returnGeometry=False)

            activities = []
            for a in mt_activities_to_process:
                attr = a.attributes
                contractor, activity_id, activity_name, review_status = attr['contractor_0'], attr['activity_id'], attr['activity_name'], attr['contractor_review_status']
                identifier = f"{contractor=}; {activity_id=}; {activity_name=}; {review_status=}"
                activities.append(identifier)

            parameters[pNum['activity']].filter.list = activities
            parameters[pNum['search']].value = False

        return
        # resourceType = pValues['resourceType']
        fgdb = r'{}'.format(pValuesAsText['fgdb'])
        mtResourceLyr = pValues['mtResourceLyr']

        if fgdb and not mtResourceLyr:
            # check to verify that the fgdb is in the proper folder
            if '\\Import_from_MT\\' not in fgdb:
                return

            arcpy.env.workspace = fgdb
            fc_name = arcpy.ListFeatureClasses()[0]
            fc_path = os.path.join(fgdb,fc_name)
            addDatasetToAPRX(aprxMap_='CURRENT',dataset=fc_path,layerName_=fc_name,layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery='')
            parameters[pNum['mtResourceLyr']].value = fc_name

            # get Resource Type from Feature Class Name
            parameters[pNum['resourceType']].value = FEATURE_LAYER_NAME_RESOURCE_TYPE_DICT.get(fc_name)

        return

    def updateMessages(self, parameters):
        """ Analyze Differences for Managed Task Deliverables """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        # pValues = paramDictValues(parameters) # k:name;v:value
        # pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)
        # resourceType = pValues['resourceType']

        return 
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        pValuesAsText = paramDictValuesAsText(parameters)

        fgdb = r'{}'.format(pValuesAsText['fgdb'])
        if fgdb and '\\Import_from_MT\\' not in fgdb:
            parameters[pNum['fgdb']].setErrorMessage("FGDB is not in the Activity's Import_from_MT subfolder")

        gisURL = checkGIS_URL()
        if gisURL != TVA_AGOL_URL:
            parameters[0].setErrorMessage('Currently signed into {} Portal. Sign in to ArcGIS Online: {}'.format(gisURL,TVA_AGOL_URL))
            return

        resourceLyr = pValues['resourceLyr']

        # check to make sure user is editing a personal version
        parameters,validVersion_TF = checkVersion(parameters,pNum,lyr=resourceLyr,paramName='resourceLyr')
        if not validVersion_TF:
            return

        return

    def execute(self, parameters, messages):
        """ Analyze Differences for Managed Task Deliverables """
        # Create GIS Connection to AGOL / Portal
        gis = connectToAGOL()

        pNum = paramDictNum(parameters) # k:name;v:number
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)

        activity = pValues['activity'] # "contractor='TVAR'; activity_id='128529793407-AS-1'; activity_name='TVAR Test Archaeological Survey'; review_status='Not Started'"
        attributes = activity.split('; ') # ["contractor='TVAR'", "activity_id='128529793407-AS-1'", "activity_name='TVAR Test Archaeological Survey'", "review_status='Not Started'"]
        attrDict = {}
        for attr in attributes:
            name,value = attr.split('=')
            value = value.strip("'")
            attrDict[name] = value
        msg(attrDict) # {'contractor': 'TVAR', 'activity_id': '128529793407-AS-1', 'activity_name': 'TVAR Test Archaeological Survey', 'review_status': 'Not Started'}

        activityName = attrDict['activity_name']
        contractor = attrDict['contractor']
        activityID = attrDict['activity_id']
        projectID = activityID.split('-')[0]

        gis = connectToAGOL()
        mt_master_feature_layer_items = gis.content.search('MT Master',item_type='Feature Layer Collection',max_items=100)

        sync_gis_fields = ['resource_id', 'res_state_id', 'resource_name', 'cemetery_number', 'legacy_object_id', 'field_site_number', 'resource_type', 'other_resource_type', 'activity_id', 'activity_type', 'project_id']


        formType = 'Resource'
        resourceTypes = ['Archaeological Resources', 'Historic Architectural Resources', 'Objects', 'Districts and Facilities', 'Cultural Landscapes', 'Cemeteries', 'Other Resources']

        
        for resourceType in resourceTypes:
            mtMasterTitle = resourceType + ' - MT Master'
            contractorViewTitle = resourceType + ' - ' + contractor + ' View'
            msg(contractorViewTitle)
            mt_featureLyrItem = getItemFromTitle(gis,contractorViewTitle,item_type='Feature Layer')
            mtMaster_formItem = getItemFromTitle(gis,mtMasterTitle,item_type='Form')
            mtMaster_formItemID = mtMaster_formItem.id
            break # TODO - remove after testing




        msg(f'\nQuerying {contractor} layers for {resourceType} associated with {activityName}')
        # identify records that have the corresponding activity name and contractor
        mt_lyr = mt_featureLyrItem.layers[0]

        if activityName:
            mt_query = f"contractor_0 = '{contractor}' and (tva_assigned_activity_name_0 = '{activityName}' or mt_assigned_activity_name_0 = '{activityName}' or activity_name = '{activityName}')"
        else:
            mt_query = f"contractor_0 = '{contractor}'"

        mt_features_to_process = getFeaturesFromLyr(lyr=mt_lyr, query=mt_query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0', 'mt_assigned_activity_name_0'],returnGeometry=False)

        mt_globalIDList = [x.attributes['globalid'] for x in mt_features_to_process]

        mt_lyrs = mt_featureLyrItem.layers + mt_featureLyrItem.tables
        mt_lyrGlobalIDDict, mt_lyrFeaturesDict, mt_lyrMissingGlobalIDDict, mt_baseGlobalIDDict, mt_globalIDobjectIDDict = getLyrGlobalIDandFeaturesDict(mt_lyrs,mt_globalIDList,formType)
        newMTFeatures = [x for x in mt_lyrFeaturesDict[0] if not x.attributes['orig_globalid_0']]
        oldMTFeatures = [x for x in mt_lyrFeaturesDict[0] if x.attributes['orig_globalid_0']]
        oldMT_agolGlobalIDList = [x.attributes['orig_globalid_0'] for x in oldMTFeatures]

        msg(f'\nQuerying TVA AGOL for {resourceType} associated with {activityName}')
        agolLayerTitle = resourceType
        msg(agolLayerTitle)
        agol_formItem = getItemFromTitle(gis,agolLayerTitle,item_type='Form')
        agol_formItemID = agol_formItem.id
        agol_featureLyrItem = getItemFromTitle(gis,agolLayerTitle,item_type='Feature Layer')
        agol_lyrs = agol_featureLyrItem.layers + agol_featureLyrItem.tables

        # oldMT_agolGlobalIDList = ['{74CFD91F-4654-4AAC-9E39-E2258B4E6B8C}'] # TODO remove
        # # {1c8ad498-4e47-4a94-ad0b-8a281ebf5cd2}
        # # parent {5c26d326-97a5-4e14-8e5a-21e421bd70ca}
        # if not oldMT_agolGlobalIDList:
        #     agol_query 
        # agol_query = "globalid in ('{}')".format("','".join(oldMT_agolGlobalIDList))
        agol_lyrGlobalIDDict, agol_lyrFeaturesDict, agol_lyrMissingGlobalIDDict, agol_baseGlobalIDDict, agol_globalIDobjectIDDict = getLyrGlobalIDandFeaturesDict(agol_lyrs,oldMT_agolGlobalIDList,formType)


        ### Logic to generate report of differences
        mt_baseGlobalIDSyncAttrDict = {}
        agol_baseGlobalIDSyncAttrDict = {}
        different_values_all_layers = []
        same_values_all_layers = []
        new_records_all_layers = []
        for lyrId,mt_features in mt_lyrFeaturesDict.items():
            lyrName = mt_lyrs[lyrId].properties.name
            agol_fields = [x['name'].lower() for x in agol_lyrs[lyrId].properties['fields']]
            mt_fields = [x['name'].lower() for x in mt_lyrs[lyrId].properties['fields']]
            msg(f'\n\n{lyrId} ({lyrName})')
            if not mt_features:
                msg(f'No MT features in {lyrId}')
                continue
            agol_features = agol_lyrFeaturesDict[lyrId]
            if not agol_features:
                msg(f'No AGOL features in {lyrId}')
                # continue

            # agol_fields = [field for field in agol_features[0].attributes.keys()]
            # mt_fields = [field for field in mt_features[0].attributes.keys()]
            common_fields = [field for field in mt_fields if field.lower() in agol_fields and field.lower() not in ('objectid','creationdate','creator','editdate','editor','globalid','parentglobalid','shape__area','shape__length')]

            agol_different_values_dict = {}
            changed_GlobalIDs = []
            unchanged_GlobalIDs = []
            sync_changed_GlobalIDs = []
            features_to_add = []
            for mt_feature in mt_features:
                msg('---')
                different_values = []
                mt_globalid = mt_feature.attributes['globalid']
                mt_baseGlobalID = mt_baseGlobalIDDict.get(mt_globalid)
                agol_globalid = mt_feature.attributes[f'orig_globalid_{lyrId}']
                agol_baseGlobalID = agol_baseGlobalIDDict.get(agol_globalid)

                # generate form URL hyperlinks      
                mt_form_url = generateFormURL(mtMaster_formItemID,mt_baseGlobalID)

                if lyrId == 0:
                    mt_baseGlobalIDSyncAttrDict[mt_globalid] = {k:v for k,v in mt_feature.attributes.items() if k in sync_gis_fields}

                mt_resourceID = mt_feature.attributes['resource_id'] if 'resource_id' in common_fields else mt_baseGlobalIDSyncAttrDict[mt_baseGlobalID].get('resource_id','')

                mt_res_state_id = mt_feature.attributes['res_state_id'] if 'res_state_id' in common_fields else mt_baseGlobalIDSyncAttrDict[mt_baseGlobalID].get('res_state_id','')
                

                mt_resource_name = mt_feature.attributes['resource_name'] if 'resource_name' in common_fields else mt_baseGlobalIDSyncAttrDict[mt_baseGlobalID].get('resource_name','')

                mt_parentglobalid = mt_feature.attributes['parentglobalid'] if 'parentglobalid' in mt_fields else ''

                if not agol_globalid:
                    msg(f'No AGOL feature expected for MT {mt_globalid}')
                    features_to_add.append(mt_feature)
                    new_records_all_layers.append([resourceType,lyrId,lyrName,mt_resourceID,mt_resource_name,mt_res_state_id,mt_baseGlobalID,mt_globalid,mt_parentglobalid,mt_form_url])
                    # TODO - These are new features; how do we want to report them???
                    continue

                # TODO - any need to check for MT related table records that may have been deleted?

                agol_form_url = generateFormURL(agol_formItemID,agol_baseGlobalID)
                #mt_features_selection = [x for x in mt_features if x.attributes[f'orig_globalid_{lyrId}'] == agol_globalid]

                agol_features_selection = [x for x in agol_features if x.attributes['globalid'] == agol_globalid]
                if not agol_features_selection:
                    wrn(f'No AGOL feature for {agol_globalid}')
                    continue
                if len(agol_features_selection) > 1:
                    wrn(f'More than 1 AGOL feature for {agol_globalid}')
                    continue

                agol_feature = agol_features_selection[0]
                
                agol_resourceID = agol_feature.attributes['resource_id'] if 'resource_id' in common_fields else agol_baseGlobalIDSyncAttrDict[agol_baseGlobalID].get('resource_id','') #agol_resourceID = ''

                agol_res_state_id = agol_feature.attributes['res_state_id'] if 'res_state_id' in common_fields else agol_baseGlobalIDSyncAttrDict[agol_baseGlobalID].get('res_state_id','') #agol_res_state_id = ''

                agol_resource_name = agol_feature.attributes['resource_name'] if 'resource_name' in common_fields else agol_baseGlobalIDSyncAttrDict[agol_baseGlobalID].get('resource_name','') #agol_resource_name = ''

                agol_parentglobalid = agol_feature.attributes['parentglobalid'] if 'parentglobalid' in agol_fields else ''

                #msg(agol_feature.geometry)
                #msg(mt_feature.attributes)
                if lyrId == 0:
                    agol_baseGlobalIDSyncAttrDict[agol_globalid] = {k:v for k,v in agol_feature.attributes.items() if k in sync_gis_fields}
                    agol_geo = agol_feature.geometry
                    mt_geo = mt_feature.geometry
                    if agol_geo != mt_geo:
                        #msg('Different Geometry for {}'.format(agol_globalid))
                        sync_changed_GlobalIDs.append(agol_globalid)
                        different_values += [(resourceType,lyrId,lyrName,'Geometry','','',agol_resourceID,agol_resource_name,agol_res_state_id,agol_baseGlobalID,mt_baseGlobalID,agol_globalid,'',mt_globalid,'',agol_form_url,mt_form_url)]
                    else:
                        msg('Same Geometry for {}'.format(agol_globalid))
                    
                    inspect_geometry = False
                    if inspect_geometry:
                        pass
                        # if agol_geo:
                        #     agol_geo_rings = agol_geo['rings']
                        # agol_geo_rings = agol_geo['rings'] if agol_geo else [[]]
                        # mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

                        # if agol_geo != mt_geo:
                        #     msg('Different Geometry for {}'.format(agol_globalid))
                        # else:
                        #     msg('Same Geometry for {}'.format(agol_globalid))
                        # mt_geo_json = managed_task_dict['{' + agol_globalid + '}']['SHAPE@JSON']
                        # mt_geo = json.loads(mt_geo_json) if mt_geo_json else None
                        # # if mt_geo:
                        # #     mt_geo_rings = mt_geo['rings']
                        # mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

                        # agol_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in agol_geo_rings for coord in ring]
                        # mt_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in mt_geo_rings for coord in ring]

                        # if sorted(agol_coord_lists) != sorted(mt_coord_lists):
                        #     msg('Different Geometry for {}'.format(agol_globalid))
                        #     msg(f'{agol_coord_lists=}\n{mt_coord_lists=}')
                        #     different_values += [('','','Geometry','','')]
                        #     if sync_gis_fields:
                        #         sync_changed_GlobalIDs.append(agol_globalid)
                        # else:
                        #     msg('Same Geometry for {}'.format(agol_globalid))

                # msg(f'{agol_resourceID},{mt_resourceID}')
                # msg(f'{agol_res_state_id},{mt_res_state_id}')

                #different_attribute_values = [(agol_resourceID,agol_res_state_id,field,agol_feature.attributes[field],managed_task_dict['{' + agol_globalid + '}'][field]) for field in common_fields if tryUpper(removeBrackets(agol_feature.attributes[field])) != tryUpper(removeBrackets(datetimeUTCToEpochMilli(managed_task_dict['{' + agol_globalid + '}'][field])))] # must attempt to convert datetime to epoch in order to compare feature class to AGOL values; must convert globalids to uppercase; must remove '{' and '}' for parentglobalid
                different_attribute_values = [(resourceType,lyrId,lyrName,field,agol_feature.attributes[field],mt_feature.attributes[field],agol_resourceID,agol_resource_name,agol_res_state_id,agol_baseGlobalID,mt_baseGlobalID,agol_globalid,agol_parentglobalid,mt_globalid,mt_parentglobalid,agol_form_url,mt_form_url) for field in common_fields if agol_feature.attributes[field] != mt_feature.attributes[field]]

                if different_attribute_values:
                    #msg('Different Attributes for {}'.format(agol_globalid))
                    pass
                else:
                    msg('Same Attributes for {}'.format(agol_globalid))

                different_values += different_attribute_values

                if different_values:
                    different_values_all_layers += different_values
                    #msg('Different Values for {}:\nField\tAGOL\tManaged Task\n{}'.format(agol_globalid,"\n".join(['{}\t{}\t{}\t{}\t{}'.format(resourceID,resStateID,field,agol_value,mt_value) for resourceID,resStateID,field,agol_value,mt_value in different_values])))
                    if 'resource_id' in common_fields:
                        msg('Different Values for {} (Resource ID: {}, Resource State ID: {}):'.format(agol_globalid,agol_resourceID,agol_res_state_id))
                    elif 'parentglobalid' in agol_fields:
                        msg('Different Values for {} (Parent GlobalID: {}):'.format(agol_globalid,agol_parentglobalid))
                    # for d in different_values:
                    #     msg(f'{len(d)=}: {d}')
                    for resourceType,lyrId,lyrName,field,agol_value,mt_value,resourceID,resourceName,resStateID,agol_baseGlobalID,mt_baseGlobalID,agol_globalid,agol_parentglobalid,mt_globalid,mt_parentglobalid,agol_form_url,mt_form_url in different_values:
                        if field == 'Geometry':
                            msg(f'\t{field}')
                        msg(f'\t{field}\t{agol_value=}\t{mt_value=}')
                    agol_different_values_dict[agol_globalid] = different_values
                    changed_GlobalIDs.append(agol_globalid)
                    # determine if field differences apply when syncing with GIS
                    if sync_gis_fields:
                        sync_value_changes = [x for x in different_values if x[2] in sync_gis_fields]
                        if sync_value_changes:
                            sync_changed_GlobalIDs.append(agol_globalid)
                else:
                    same_values_all_layers.append([resourceType,lyrId,lyrName,agol_resourceID,agol_resource_name,agol_res_state_id,agol_baseGlobalID,mt_baseGlobalID,agol_globalid,agol_parentglobalid,mt_globalid,mt_parentglobalid,agol_form_url,mt_form_url])
                    unchanged_GlobalIDs.append(agol_globalid)
                    msg('Same Values for {}'.format(agol_globalid))

            sync_changed_GlobalIDs = list(set(sync_changed_GlobalIDs))

        import pandas
        df_dict = {}
        existing_values_all_layers = different_values_all_layers + same_values_all_layers
        df_diff = pandas.DataFrame(different_values_all_layers)
        df_diff.columns = ['Resource Type','Layer ID','Layer Name','Field','AGOL Value','MT Value','Resource ID','Resource Name','Resource State ID','AGOL Base GlobalID','MT Base GlobalID','AGOL GlobalID','AGOL Parent GlobalID','MT GlobalID','MT Parent GlobalID','AGOL Form URL','MT Form URL']
        df_diff['Status'] = 'Updated Record'
        #df_dict['updates'] = df_diff

        df_same = pandas.DataFrame(same_values_all_layers)
        df_same.columns = ['Resource Type','Layer ID','Layer Name','Resource ID','Resource Name','Resource State ID','AGOL Base GlobalID','MT Base GlobalID','AGOL GlobalID','AGOL Parent GlobalID','MT GlobalID','MT Parent GlobalID','AGOL Form URL','MT Form URL']
        df_same['Status'] = 'No Update'
        #df_dict['same'] = df_same


        df_new = pandas.DataFrame(new_records_all_layers)
        df_new.columns = ['Resource Type','Layer ID','Layer Name','Resource ID','Resource Name','Resource State ID','MT Base GlobalID','MT GlobalID','MT Parent GlobalID','MT Form URL']
        df_new['Status'] = 'New Record'
        #df_dict['new'] = df_new

        def moveColumnToIndex(df,columnName,idx):
            # moves column to the specified index in the data frame
            column = df.pop(columnName)
            df.insert(idx,columnName,column)
            return df

        def style_specific_cell(x):
            global new_cells
            global updated_cells
            new_color = 'background-color: lightblue'
            updated_color = 'background-color: lightgreen'
            df1 = pandas.DataFrame('', index=x.index, columns=x.columns)
            for cell in new_cells:
                row,col = cell
                df1.iloc[row,col] = new_color
            for cell in updated_cells:
                row,col = cell
                df1.iloc[row,col] = updated_color                
            return df1

        def format_text(val):
            # function to return formatting for highlighting empty cells and setting hyperlinks
            # https://stackoverflow.com/questions/75221463/highlighting-excel-cells-based-on-the-cell-values-using-pandas-dataframe-and-xls
            # https://www.analyticsvidhya.com/blog/2021/06/style-your-pandas-dataframe-and-make-it-stunning/
            val = str(val)
            if val == 'New Record':
                return 'background-color: lightblue'
            elif val == 'Updated Record':
                return 'background-color: lightgreen'
            elif val.startswith('https://survey123.arcgis.com'):
                return 'color: blue'
            return

        result = pandas.concat([df_diff, df_same, df_new], ignore_index=True, sort=False)
        statusColumnIndex = result.columns.tolist().index('Layer Name') + 1
        result = moveColumnToIndex(result,'Status',statusColumnIndex)
        df_dict['Result'] = result
        result.sort_values(['Resource Type','MT Base GlobalID','Layer ID','Status'],inplace=True)

        tempFolder = createFolder(r'c:\temp')
        outputXLSX = os.path.join(tempFolder,'Change_Report_for_{}_{}.xlsx'.format(activityName,'{:%Y%m%d_%H%M%S}'.format(datetime.datetime.now())))
        #df.to_excel(outputXLSX,index=False)

        new_cells = set()
        updated_cells = set()
        columnCount = len(result.columns)
        result_values = result.values.tolist()
        final_values = []
        previous_globalid = None
        globalid = None
        globalidColumnIndex = result.columns.tolist().index('MT Base GlobalID')
        row_idx = 0
        for row in result_values:
            globalid = row[globalidColumnIndex]
            status = row[statusColumnIndex]
            if not previous_globalid:
                previous_globalid = globalid
            if globalid != previous_globalid:
                final_values.append([''*columnCount])
                row_idx+=1
            final_values.append(row)
            previous_globalid = globalid
            if status == 'New':
                new_cell_coords = set([(row_idx,i) for i in range(columnCount)])
                new_cells.update(new_cell_coords)                
            elif status == 'Update':
                updated_cell_coords = set([(row_idx,i) for i in range(columnCount)])
                updated_cells.update(updated_cell_coords)
            row_idx+=1
        msg(f'{new_cells=}')
        msg(f'{updated_cells=}')

        df_final = pandas.DataFrame(final_values)
        df_final.columns = result.columns
        df_final = df_final.style.applymap(format_text)
        df_final.to_excel(outputXLSX, index=False)

        #df_styled = df_final.style.apply(style_specific_cell, axis = None)
        #df_styled.to_excel(outputXLSX, index = False)
        #df_styled.to_excel(xlsx_path, index = False)
        #df_styled = df_final.style.apply(style_specific_cell, axis = None)
        #df_styled = df_final
        #df_dict['Results'] = df_styled
        # xlsx_path = outputXLSX.replace('.xlsx','_styled.xlsx')
        # df_styled.to_excel(xlsx_path, index = False)
        # os.startfile(xlsx_path)

        #saveSheets(df_dict, outputXLSX, index=False, header=True)
        os.startfile(outputXLSX)

        return

 





        return








        fgdb = r'{}'.format(pValuesAsText['fgdb'])
        resourceType = pValues['resourceType']
        mt_resourceLyr = pValues['mtResourceLyr']
        resourceLyr = pValues['resourceLyr']
        # prevResourceLyr = pValues['prevResourceLyr']

        gis_wkid = arcpy.Describe(resourceLyr).SpatialReference.factoryCode
        mt_wkid = arcpy.Describe(mt_resourceLyr).SpatialReference.factoryCode

        gis_fields = [f.name.lower() for f in arcpy.ListFields(resourceLyr)]
        msg(gis_fields)
        mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_resourceLyr)]
        msg(mt_fields)
        common_fields = [x for x in mt_fields if x in gis_fields]
        msg('common fields: {}'.format(common_fields))

        # clear map selections, if any
        arcpy.SelectLayerByAttribute_management(mt_resourceLyr,'CLEAR_SELECTION')
        arcpy.SelectLayerByAttribute_management(resourceLyr,'CLEAR_SELECTION')

        globalIDIndex = common_fields.index('globalid')
        managed_task_dict = {}
        mt_full_GlobalID_list = []
        mt_full_gis_globalid_list = []
        mt_GlobalID_list_with_gis_globalid = []
        mt_fields = common_fields + ['SHAPE@JSON','gis_globalid']
        with arcpy.da.SearchCursor(mt_resourceLyr,mt_fields) as cursor:
            for row in cursor:
                gis_globalid = row[-1]
                globalID = row[globalIDIndex]
                # if gis_globalid:
                #     if gis_globalid in mt_full_gis_globalid_list:
                #         wrn('Multiple instances of gis_globalid {}'.format(gis_globalid))
                #     elif '{' in gis_globalid:
                #         managed_task_dict[gis_globalid] = dict(zip(mt_fields,row))
                #         managed_task_dict[globalID] = dict(zip(mt_fields,row))
                #         mt_GlobalID_list_with_gis_globalid.append(globalID)
                #     else:
                #         mt_GlobalID_list_with_gis_globalid.append(globalID)
                #     mt_full_gis_globalid_list.append(gis_globalid)
                if gis_globalid:
                    if gis_globalid in mt_full_gis_globalid_list:
                        wrn('Multiple instances of gis_globalid {}'.format(gis_globalid))
                    elif '{' in gis_globalid:
                        managed_task_dict[gis_globalid] = dict(zip(mt_fields,row))
                        managed_task_dict[globalID] = dict(zip(mt_fields,row))
                        mt_GlobalID_list_with_gis_globalid.append(globalID)
                    mt_GlobalID_list_with_gis_globalid.append(globalID)
                    mt_full_gis_globalid_list.append(gis_globalid)
                else:
                    managed_task_dict[globalID] = dict(zip(mt_fields,row))
                mt_full_GlobalID_list.append(globalID)

        gis_fields = common_fields + ['SHAPE@JSON','globalid']
        msg('gis_fields = {}'.format(gis_fields))

        # use gis_globalid field to determine records to select in GIS
        mt_gis_globalid_list = list(managed_task_dict.keys())
        msg('Length of managed task dict: {}'.format(len(managed_task_dict)))

        msg('mt_gis_globalid_list: {}: {}'.format(len(mt_gis_globalid_list),mt_gis_globalid_list))
        msg('mt_full_gis_globalid_list: {}: {}'.format(len(mt_full_gis_globalid_list),mt_full_gis_globalid_list))
        msg('mt_full_GlobalID_list: {}: {}'.format(len(mt_full_GlobalID_list),mt_full_GlobalID_list))

        # globalid_list_for_resource_query = list(set(mt_full_gis_globalid_list + mt_full_GlobalID_list))
        globalid_list_for_resource_query = list(set(mt_gis_globalid_list + mt_full_GlobalID_list))
        msg('globalid_list_for_resource_query: {}: {}'.format(len(globalid_list_for_resource_query),globalid_list_for_resource_query))

        globalid_to_form_record_globalid_dict = {}
        if globalid_list_for_resource_query:
            query = "GlobalID in ('{0}')".format("','".join(globalid_list_for_resource_query))
            msg('Query for Resource Layer: {}'.format(query))
            gis_dict = {}
            with arcpy.da.SearchCursor(resourceLyr,gis_fields,where_clause=query) as cursor:
                for row in cursor:
                    gis_dict[row[-1]] = dict(zip(gis_fields,row))

            query = "GlobalID in ('{0}') or form_record_globalid in ('{0}')".format("','".join(globalid_list_for_resource_query))
            msg('Query for Resource Layer: {}'.format(query))
            with arcpy.da.SearchCursor(resourceLyr,['globalid','form_record_globalid'],where_clause=query) as cursor:
                for row in cursor:
                    globalid_to_form_record_globalid_dict[row[0]] = tryUpper(addBrackets(row[1]))
            gis_GlobalID_list = list(globalid_to_form_record_globalid_dict.keys())

        else:
            gis_dict = {}
            gis_GlobalID_list = []

        # get full value lists
        gis_full_globalid_list = returnDistinctFieldValues(resourceLyr,'globalId',"resource_type = '{}'".format(resourceType))
        if not gis_full_globalid_list:
            gis_full_globalid_list = []
        gis_full_form_record_globalid_list = returnDistinctFieldValues(resourceLyr,'form_record_globalid',"resource_type = '{}'".format(resourceType))
        if gis_full_form_record_globalid_list:
            gis_full_form_record_globalid_list = [tryUpper(addBrackets(x)) for x in gis_full_form_record_globalid_list]
        else:
            gis_full_form_record_globalid_list = []
        matches_by_gis_globalid = [globalid_to_form_record_globalid_dict[x] for x in mt_full_gis_globalid_list if x in gis_full_globalid_list]
        matches_by_form_record_globalid = [x for x in mt_full_GlobalID_list if x in gis_full_form_record_globalid_list]
        match_by_both = [x for x in matches_by_gis_globalid if x in matches_by_form_record_globalid]
        match_by_only_gis_globalid = [x for x in matches_by_gis_globalid if x not in matches_by_form_record_globalid]
        match_by_only_form_record_globalid = [x for x in matches_by_form_record_globalid if x not in matches_by_gis_globalid]
        no_match = [x for x in mt_full_GlobalID_list if x not in match_by_both and x not in match_by_only_gis_globalid and x not in match_by_only_form_record_globalid]

        #msg('gis_full_globalid_list: {}:{}'.format(len(gis_full_globalid_list),gis_full_globalid_list))
        msg('mt_full_gis_globalid_list: {}:{}'.format(len(mt_full_gis_globalid_list),mt_full_gis_globalid_list))
        #msg('gis_full_form_record_globalid_list: {}:{}'.format(len(gis_full_form_record_globalid_list),gis_full_form_record_globalid_list))
        msg('mt_full_GlobalID_list: {}:{}'.format(len(mt_full_GlobalID_list),mt_full_GlobalID_list))

        msg('matches_by_gis_globalid: {}:{}'.format(len(matches_by_gis_globalid),matches_by_gis_globalid))
        msg('matches_by_form_record_globalid: {}:{}'.format(len(matches_by_form_record_globalid),matches_by_form_record_globalid))

        msg('match_by_both: {}:{}'.format(len(match_by_both),match_by_both))
        msg('match_by_only_gis_globalid: {}:{}'.format(len(match_by_only_gis_globalid),match_by_only_gis_globalid))
        msg('match_by_only_form_record_globalid: {}:{}'.format(len(match_by_only_form_record_globalid),match_by_only_form_record_globalid))
        msg('no_match: {}:{}'.format(len(no_match),no_match))

        gis_mt_different_values_dict = compare_GIS_and_MT_Resources(gis_dict,gis_fields,gis_wkid,managed_task_dict,mt_fields,mt_wkid)

        plain_msg('\n\ngis_mt_different_values_dict = {}'.format(gis_mt_different_values_dict))
        plain_msg('\nglobalid_to_form_record_globalid_dict = {}'.format(globalid_to_form_record_globalid_dict))

        form_record_globalid_to_globalid_dict = {v:k for k,v in globalid_to_form_record_globalid_dict.items()}
        plain_msg('\nform_record_globalid_to_globalid_dict = {}'.format(form_record_globalid_to_globalid_dict))

        form_record_globalid_mt_different_values_dict = {globalid_to_form_record_globalid_dict[globalid]:valueDict for globalid,valueDict in gis_mt_different_values_dict.items()}
        plain_msg('\n\nform_record_globalid_mt_different_values_dict = {}'.format(form_record_globalid_mt_different_values_dict))

        form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=resourceType,formType='Resource',gis=gis)
        # ### TODO*** - change after testing -
        # feature_layer_item_id = 'fa6d5bc6f21e49c8b5cb6cc16e780ff7'
        # form_item_id = '23483e7c8ac24870b8af717777ada78b'
        lyrs = returnAllLayerObjectsFromAGOLItem(gis,item_id=feature_layer_item_id)

        arcpy.env.workspace = fgdb
        msg('Workspace: {}'.format(arcpy.env.workspace))

        master_different_values_dict = {}
        master_new_records_dict = {}

        # get feature class
        #fcList = arcpy.ListFeatureClasses()
        #msg('FC List: {}'.format(fcList))
        # fc = os.path.join(fgdb,fcList[0])
        mt_fc_name = mt_resourceLyr.name
        mt_fc = os.path.join(fgdb,mt_fc_name) # r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb\cultural_landscape_poly"
        msg('Managed Task FC: {}'.format(mt_fc))
        #mt_fc_name = os.path.basename(mt_fc)

        mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_fc)]
        globalIDIndex = mt_fields.index('globalid')
        managed_task_dict = {}
        with arcpy.da.SearchCursor(mt_fc,mt_fields + ['SHAPE@JSON']) as cursor:
            for row in cursor:
                managed_task_dict[row[globalIDIndex].upper()] = dict(zip(mt_fields + ['SHAPE@JSON'],row))

        mt_globalIDList = [key for key in managed_task_dict.keys()]
        msg('Managed Task Record Count: {}'.format(len(managed_task_dict)))
        msg(mt_globalIDList)
        #msg('Managed Task Dict: {}'.format(managed_task_dict))
        query = "globalid in ('{}')".format("','".join(mt_globalIDList))
        lyr = [lyr for lyr in lyrs if lyr.properties.name == mt_fc_name][0]
        agol_features = getFeaturesFromLyr(lyr, query=query, outputFields='*', returnGeometry=True)
        msg('AGOL Features Count: {}'.format(len(agol_features)))
        #msg('AGOL Features: {}'.format(agol_features))
        agol_globalIDList = [agol_feature.attributes['globalid'] for agol_feature in agol_features]
        msg(agol_globalIDList)
        agol_different_values_dict, sync_changed_GlobalIDs = compare_AGOL_and_MT_Resources(agol_features,managed_task_dict,mt_fields,COMMON_RESOURCE_FIELDS_FOR_SYNC)

        if len(mt_globalIDList) != len(agol_globalIDList):
            msg('Counts do not match between MT and AGOL')
            # msg(mt_globalIDList)
            # msg(agol_globalIDList)
            # mtNotAGOL = [x for x in mt_globalIDList if x not in agol_GlobalIDList]
            # agolNotMT = [x for x in agol_GlobalIDList if x not in mt_globalIDList]
            # wrn('In MT but not AGOL: {}'.format(mtNotAGOL))
            # wrn('In AGOL but not MT: {}'.format(agolNotMT))
        plain_msg('\nagol_different_values_dict = {}'.format(agol_different_values_dict))


        # compare AGOL and MT and GIS
        def compare_AGOL_MT_GIS_dict(agol_different_values_dict,gis_mt_different_values_dict,globalid_to_form_record_globalid_dict,form_record_globalid_to_globalid_dict,managed_task_dict,agol_features,gis_dict):
            agol_mt_gis_values_dict = {} # new master dict that also contains gis value
            gis_mt_different_values_form_globalid_dict = {globalid_to_form_record_globalid_dict[k]:v for k,v in gis_mt_different_values_dict.items()}
            agol_different_values_dict = {addBrackets(k):v for k,v in agol_different_values_dict.items()}
            different_globalids = list(set(list(gis_mt_different_values_form_globalid_dict.keys()) + list(agol_different_values_dict.keys())))
            for globalid in different_globalids:
                agol_fields_diff = [x[2] for x in agol_different_values_dict.get(globalid,[])]
                gis_fields_diff = [x[0] for x in gis_mt_different_values_form_globalid_dict.get(globalid,[])]
                fields_diff = list(set(agol_fields_diff + gis_fields_diff))
                agol_mt_val_lists = agol_different_values_dict.get(globalid,None)
                gis_mt_val_lists = gis_mt_different_values_form_globalid_dict.get(globalid,None)
                single_agol_mt_gis_values_list = []
                for field in fields_diff:
                    # set default as False for each value
                    agol_val = False
                    mt_val = False
                    gis_val = False
                    res_id = False
                    res_state_id = False
                    if agol_mt_val_lists:
                        agol_mt_val_list = [x for x in agol_mt_val_lists if x[2] == field]
                        if agol_mt_val_list:
                            msg(agol_mt_val_list)
                            res_id = agol_mt_val_list[0][0]
                            res_state_id = agol_mt_val_list[0][1]
                            agol_val = agol_mt_val_list[0][3]
                            mt_val = agol_mt_val_list[0][4]
                    if gis_mt_val_lists:
                        gis_mt_val_list = [x for x in gis_mt_val_lists if x[0] == field]
                        if gis_mt_val_list:
                            msg(gis_mt_val_list)
                            gis_val = gis_mt_val_list[0][1]
                            mt_val = gis_mt_val_list[0][2]

                    # assume that if it is not in the change dictionary, it matches the MT value
                    if agol_val == False:
                        agol_val = mt_val

                    gis_dict_values = gis_dict.get(form_record_globalid_to_globalid_dict.get(globalid,None),{})
                    msg('gis_dict_values: {}'.format(gis_dict_values))
                    if res_id == False: # if not in AGOL, try GIS value
                        res_id = gis_dict_values.get('resource_id',None)

                    if res_state_id == False: # if not in AGOL, try GIS value
                        res_state_id = gis_dict_values.get('res_state_id',None)

                    if gis_val == False:
                        if field in gis_fields:
                            gis_val = gis_dict_values.get(field,None)
                        elif field == 'gis_globalid':
                            msg('globalid to search for: {}'.format(globalid))
                            gis_val = form_record_globalid_to_globalid_dict.get(globalid,'???')
                        elif field == 'Geometry':
                            gis_val = None
                        else:
                            gis_val = 'n/a'

                    single_agol_mt_gis_values_list.append((res_id,res_state_id,field,agol_val,mt_val,gis_val))

                agol_mt_gis_values_dict[globalid] = single_agol_mt_gis_values_list
            return agol_mt_gis_values_dict


        agol_mt_gis_values_dict = compare_AGOL_MT_GIS_dict(agol_different_values_dict,gis_mt_different_values_dict,globalid_to_form_record_globalid_dict,form_record_globalid_to_globalid_dict,managed_task_dict,agol_features,gis_dict)
        msg('agol_mt_gis_values_dict = {}'.format(agol_mt_gis_values_dict))

        master_different_values_dict[mt_fc_name] = agol_mt_gis_values_dict
        agolGlobalIDList = [x.attributes['globalid'].upper() for x in agol_features]
        notInAGOL = [x for x in mt_globalIDList if removeBrackets(x) not in agolGlobalIDList]
        msg(f'{notInAGOL=}')
        #if notInAGOL:

        # find if any of those missing in AGOL have a gis_globalid populated. If so, that record may have been removed from the TVA main survey
        mt_globalid_missing_from_AGOL_with_gis_globalid = [globalid for globalid,valueDict in managed_task_dict.items() if globalid in notInAGOL and valueDict['gis_globalid']]
        new_mt_fc = [globalid for globalid in notInAGOL if globalid not in mt_globalid_missing_from_AGOL_with_gis_globalid]
        master_new_records_dict[mt_fc_name] = new_mt_fc

        # repeat process for tables
        arcpy.env.workspace = fgdb
        msg('Workspace: {}'.format(arcpy.env.workspace))
        tblList = arcpy.ListTables()
        msg('Tbl List: {}'.format(tblList))
        tblList = [x for x in tblList if os.path.basename(x) != 'GDB_ServiceItems']
        #parentGlobalIDSet = set()
        parentGlobalIDDict = {}
        for mt_tbl in tblList:
            mt_tbl_name = os.path.basename(mt_tbl)
            msg('Table: {}'.format(mt_tbl))
            mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_tbl)]
            globalIDIndex = mt_fields.index('globalid')
            managed_task_dict = {}
            with arcpy.da.SearchCursor(mt_tbl,mt_fields) as cursor:
                for row in cursor:
                    managed_task_dict[row[globalIDIndex].upper()] = dict(zip(mt_fields,row))
            # tblParentGlobalIDSet = returnDistinctFieldValues(mt_tbl,'ParentGlobalID')
            # if tblParentGlobalIDSet:
            #     msg('Table: {}; ParentGlobalIDList: {}:{}'.format(mt_tbl_name,len(tblParentGlobalIDSet),tblParentGlobalIDSet))
            # else:
            #     msg('No records for {}'.format(mt_tbl_name))
            # parentGlobalIDSet.update(tblParentGlobalIDSet)
            msg('Managed Task Records: {}'.format(len(managed_task_dict)))
            msg('Managed Task Dict: {}'.format(managed_task_dict))
            parentGlobalIDDict.update({k:v['parentglobalid'] for k,v in managed_task_dict.items()})
            msg('Parent GlobalID Dict: {}'.format(parentGlobalIDDict))
            if managed_task_dict:
                mt_globalIDList = [key for key in managed_task_dict.keys()]
                lyr = [lyr for lyr in lyrs if lyr.properties.name == mt_tbl_name][0]
                query = "globalid in ('{}')".format("','".join(list(mt_globalIDList)))
                agol_features = getFeaturesFromLyr(lyr, query=query, outputFields='*', returnGeometry=False)
                msg('AGOL Features: {}'.format(len(agol_features)))
                agol_different_values_dict, _ = compare_AGOL_and_MT_Resources(agol_features,managed_task_dict,mt_fields,None)
                agolGlobalIDList = [x.attributes['globalid'].upper() for x in agol_features]
                new_mt = [x for x in mt_globalIDList if x.replace('{','').replace('}','') not in agolGlobalIDList]
                master_new_records_dict[mt_tbl_name] = new_mt
            else:
                agol_different_values_dict = {}
                master_new_records_dict[mt_tbl_name] = []
            master_different_values_dict[mt_tbl_name] = agol_different_values_dict


        msg('Differences:\n{}'.format(master_different_values_dict))
        msg('New:\n{}'.format(master_new_records_dict))

        new_and_updated_mt_globalids = []

        df_values = []
        query_df_diff_values = []
        cultural_resources_difference_query = ''
        # list out differences
        for layer,layerDict in master_different_values_dict.items():
            cultural_resources_difference_global_ids = []
            difference_globalidList = []
            for globalid,valueLists in layerDict.items():
                form_url = generateFormURL(form_item_id,globalid)
                globalid = addBrackets(globalid)
                new_and_updated_mt_globalids.append(globalid)
                for valueList in valueLists:
                    df_value_list = ['difference',layer,globalid] + list(valueList) + [form_url]
                    df_values.append(df_value_list)
                    if valueList[-1] != 'n/a':
                        cultural_resources_difference_global_ids.append(globalid)
                difference_globalidList.append(globalid)
            if difference_globalidList:
                query = "globalid in ('{}')".format("','".join(difference_globalidList))
                msg("Different {}\n{}".format(layer,query))
                query_df_diff_values.append(['selection query for records with differences',layer,query])
                df_values.append(["","","","","","","","",""])

                if layer == mt_fc_name:
                    cultural_resources_difference_query = "form_record_globalid in ('{}')".format("','".join(cultural_resources_difference_global_ids))

        # finding layer with the map
        aprxMap = returnMapObject()
        lyrList = getValidLayers(aprxMap,searchLyrName=resourceLyr)
        rLyr = lyrList[0]
        desc = arcpy.Describe(rLyr)
        rLyrDataSource = desc.catalogPath
        msg('Resource Layer Data Source: {}'.format(rLyrDataSource))


        mt_lyrList = getValidLayers(aprxMap,searchLyrName=mt_resourceLyr)
        mt_rLyr = mt_lyrList[0]
        desc = arcpy.Describe(mt_rLyr)
        mt_rLyrDataSource = desc.catalogPath
        msg('Managed Task Resource Layer Data Source: {}'.format(mt_rLyrDataSource))

        # add query for new records
        query_df_new_values = []
        for layer,newList in master_new_records_dict.items():
            new_globalidList = []
            for globalid in newList:
                # df_value_list = ['new record',layer,globalid,'','','','']
                # df_values.append(df_value_list)
                globalid = addBrackets(globalid)
                new_globalidList.append(globalid)
                new_and_updated_mt_globalids.append(globalid)
            if new_globalidList:
                query = "globalid in ('{}')".format("','".join(new_globalidList))
                msg("New {}\n{}".format(layer,query))
                query_df_new_values.append(['selection query for new records',layer,query])

        update_summary_df_values = []
        for lyr in [mt_fc_name] + tblList:
            diff_count = len(master_different_values_dict.get(lyr,0))
            new_count = len(master_new_records_dict.get(lyr,0))
            append_status = 'Y' if diff_count or new_count else ''
            update_summary_df_values.append([resourceType,lyr,diff_count,new_count,append_status])


        all_new_and_updated_mt_globalids = new_and_updated_mt_globalids + [parentGlobalIDDict.get(globalid,globalid) for globalid in new_and_updated_mt_globalids]
        all_new_and_updated_mt_globalids = list(set(all_new_and_updated_mt_globalids))
        if all_new_and_updated_mt_globalids:
            # adding in subset of new/updated Resources from MT layer
            all_new_and_updated_mt_globalids_query = "globalid in ('{}')".format("','".join(all_new_and_updated_mt_globalids))
            try:
                addDatasetToAPRX(aprxMap_=aprxMap,dataset=mt_rLyrDataSource,layerName_='MT {} - To Append to GIS'.format(resourceType),layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery=all_new_and_updated_mt_globalids_query)
            except OSError:
                err('!!!Active Map must be open. Open or import a map to continue!!!')
                return
            # adding in subset of Resources from Enterprise GIS layer
            all_new_and_updated_mt_globalids_query = "form_record_globalid in ('{}')".format("','".join(all_new_and_updated_mt_globalids))
            try:
                addDatasetToAPRX(aprxMap_=aprxMap,dataset=rLyrDataSource,layerName_='{} with Survey Form Differences - To Export Report'.format(resourceType),layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery=all_new_and_updated_mt_globalids_query)
            except OSError:
                err('!!!Active Map must be open. Open or import a map to continue!!!')
                return

        if cultural_resources_difference_query:
            try:
                addDatasetToAPRX(aprxMap_=aprxMap,dataset=rLyrDataSource,layerName_='{} with GIS Differences - To Backup and Update'.format(resourceType),layerFile_='',replaceExisting_='REMOVE',visibility_=True,groupLayer_='',defQuery=cultural_resources_difference_query)
            except OSError:
                err('!!!Active Map must be open. Open or import a map to continue!!!')
                return

        # # add section for values that apply to syncing with GIS
        # df_values.append(["","","","","","","",""])
        # if sync_changed_GlobalIDs:
        #     # change_query = "globalid in ('{}')".format("','".join(sync_changed_GlobalIDs))
        #     change_query = "globalid in ('{" + "}','{".join(sync_changed_GlobalIDs) + "}')"
        # else:
        #     change_query = ""
        # df_values.append(['Sync selection query for records with differences',mt_fc_name,change_query,"","","","","",""])
        # if new_mt_fc:
        #     new_query = "globalid in ('{}')".format("','".join(new_mt_fc))
        # else:
        #     new_query = ""
        # df_values.append(['Sync selection query for new records',mt_fc_name,new_query,"","","","","",""])

        msg(df_values)
        for x in df_values:
            plain_msg('{}:{}'.format(len(x),x))
        #msg('Copied to clipboard: {}'.format("\n".join([str(x) for x in df_values])))


        df_dict = {}
        columns = ['change_type','layer','unique_form_record_globalid','resource_id','resource_state_id','field','current_tva_agol_value','managed_task_value','gis_value','form_record_url']
        if not df_values:
            df_values.append([""]*len(columns))
        df = pandas.DataFrame(df_values)
        df.columns = columns
        df_dict['Differences'] = df

        columns = ['change_type','layer','query']
        if not query_df_diff_values:
            query_df_diff_values.append([""]*len(columns))
        df = pandas.DataFrame(query_df_diff_values)
        df.columns = columns
        df_dict['Difference Queries'] = df

        columns = ['change_type','layer','query']
        if not query_df_new_values:
            query_df_new_values.append([""]*len(columns))
        df = pandas.DataFrame(query_df_new_values)
        df.columns = columns
        df_dict['New Record Queries'] = df

        columns = ['Resource_Type','Layer_Name','Difference_Count','New_Count','Need to Update AGOL Form?']
        if not update_summary_df_values:
            update_summary_df_values.append([""]*len(columns))
        df = pandas.DataFrame(update_summary_df_values)
        df.columns = columns
        df_dict['Update Summary'] = df


        tempFolder = createFolder(r'c:\temp')
        outputXLSX = os.path.join(tempFolder,'Change_Report_for_{}_{}_{}.xlsx'.format(resourceType,os.path.basename(fgdb).replace('.gdb',''),'{:%Y%m%d_%H%M%S}'.format(datetime.datetime.utcnow())))

        saveSheets(df_dict, outputXLSX, index=False, header=True)
        os.startfile(outputXLSX)

        return

class placeholderExportReport(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "3. Export Survey123 Reports"
        self.description = "Export Survey123 Reports"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):
        return exportSurvey123Report.getParameterInfo(self)

    def isLicensed(self):
        return True

    def updateParameters(self, parameters):
        return exportSurvey123Report.updateParameters(self, parameters)

    def updateMessages(self, parameters):
        return exportSurvey123Report.updateMessages(self, parameters)

    def execute(self, parameters, messages):
        return exportSurvey123Report.execute(self, parameters, messages)

''' Backup Resources into Cultural Resources Previous Boundaries '''
class copyResourcesIntoPreviousBoundaries(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "2. Backup Resources into Cultural Resources Previous Boundaries"
        self.description = "Backup Resources into Cultural Resources Previous Boundaries"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):
        """ Backup Resources into Cultural Resources Previous Boundaries """
        '''
            Defines the tool's parameters
        '''
        # fgdbParam = arcpy.Parameter(
        #     displayName='Managed Task FGDB',
        #     name='fgdb', ### Name the same as parameter in main script
        #     datatype='DEWorkspace',
        #     parameterType='Required',
        #     direction='Input')
        # fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb"

        resourceLyrParam = arcpy.Parameter(
            displayName='Resource Layer To Backup',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        #resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        prevResourceLyrParam = arcpy.Parameter(
            displayName='Resource Previous Boundary Layer',
            name='prevResourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        prevResourceLyrParam.value = resourcePrevBoundariesLyrParamDefault # set default cultural resources layer

        # mtResourceLyrParam = arcpy.Parameter(
        #     displayName='Managed Task Resource Layer',
        #     name='mtResourceLyr', ### Name the same as parameter in main script
        #     datatype='GPFeatureLayer',
        #     parameterType='Required',
        #     direction='Input')
        # #mtResourceLyrParam.value = mtResourceLyrParamDefault # set default cultural resources layer

        noteParam = arcpy.Parameter(
            displayName='Note reason for backup',
            name='note', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Required',
            direction='Input')

        refreshVersioningParam = arcpy.Parameter(
            displayName='Confirm Refresh of Version',
            name='refreshVersioning', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        refreshVersioningParam.value = False

        params = [resourceLyrParam,prevResourceLyrParam,noteParam,refreshVersioningParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Backup Resources into Cultural Resources Previous Boundaries """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # generate pNum and pValue dictionaries for parameters
        # pValues = paramDictValues(parameters) # k:name;v:value
        # pNum = paramDictNum(parameters) # k:name;v:number
        # # pValuesAsText = paramDictValuesAsText(parameters)
        # # resourceType = pValues['resourceType']
        # #selectNewRecords_TF = pValues['selectNewRecords_TF']
        # resourceLyr = pValues['resourceLyr']

        # if not resourceLyr:
        #     validLayerList = getValidLayers()
        #     dropdownList = [x.name for x in validLayerList if 'To Backup and Update' in x.name] # TODO - update if layer name changes
        #     parameters[pNum['resourceLyr']].filter.list = dropdownList

        return

    def updateMessages(self, parameters):
        """ Backup Resources into Cultural Resources Previous Boundaries """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)

        resourceLyr = pValues['resourceLyr']
        prevResourceLyr = pValues['prevResourceLyr']
        refreshVersioning = pValues['refreshVersioning']
        note = pValues['note']

        # check to make sure user is editing a personal version
        if not resourceLyr:
            return
        parameters,validVersion_TF = checkVersion(parameters,pNum,lyr=resourceLyr,paramName='resourceLyr')
        if not validVersion_TF:
            return

        # # verify selection on lyr
        # oidList = getOIDListFromLyrSelection(mt_resourceLyr)
        # if not oidList:
        #     parameters[pNum['mtResourceLyr']].setWarningMessage('No Features selected.')

        if note and len(note) > 50:
            parameters[pNum['note']].setWarningMessage('Note must be less than 50 characters..')

        if resourceLyr and prevResourceLyr and not refreshVersioning:
            parameters[pNum['refreshVersioning']].setWarningMessage('Must confirm that Version has been Refreshed through Versioning Toolbar.')
        return

    def execute(self, parameters, messages):
        """ Backup Resources into Cultural Resources Previous Boundaries """
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        resourceLyr = pValues['resourceLyr']
        prevResourceLyr = pValues['prevResourceLyr']
        refreshVersioning = pValues['refreshVersioning']
        note = pValues['note']

        if not refreshVersioning:
            wrn('Must confirm that you have refreshed the version through the Versioning Toolbar')
            return

        resource_fields = [f.name.lower() for f in arcpy.ListFields(resourceLyr)]
        msg(resource_fields)
        prevResource_fields = [f.name.lower() for f in arcpy.ListFields(prevResourceLyr)]
        msg(prevResource_fields)
        common_fields = [x for x in prevResource_fields if x in resource_fields if x not in ('globalid','objectid','shape','shape.starea()', 'shape.stlength()')]
        msg('common fields: {}'.format(common_fields))


        search_fields = common_fields + ['SHAPE@JSON']
        insert_fields = common_fields + ['SHAPE@JSON','previous_resource_class']
        resourceID_index = search_fields.index('resource_id')

        editSession = startEditingLyr(prevResourceLyr)
        with arcpy.da.InsertCursor(prevResourceLyr,insert_fields) as i_cursor:
            with arcpy.da.SearchCursor(resourceLyr, search_fields) as s_cursor:
                for row in s_cursor:
                    #msg('row: {}'.format(row))
                    new_row = list(row) + [note]
                    #msg('new row: {}'.format(new_row))
                    #i_cursor.insertRow(new_row)
                    resourceID = row[resourceID_index]
                    msg('Resource ID {} copied into Cultural Resources Previous Boundaries'.format(resourceID))
        stopEditing(editSession)

class placeholderManuallyUpdateResources(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "4. (Manually update existing GIS Cultural Resources to match Managed Task)"
        self.description = "Manually update existing GIS Cultural Resources to match Managed Task"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):

        resourceTypeParam = arcpy.Parameter(
            displayName="Make sure Managed Task Feature Class and GIS Features are in Sync prior to updating Survey123 Form",
            name='resourceType',
            datatype='GPString',
            parameterType='Optional',
            direction='Output')
        #resourceTypeParam.filter.type = 'ValueList'
        #resourceTypeParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])

        resourceLyrParam = arcpy.Parameter(
            displayName='Layers',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input',
            multiValue=True)
        #resourceLyrParam.filter.type = 'ValueList'

        projectIDParam = arcpy.Parameter(
            displayName="Project ID",
            name='projectID',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        projectIDParam.filter.type = 'ValueList'
        projectIDParam.filter.list = ['project1','project2']
        return [resourceTypeParam]

    def isLicensed(self):
        return True

    def updateParameters(self, parameters):
        # pValues = paramDictValues(parameters) # k:name;v:value
        # pNum = paramDictNum(parameters) # k:name;v:number
        # allLayers = getValidLayers()
        # #lyrsOfInterest = [lyr for lyr in allLayers if 'tx_siting' in arcpy.Describe(lyr).catalogPath.lower()]
        # lyrsOfInterest = [lyr.name for lyr in allLayers if 'cultural_resource' in arcpy.Describe(lyr).catalogPath.lower()]
        # parameters[pNum['resourceLyr']].filter.list = lyrsOfInterest
        # parameters[pNum['resourceLyr']].values = lyrsOfInterest

        # for lyr in lyrsOfInterest:
        #     #arcpy.SelectLayerByAttribute_management(lyr,selection_type='NEW_SELECTION',where_clause="created_user = '{}'".format(pcUser.upper()))
        #     arcpy.SelectLayerByAttribute_management(lyr,selection_type='NEW_SELECTION',where_clause="created_user = '{}'".format('JWALL0'))
        return parameters

    def execute(self, parameters, messages):
        return

''' Append New Managed Task Records into the Enterprise GDB '''
class appendNewMTtoGIS(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "5. Append New Managed Task Records into the Enterprise GDB"
        self.description = "Tool to Append New Managed Task Records into the Enterprise GDB"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):
        """ Append New Managed Task Records into the Enterprise GDB """
        '''
            Defines the tool's parameters
        '''
        # fgdbParam = arcpy.Parameter(
        #     displayName='Managed Task FGDB',
        #     name='fgdb', ### Name the same as parameter in main script
        #     datatype='DEWorkspace',
        #     parameterType='Required',
        #     direction='Input')
        # fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb"

        resourceTypeParam = arcpy.Parameter(
            displayName='Resource Type',
            name='resourceType',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        resourceTypeParam.filter.type = 'ValueList'
        resourceTypeParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])
        #resourceTypeParam.value = 'Cultural_Landscapes'
        #resourceTypeParam.value = 'Objects'

        resourceLyrParam = arcpy.Parameter(
            displayName='Resource Layer',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        mtResourceLyrParam = arcpy.Parameter(
            displayName='Managed Task Resource Layer',
            name='mtResourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        #mtResourceLyrParam.value = mtResourceLyrParamDefault # set default cultural resources layer

        queryForNewRecordsParam = arcpy.Parameter(
            displayName='Query for records to Append to GIS',
            name='queryForNewRecords', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Optional',
            direction='Input')
        #queryForNewRecordsParam.value = "globalid in ('{99E27D42-0DFB-4187-8F1A-D336EC384177}')"

        # selectNewRecords_TFParam = arcpy.Parameter(
        #     displayName='Select new records from Managed Task?',
        #     name='selectNewRecords_TF', ### Name the same as parameter in main script
        #     datatype='GPBoolean',
        #     parameterType='Optional',
        #     direction='Input')
        # selectNewRecords_TFParam.value = False

        refreshVersioningParam = arcpy.Parameter(
            displayName='Confirm Refresh of Version',
            name='refreshVersioning', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        refreshVersioningParam.value = False

        params = [resourceLyrParam,mtResourceLyrParam,resourceTypeParam,refreshVersioningParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Append New Managed Task Records into the Enterprise GDB """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)
        # resourceType = pValues['resourceType']
        #selectNewRecords_TF = pValues['selectNewRecords_TF']
        mt_resourceLyr = pValues['mtResourceLyr']
        resourceLyr = pValues['resourceLyr']

        if not mt_resourceLyr:
            #pNum['selectNewRecords_TF'].value = False
            return

        # oidList = getOIDListFromLyrSelection(mt_resourceLyr)
        # #if not oidList and selectNewRecords_TF:
        # if selectNewRecords_TF:
        #     #arcpy.SelectLayerByAttribute_management(mt_resourceLyr,"NEW_SELECTION","gis_globalid is null and resource_id is null","INVERT")
        #     arcpy.SelectLayerByAttribute_management(mt_resourceLyr,"NEW_SELECTION","1=2","INVERT")
        #     parameters[pNum['selectNewRecords_TF']].value = False

        #     arcpy.SelectLayerByAttribute_management(resourceLyr,"NEW_SELECTION","ACTIVITY_ID = '20188688792-AS'")

        # # Check GIS Connection to AGOL / Portal
        # gisURL = checkGIS_URL()
        # if gisURL != TVA_AGOL_URL:
        #     return

        return

    def updateMessages(self, parameters):
        """ Append New Managed Task Records into the Enterprise GDB """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)

        resourceLyr = pValues['resourceLyr']
        mt_resourceLyr = pValues['mtResourceLyr']
        refreshVersioning = pValues['refreshVersioning']

        gisURL = checkGIS_URL()
        if gisURL != TVA_AGOL_URL:
            parameters[0].setErrorMessage('Currently signed into {}. Sign in to ArcGIS Online: {}'.format(gisURL,TVA_AGOL_URL))
            return

        # check to make sure user is editing a personal version
        parameters,validVersion_TF = checkVersion(parameters,pNum,lyr=resourceLyr,paramName='resourceLyr')
        if not validVersion_TF:
            return

        # # verify selection on lyr
        # oidList = getOIDListFromLyrSelection(mt_resourceLyr)
        # if not oidList:
        #     parameters[pNum['mtResourceLyr']].setWarningMessage('No Features selected.')

        if resourceLyr and mt_resourceLyr and not refreshVersioning:
            parameters[pNum['refreshVersioning']].setWarningMessage('Must confirm that Version has been Refreshed through Versioning Toolbar.')

        return

    def execute(self, parameters, messages):
        """ Append New Managed Task Records into the Enterprise GDB """
        # Create GIS Connection to AGOL / Portal
        gis = connectToAGOL()

        pNum = paramDictNum(parameters) # k:name;v:number
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        # fgdb = r'{}'.format(pValuesAsText['fgdb'])
        mt_resourceLyr = pValues['mtResourceLyr']
        resourceType = pValues['resourceType']
        resourceLyr = pValues['resourceLyr']
        refreshVersioning = pValues['refreshVersioning']

        if not refreshVersioning:
            wrn('Must confirm that you have refreshed the version through the Versioning Toolbar')
            return

        gis_wkid = arcpy.Describe(resourceLyr).SpatialReference.factoryCode
        mt_wkid = arcpy.Describe(mt_resourceLyr).SpatialReference.factoryCode

        gis_fields = [f.name.lower() for f in arcpy.ListFields(resourceLyr)]
        msg(gis_fields)
        mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_resourceLyr)]
        msg(mt_fields)
        common_fields = [x for x in mt_fields if x in gis_fields if x not in ('globalid','objectid','shape')]
        msg('common fields: {}'.format(common_fields))


        # check to see if it exists in resource lyr
        # get globalid selection from MT
        mt_globalid_list = returnDistinctFieldValues(mt_resourceLyr,'globalid')
        msg('mt_globalid_list: {}'.format(mt_globalid_list))
        form_record_globalid_query = "form_record_globalid in ('{}')".format("','".join(mt_globalid_list))
        msg('form_record_globalid_query: {}'.format(form_record_globalid_query))

        # search for MT globalid in form_record_globalid in resource lyr
        gis_match_form_record_globalid_list = returnDistinctFieldValues(resourceLyr,'form_record_globalid',form_record_globalid_query)
        msg('gis_match_form_record_globalid_list: {}'.format(gis_match_form_record_globalid_list))

        # warn if any of the selected records already exist in the resource lyr
        if gis_match_form_record_globalid_list:
            wrn('Selected Records already in GIS Dataset: {}'.format(gis_match_form_record_globalid_list))

        # determine records to append
        records_to_append = [x for x in mt_globalid_list if x not in (globalid.upper() for globalid in gis_match_form_record_globalid_list)]
        msg('records_to_append: {}'.format(records_to_append))


        if not records_to_append:
            wrn('No records to append.')
            return

        records_to_append_globalid_query = "globalid in ('{}')".format("','".join(records_to_append))
        msg('records_to_append_globalid_query: {}'.format(records_to_append_globalid_query))

        form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=resourceType,formType='Resource',gis=gis)

        search_fields = common_fields + ['globalid', 'SHAPE@JSON']
        insert_fields = common_fields + ['form_record_globalid', 'SHAPE@JSON', 'form_item_id', 'feature_layer_item_id', 'form_url']


        def calcResourceID(resourceType):
            resourceID = getEpochTimeMilli()
            resourceID = resourceID + '-' + RESOURCE_TYPE_ABBREV_DICT[resourceType]
            msg('New ResourceID: {}'.format(resourceID))
            return resourceID

        resourceIDIndex = search_fields.index('resource_id')

        editSession = startEditingLyr(resourceLyr)
        with arcpy.da.InsertCursor(resourceLyr,insert_fields) as i_cursor:
            with arcpy.da.SearchCursor(mt_resourceLyr, search_fields, records_to_append_globalid_query) as s_cursor:
                for row in s_cursor:
                    msg('row: {}'.format(row))
                    shapeJSON = row[-1]
                    form_record_global_id = row[-2]
                    # convert shapeJSON if necessary
                    if gis_wkid != mt_wkid:
                        newShapeJSON = str(getGeometryDictFromShapeJSON(shapeJSON,target_WKID=gis_wkid))
                    else:
                        newShapeJSON = shapeJSON
                    new_row = list(row[:-1]) + [newShapeJSON, form_item_id, feature_layer_item_id, generateFormURL(form_item_id,form_record_global_id)]
                    # populate new resource ID
                    current_resource_id = new_row[resourceIDIndex]
                    # calculate new Resource ID
                    if not current_resource_id:
                        new_resource_id = calcResourceID(resourceType)
                        new_row[resourceIDIndex] = new_resource_id
                        msg('New Resource ID created for {}: {}'.format(form_record_global_id,new_resource_id))
                        time.sleep(1)
                    msg('new row: {}'.format(new_row))
                    i_cursor.insertRow(new_row)
                    msg('Row inserted')
        stopEditing(editSession)

        return

''' Populate Managed Task Resource ID and GIS GlobalID for New Records '''
class populateMTGISGlobalID(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "6. Populate Managed Task Resource ID and GIS GlobalID for New Records"
        self.description = "Tool Populates Managed Task Resource ID and GIS GlobalID for New Records"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):
        """ Populate Managed Task Resource ID and GIS GlobalID for New Records """
        """
            Defines the tool's parameters
        """
        # fgdbParam = arcpy.Parameter(
        #     displayName='Managed Task FGDB',
        #     name='fgdb', ### Name the same as parameter in main script
        #     datatype='DEWorkspace',
        #     parameterType='Required',
        #     direction='Input')
        # fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb"

        resourceTypeParam = arcpy.Parameter(
            displayName='Resource Type',
            name='resourceType',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        resourceTypeParam.filter.type = 'ValueList'
        resourceTypeParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])
        #resourceTypeParam.value = 'Cultural_Landscapes'

        resourceLyrParam = arcpy.Parameter(
            displayName='Resource Layer',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        mtResourceLyrParam = arcpy.Parameter(
            displayName='Managed Task Resource Layer',
            name='mtResourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        #mtResourceLyrParam.value = mtResourceLyrParamDefault # set default cultural resources layer

        queryForNewRecordsParam = arcpy.Parameter(
            displayName='Query for records to Append to GIS',
            name='queryForNewRecords', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Optional',
            direction='Input')
        #queryForNewRecordsParam.value = "globalid in ('{99E27D42-0DFB-4187-8F1A-D336EC384177}')"

        # selectNewRecords_TFParam = arcpy.Parameter(
        #     displayName='Select new records from Managed Task?',
        #     name='selectNewRecords_TF', ### Name the same as parameter in main script
        #     datatype='GPBoolean',
        #     parameterType='Optional',
        #     direction='Input')
        # selectNewRecords_TFParam.value = False

        refreshVersioningParam = arcpy.Parameter(
            displayName='Confirm Refresh of Version',
            name='refreshVersioning', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        refreshVersioningParam.value = False

        params = [resourceLyrParam,mtResourceLyrParam,resourceTypeParam,refreshVersioningParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Populate Managed Task Resource ID and GIS GlobalID for New Records """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # # generate pNum and pValue dictionaries for parameters
        # pValues = paramDictValues(parameters) # k:name;v:value
        # pNum = paramDictNum(parameters) # k:name;v:number
        # # pValuesAsText = paramDictValuesAsText(parameters)
        # # resourceType = pValues['resourceType']
        # selectNewRecords_TF = pValues['selectNewRecords_TF']
        # mt_resourceLyr = pValues['mtResourceLyr']

        # if not mt_resourceLyr:
        #     pNum['selectNewRecords_TF'].value = False
        #     return

        # oidList = getOIDListFromLyrSelection(mt_resourceLyr)
        # if not oidList and selectNewRecords_TF:
        #     arcpy.SelectLayerByAttribute_management(mt_resourceLyr,"NEW_SELECTION","gis_globalid is null and resource_id is null")
        #     pNum['selectNewRecords_TF'].value = False
        return

    def updateMessages(self, parameters):
        """ Populate Managed Task Resource ID and GIS GlobalID for New Records """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)

        resourceLyr = pValues['resourceLyr']
        mt_resourceLyr = pValues['mtResourceLyr']
        refreshVersioning = pValues['refreshVersioning']

        # check to make sure user is editing a personal version
        parameters,validVersion_TF = checkVersion(parameters,pNum,lyr=resourceLyr,paramName='resourceLyr')
        if not validVersion_TF:
            return

        # verify selection on lyr
        oidList = getOIDListFromLyrSelection(mt_resourceLyr)
        # if not oidList:
        #     parameters[pNum['mtResourceLyr']].setWarningMessage('No Features selected.')
        if oidList:
            parameters[pNum['mtResourceLyr']].setWarningMessage('{} Features selected.'.format(len(oidList)))

        if resourceLyr and mt_resourceLyr and not refreshVersioning:
            parameters[pNum['refreshVersioning']].setWarningMessage('Must confirm that Version has been Refreshed through Versioning Toolbar.')

        return

    def execute(self, parameters, messages):
        """ Populate Managed Task Resource ID and GIS GlobalID for New Records """
        # Create GIS Connection to AGOL / Portal
        gis = connectToAGOL()

        pNum = paramDictNum(parameters) # k:name;v:number
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        # fgdb = r'{}'.format(pValuesAsText['fgdb'])
        mt_resourceLyr = pValues['mtResourceLyr']
        resourceType = pValues['resourceType']
        resourceLyr = pValues['resourceLyr']
        refreshVersioning = pValues['refreshVersioning']

        if not refreshVersioning:
            wrn('Must confirm that you have refreshed the version through the Versioning Toolbar')
            return

        # check to see if it exists in resource lyr
        # get globalid selection from MT
        # filter to those where resource id has not been populated and gis_global_id has not been populated
        mt_globalid_list = returnDistinctFieldValues(mt_resourceLyr,'globalid','resource_id is null and gis_globalid is null')
        if not mt_globalid_list:
            wrn('No records found in {} matching criteria "resource_id is null and gis_globalid is null"'.format(mt_resourceLyr.name))
            return
        msg('mt_globalid_list: {}'.format(mt_globalid_list))
        form_record_globalid_query = "form_record_globalid in ('{}')".format("','".join(mt_globalid_list))
        msg('form_record_globalid_query: {}'.format(form_record_globalid_query))

        # search for MT globalid in form_record_globalid in resource lyr
        gis_match_form_record_globalid_list = returnDistinctFieldValues(resourceLyr,'form_record_globalid',form_record_globalid_query)
        msg('gis_match_form_record_globalid_list: {}'.format(gis_match_form_record_globalid_list))

        # determine records to update in the MT feature class
        records_to_update = [x for x in mt_globalid_list if x in (globalid.upper() for globalid in gis_match_form_record_globalid_list)]
        msg('records_to_update: {}'.format(records_to_update))

        # warn if any of the selected records do not exist in the resource lyr
        if not records_to_update or sorted(records_to_update) != sorted(mt_globalid_list):
            records_missing = [x for x in mt_globalid_list if x not in gis_match_form_record_globalid_list]
            wrn('Records missing from GIS: {}. Run Append tool for those records.'.format(records_missing))
            return

        records_to_update_globalid_query = "globalid in ('{}')".format("','".join(records_to_update))
        msg('records_to_update_globalid_query: {}'.format(records_to_update_globalid_query))

        records_to_update_form_record_globalid_query = "form_record_globalid in ('{}')".format("','".join(records_to_update))
        msg('records_to_update_form_record_globalid_query: {}'.format(records_to_update_form_record_globalid_query))

        form_record_globalid_to_gis_globalid_dict = {}
        form_record_globalid_to_resource_id_dict = {}
        with arcpy.da.SearchCursor(resourceLyr, ['form_record_globalid','globalid','resource_id'], records_to_update_form_record_globalid_query) as s_cursor:
            for row in s_cursor:
                form_record_globalid_to_gis_globalid_dict[row[0]] = row[1]
                form_record_globalid_to_resource_id_dict[row[0]] = row[2]


        # for globalid, gis_globalid in form_record_globalid_to_gis_globalid_dict.items():
        #     resource_id = form_record_globalid_to_resource_id_dict[globalid]
        #     #arcpy.management.SelectLayerByAttribute(mt_resourceLyr,"NEW_SELECTION","globalid = '{}'".format(globalid))
        #     msg('Attempting to populate {} with {} and {} using Calculate Field'.format(globalid,gis_globalid,resource_id))
        #     where_clause = "globalid = '{}'".format(globalid)
        #     arcpy.management.CalculateFields(mt_resourceLyr, "PYTHON3",
        #         [["gis_globalid", "'{}'".format(gis_globalid), where_clause],
        #         ["resource_id", "'{}'".format(resource_id), where_clause]])

        msg('form_record_globalid_to_gis_globalid_dict: {}'.format(form_record_globalid_to_gis_globalid_dict))
        desc = arcpy.Describe(mt_resourceLyr)
        versioned = desc.isVersioned
        msg('Versioned? {}'.format(versioned))
        mt_resourceLyr_path = desc.catalogPath
        workspace = os.path.dirname(desc.catalogPath)
        removeLayers(mt_resourceLyr)
        editSession = startEditSession(workspace,versioned)
        msg('Starting update cursor...')
        with arcpy.da.UpdateCursor(mt_resourceLyr_path, ['globalid','gis_globalid','resource_id'], records_to_update_globalid_query) as u_cursor:
            for row in u_cursor:
                globalid = row[0]
                # populate GIS Global ID
                msg('Attempting to populate {} with {} and {}'.format(globalid,row[1],row[2]))
                row[1] = form_record_globalid_to_gis_globalid_dict[globalid]
                # populate Resource ID
                row[2] = form_record_globalid_to_resource_id_dict[globalid]
                u_cursor.updateRow(row)
                msg('gis_global id populated for {}: {}'.format(globalid,row[1]))
                msg('resource_id populated for {}: {}'.format(globalid,row[2]))
                #time.sleep(1)
        stopEditing(editSession)
        msg('Attempting to add layer to map...')
        addDatasetToAPRX(dataset=mt_resourceLyr_path)


        msg('Managed Task Records Updated with GIS GlobalID and Resource ID')

        # insert cursor
        # with spatial reference transformation
        # map globalid to form_record_globalid
        # calc form_item_id, feature_layer_item_id, form_url

        return

''' Append Managed Task Deliverables to AGOL '''
class appendMTtoAGOL(object):
    def __init__(self):
        '''
            Initializes the tool class
        '''
        self.label = "7. Append Managed Task Deliverables to AGOL"
        self.description = "Tool to Append Managed Task Deliverables to AGOL"
        self.canRunInBackground = False
        self.category = "Managed Task Import"

    def getParameterInfo(self):
        """ Append Managed Task Deliverables to AGOL """
        '''
            Defines the tool's parameters
        '''
        fgdbParam = arcpy.Parameter(
            displayName='Managed Task FGDB',
            name='fgdb', ### Name the same as parameter in main script
            datatype='DEWorkspace',
            parameterType='Required',
            direction='Input')
        #fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\CRMS\Managed_Task_Resources\Received\Cultural_Landscapes_Test__Managed_Task\e2abddb5-c3f9-48eb-99aa-ddb0ca25c029.gdb"
        # fgdbParam.value = r"\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\Projects\70568\70568-AS-1\MT\Import_from_MT\Historic_Architectural_Resources_70568-AS-1\2bd66427-6ea8-4c6a-8823-7a6e8381fff7.gdb"
        # \\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\Projects\21573447673_26a Kim-Geyer\21573447673-AS-1\Export_to_MT
        resourceTypeParam = arcpy.Parameter(
            displayName='Resource Type',
            name='resourceType',
            datatype='GPSTRING',
            parameterType='Required',
            direction='Input')
        resourceTypeParam.filter.type = 'ValueList'
        resourceTypeParam.filter.list = sorted([k for k in RESOURCE_TYPE_DOMAIN_DICT.keys()])
        #resourceTypeParam.value = 'Cultural_Landscapes'

        resourceLyrParam = arcpy.Parameter(
            displayName='Resource Layer',
            name='resourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        resourceLyrParam.value = resourceLyrParamDefault # set default cultural resources layer

        # prevResourceLyrParam = arcpy.Parameter(
        #     displayName='Resource Previous Boundary Layer',
        #     name='prevResourceLyr', ### Name the same as parameter in main script
        #     datatype='GPFeatureLayer',
        #     parameterType='Required',
        #     direction='Input')
        # prevResourceLyrParam.value = resourcePrevBoundariesLyrParamDefault # set default cultural resources layer

        mtResourceLyrParam = arcpy.Parameter(
            displayName='Managed Task Resource Layer',
            name='mtResourceLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')
        #mtResourceLyrParam.value = mtResourceLyrParamDefault # set default cultural resources layer

        refreshVersioningParam = arcpy.Parameter(
            displayName='Confirm Refresh of Version and Reload of Tool',
            name='refreshVersioning', ### Name the same as parameter in main script
            datatype='GPBoolean',
            parameterType='Optional',
            direction='Input')
        refreshVersioningParam.value = False

        # params = [resourceLyrParam,prevResourceLyrParam,mtResourceLyrParam,queryForDifferencesParam,queryForNewRecordsParam]

        params = [resourceLyrParam,mtResourceLyrParam,fgdbParam,resourceTypeParam,refreshVersioningParam]

        return params

    def isLicensed(self):
        '''
            Returns whether the tool is licensed.
        '''
        return True

    def updateParameters(self, parameters):
        """ Append Managed Task Deliverables to AGOL """
        '''
            Called each time the user changes a parameter on the tool dialog box.
            After returning from Update Parameters, geoprocessing calls its internal validation routine.
        '''
        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        pValuesAsText = paramDictValuesAsText(parameters)
        # resourceType = pValues['resourceType']
        fgdb = r'{}'.format(pValuesAsText['fgdb'])
        mtResourceLyr = pValues['mtResourceLyr']

        if not mtResourceLyr:
            return

        catalogPath = arcpy.Describe(mtResourceLyr).catalogPath
        fgdb = os.path.dirname(catalogPath)
        # check to verify that the fgdb is in the proper folder
        if '\\Import_from_MT\\' not in fgdb:
            return

        parameters[pNum['fgdb']].value = fgdb

        arcpy.env.workspace = fgdb
        fc_name = arcpy.ListFeatureClasses()[0]
        #fc_path = os.path.join(fgdb,fc_name)

        # get Resource Type from Feature Class Name
        parameters[pNum['resourceType']].value = FEATURE_LAYER_NAME_RESOURCE_TYPE_DICT.get(fc_name)

        return

    def updateMessages(self, parameters):
        """ Append Managed Task Deliverables to AGOL """
        '''
            Called after returning from the internal validation routine.
            You can examine the messages created from internal validation and change them if desired.
        '''
        # generate pNum and pValue dictionaries for parameters
        # pValues = paramDictValues(parameters) # k:name;v:value
        # pNum = paramDictNum(parameters) # k:name;v:number
        # pValuesAsText = paramDictValuesAsText(parameters)
        # resourceType = pValues['resourceType']

        # generate pNum and pValue dictionaries for parameters
        pValues = paramDictValues(parameters) # k:name;v:value
        pNum = paramDictNum(parameters) # k:name;v:number
        pValuesAsText = paramDictValuesAsText(parameters)
        resourceLyr = pValues['resourceLyr']
        mtResourceLyr = pValues['mtResourceLyr']

        refreshVersioning = pValues['refreshVersioning']
        if not refreshVersioning:
            wrn('Must confirm that you have refreshed the version through the Versioning Toolbar')
            return

        if mtResourceLyr:
            catalogPath = arcpy.Describe(mtResourceLyr).catalogPath
            fgdb = os.path.dirname(catalogPath)
            if fgdb and '\\Import_from_MT\\' not in fgdb:
                parameters[pNum['mtResourceLyr']].setErrorMessage(f"Source FGDB is not in the Activity's Import_from_MT subfolder. Current source: {fgdb}")

        gisURL = checkGIS_URL()
        if gisURL != TVA_AGOL_URL:
            parameters[0].setErrorMessage('Currently signed into {} Portal. Sign in to ArcGIS Online: {}'.format(gisURL,TVA_AGOL_URL))
            return


        # check to make sure user is editing a personal version
        parameters,validVersion_TF = checkVersion(parameters,pNum,lyr=resourceLyr,paramName='resourceLyr')
        if not validVersion_TF:
            return

        return

    def execute(self, parameters, messages):
        """ Append Managed Task Deliverables to AGOL """
        # Create GIS Connection to AGOL / Portal
        gis = connectToAGOL()

        pNum = paramDictNum(parameters) # k:name;v:number
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        fgdb = r'{}'.format(pValuesAsText['fgdb'])
        resourceType = pValues['resourceType']
        mt_resourceLyr = pValues['mtResourceLyr']
        resourceLyr = pValues['resourceLyr']
        # prevResourceLyr = pValues['prevResourceLyr']


        # clear map selections, if any
        arcpy.SelectLayerByAttribute_management(mt_resourceLyr,'CLEAR_SELECTION')
        arcpy.SelectLayerByAttribute_management(resourceLyr,'CLEAR_SELECTION')
        #resourceType = 'Archaeological_Resources'
        form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=resourceType,formType='Resource',gis=gis)
        # ### TODO*** - change after testing -
        # feature_layer_item_id = 'fa6d5bc6f21e49c8b5cb6cc16e780ff7'
        # form_item_id = '23483e7c8ac24870b8af717777ada78b'
        lyrs = returnAllLayerObjectsFromAGOLItem(gis,item_id=feature_layer_item_id)

        mt_path = arcpy.Describe(mt_resourceLyr).catalogPath
        msg('Catalog Path: {}'.format(mt_path))

        arcpy.env.workspace = fgdb
        msg('Workspace: {}'.format(arcpy.env.workspace))

        # get list of layers in GDB
        tblList = arcpy.ListTables()
        msg('Tbl List: {}'.format(tblList))
        tblList = [x for x in tblList if os.path.basename(x) != 'GDB_ServiceItems']
        gdbServiceItemsTable = os.path.join(fgdb,'GDB_ServiceItems')

        lyrId_lyrName_dict = {}
        with arcpy.da.SearchCursor(gdbServiceItemsTable,['ItemId','DatasetName'],where_clause="ItemType=1") as cursor:
            for row in cursor:
                lyrId_lyrName_dict[row[0]] = row[1]
        lyrIdList = sorted(lyrId_lyrName_dict.keys())

        msg(lyrId_lyrName_dict)

        msg('Managed Task FC: {}'.format(mt_path))
        #mt_fc_name = os.path.basename(mt_fc)

        """
        orig_agol_globalid > gis_globalid
        gis_globalid > new_agol_globalid
        orig_agol_globalid > new_agol_globalid (parent_globalid > new_parent_globalid)
        """

        """
        verify uppercase vs lowercase globalids and gis_globalids

        SDE Attribute Table
        GlobalID - upper
        FORM_RECORD_GLOBALID - upper

        AGOL Service Layer in Pro
        globalid - lower
        gis_globalid - upper

        AGOL Service in Browser
        globalid - lower
        gis_globalid - upper
        parentGlobalID - lower

        MT FC
        GlobalID - upper
        gis_globalid - upper

        globalid = '{68b9562f-964c-4885-8813-e03cde46236c}'
        globalid = '{68B9562F-964C-4885-8813-E03CDE46236C}'

        globalid = '{005FAA54-1818-43F9-8536-23FBAB3533C8}'
        globalid = '{3CBDDE9E-AB8E-4954-8730-57E04390331E}'

        we want the parentGlobalID to be lowercase when recalculating.
        """

        # get list of MT globalids
        # get initial dictionary for local MT globalid and gis_globalid
        orig_globalid_gis_globalid_dict = {}
        mt_gis_globalid_list = []
        mt_globalid_list = []
        with arcpy.da.SearchCursor(mt_path,['globalid','gis_globalid']) as cursor: # TODO - verify with Jacob that it is ok to attempt to update all records, disregarding any selections
            orig_globalid_gis_globalid_dict = {row[0].upper():row[1].upper() for row in cursor}

        mt_globalid_list = list(orig_globalid_gis_globalid_dict.keys())
        mt_gis_globalid_list = list(orig_globalid_gis_globalid_dict.values())
        msg(f'{orig_globalid_gis_globalid_dict=}')
        # {'{68B9562F-964C-4885-8813-E03CDE46236C}': '{CE8DEDBC-8313-4C58-B0D9-9952E40BDF69}',
        # '{83C785C7-5123-4915-892E-36CB39C52A22}': '{11C31641-D3D9-4AFF-B085-6DBD675EDBE8}',
        # '{55903CEC-AF86-4906-B092-9B92FC279573}': '{5D418E2E-1671-410E-AD32-067BE7651D68}'}

        # check to ensure that all items have a gis_globalid
        missing_gis_globalids = [k for k,v in orig_globalid_gis_globalid_dict.items() if not v]
        if missing_gis_globalids:
            err(f'One or more records is missing a gis_globalid: {missing_gis_globalids}')
            return

        # append GIS features
        agol_lyr = lyrs[0].url
        base_url = os.path.dirname(agol_lyr)

        # create dictionary to store relationship between original local globalid and new globalid for new records
        orig_globalid_new_globalid_dict = {}
        gis_globalid_new_globalid_dict = {}
        # for each related table, get new globalid and update parentglobalid as appropriate
        orig_globalid_parent_globalid_dict = {}

        make_updates_TF = True

        for layer_id in lyrIdList:
            agol_path = base_url + '/' + str(layer_id)
            tbl_name = lyrId_lyrName_dict[layer_id]
            mt_path = os.path.join(fgdb,tbl_name)

            mt_fields = [f for f in arcpy.ListFields(mt_path)]
            mt_field_names = [f.name.lower() for f in mt_fields]
            date_field_names = [f.name.lower() for f in mt_fields if f.type == 'DATE']
            agol_field_names = [f.name.lower() for f in arcpy.ListFields(agol_path)]
            common_fields = [x for x in mt_field_names if x in agol_field_names]

            # get list of AGOL objectids, globalids and gis_globalids
            agolLyr = returnLyrObjectFromAGOLItem(gis,item_id=feature_layer_item_id,relativeLyrNum=layer_id)

            """
            #code to generate state and county fields - should already by done by MT, but might include in post-processing
            county_state_list, stcntyfips_list, state_list, state_county_list = getCountyList(counties_sde,resourceLyr)
            spatial_analysis_dict = {"state_name": ",".join(sorted(state_list)), "county": ",".join(sorted(state_county_list))}
            """
            msg(f'Processing {mt_path}')
            if layer_id == 0: # first layer, spatial
                agolQuery = "gis_globalid in ('{}')".format("','".join(mt_gis_globalid_list))
                existingFeatures = getFeaturesFromLyr(agolLyr,query=agolQuery,outputFields=['globalid','objectid','gis_globalid'],returnGeometry=False)
                found_gis_globalids = [feature.attributes['gis_globalid'].upper() for feature in existingFeatures]
                agol_gis_globalid_objectid_dict = {feature.attributes['gis_globalid'].upper():feature.attributes['objectid'] for feature in existingFeatures}
                agol_gis_globalid_globalid_dict = {feature.attributes['gis_globalid'].upper():feature.attributes['globalid'].upper() for feature in existingFeatures}

                msg(f'{len(found_gis_globalids)} found in AGOL: {found_gis_globalids}')

                new_or_updated_record_count = 0
                with arcpy.da.SearchCursor(mt_path,common_fields + ['SHAPE@JSON']) as cursor:
                    for row in cursor:
                        shapeJSON = row[-1]
                        geometryDict = getGeometryDictFromShapeJSON(shapeJSON,target_WKID=SURVEY123_WKID)
                        #attributes_dict = {field:datetimeUTCToEpochMilli(row[i]) if field in date_field_names else field:row[i] for i,field in common_fields} # TODO - may need to exclude shape area and shape length
                        attributes_dict = {field:datetimeUTCToEpochMilli(row[i]) for i,field in enumerate(common_fields)} # TODO - may need to exclude shape area and shape length
                        gis_globalid = attributes_dict['gis_globalid'].upper()
                        mt_globalid = attributes_dict['globalid'].upper()

                        found_count = found_gis_globalids.count(gis_globalid)

                        form_globalid = '' # TODO - remove after testing
                        if found_count > 1:
                            wrn(f'Multiple records exist in AGOL for gis_globalid: {gis_globalid}')
                            continue

                        elif found_count == 1:
                            # update record
                            msg(f'Updating record {mt_globalid}...')
                            attributes_dict['objectid'] = agol_gis_globalid_objectid_dict[gis_globalid]
                            attributes_dict['globalid'] = agol_gis_globalid_globalid_dict[gis_globalid]
                            msg(f'Attributes Dict: {attributes_dict}')
                            updates_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
                            if make_updates_TF:
                                form_globalid, form_objectid = updateFormRecord(agolLyr,updates_dict_list)
                                msg(f'Updated GlobalID: {form_globalid}')
                                orig_globalid_new_globalid_dict[mt_globalid] = form_globalid.upper()
                                new_or_updated_record_count+=1

                        else:
                            # add record
                            msg(f'Adding new record for {mt_globalid}...')
                            del attributes_dict['objectid']
                            del attributes_dict['globalid']
                            msg(f'Attributes Dict: {attributes_dict}')
                            adds_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
                            if make_updates_TF:
                                form_globalid, form_objectid = addFormRecord(agolLyr,adds_dict_list)
                                msg(f'New GlobalID: {form_globalid}')
                                orig_globalid_new_globalid_dict[mt_globalid] = form_globalid.upper()
                                new_or_updated_record_count

                        gis_globalid_new_globalid_dict[gis_globalid] = form_globalid.upper()



                # 64176AR #

                # update update form_record_globalid, form_url in SDE Cultural Resources layer if difference exists
                if not new_or_updated_record_count:
                    continue

                msg('Updating form record globalid in Resource feature class.')
                editSession = startEditingLyr(resourceLyr)
                where_clause="GlobalID IN ('{}')".format("','".join(gis_globalid_new_globalid_dict.keys()))
                with arcpy.da.UpdateCursor(resourceLyr,['GlobalID','form_record_globalid','form_item_id','form_url'],where_clause=where_clause) as cursor:
                    for row in cursor:
                        globalid = row[0].upper()
                        form_record_globalid = row[1].upper()
                        new_form_record_globalid = gis_globalid_new_globalid_dict.get(globalid)
                        if not new_form_record_globalid:
                            wrn(f'No GIS record found for globalid {globalid}')
                            continue
                        if form_record_globalid == new_form_record_globalid:
                            msg(f'No update needed in GIS for globalid {globalid}')
                            continue
                        msg(f'Update needed in GIS for globalid {globalid}')
                        form_item_id = row[2]
                        new_form_url = generateFormURL(form_item_id,new_form_record_globalid)
                        row[1] = new_form_record_globalid
                        row[3] = new_form_url
                        msg('Updating {} to {}'.format(form_record_globalid,new_form_record_globalid))
                        cursor.updateRow(row)
                stopEditing(editSession)

                continue


            # process for related tables
            with arcpy.da.SearchCursor(mt_path,['globalid','parentglobalid']) as cursor:
                for row in cursor:
                    orig_globalid_parent_globalid_dict[row[0].upper()] = row[1].upper()

            if not orig_globalid_parent_globalid_dict:
                wrn(f'No records for {mt_path} layer')
                continue

            agolQuery = "globalid in ('{}')".format("','".join(orig_globalid_parent_globalid_dict.values()))
            # TODO - account for cases where there are no records
            existingFeatures = getFeaturesFromLyr(agolLyr,query=agolQuery,outputFields=['globalid','objectid','parentglobalid'],returnGeometry=False)
            found_globalids = [feature.attributes['globalid'].upper() for feature in existingFeatures]
            agol_globalid_objectid_dict = {feature.attributes['globalid'].upper():feature.attributes['objectid'] for feature in existingFeatures}
            #agol_gis_globalid_globalid_dict = {feature.attributes['gis_globalid'].upper():feature.attributes['globalid'].upper() for feature in existingFeatures}


            with arcpy.da.SearchCursor(mt_path,common_fields) as cursor:
                for row in cursor:
                    #attributes_dict = {field:datetimeUTCToEpochMilli(row[i]) if field in date_field_names else field:row[i] for i,field in common_fields}
                    attributes_dict = {field:datetimeUTCToEpochMilli(row[i]) for i,field in common_fields}
                    mt_parentglobalid = attributes_dict['parentglobalid'].upper()
                    mt_globalid = attributes_dict['globalid'].upper()

                    found_count = found_globalids.count(mt_globalid)

                    if found_count > 1:
                        wrn(f'Multiple records exist in AGOL for globalid: {mt_globalid}')
                        continue

                    elif found_count == 1:
                        # update record
                        msg('Updating record {mt_globalid}...')
                        attributes_dict['objectid'] = agol_globalid_objectid_dict[mt_globalid]
                        attributes_dict['globalid'] = mt_globalid
                        msg('Attributes Dict: {}'.format(attributes_dict))
                        updates_dict_list = [{"attributes":attributes_dict}]
                        if make_updates_TF:
                            form_globalid, form_objectid = updateFormRecord(agolLyr,updates_dict_list)
                            msg(f'Updated GlobalID: {form_globalid}')
                            orig_globalid_new_globalid_dict[mt_globalid] = form_globalid.upper()

                    else:
                        # add record
                        msg('Updating record for {mt_globalid}...')
                        del attributes_dict['objectid']
                        del attributes_dict['globalid']
                        new_parentglobalid = orig_globalid_new_globalid_dict.get(mt_parentglobalid)
                        if not new_parentglobalid:
                            wrn(f'No new parentglobalid record found for globalid {mt_globalid}, parentglobalid {mt_parentglobalid}')
                            continue
                        attributes_dict['parentglobalid'] = new_parentglobalid # update parentglobalid in cases where new records are created
                        msg('Attributes Dict: {}'.format(attributes_dict))
                        adds_dict_list = [{"attributes":attributes_dict}]
                        if make_updates_TF:
                            form_globalid, form_objectid = addFormRecord(agolLyr,adds_dict_list)
                            msg(f'New GlobalID: {form_globalid}')
                            orig_globalid_new_globalid_dict[mt_globalid] = form_globalid.upper()



# # lookup new globalid values based on gis_globalid
# gis_globalid_list = orig_globalid_gis_globalid_dict.values()
# where_clause = "gis_globalid in ('{}')".format("','".join(gis_globalid_list))
# gis_globalid_new_globalid_dict = {}
# with arcpy.da.SearchCursor(agol_path,['gis_globalid','globalid'],where_clause=where_clause) as cursor:
#     gis_globalid_new_globalid_dict = {row[0]:row[1].upper() for row in cursor}

# {'{CE8DEDBC-8313-4C58-B0D9-9952E40BDF69}': '{68B9562F-964C-4885-8813-E03CDE46236C}',
#  '{11C31641-D3D9-4AFF-B085-6DBD675EDBE8}': '{E1F501B3-FEC8-4FF4-9C53-17A3BB3A29E0}',
#  '{5D418E2E-1671-410E-AD32-067BE7651D68}': '{10AC0683-E821-49D1-A34A-4DFEDBBE5B0D}'}

# # create dictionary to store relationship between original local globalid and new globalid for new records
# orig_globalid_new_globalid_dict = {}
# for orig_globalid,gis_globalid in orig_globalid_gis_globalid_dict.items():
#     new_globalid = gis_globalid_new_globalid_dict[gis_globalid]
#     orig_globalid_new_globalid_dict[orig_globalid] = new_globalid

# {'{68B9562F-964C-4885-8813-E03CDE46236C}': '{68B9562F-964C-4885-8813-E03CDE46236C}',
#  '{83C785C7-5123-4915-892E-36CB39C52A22}': '{E1F501B3-FEC8-4FF4-9C53-17A3BB3A29E0}',
#  '{55903CEC-AF86-4906-B092-9B92FC279573}': '{10AC0683-E821-49D1-A34A-4DFEDBBE5B0D}'}





# ############################
# else:
#     # warning that there are multiple
#     wrn('Multiple records existExisting Features Attributes: {}'.format([feature.attributes for feature in existingFeatures]))
#     # TODO - ask how they want to handle multiple Survey123 records with the same gis_globalid
#     existingFeature = existingFeatures[0]
#     # existingFeatureAttributes = existingFeature['attributes']
#     # existingGlobalID = existingFeatureAttributes['globalid']
#     # existingObjectid = existingFeatureAttributes['objectid']
#     existingGlobalID = existingFeature.attributes['globalid']
#     existingObjectid = existingFeature.attributes['objectid']
#     existingFeatureCount = len(existingFeatures)
#     if existingFeatureCount > 1:
#         msg('Existing Features Attributes: {}'.format([feature.attributes for feature in existingFeatures]))
#         multipleGlobalIDs = [feature.attributes['globalid'] for feature in existingFeatures]
#         #form_url = generateFormURL(form_item_id,existingGlobalID)
#         #wrn('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,multipleGlobalIDs,existingObjectid,existingGlobalID,form_url))
#         form_url_list = [generateFormURL(form_item_id,globalid) for globalid in multipleGlobalIDs]
#         err('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,"\n".join(form_url_list)))
#     msg('Survey Record for GIS GlobalID {} already exists in {}: ObjectID: {}, GlobalID: {}'.format(gis_globalid,itemTitle,existingObjectid,existingGlobalID))
#     updateExistingRecord = True # TODO - make parameter to determine if update is needed
#     if updateExistingRecord:
#         form_globalid = updateMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,existingGlobalID,existingObjectid,gis_globalid,attribute_dict,shapeJSON=shapeJSON)
#     else:
#         msg('No change to existing record')
#         form_globalid = existingGlobalID
#     # update GIS dataset based on record in Survey123
#     form_url = generateFormURL(form_item_id,form_globalid)
#     calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
#     new_record_TF = False
# else:
#     form_globalid = addMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict,shapeJSON=shapeJSON)
#     msg('Survey Record for GIS GlobalID {} does not yet exist in {}.'.format(gis_globalid,itemTitle))
#     if isinstance(form_globalid,str): # check to see if the resulting globalid is a string
#         form_url = generateFormURL(form_item_id,form_globalid)
#         calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
#         new_record_TF = True
#     else:
#         err('Error creating new Survey123 record: {}'.format(form_globalid))
#         form_url = None
#         form_globalid = None
#         new_record_TF = False











# existingFeatures = getFeaturesFromLyr(mainSurveyLyr,query=agolQuery,outputFields=['globalid','objectid','gis_globalid'],returnGeometry=False)

# with arcpy.EnvManager(preserveGlobalIds=False): # note that preserveGlobalIds does not work when appending into AGOL
#     arcpy.management.Append(
#         inputs=mt_resourceLyr,
#         target=agol_path, # "https://services.arcgis.com/w8auYAijfGK1Mydj/arcgis/rest/services/service_e5c4c15f908a40ce8c1a2ca9475a2711/FeatureServer/0"
#         schema_type="TEST",
#         field_mapping=None,
#         subtype="",
#         expression="",
#         match_fields="globalid globalid",
#         update_geometry="UPDATE_GEOMETRY"
#     )



# def addMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict=None,shapeJSON=None):
#     if not shapeJSON:
#         fieldValues, shapeJSON = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
#     else:
#         fieldValues, _ = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
#     # msg('shapeJSON: {}'.format(shapeJSON))
#     msg('{} >>> {}'.format(sde_fields,fieldValues))
#     if survey123_fields:
#         fieldValuesDict = dict(zip(survey123_fields,fieldValues))
#     else:
#         fieldValuesDict = {}
#     gis_globalid = returnDistinctFieldValueForOneRecord(mapLyr,'GLOBALID')
#     fieldValuesDict['gis_globalid'] = gis_globalid # .replace('{','').replace('}','') # added GLOBALID in order to populate gis_globalid in survey123
#     currentDateEpoch = getEpochTime()
#     msg(currentDateEpoch)
#     fieldValuesDict['begin_date'] = currentDateEpoch * 1000
#     if attribute_dict:
#         attributes_dict = {**spatial_analysis_dict, **fieldValuesDict, **attribute_dict} # combined spatial analysis dict with field values dict and any custom attribute dictionary provided
#     else:
#         attributes_dict = {**spatial_analysis_dict, **fieldValuesDict} # combined spatial analysis dict with field values dict
#     mainSurveyLyrName = mainSurveyLyr.properties.name
#     msg('Attributes Dict for {}: {}'.format(mainSurveyLyrName,attributes_dict))
#     geometryDict = getGeometryDictFromShapeJSON(shapeJSON,target_WKID=SURVEY123_WKID)
#     adds_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
#     form_globalid, form_objectid = addFormRecord(mainSurveyLyr,adds_dict_list)
#     return form_globalid # returns the global id of the new survey123 form record

# def updateMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,existingGlobalID,existingObjectid,gis_globalid,attribute_dict=None,shapeJSON=None):
#     # if a shape is provided, use that for the geometry (in cases where updating based on another feature); otherwise, use the geometry from the main record
#     if not shapeJSON:
#         fieldValues, shapeJSON = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
#     else:
#         fieldValues, _ = getFieldValuesAndShapeJSON(mapLyr,sde_fields)
#     if survey123_fields:
#         fieldValuesDict = dict(zip(survey123_fields,fieldValues))
#     else:
#         fieldValuesDict = {}
#     fieldValuesDict['gis_globalid'] = gis_globalid # added GLOBALID in order to populate gis_globalid in survey123
#     if attribute_dict:
#         attributes_dict = {**spatial_analysis_dict, **fieldValuesDict, **attribute_dict} # combined spatial analysis dict with field values dict and any custom attribute dictionary provided
#     else:
#         attributes_dict = {**spatial_analysis_dict, **fieldValuesDict} # combined spatial analysis dict with field values dict
#     attributes_dict['globalId'] = existingGlobalID
#     attributes_dict['objectid'] = existingObjectid
#     msg('Attributes Dict: {}'.format(attributes_dict))
#     geometryDict = getGeometryDictFromShapeJSON(shapeJSON,target_WKID=SURVEY123_WKID)
#     updates_dict_list = [{"geometry":geometryDict,"attributes":attributes_dict}]
#     form_globalid, form_objectid = updateFormRecord(mainSurveyLyr,updates_dict_list)
#     return form_globalid # returns the global id of the new survey123 form record

# def createOrUpdateSurvey123Record(mainSurveyLyr,mapLyr,gis_globalid,itemTitle,form_item_id,feature_layer_item_id,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict=None,agolQuery=None,shapeJSON=None):
#     """
#     Creates a new Survey123 Record or Updates an Existing Records

#     Args:
#         mainSurveyLyr (AGOL Feature Layer Object): [description]
#         mapLyr (ArcGIS Pro Map Layer): [description]
#         gis_globalid (str): [description]
#         itemTitle (str): [description]
#         form_item_id (str): [description]
#         feature_layer_item_id (str): [description]
#         sde_fields (list): [description]
#         survey123_fields (list): [description]
#         spatial_analysis_dict (dict): [description]
#         attribute_dict (dict, optional): [description]. Defaults to None.
#         agolQuery (str, optional): [description]. Defaults to None.

#     Returns:
#         [type]: [description]
#         form_globalid,form_url,new_record_TF
#         (str,str,boolean)
#     """
#     #existingFeatures = getFeaturesFromLyr(mainSurveyLyr,query="gis_globalid = '{}'".format(gis_globalid),outputFields='*')
#     if not agolQuery:
#         agolQuery = "gis_globalid = '{}'".format(gis_globalid)
#     existingFeatures = getFeaturesFromLyr(mainSurveyLyr,query=agolQuery,outputFields=['globalid','objectid','gis_globalid'],returnGeometry=False)
#     if existingFeatures:
#         msg('existingFeatures: {}'.format(existingFeatures))
#         # TODO - ask how they want to handle multiple Survey123 records with the same gis_globalid
#         existingFeature = existingFeatures[0]
#         # existingFeatureAttributes = existingFeature['attributes']
#         # existingGlobalID = existingFeatureAttributes['globalid']
#         # existingObjectid = existingFeatureAttributes['objectid']
#         existingGlobalID = existingFeature.attributes['globalid']
#         existingObjectid = existingFeature.attributes['objectid']
#         existingFeatureCount = len(existingFeatures)
#         if existingFeatureCount > 1:
#             msg('Existing Features Attributes: {}'.format([feature.attributes for feature in existingFeatures]))
#             multipleGlobalIDs = [feature.attributes['globalid'] for feature in existingFeatures]
#             #form_url = generateFormURL(form_item_id,existingGlobalID)
#             #wrn('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,multipleGlobalIDs,existingObjectid,existingGlobalID,form_url))
#             form_url_list = [generateFormURL(form_item_id,globalid) for globalid in multipleGlobalIDs]
#             err('{} Survey123 Record with gis_globalid: {} in {}: {}.\nUpdating first record found: ObjectID: {}, GlobalID: {}:\n{}'.format(existingFeatureCount,gis_globalid,itemTitle,"\n".join(form_url_list)))
#         msg('Survey Record for GIS GlobalID {} already exists in {}: ObjectID: {}, GlobalID: {}'.format(gis_globalid,itemTitle,existingObjectid,existingGlobalID))
#         updateExistingRecord = True # TODO - make parameter to determine if update is needed
#         if updateExistingRecord:
#             form_globalid = updateMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,existingGlobalID,existingObjectid,gis_globalid,attribute_dict,shapeJSON=shapeJSON)
#         else:
#             msg('No change to existing record')
#             form_globalid = existingGlobalID
#         # update GIS dataset based on record in Survey123
#         form_url = generateFormURL(form_item_id,form_globalid)
#         calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
#         new_record_TF = False
#     else:
#         form_globalid = addMainSurveyLyrRecord(mapLyr,mainSurveyLyr,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict,shapeJSON=shapeJSON)
#         msg('Survey Record for GIS GlobalID {} does not yet exist in {}.'.format(gis_globalid,itemTitle))
#         if isinstance(form_globalid,str): # check to see if the resulting globalid is a string
#             form_url = generateFormURL(form_item_id,form_globalid)
#             calcFields(mapLyr,[formItemIDField,formRecordGlobalIDField,featureLayerItemIDField,formURLField],[form_item_id,form_globalid,feature_layer_item_id,form_url])
#             new_record_TF = True
#         else:
#             err('Error creating new Survey123 record: {}'.format(form_globalid))
#             form_url = None
#             form_globalid = None
#             new_record_TF = False

#     return form_globalid,form_url,new_record_TF

# ### Generate Resource Attribute Dict
# sde_fields = [resourceTypeField,otherResourceTypeField,resourceNameField,resourceStateIDField,cemeteryNumberField,legacyObjectIDField,fieldSiteNumberField,resourceIDField]
# survey123_fields = [x.lower() for x in sde_fields] # TODO - this assumes that the field names are the same
# resourceType,otherResourceType,resourceName,resourceStateID,cemeteryNumber,legacyObjectID,fieldSiteNumber,resourceID
# survey123_field_values = [resourceType,otherResourceType,resourceName,resourceStateID,cemeteryNumber,legacyObjectID,fieldSiteNumber,resourceID]
# resourceAttributeDict = dict(zip(survey123_fields,survey123_field_values))

# ### Combine resource attribute dict and activity attribute dict
# attribute_dict = {**resourceAttributeDict, **activityAttributeDict}

# # Set contributing resource Yes/No/Unknown values
# if hasContributingResources and hasContributingResources != 'Keep Existing Value':
#     attribute_dict['has_contributing_resources'] = hasContributingResources
# if hasNonContributingResources and hasNonContributingResources != 'Keep Existing Value':
#     attribute_dict['has_noncontributing_resources'] = hasNonContributingResources
# if contributesToOtherResources and contributesToOtherResources != 'Keep Existing Value':
#     attribute_dict['contributes_to_other_resources'] = contributesToOtherResources



# # try:
# #     arcpy.edit.Generalize(resourceLyr, GENERALIZE_THRESHOLD) # TODO!!! - remove if curve ring conversion is fixed
# # except Exception as e:
# #     wrn('Unable to generalize: {}'.format(e))

# sde_fields,survey123_fields = None,None # clear out prior to running since attributes are all handled through attribute_dict
# form_globalid,form_url,new_record_TF = createOrUpdateSurvey123Record(mainSurveyLyr,resourceLyr,gis_globalid,itemTitle,form_item_id,feature_layer_item_id,sde_fields,survey123_fields,spatial_analysis_dict,attribute_dict)






# addFormRecord(lyr,adds_dict_list)
# updateFormRecord(lyr,updates_dict_list)


# mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_fc)]
# tbl = r'memory\dates'
# fields = [f.name for f in arcpy.ListFields(tbl) if f.name.lower() not in ('objectid')]
# globalIDIndex = fields.index('globalid')
# managed_task_dict = {}
# with arcpy.da.SearchCursor(tbl,fields) as cursor:
#     for row in cursor:
#         managed_task_dict[row[globalIDIndex].upper()] = dict(zip(fields,row))

# updates_dict_list = [{'globalid': '{95c237b7-32fd-411f-8494-f187d20d4508}', 'significant_dates': datetime.datetime(2024, 5, 9, 16, 0), 'significant_date_description': 'Test date added through the web form.', 'parentglobalid': '{68b9562f-964c-4885-8813-e03cde46236c}', 'CreationDate': datetime.datetime(2024, 5, 11, 0, 33, 59, 522000), 'Creator': 'japarkhurst', 'EditDate': datetime.datetime(2024, 5, 11, 0, 33, 59, 522000), 'Editor': 'japarkhurst'}]

# ############################
# with arcpy.da.SearchCursor(mt_resourceLyr,['globalid','gis_globalid']) as cursor:
#     orig_globalid_gis_globalid_dict = {row[0]:row[1] for row in cursor}
#         # append/update to AGOL
#         mainSurveyLyr = returnLyrObjectFromAGOLItem(gis,item_id=itemId,relativeLyrNum=0)
#         survey_lyr_url = mainSurveyLyr.url
#         base_url = os.path.dirname(survey_lyr_url)
#         fc_fields = [f.name for f in arcpy.ListFields(fc)]

#         with arcpy.da.SearchCursor(mt_resourceLyr,field_names=['globalid'],where_clause='1=1') as cursor:
#             valueList = [list(row) for row in cursor]

#         layer_id = 0
#         with arcpy.EnvManager(preserveGlobalIds=False):
#             arcpy.management.Append(
#                 inputs=mt_resourceLyr,
#                 target=base_url + '/' + layer_id, # "https://services.arcgis.com/w8auYAijfGK1Mydj/arcgis/rest/services/service_e5c4c15f908a40ce8c1a2ca9475a2711/FeatureServer/0"
#                 schema_type="TEST",
#                 field_mapping=None,
#                 subtype="",
#                 expression="",
#                 match_fields="globalid globalid",
#                 update_geometry="UPDATE_GEOMETRY"
#             )


#         #

#         globalIDIndex = common_fields.index('globalid')
#         managed_task_dict = {}
#         mt_full_GlobalID_list = []
#         mt_full_gis_globalid_list = []
#         mt_GlobalID_list_with_gis_globalid = []
#         mt_fields = common_fields + ['SHAPE@JSON','gis_globalid']
#         orig_mt_globalid_gis_globalid_dict = {}
#         mt_fields = [f.name for f in arcpy.ListFields(mt_resourceLyr) if f.name.lower() not in ('shape_length', 'shape_area', 'objectid')] + ['SHAPE@JSON']
#         globalID_index = mt_fields.index('globalid')
#         gis_globalID_index = mt_fields.index('gis_globalid')
#         with arcpy.da.SearchCursor(mt_resourceLyr,mt_fields) as cursor:
#             for row in cursor:
#                 gis_globalID = row[gis_globalID_index]
#                 globalID = row[globalID_index]
#                 orig_mt_globalid_gis_globalid_dict[globalID] = gis_globalID
#                 if gis_globalid:
#                     if gis_globalid in mt_full_gis_globalid_list:
#                         wrn('Multiple instances of gis_globalid {}'.format(gis_globalid))
#                     elif '{' in gis_globalid:
#                         managed_task_dict[gis_globalid] = dict(zip(mt_fields,row))
#                         managed_task_dict[globalID] = dict(zip(mt_fields,row))
#                         mt_GlobalID_list_with_gis_globalid.append(globalID)
#                     mt_GlobalID_list_with_gis_globalid.append(globalID)
#                     mt_full_gis_globalid_list.append(gis_globalid)
#                 else:
#                     managed_task_dict[globalID] = dict(zip(mt_fields,row))
#                 mt_full_GlobalID_list.append(globalID)

#         # for globalID,valueDict in managed_task_dict.items():


#         # gis_fields = common_fields + ['SHAPE@JSON','globalid']
#         # msg('gis_fields = {}'.format(gis_fields))

#         # # use gis_globalid field to determine records to select in GIS
#         # mt_gis_globalid_list = list(managed_task_dict.keys())
#         # msg('Length of managed task dict: {}'.format(len(managed_task_dict)))

#         # msg('mt_gis_globalid_list: {}: {}'.format(len(mt_gis_globalid_list),mt_gis_globalid_list))
#         # msg('mt_full_gis_globalid_list: {}: {}'.format(len(mt_full_gis_globalid_list),mt_full_gis_globalid_list))
#         # msg('mt_full_GlobalID_list: {}: {}'.format(len(mt_full_GlobalID_list),mt_full_GlobalID_list))

#         # # globalid_list_for_resource_query = list(set(mt_full_gis_globalid_list + mt_full_GlobalID_list))
#         # globalid_list_for_resource_query = list(set(mt_gis_globalid_list + mt_full_GlobalID_list))
#         # msg('globalid_list_for_resource_query: {}: {}'.format(len(globalid_list_for_resource_query),globalid_list_for_resource_query))

#         # globalid_to_form_record_globalid_dict = {}
#         # if globalid_list_for_resource_query:
#         #     query = "GlobalID in ('{0}')".format("','".join(globalid_list_for_resource_query))
#         #     msg('Query for Resource Layer: {}'.format(query))
#         #     gis_dict = {}
#         #     with arcpy.da.SearchCursor(resourceLyr,gis_fields,where_clause=query) as cursor:
#         #         for row in cursor:
#         #             gis_dict[row[-1]] = dict(zip(gis_fields,row))

#         #     query = "GlobalID in ('{0}') or form_record_globalid in ('{0}')".format("','".join(globalid_list_for_resource_query))
#         #     msg('Query for Resource Layer: {}'.format(query))
#         #     with arcpy.da.SearchCursor(resourceLyr,['globalid','form_record_globalid'],where_clause=query) as cursor:
#         #         for row in cursor:
#         #             globalid_to_form_record_globalid_dict[row[0]] = tryUpper(addBrackets(row[1]))
#         #     gis_GlobalID_list = list(globalid_to_form_record_globalid_dict.keys())

#         # else:
#         #     gis_dict = {}
#         #     gis_GlobalID_list = []

#         # # get full value lists
#         # gis_full_globalid_list = returnDistinctFieldValues(resourceLyr,'globalId',"resource_type = '{}'".format(resourceType))
#         # if not gis_full_globalid_list:
#         #     gis_full_globalid_list = []
#         # gis_full_form_record_globalid_list = returnDistinctFieldValues(resourceLyr,'form_record_globalid',"resource_type = '{}'".format(resourceType))
#         # if gis_full_form_record_globalid_list:
#         #     gis_full_form_record_globalid_list = [tryUpper(addBrackets(x)) for x in gis_full_form_record_globalid_list]
#         # else:
#         #     gis_full_form_record_globalid_list = []
#         # matches_by_gis_globalid = [globalid_to_form_record_globalid_dict[x] for x in mt_full_gis_globalid_list if x in gis_full_globalid_list]
#         # matches_by_form_record_globalid = [x for x in mt_full_GlobalID_list if x in gis_full_form_record_globalid_list]
#         # match_by_both = [x for x in matches_by_gis_globalid if x in matches_by_form_record_globalid]
#         # match_by_only_gis_globalid = [x for x in matches_by_gis_globalid if x not in matches_by_form_record_globalid]
#         # match_by_only_form_record_globalid = [x for x in matches_by_form_record_globalid if x not in matches_by_gis_globalid]
#         # no_match = [x for x in mt_full_GlobalID_list if x not in match_by_both and x not in match_by_only_gis_globalid and x not in match_by_only_form_record_globalid]

#         # #msg('gis_full_globalid_list: {}:{}'.format(len(gis_full_globalid_list),gis_full_globalid_list))
#         # msg('mt_full_gis_globalid_list: {}:{}'.format(len(mt_full_gis_globalid_list),mt_full_gis_globalid_list))
#         # #msg('gis_full_form_record_globalid_list: {}:{}'.format(len(gis_full_form_record_globalid_list),gis_full_form_record_globalid_list))
#         # msg('mt_full_GlobalID_list: {}:{}'.format(len(mt_full_GlobalID_list),mt_full_GlobalID_list))

#         # msg('matches_by_gis_globalid: {}:{}'.format(len(matches_by_gis_globalid),matches_by_gis_globalid))
#         # msg('matches_by_form_record_globalid: {}:{}'.format(len(matches_by_form_record_globalid),matches_by_form_record_globalid))

#         # msg('match_by_both: {}:{}'.format(len(match_by_both),match_by_both))
#         # msg('match_by_only_gis_globalid: {}:{}'.format(len(match_by_only_gis_globalid),match_by_only_gis_globalid))
#         # msg('match_by_only_form_record_globalid: {}:{}'.format(len(match_by_only_form_record_globalid),match_by_only_form_record_globalid))
#         # msg('no_match: {}:{}'.format(len(no_match),no_match))

#         # gis_mt_different_values_dict = compare_GIS_and_MT_Resources(gis_dict,gis_fields,gis_wkid,managed_task_dict,mt_fields,mt_wkid)

#         # plain_msg('\n\ngis_mt_different_values_dict = {}'.format(gis_mt_different_values_dict))
#         # plain_msg('\nglobalid_to_form_record_globalid_dict = {}'.format(globalid_to_form_record_globalid_dict))

#         # form_record_globalid_to_globalid_dict = {v:k for k,v in globalid_to_form_record_globalid_dict.items()}
#         # plain_msg('\nform_record_globalid_to_globalid_dict = {}'.format(form_record_globalid_to_globalid_dict))

#         # form_record_globalid_mt_different_values_dict = {globalid_to_form_record_globalid_dict[globalid]:valueDict for globalid,valueDict in gis_mt_different_values_dict.items()}
#         # plain_msg('\n\nform_record_globalid_mt_different_values_dict = {}'.format(form_record_globalid_mt_different_values_dict))





#         mt_fields = [f.name.lower() for f in arcpy.ListFields(mt_fc)]
#         globalIDIndex = mt_fields.index('globalid')
#         managed_task_dict = {}
#         with arcpy.da.SearchCursor(mt_fc,mt_fields + ['SHAPE@JSON']) as cursor:
#             for row in cursor:
#                 managed_task_dict[row[globalIDIndex].upper()] = dict(zip(mt_fields + ['SHAPE@JSON'],row))

#         mt_globalIDList = [key for key in managed_task_dict.keys()]
#         msg('Managed Task Record Count: {}'.format(len(managed_task_dict)))
#         msg(mt_globalIDList)
#         #msg('Managed Task Dict: {}'.format(managed_task_dict))
#         query = "globalid in ('{}')".format("','".join(mt_globalIDList))
#         lyr = [lyr for lyr in lyrs if lyr.properties.name == mt_fc_name][0]
#         agol_features = getFeaturesFromLyr(lyr, query=query, outputFields='*', returnGeometry=True)
#         msg('AGOL Features Count: {}'.format(len(agol_features)))
#         #msg('AGOL Features: {}'.format(agol_features))
#         agol_globalIDList = [agol_feature.attributes['globalid'] for agol_feature in agol_features]
#         msg(agol_globalIDList)



class UpdateAGOLLists(object):
    def __init__(self):
        '''
            Initializes the tool class
            - canRunInBackground - used only for ArcGIS Desktop; it has no effect in ArcGIS Pro
                ie, True or False
            - category - The name of the toolset in which the tool is located. A toolset is a way to organize tools within a toolbox
            - description - The description for the tool
            - label - display name for the tool as shown in the Geoprocessing pane
            - stylesheet - to change the default style sheet used for the tool
        '''
        self.label = "Update AGOL Lists"
        self.description = "Update AGOL Contractor and Activity Lists"
        self.canRunInBackground = False
        self.category = ''


    def getParameterInfo(self):
        """
            Defines the tool's parameters
        """

        contractor = arcpy.Parameter(
            displayName='Contractor',
            name='contractor', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Optional',
            direction='Input')

        activityName = arcpy.Parameter(
            displayName='Activity Name',
            name='activityName', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Required',
            direction='Input')

        contractor.value = 'Test Contractor Jacob' # Define default for parameter
        activityName.value = 'Test Activity Name Jacob' # Define default for parameter
        params = [contractor,activityName]
        return params

    def isLicensed(self):
        '''Set whether tool is licensed to execute.'''
        return True

    # Sample Validation
    def updateParameters(self, parameters):
        '''Modify the values and properties of parameters before internal
        validation is performed.  This method is called whenever a parameter
        has been changed.'''
        pValues = paramDictValues(parameters)
        pNum = paramDictNum(parameters)
        return

    def updateMessages(self, parameters):
        '''Modify the messages created by internal validation for each tool
        parameter.  This method is called after internal validation.'''
        return

    # Sample parameter call
    def execute(self, parameters, messages):
        '''The source code of the tool.'''
        # User Parameters
        # See GetParameterInfo function for default values
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        contractor = pValues['contractor']
        activityName = pValues['activityName']

        # aprx = arcpy.mp.ArcGISProject("CURRENT")

        from arcgis import GIS
        import pandas
        gis = GIS("PRO")

        updateContractorAndActivityLists(contractor=contractor,activityName=activityName,gis=gis)

        return

class testTool(object):
    def __init__(self):
        '''
            Initializes the tool class
            - canRunInBackground - used only for ArcGIS Desktop; it has no effect in ArcGIS Pro
                ie, True or False
            - category - The name of the toolset in which the tool is located. A toolset is a way to organize tools within a toolbox
            - description - The description for the tool
            - label - display name for the tool as shown in the Geoprocessing pane
            - stylesheet - to change the default style sheet used for the tool
        '''
        self.label = "Test Tool"
        self.description = "This tool does xyz"
        self.canRunInBackground = False
        self.category = ''


    def getParameterInfo(self):
        """
            Defines the tool's parameters
        """

        sampleParam = arcpy.Parameter(
            displayName='Sample Parameter',
            name='sample', ### Name the same as parameter in main script
            datatype='GPString',
            parameterType='Optional',
            direction='Input')

        mapLyrParam = arcpy.Parameter(
            displayName='Map Layer',
            name='mapLyr', ### Name the same as parameter in main script
            datatype='GPFeatureLayer',
            parameterType='Required',
            direction='Input')

        sampleParam.value = 'Default Value' # Define default for parameter
        params = [sampleParam,mapLyrParam]
        return params

    def isLicensed(self):
        '''Set whether tool is licensed to execute.'''
        return True

    # Sample Validation
    def updateParameters(self, parameters):
        '''Modify the values and properties of parameters before internal
        validation is performed.  This method is called whenever a parameter
        has been changed.'''
        pValues = paramDictValues(parameters)
        pNum = paramDictNum(parameters)
        return

    def updateMessages(self, parameters):
        '''Modify the messages created by internal validation for each tool
        parameter.  This method is called after internal validation.'''
        return

    # Sample parameter call
    def execute(self, parameters, messages):
        '''The source code of the tool.'''
        # User Parameters
        # See GetParameterInfo function for default values
        pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
        pValuesAsText = paramDictValuesAsText(parameters)
        sample = pValues['sample']
        mapLyr = pValues['mapLyr']

        catalogPath = arcpy.Describe(mapLyr).catalogPath
        msg('Catalog Path: {}'.format(catalogPath))

        # aprx = arcpy.mp.ArcGISProject("CURRENT")
        # activeMapView = aprx.activeView
        # activeMapView.zoomToAllLayers()

        # calcFields(mapLyr,[otherResourceTypeField],['test'])

        # aprxMap = returnMapObject()
        # for name,url in lyrNameLyrURL_dict.items():
        #     lyr = aprxMap.addDataFromPath(url) # add layer to map
        #     lyrName = lyr.name
        #     msg("'{}':'{}'".format(name,lyrName))
        #     break

        return



# 'C:\\Users\\japarkhu\\AppData\\Local\\Temp\\3\\ArcGISProTemp12860\\Section_106_Polygon_Test_2_report_aeb50a.docx'
# out = 'C:\\Users\\japarkhu\\AppData\\Local\\Temp\\3\\ArcGISProTemp12860\\Section_106_Polygon_Test_2_report_aeb50a.docx'
# print(out)
# C:\Users\japarkhu\AppData\Local\Temp\3\ArcGISProTemp12860\Section_106_Polygon_Test_2_report_aeb50a.docx

# #print(letslearn.reports)
# for x in letslearn.reports: #finds all Microsoft Word doc files in AGO with you as their owner and "Survey 123" in the tags. Could upload a Word document that wasn't a generated report, add that tag, and it would be found with this
#    #print(x.description) remove "#" if you want to see them listed
#    if x.description: #if it exists.. (some of my old generated reports had null descriptions)
#       if 'id for survey form item in AGO' in x.description: #fill in 1 string here
#          id = x.id #get id of that Word doc
#          data_item = gis.content.get(id) #get Word doc as an item from AGO
#          data_item.download(save_path = r'C:\where\ever\you\want\it') #fill in 1 string here; downloads that Word doc
#          data_item.delete() #deletes that Word doc in AGO


# from arcgis.gis import GIS
# from arcgis.apps.survey123._survey import SurveyManager, Survey
# gis = GIS("https://tva.maps.arcgis.com/home",'python_tva','tvagis00',verify_cert=False)
# survey_mgr = SurveyManager(gis)
# letslearn = survey_mgr.get("ead814540c3a406da85e49d978b932dc") #fill in 1 string here
# print(letslearn.report_templates)
# [<Item title:"Section 106 Polygon Test 2_sampleTemplateIndividual" type:Microsoft Word owner:japarkhurst>]
# choice = letslearn.report_templates[0] #the print template you want to use (0 = first); fill in 1 number here
# print(choice)
# <Item title:"Section 106 Polygon Test 2_sampleTemplateIndividual" type:Microsoft Word owner:japarkhurst>

# #print(letslearn.reports)
# for x in letslearn.reports: #finds all Microsoft Word doc files in AGO with you as their owner and "Survey 123" in the tags. Could upload a Word document that wasn't a generated report, add that tag, and it would be found with this
#    #print(x.description) remove "#" if you want to see them listed
#    if x.description: #if it exists.. (some of my old generated reports had null descriptions)
#       if 'id for survey form item in AGO' in x.description: #fill in 1 string here
#          id = x.id #get id of that Word doc
#          data_item = gis.content.get(id) #get Word doc as an item from AGO
#          data_item.download(save_path = r'C:\where\ever\you\want\it') #fill in 1 string here; downloads that Word doc
#          data_item.delete() #deletes that Word doc in AGO

#


### SAMPLE TOOLS ###
### See URL for more info: https://pro.arcgis.com/en/pro-app/arcpy/geoprocessing_and_python/creating-a-new-python-toolbox.htm

# class Toolbox(object):
#     def __init__(self):
#         '''
#             Define the toolbox (the name of the toolbox is the name of the .pyt file).
#         '''
#         self.label = "Toolbox"
#         self.alias = ""

#         # List of tool classes associated with this toolbox
#         toolList = []
#         toolList.append(toolName)
#         self.tools = toolList

# class toolName(object):
#     def __init__(self):
#         '''
#             Initializes the tool class
#             - canRunInBackground - used only for ArcGIS Desktop; it has no effect in ArcGIS Pro
#                 ie, True or False
#             - category - The name of the toolset in which the tool is located. A toolset is a way to organize tools within a toolbox
#             - description - The description for the tool
#             - label - display name for the tool as shown in the Geoprocessing pane
#             - stylesheet - to change the default style sheet used for the tool
#         '''
#         self.label = "Tool Name"
#         self.description = "This tool does xyz"
#         self.canRunInBackground = False
#         self.category = ''

#     # sample parameter definition
#     def getParameterInfo(self):
#         '''
#             Define parameter definitions
#             - displayName - The parameter name as displayed on the tool's dialog box
#                 - ie, "Survey Type"
#             - name - The parameter name as shown in the tool's syntax in Python
#                 - ie, "surveyType"
#             - datatype - The data type being used in the parameter
#                 - ie, see URL: https://pro.arcgis.com/en/pro-app/arcpy/geoprocessing_and_python/defining-parameter-data-types-in-a-python-toolbox.htm
#             - parameterType
#                 - ie, "Required", "Optional", or "Derived"
#             - direction - defines whether the parameter is an input to the tool or an output of the tool
#                 - ie, "Input"
#         '''
#         # Define List of Parameters to be displayed in the tool

#         sampleParam = arcpy.Parameter(
#             displayName='Sample Parameter',
#             name='sample', ### Name the same as parameter in main script
#             datatype='GPString',
#             parameterType='Required',
#             direction='Input')

#         sampleParam.value = 'Default Value' # Define default for parameter
#         params = [sampleParam]
#         return params

#     def isLicensed(self):
#         '''Set whether tool is licensed.'''
#         return True

#     # Sample Validation
#     def updateParameters(self, parameters):
#         '''Modify the values and properties of parameters before internal
#         validation is performed.  This method is called whenever a parameter
#         has been changed.'''
#         pValues = paramDictValues(parameters)
#         pNum = paramDictNum(parameters)
#         return

#     def updateMessages(self, parameters):
#         '''Modify the messages created by internal validation for each tool
#         parameter.  This method is called after internal validation.'''
#         return

#     # Sample parameter call
#     def execute(self, parameters, messages):
#         '''The source code of the tool.'''
#         # User Parameters
#         # See GetParameterInfo function for default values
#         pValues = paramDictValues(parameters) # Generate Dictionary of Parameter Values
#         pValuesAsText = paramDictValuesAsText(parameters)
#         sample = pValues['sample']
#         return
# '''



"""
Compare Fields for resource and finding and effects activity forms
resource_type = 'Archaeological Resource'

arch_res_id = '2f2818068968417aadddb6b017a9e5f2'
arch_res_item = gis.content.get(arch_res_id)
arch_res_layer = arch_res_item.layers[0]
arch_res_fields = [field['name'].lower().replace(' ','_') for field in arch_res_layer.properties.fields]

suffix = RESOURCE_FIELD_SUFFIX_DICT[resource_type.lower().replace(' ','_')]
excludedFields = ['creationdate', 'creator', 'editdate', 'editor']


fe_arch_res_fields = []
for tbl in lyr.tables:
    #print(tbl.properties.name)
    if tbl.properties.name == resource_type.lower().replace(' ','_'):
        fields = tbl.properties.fields
        print(fields)
        for field in fields:
            print(field['name'])
            fieldName = field['name'].lower().replace(' ','_')
            fe_arch_res_fields.append(fieldName)

arch_res_fields = [f for f in arch_res_fields if f not in excludedFields]
fe_arch_res_fields = [f for f in fe_arch_res_fields if f not in excludedFields]

field_map_dict = {}
for field in fe_arch_res_fields:
    if field in arch_res_fields:
        field_map_dict[field] = field
    elif field.endswith(suffix):
        mod_field = field[:field.rfind(suffix)]
        if mod_field in arch_res_fields:
            field_map_dict[mod_field] = field

source_res_fields, target_fe_fields = zip(*field_map_dict.items())
"""





def getFieldPicklistDict(form_xlsx):
    import pandas
    #form_xlsx = r'C:\Users\japarkhu\ArcGIS\My Survey Designs\Objects_21068072150-AS-1_GA_NC_SC v2\Objects_21068072150-AS-1_GA_NC_SC v2.xlsx'

    survey_df = pandas.read_excel(form_xlsx,sheet_name='survey')
    field_to_domain_Dict = {}
    for row in survey_df.values.tolist():
        field_type =  str(row[0])
        field_name = row[1]
        if field_type.startswith('select_multiple ') or field_type.startswith('select_one '):
            domain_name = field_type.split(' ')[1]
            field_to_domain_Dict[field_name] = domain_name

    choices_df = pandas.read_excel(form_xlsx,sheet_name='choices')
    domain_names = list(set(choices_df['list_name'].values.tolist()))
    domain_value_dict_Dict = {x:{} for x in domain_names}
    for row in choices_df.values.tolist():
        domain_name =  row[0]
        name = row[1]
        label = row[2]
        domain_value_dict_Dict[domain_name].update({name:label})

    field_to_values_Dict = {}
    for field,domain_name in field_to_domain_Dict.items():
        domain_values_dict = domain_value_dict_Dict.get(domain_name)
        field_to_values_Dict[field] = domain_values_dict

    return field_to_values_Dict




def msg(txt):
    print(txt)
    arcpy.AddMessage(txt)

def getItemFromTitle(gis,title,item_type="Feature Layer"):
    items = gis.content.search(title,item_type,max_items=1000)
    items = [item for item in items if item.title == title]
    if len(items) != 1:
        wrn(f'Unable to uniquely find "{title}": {items}')
        return
    return items[0]

def getFeatureLayerIDFromFormItem(form_item):
    related_items = form_item.related_items(rel_type='Survey2Service',direction='forward')
    if len(related_items) == 1:
        related_item = related_items[0]
        return related_item.id
    elif len(related_items) > 1:
        wrn('More than 1 related feature layer for form id {}'.format(form_item))
        return None
    else:
        wrn('No related feature layer for form id {}'.format(form_item))
        return None

def getFeaturesFromLyr(lyr,query='1=1',outputFields='*',returnGeometry=True):
    """ Function to return the features (in list of dictionaries) from a feature service layer """
    lyrName = lyr.properties.name

    msg('Query for getting features from {}: {} with {} output fields requested'.format(lyrName,query,outputFields))
    fset = lyr.query(where=query,out_fields=",".join(outputFields),return_geometry=returnGeometry)
    # features_json = fset.features
    # features = json.loads(str(features_json))
    features = fset.features
    if not features:
        msg('No features found in {}'.format(lyrName))
    return features



# # user entered parameter
# contractorAndActivityName = "TVAR : MT  Tool Oct 31 Testing Survey"

# contractor = 'TVAR'
# activityName = 'MT  Tool Oct 31 Testing Survey'

# # formTitle = 'Managed Task Activities'
# # mtFormTitle = formTitle + ' - MT Master'
# # formItem = getItemFromTitle(gis,mtFormTitle,item_type='Form')
# # featureLyrItemId = getFeatureLayerIDFromFormItem(formItem)
# # featureLyrItem = gis.content.get(featureLyrItemId)

# featureLyrTitle = f'Managed Task Activities - {contractor} View'
# featureLyrItem = getItemFromTitle(gis,featureLyrTitle,item_type='Feature Layer')
# lyr = featureLyrItem.layers[0]
# activitiesReadyForReview = getFeaturesFromLyr(lyr=lyr, query="contractor_review_status = 'Ready for TVA Review'", outputFields=['contractor_0', 'activity_name', 'contractor_review_status', 'activity_id'],returnGeometry=False)

# resourceTypes = ['Archaeological Resources', 'Historic Architectural Resources', 'Objects', 'Districts and Facilities', 'Cultural Landscapes', 'Cemeteries', 'Other Resources']

# for resourceType in resourceTypes:
#     title = resourceType + ' - ' + contractor + ' View'
#     print(title)
#     featureLyrItem = getItemFromTitle(gis,title,item_type='Feature Layer')

#     # identify records that have the corresponding activity name and contractor
#     lyr = featureLyrItem.layers[0]
#     query = f"contractor_0 = '{contractor}' and (tva_assigned_activity_name_0 = '{activityName}' or activity_name = '{activityName}')"
#     existingFeatures = getFeaturesFromLyr(lyr=lyr, query=query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0'],returnGeometry=False)


#     from collections import Counter
#     stats = Counter([x.attributes['contractor_review_status'] for x in existingFeatures])
#     print(stats)

#     query = f"contractor_0 = '{contractor}' and tva_assigned_activity_name_0 is null and activity_name is null"
#     featuresWithNoActivityName = getFeaturesFromLyr(lyr=lyr, query=query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0','Creator'],returnGeometry=False)
#     print(featuresWithNoActivityName)
#     break # for testing, only evaluate Archaeologicaly Resources


# # updates = [{'attributes':{'objectid':_.attributes['objectid'],'tva_assigned_activity_name_0':'77556-AS-1'}} for _ in fset.features]
# # lyr.edit_features(updates=updates)


# resourceTypes = ['Archaeological Resources', 'Historic Architectural Resources', 'Objects', 'Districts and Facilities', 'Cultural Landscapes', 'Cemeteries', 'Other Resources']

# for resourceType in resourceTypes:
#     title = resourceType + ' - ' + contractor + ' View'
#     print(title)
#     mt_featureLyrItem = getItemFromTitle(gis,title,item_type='Feature Layer')

#     # identify records that have the corresponding activity name and contractor
#     lyr = mt_featureLyrItem.layers[0]
#     query = f"contractor_0 = '{contractor}' and (tva_assigned_activity_name_0 = '{activityName}' or activity_name = '{activityName}')"
#     existingFeatures = getFeaturesFromLyr(lyr=lyr, query=query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0'],returnGeometry=False)


#     from collections import Counter
#     stats = dict(Counter([x.attributes['contractor_review_status'] for x in existingFeatures]))
#     print(stats)
#     ready = stats['Ready for TVA Review']
#     noStatus = stats[None]
#     inProgress = stats['In Progress']

#     #if noStatus:
#         #f'{len(noStatus)} records have no contractor status set'
#     #if inProgress:
#         #f'{len(inProgress)} records are set to In Progress'}

#     query = f"contractor_0 = '{contractor}' and tva_assigned_activity_name_0 is null and activity_name is null"
#     featuresWithNoActivityName = getFeaturesFromLyr(lyr=lyr, query=query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0','Creator'],returnGeometry=False)
#     print(featuresWithNoActivityName)
#     break # for testing, only evaluate Archaeologicaly Resources

# def msg(txt):
#     print(txt)

# def wrn(txt):
#     print(txt)



# mt_lyrs = mt_featureLyrItem.layers + mt_featureLyrItem.tables
# mt_globalIDList = [_.attributes['globalid'] for _ in existingFeatures]
# mt_lyrGlobalIDDict, mt_lyrFeaturesDict = getLyrGlobalIDandFeaturesDict(mt_lyrs,mt_globalIDList)
# msg('---')
# msg(mt_lyrGlobalIDDict)
# #msg(lyrFeaturesDict)

# #mt_globalIDList = [_.attributes['globalid'] for _ in existingFeatures] # get existingFeatures from contractor/activity input values

# mt_to_agol_globalIDDict = {_.attributes['globalid']:_.attributes['orig_globalid_0'] for _ in mt_lyrFeaturesDict[0]}
# agol_globalIDList = list(_ for _ in mt_to_agol_globalIDDict.values() if _)

# formType = 'Resource'
# formTitle = 'Archaeological_Resources'
# form_item_id, feature_layer_item_id = getFormAndFeatureLayerItemIDs(ITEM_ID_REFERENCE_LIST,itemTitle=formTitle,formType=formType,gis=gis)
# agol_featureLyrItem = gis.content.get(feature_layer_item_id)
# agol_lyrs = agol_featureLyrItem.layers + agol_featureLyrItem.tables


# agol_lyrGlobalIDDict, agol_lyrFeaturesDict = getLyrGlobalIDandFeaturesDict(agol_lyrs,agol_globalIDList)
# msg('---')
# msg(agol_lyrGlobalIDDict)

# """
# 1 hour	3600 seconds
# 1 day	86400 seconds
# 1 week	604800 seconds
# 1 month (30.44 days) 	2629743 seconds
# 1 year (365.24 days) 	 31556926 seconds
# """

# def epochToDatetimeUTC(epochObject):
#     if isinstance(epochObject,(int,float)):
#         if len(str(epochObject)) == 13:
#             return datetime.datetime.fromtimestamp(int(epochObject/1000), datetime.timezone.utc)
#         elif len(str(epochObject)) == 15:
#             return datetime.datetime.fromtimestamp(int(epochObject/100000), datetime.timezone.utc)
#         elif len(str(epochObject)) == 12:
#             return datetime.datetime.fromtimestamp(int(epochObject/100), datetime.timezone.utc)
#         else:
#             return datetime.datetime.fromtimestamp(int(epochObject), datetime.timezone.utc)
#     else:
#         return epochObject

# for lyr_id,mt_feature_list in mt_lyrFeaturesDict.items():
#     #mt_gid_list = [f['attributes']['globalid'] for f in mt_feature_list]
#     mt_gid_dict = {f.attributes['globalid']:f for f in mt_feature_list}
#     agol_features = agol_lyrFeaturesDict.get(lyr_id)
#     agol_gid_dict = {f.attributes['globalid']:f for f in agol_features}

#     excludeFromComparison = ['objectid','globalid','Shape__Length','CreationDate','Creator','EditDate','Editor'] + ['Shape__Area'] # TODO - decide whether we want to compare Shape Area
#     for mt_gid,mt_feature in mt_gid_dict.items():
#         agol_gid = mt_to_agol_globalIDDict.get(mt_gid)
#         if not agol_gid:
#             print(f'\nNew record (globalid {mt_gid=}')
#         else:
#             print(f'\nCommon record (globalid {mt_gid=}/{agol_gid=}')
#             agol_feature = agol_gid_dict.get(agol_gid)
#             agol_values = agol_feature.attributes
#             fields = [field for field in agol_values.keys() if field not in excludeFromComparison]
#             mt_values = mt_feature.attributes

#             for field in fields:
#                 agol_value = agol_values[field]
#                 mt_value = mt_values.get(field)
#                 if not mt_value:
#                     continue # value not in MT export (Sacred Sites, TCP)
#                 if agol_value == mt_value:
#                     #print(f'Same value for {field}:{agol_value}')
#                     pass
#                 else:
#                     mt_date = epochToDatetimeUTC(mt_value)
#                     if mt_date != mt_value:
#                         if abs(agol_value-mt_value) <= 43200000: # if epoch is within 12 hours, ignore as conflict
#                             print('Within 12 hrs')
#                         else:
#                             print(f'{mt_date=}')
#                             agol_date = epochToDatetimeUTC(agol_value)
#                             print(f'{agol_date=}')


#                     print(f'Different value for {field}\n{mt_value=}\n{agol_value=}')

#                 # compare attributes
#                 # compare geometry
#                 mt_geometry = mt_feature.geometry
#                 agol_geometry = agol_feature.geometry
#         break

# agol_dict = mt_to_agol_globalIDDict
# # # - update getFeatureLayerIDFromFormID to return the item, not the ID

# mt_fields = [f for f in arcpy.ListFields(mt_path)]
# mt_field_names = [f.name.lower() for f in mt_fields]
# date_field_names = [f.name.lower() for f in mt_fields if f.type == 'DATE']
# agol_field_names = [f.name.lower() for f in arcpy.ListFields(agol_path)]
# common_fields = [x for x in mt_field_names if x in agol_field_names]


# Common record (globalid mt_gid='ccd2af5a-1b4b-4a1a-9554-c726068881b8'/agol_gid='f056e6eb-64c9-4e24-ac7c-b75b0b5895c4'
# Different value for recorded_date
# mt_value=253400547600000
# agol_value=253400572800000
# Different value for date_of_assessment_
# mt_value=1233334800000
# agol_value=1233291600000
# Different value for has_contributing_resources
# mt_value='No'
# agol_value=None
# Different value for has_noncontributing_resources
# mt_value='No'
# agol_value=None
# Different value for contributes_to_other_resources
# mt_value='No'
# agol_value=None

# # record from AGOL
# Recorded Date	12/12/9999
# Date of Assessment	1/30/2009 5:00:00 AM

# # record from MT
# Recorded Date	12/11/9999, 5:00 PM
# Date of Assessment	1/30/2009, 5:00 PM




# Common record (globalid mt_gid='68c2d3c2-8563-49b3-813b-547848a4b35b'/agol_gid='01757e56-09d5-4f03-b33a-7d3a6351ec36'
# Different value for recorded_date
# mt_value=253400547600000
# agol_value=253400572800000
# Different value for location_reliability
# mt_value='Accurate'
# agol_value=None
# Different value for date_of_assessment_
# mt_value=-30610249200000
# agol_value=-30610224000000

# Common record (globalid mt_gid='6dac22c9-cc7a-469a-bcc1-95fef8922256'/agol_gid='0e4e0fb7-22cd-4d59-979a-e5f90c4639da'
# Different value for recorded_date
# mt_value=253400547600000
# agol_value=253400572800000
# Different value for date_of_assessment_
# mt_value=1233248400000
# agol_value=1233205200000

# Common record (globalid mt_gid='00cbde78-a076-4976-b97e-b68fc7bb9f1a'/agol_gid='6569d768-b885-4c85-80d9-5256e8bd4a47'

# Common record (globalid mt_gid='7590316f-218c-4a86-8a23-1b974c2f5069'/agol_gid='6ec5c550-05aa-4fec-ac8e-7ef17f805c46'

# Common record (globalid mt_gid='ccd2af5a-1b4b-4a1a-9554-c726068881b8'/agol_gid='f056e6eb-64c9-4e24-ac7c-b75b0b5895c4'
# Different value for recorded_date
# mt_value=253400547600000
# agol_value=253400572800000
# Different value for date_of_assessment_
# mt_value=1233334800000
# agol_value=1233291600000
# Different value for has_contributing_resources
# mt_value='No'
# agol_value=None
# Different value for has_noncontributing_resources
# mt_value='No'
# agol_value=None
# Different value for contributes_to_other_resources
# mt_value='No'
# agol_value=None

# New record (globalid mt_gid='3e176920-4c00-4781-b591-05353e7025bf'

# New record (globalid mt_gid='4da2d38f-70f8-4da8-bc17-3ebf4ecf468c'

# test comment on MT GlobalID: {C131CD9C-F6C3-4442-B99B-C8F2AB8D1C18}
# original AGOL GlobalID: c83613b9-bc38-4516-85a7-5a20a1572c4e


# mt_globalIDList = ['68c2d3c2-8563-49b3-813b-547848a4b35b',
#  '6dac22c9-cc7a-469a-bcc1-95fef8922256',
#  '00cbde78-a076-4976-b97e-b68fc7bb9f1a',
#  '7590316f-218c-4a86-8a23-1b974c2f5069',
#  'ccd2af5a-1b4b-4a1a-9554-c726068881b8',
#  '3e176920-4c00-4781-b591-05353e7025bf',
#  '4da2d38f-70f8-4da8-bc17-3ebf4ecf468c'] + ['c131cd9c-f6c3-4442-b99b-c8f2ab8d1c18']


# https://survey123.arcgis.com/surveys/dc46decbb01246199ff62dea3bbf58c2/analyze?position=0.activity_name&filter=(0.activity_name%20is%20%22MT%20%20Tool%20Oct%2031%20Testing%20Survey%22%20OR%200.tva_assigned_activity_name_0%20is%20%22MT%20%20Tool%20Oct%2031%20Testing%20Survey%22)&navigation=open:true&hideFields=0:contractor_0,resource_id,res_state_id,resource_name,cemetery_number,legacy_object_id,field_site_number,resource_type,other_resource_type,activity_id,activity_type,project_id,record_status,state_name,county,isolated_find,ownership,resource_description,full_resource_description,bank,physiography,topographic_situation,topographic_situation_other,soil_series,property_type,property_name,recorded_in_the_field_by,recorded_date,shoreline_condition_conducted,elevation_ft_amsl,slope,aspect,length_meters,width_meters,depth_meters,height_meters,distance_to_water_meters,location_reliability,archaeological_field_methodolog,management_recommendation_,management_recommend_other,repository_storage_location,boundary_delineation_method,site_type,cultural_affiliation,period_of_historic_occupation,features_present,current_land_use,type_of_disturbance,digitization_method,date_of_assessment_,when_will_project_be_ready,stabilization_funded,approximate_linear_ft,bank_height_protected_ft,heritage_review_started,corps_permit,arpa_damage_type,presence_of_human_remains_sc,archaeological_protection_needs,archaeological_protection_feasi,presence_of_human_remains_apnr,eligibility_status_apnr,archaeological_features_present,arpa_damage_severity,erosion_severity,presence_of_cutbank_toe,t_and_e_species_present,condition_needed_for_installati,accessibility,partnership_potential,bank_height_feet,comments_archaeological_protect,nrhp_aspects_of_integrity,nrhp_recommendation,nrhp_determination,nrhp_listing_date,nrhp_concurrence,applicable_nrhp_criteria,nrhp_criteria_consideration,areas_of_significance,areas_of_significance_other,periods_of_significance,statement_of_significance,nrhp_boundary_justification,has_contributing_resources,has_noncontributing_resources,contributes_to_other_resources,comment,gis_globalid,orig_globalid_0,tva_assigned_activity_name_0;1:contractor_1,significant_dates,significant_date_description,orig_globalid_1,orig_parentglobalid_1,tva_assigned_activity_name_1;2:contractor_2,resource_id_cr,res_state_id_cr,resource_name_cr,cemetery_number_cr,legacy_object_id_cr,field_site_number_cr,resource_type_cr,other_resource_type_cr,orig_globalid_2,orig_parentglobalid_2,tva_assigned_activity_name_2;3:contractor_3,resource_id_ncr,res_state_id_ncr,resource_name_ncr,cemetery_number_ncr,legacy_object_id_ncr,field_site_number_ncr,resource_type_ncr,other_resource_type_ncr,orig_globalid_3,orig_parentglobalid_3,tva_assigned_activity_name_3;4:contractor_4,resource_id_rct,res_state_id_rct,resource_name_rct,cemetery_number_rct,legacy_object_id_rct,field_site_number_rct,resource_type_rct,other_resource_type_rct,description_of_resources_contri,orig_globalid_4,orig_parentglobalid_4,tva_assigned_activity_name_4

# inputs = [('ENVIRONMENT.GIS.ICD_ARCH_SITE_DATA','SHORE_DT_OF_ASSESMENT'),
# ('ENVIRONMENT.GIS.ICD_ARCH_SITE_DATA','RECORDER_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','CONTRACTOR_DOCUMENT_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','DOCUMENT_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','CONTRACTOR_END_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','SURV_END_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','CONTRACTOR_START_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY_FIELD_RECON','SURV_BEGIN_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_SURVEY','CREATED_DT'),
# ('ENVIRONMENT.GIS.ICD_SITE_MONITORING','ASSESSMENT_DT')]

# def saveSheets(dfDict, xlsxFile, index=False, header=True):
#     """
#     Save a dictionary of dataframes to an excel file, with each dataframe as a separate page
#     dfDict is a dictionary where the sheet name is the key and the dataframe is the value
#     # must be run in python 3.x
#     """
#     from pandas import ExcelWriter
#     import warnings
#     warnings.simplefilter(action='ignore', category=UserWarning)
#     writer = ExcelWriter(xlsxFile)
#     for sheet in dfDict:
#         dfDict[sheet].to_excel(writer, sheet, index=index, header=header)
#     writer.close()

# outputXLSX = r'\\tva\egis\EGIS-Projects\Environment\Cultural\team_Cultural_Compliance\ALIS_Inventory\Date_Value_Validation_20241114.xlsx'
# dfDict = {}
# for tblName,field in inputs:
#     out_table=fr"memory\{tblName.split('.')[-1]}_frequency_{field}"
#     print(out_table)
#     arcpy.analysis.Frequency(
#     in_table=os.path.join(r'\\tva\egis\EGIS-ARCGISSERVER\connections\Production_PUB\UDXPGISPUB@ENVIRONMENT@OSA.sde',tblName),
#     out_table=out_table,
#     frequency_fields=field,
#     summary_fields=None
#     )
#     with arcpy.da.SearchCursor(out_table,['FREQUENCY',field]) as cursor:
#         rows = [row for row in cursor]
#     df = pandas.DataFrame(rows)
#     df.columns = ['FREQUENCY',field]
#     dfDict[os.path.basename(out_table)] = df
#     print(f'Processed {tblName}, {field}')

# saveSheets(dfDict, outputXLSX, index=False, header=True)

# for dfName,df in dfDict.items():
#     df.to_clipboard(index=False)
#     print('Copied to clipboard')
#     time.sleep(5)



















"""
def compare_AGOL_and_MT_Resources(agol_features,managed_task_dict,mt_fields,sync_gis_fields):

def msg(txt):
    print(txt)

def wrn(txt):
    print(txt)

def getLyrGlobalIDandFeaturesDict(lyrs,globalIDList,formType='Resource'):
    globalIDQuery = "globalid IN ('{}')".format("','".join(globalIDList))
    parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(globalIDList))
    lyrGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
    lyrMissingGlobalIDDict = {lyr.properties.id:[] for lyr in lyrs}
    lyrFeaturesDict = {lyr.properties.id:[] for lyr in lyrs}

    outputFields = "*"
    returnGeometry = True
    if globalIDList:
        full_globalid_list = []
        lyrsToReprocess = [] # variable to hold the table layers that are not directly related to the polygon feature (i.e. related to a related table)
        maxNon0Rel = 0 # variable to hold the number of relationships that are not directly with the polygon feature
        for i,lyr in enumerate(lyrs):
            lyrId = lyr.properties.id
            if i == 0:
                agol_features = getFeaturesFromLyr(lyr=lyr, query=globalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                #msg(agol_features)
                gid_list = [x.attributes['globalid'] for x in agol_features]
                msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
                if formType == 'Activity' and not gid_list:
                    err(f'No ArcGIS Online record found for Activity ID {activityID}.')
                    exit()

                if len(globalIDList) != len(gid_list):
                    missing = [x for x in globalIDList if x.lower().strip('{').strip('}') not in gid_list]
                    err(f'{len(missing)} not found in ArcGIS Online: {missing}')
                    msg(f'{globalIDList=}')
                    msg(f'{gid_list=}')
                    lyrMissingGlobalIDDict[lyrId] += missing
                    #if formType == 'Resource':
                        #exit()
                full_globalid_list += gid_list
                lyrGlobalIDDict[lyrId] += gid_list
                lyrFeaturesDict[lyrId] += agol_features
                continue
            relationships = lyr.properties['relationships']
            non0Rel = [x for x in relationships if x['relatedTableId'] != 0] # get list of layers that have relationships other than with the main polygon features
            if non0Rel:
                lyrsToReprocess.append(lyr)

            agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
            gid_list = [x.attributes['globalid'] for x in agol_features]
            msg(f'{len(gid_list)} found in layer {lyrId} - {lyr.properties.name}')
            full_globalid_list += gid_list
            lyrGlobalIDDict[lyrId] += gid_list
            lyrFeaturesDict[lyrId] += agol_features

            if lyrsToReprocess:
                found_new_TF = True
                pass_count = 1
                while found_new_TF:
                    pass_count+=1
                    msg(f'Making Pass {pass_count} for related records...')
                    found_new_TF = False
                    for lyr in lyrsToReprocess:
                        lyrId = lyr.properties.id
                        newParentGlobalIDList = []
                        newParentGlobalIDList += full_globalid_list
                        parentGlobalIDQuery = "parentglobalid IN ('{}')".format("','".join(newParentGlobalIDList))
                        agol_features = getFeaturesFromLyr(lyr=lyr, query=parentGlobalIDQuery, outputFields=outputFields, returnGeometry=returnGeometry)
                        gid_list = [x.attributes['globalid'] for x in agol_features]
                        new_gid_list = [x for x in gid_list if x not in full_globalid_list]
                        if not new_gid_list:
                            continue
                        msg(f'{len(new_gid_list)} found during pass {pass_count} on layer {lyrId} - {lyr.properties.name}')
                        full_globalid_list += new_gid_list
                        lyrGlobalIDDict[lyrId] += new_gid_list
                        found_new_TF = True
                        lyrFeaturesDict[lyrId] += agol_features

    return lyrGlobalIDDict, lyrFeaturesDict, lyrMissingGlobalIDDict


mt_master_feature_layer_items = gis.content.search('MT Master',item_type='Feature Layer Collection',max_items=100)

contractor = 'TVAR'
activityName = '21068072150-AS-1_GA NC SC'
formType = 'Resource'
resourceTypes = ['Archaeological Resources', 'Historic Architectural Resources', 'Objects', 'Districts and Facilities', 'Cultural Landscapes', 'Cemeteries', 'Other Resources']
ready = True
for resourceType in resourceTypes:
    contractorViewTitle = resourceType + ' - ' + contractor + ' View'
    print(contractorViewTitle)
    mt_featureLyrItem = getItemFromTitle(gis,contractorViewTitle,item_type='Feature Layer')
    break

mt_lyrs = mt_featureLyrItem.layers + mt_featureLyrItem.tables


# identify records that have the corresponding activity name and contractor
mt_lyr = mt_featureLyrItem.layers[0]

mt_query = f"contractor_0 = '{contractor}' and (tva_assigned_activity_name_0 = '{activityName}' or mt_assigned_activity_name_0 = '{activityName}' or activity_name = '{activityName}')"
mt_features_to_process = getFeaturesFromLyr(lyr=mt_lyr, query=mt_query, outputFields=['objectid', 'globalid', 'contractor_0', 'activity_name', 'contractor_review_status','gis_globalid', 'orig_globalid_0','tva_assigned_activity_name_0', 'tva_assigned_activity_id_0', 'mt_assigned_activity_name_0'],returnGeometry=False)

mt_globalIDList = [x.attributes['globalid'] for x in mt_features_to_process]

mt_lyrGlobalIDDict, mt_lyrFeaturesDict, mt_lyrMissingGlobalIDDict = getLyrGlobalIDandFeaturesDict(mt_lyrs,mt_globalIDList,formType)
newMTFeatures = [x for x in mt_lyrFeaturesDict[0] if not x.attributes['orig_globalid_0']]
oldMTFeatures = [x for x in mt_lyrFeaturesDict[0] if x.attributes['orig_globalid_0']]
oldMT_agolGlobalIDList = [x.attributes['orig_globalid_0'] for x in oldMTFeatures]

agolLayerTitle = resourceType
print(agolLayerTitle)
agol_featureLyrItem = getItemFromTitle(gis,agolLayerTitle,item_type='Feature Layer')
agol_lyrs = agol_featureLyrItem.layers + agol_featureLyrItem.tables
agol_query = "globalid in ('{}')".format("','".join(oldMT_agolGlobalIDList))
agol_lyrGlobalIDDict, agol_lyrFeaturesDict, agol_lyrMissingGlobalIDDict = getLyrGlobalIDandFeaturesDict(agol_lyrs,oldMT_agolGlobalIDList,formType)

for lyrID,mt_features in mt_lyrFeaturesDict.items():
    print(f'\n\n{lyrID}')
    if not mt_features:
        print(f'No MT features in {lyrID}')
        continue
    agol_features = agol_lyrFeaturesDict[lyrID]
    if not agol_features:
        print(f'No AGOL features in {lyrID}')
        continue

    agol_fields = [field for field in agol_features[0].attributes.keys()]
    mt_fields = [field for field in mt_features[0].attributes.keys()]
    common_fields = [field for field in mt_fields if field.lower() in agol_fields and field.lower() not in ('objectid','creationdate','creator','editdate','editor','globalid')]


    agol_different_values_dict = {}
    changed_GlobalIDs = []
    unchanged_GlobalIDs = []
    sync_changed_GlobalIDs = []
    for mt_feature in mt_features:
        msg('---')
        different_values = []
        agol_globalid = mt_feature.attributes[f'orig_globalid_{lyrID}']
        mt_globalid = mt_feature.attributes[f'globalid']

        if not agol_globalid:
            msg(f'No AGOL feature expected for MT {mt_globalid}')
            continue

        #mt_features_selection = [x for x in mt_features if x.attributes[f'orig_globalid_{lyrID}'] == agol_globalid]

        agol_features_selection = [x for x in agol_features if x.attributes['globalid'] == agol_globalid]
        if not agol_features_selection:
            wrn(f'No AGOL feature for {agol_globalid}')
            continue
        if len(agol_features_selection) > 1:
            print(f'More than 1 AGOL feature for {agol_globalid}')
            continue

        agol_feature = agol_features_selection[0]

        #print(agol_feature.geometry)
        #print(mt_feature.attributes)
        if lyrID == 0:
            agol_geo = agol_feature.geometry
            mt_geo = mt_feature.geometry
            if agol_geo != mt_geo:
                msg('Different Geometry for {}'.format(agol_globalid))
                sync_changed_GlobalIDs.append(agol_globalid)
            else:
                msg('Same Geometry for {}'.format(agol_globalid))            
            # if agol_geo:
            #     agol_geo_rings = agol_geo['rings']
            # agol_geo_rings = agol_geo['rings'] if agol_geo else [[]]
            # mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

            # if agol_geo != mt_geo:
            #     msg('Different Geometry for {}'.format(agol_globalid))
            # else:
            #     msg('Same Geometry for {}'.format(agol_globalid))
            # mt_geo_json = managed_task_dict['{' + agol_globalid + '}']['SHAPE@JSON']
            # mt_geo = json.loads(mt_geo_json) if mt_geo_json else None
            # # if mt_geo:
            # #     mt_geo_rings = mt_geo['rings']
            # mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

            # agol_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in agol_geo_rings for coord in ring]
            # mt_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in mt_geo_rings for coord in ring]

            # if sorted(agol_coord_lists) != sorted(mt_coord_lists):
            #     msg('Different Geometry for {}'.format(agol_globalid))
            #     msg(f'{agol_coord_lists=}\n{mt_coord_lists=}')
            #     different_values += [('','','Geometry','','')]
            #     if sync_gis_fields:
            #         sync_changed_GlobalIDs.append(agol_globalid)
            # else:
            #     msg('Same Geometry for {}'.format(agol_globalid))

        if 'resource_id' in common_fields:
            agol_resourceID = agol_feature.attributes['resource_id']
            mt_resourceID = mt_feature.attributes['resource_id']
        else:
            agol_resourceID = ''

        if 'res_state_id' in common_fields:
            agol_res_state_id = agol_feature.attributes['res_state_id']
            mt_res_state_id = mt_feature.attributes['res_state_id']
        else:
            agol_res_state_id = ''
        # print(f'{agol_resourceID},{mt_resourceID}')
        # print(f'{agol_res_state_id},{mt_res_state_id}')

        #different_attribute_values = [(agol_resourceID,agol_res_state_id,field,agol_feature.attributes[field],managed_task_dict['{' + agol_globalid + '}'][field]) for field in common_fields if tryUpper(removeBrackets(agol_feature.attributes[field])) != tryUpper(removeBrackets(datetimeUTCToEpochMilli(managed_task_dict['{' + agol_globalid + '}'][field])))] # must attempt to convert datetime to epoch in order to compare feature class to AGOL values; must convert globalids to uppercase; must remove '{' and '}' for parentglobalid
        different_attribute_values = [(agol_resourceID,agol_res_state_id,field,agol_feature.attributes[field],mt_feature.attributes[field]) for field in common_fields if agol_feature.attributes[field] != mt_feature.attributes[field]]

        if different_attribute_values:
            msg('Different Attributes for {}'.format(agol_globalid))
        else:
            msg('Same Attributes for {}'.format(agol_globalid))

        different_values += different_attribute_values

        if different_values:
            msg('Different Values for {}:\nField\tAGOL\tManaged Task\n{}'.format(agol_globalid,"\n".join(['{}\t{}\t{}\t{}\t{}'.format(resourceID,resStateID,field,agol_value,mt_value) for resourceID,resStateID,field,agol_value,mt_value in different_values])))
            agol_different_values_dict[agol_globalid] = different_values
            changed_GlobalIDs.append(agol_globalid)
            # determine if field differences apply when syncing with GIS
            if sync_gis_fields:
                sync_value_changes = [x for x in different_values if x[2] in sync_gis_fields]
                if sync_value_changes:
                    sync_changed_GlobalIDs.append(agol_globalid)
        else:
            unchanged_GlobalIDs.append(agol_globalid)
            msg('Same Values for {}'.format(agol_globalid))

    sync_changed_GlobalIDs = list(set(sync_changed_GlobalIDs))


agol_different_values_dict = {addBrackets(k):v for k,v in agol_different_values_dict.items()}


sync_gis_fields = ['resource_id', 'res_state_id', 'resource_name', 'cemetery_number', 'legacy_object_id', 'field_site_number', 'resource_type', 'other_resource_type', 'activity_id', 'activity_type', 'project_id']














query = "globalid in ('{}')".format("','".join(mt_globalIDList))
lyr = [lyr for lyr in lyrs if lyr.properties.name == mt_fc_name][0]
agol_features = getFeaturesFromLyr(lyr, query=query, outputFields='*', returnGeometry=True)


agol_different_values_dict = {}
changed_GlobalIDs = []
unchanged_GlobalIDs = []
sync_changed_GlobalIDs = []
for agol_feature in agol_features:
    msg('---')
    different_values = []
    agol_globalid = agol_feature.attributes['globalid'].upper() # convert to upper in order to compare with MT and GIS

    common_fields = [field for field in agol_feature.attributes.keys() if field.lower() in mt_fields and field.lower() not in ('objectid','creationdate','creator','editdate','editor','globalid')]

    if 'shape' in mt_fields:
        agol_geo = agol_feature.geometry
        # if agol_geo:
        #     agol_geo_rings = agol_geo['rings']
        agol_geo_rings = agol_geo['rings'] if agol_geo else [[]]
        mt_geo_json = managed_task_dict['{' + agol_globalid + '}']['SHAPE@JSON']
        mt_geo = json.loads(mt_geo_json) if mt_geo_json else None
        # if mt_geo:
        #     mt_geo_rings = mt_geo['rings']
        mt_geo_rings = mt_geo['rings'] if mt_geo else [[]]

        agol_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in agol_geo_rings for coord in ring]
        mt_coord_lists = [[round(coord[0],COORDINATE_ROUNDING),round(coord[1],COORDINATE_ROUNDING)] for ring in mt_geo_rings for coord in ring]

        if sorted(agol_coord_lists) != sorted(mt_coord_lists):
            msg('Different Geometry for {}'.format(agol_globalid))
            msg(f'{agol_coord_lists=}\n{mt_coord_lists=}')
            different_values += [('','','Geometry','','')]
            if sync_gis_fields:
                sync_changed_GlobalIDs.append(agol_globalid)
        else:
            msg('Same Geometry for {}'.format(agol_globalid))

    if 'resource_id' in common_fields:
        agol_resourceID = agol_feature.attributes['resource_id']
        if agol_resourceID:
            agol_resourceID = agol_resourceID.upper()
    else:
        agol_resourceID = ''

    if 'res_state_id' in common_fields:
        agol_res_state_id = agol_feature.attributes['res_state_id']
    else:
        agol_res_state_id = ''

    different_attribute_values = [(agol_resourceID,agol_res_state_id,field,agol_feature.attributes[field],managed_task_dict['{' + agol_globalid + '}'][field]) for field in common_fields if tryUpper(removeBrackets(agol_feature.attributes[field])) != tryUpper(removeBrackets(datetimeUTCToEpochMilli(managed_task_dict['{' + agol_globalid + '}'][field])))] # must attempt to convert datetime to epoch in order to compare feature class to AGOL values; must convert globalids to uppercase; must remove '{' and '}' for parentglobalid

    if different_attribute_values:
        msg('Different Attributes for {}'.format(agol_globalid))
    else:
        msg('Same Attributes for {}'.format(agol_globalid))

    different_values += different_attribute_values

    if different_values:
        msg('Different Values for {}:\nField\tAGOL\tManaged Task\n{}'.format(agol_globalid,"\n".join(['{}\t{}\t{}\t{}\t{}'.format(resourceID,resStateID,field,agol_value,mt_value) for resourceID,resStateID,field,agol_value,mt_value in different_values])))
        agol_different_values_dict[agol_globalid] = different_values
        changed_GlobalIDs.append(agol_globalid)
        # determine if field differences apply when syncing with GIS
        if sync_gis_fields:
            sync_value_changes = [x for x in different_values if x[2] in sync_gis_fields]
            if sync_value_changes:
                sync_changed_GlobalIDs.append(agol_globalid)
    else:
        unchanged_GlobalIDs.append(agol_globalid)
        msg('Same Values for {}'.format(agol_globalid))

sync_changed_GlobalIDs = list(set(sync_changed_GlobalIDs))
agol_different_values_dict = {addBrackets(k):v for k,v in agol_different_values_dict.items()}
"""
